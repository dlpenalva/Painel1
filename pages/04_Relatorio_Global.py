from io import BytesIO
import re
from pathlib import Path
from datetime import datetime
from zoneinfo import ZoneInfo

import pandas as pd
import streamlit as st


st.set_page_config(page_title="Análises de Reajustes - Relatório Global", layout="wide")




# >>> UX_ADITIVOS_25_COMPACTO
def aplicar_css_aditivos25_compacto():
    st.markdown(
        """
        <style>
        div[data-testid="stMetric"] {
            min-height: 72px;
            padding: 8px 10px;
        }
        div[data-testid="stMetricValue"] {
            font-size: clamp(0.95rem, 1.55vw, 1.28rem) !important;
            line-height: 1.12 !important;
            white-space: normal !important;
            overflow-wrap: anywhere !important;
            word-break: normal !important;
        }
        div[data-testid="stMetricLabel"] p {
            font-size: clamp(0.70rem, 1.00vw, 0.86rem) !important;
            line-height: 1.15 !important;
            white-space: normal !important;
        }
        .aditivos25-ux-note {
            font-size: 0.86rem;
            color: #475569;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
# <<< UX_ADITIVOS_25_COMPACTO
from _ui_utils import render_marca_topo, render_aviso_privacidade

def moeda(valor):
    try:
        valor = round(float(valor), 2)
    except Exception:
        valor = 0.0
    return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def texto_seguro(valor, padrao="Não"):
    if valor is None:
        return padrao
    try:
        if pd.isna(valor):
            return padrao
    except Exception:
        pass
    texto = str(valor).strip()
    if texto.lower() in ["", "nan", "none", "nat", "<na>"]:
        return padrao
    return texto


def moeda_ou_texto(valor):
    """Formata números como moeda e preserva textos explicativos."""
    if isinstance(valor, str):
        texto = valor.strip()
        if texto and not any(ch.isdigit() for ch in texto):
            return texto
        if texto and any(ch.isdigit() for ch in texto):
            # Permite números em formato brasileiro, preservando textos mistos não numéricos.
            limpo = texto.replace("R$", "").replace(".", "").replace(",", ".").strip()
            try:
                return moeda(float(limpo))
            except Exception:
                return texto
    try:
        return moeda(float(valor))
    except Exception:
        return "" if pd.isna(valor) else str(valor)


def percentual(valor, casas=2):
    try:
        valor = float(valor)
    except Exception:
        valor = 0.0
    return f"{valor * 100:.{casas}f}%".replace(".", ",")


def fator_fmt(valor):
    try:
        valor = float(valor)
    except Exception:
        valor = 1.0
    return f"{valor:.4f}".replace(".", ",")



def _normalizar_modo_apuracao(valor):
    texto = texto_seguro(valor, "")
    texto = texto.lower()
    mapa = str.maketrans("áàâãéêíóôõúç", "aaaaeeiooouc")
    texto = texto.translate(mapa)
    return re.sub(r"\s+", " ", texto).strip()


def eh_modo_reduzido_itens(res):
    if not isinstance(res, dict):
        return False
    modo = _normalizar_modo_apuracao(res.get("modo_apuracao", ""))
    return "reduzido" in modo and ("item" in modo or "estoque" in modo)


def eh_modo_consumo_itens_ciclo(res):
    if not isinstance(res, dict):
        return False
    modo = _normalizar_modo_apuracao(res.get("modo_apuracao", ""))
    return "consumo" in modo and ("item" in modo or "iten" in modo) and "ciclo" in modo


def valor_retroativo_consumo_itens_ciclo(res):
    if not isinstance(res, dict):
        return 0.0
    for chave in ["valor_retroativo_consumo_itens_ciclo", "retroativo_consumo_itens_ciclo"]:
        if chave in res:
            try:
                return float(res.get(chave) or 0)
            except Exception:
                return 0.0
    return valor_retroativo_estimado_itens(res)


def df_retroativo_consumo_itens_ciclo(res):
    if not isinstance(res, dict):
        return pd.DataFrame()
    df = res.get("df_retroativo_itemizado_por_ciclo")
    if isinstance(df, pd.DataFrame):
        return df
    return pd.DataFrame()


def valor_retroativo_estimado_itens(res):
    if not isinstance(res, dict):
        return 0.0
    for chave in [
        "valor_retroativo_estimado_itens_estoque",
        "retroativo_estimado_itens_estoque",
        "valor_retroativo_itens_estoque",
        "retroativo_itens_estoque",
    ]:
        if chave in res:
            try:
                return float(res.get(chave) or 0)
            except Exception:
                return 0.0
    return 0.0


def df_retroativo_estimado_itens(res):
    if not isinstance(res, dict):
        return pd.DataFrame()
    df = res.get("df_retroativo_estimado_itens_estoque")
    if isinstance(df, pd.DataFrame):
        return df
    return pd.DataFrame()


def indicadores_executivos_relatorio(res):
    if eh_modo_consumo_itens_ciclo(res):
        return [
            ["Indicador", "Valor"],
            ["Modo de apuração", texto_seguro(res.get("modo_apuracao"), "Consumo por Itens/Ciclo")],
            ["Valor original", moeda(res.get("valor_original_contrato", 0))],
            ["Valor Total Atualizado do Contrato", moeda(_m20_componentes_vta_documentos(res).get("total", res.get("valor_atualizado_contrato", res.get("valor_global_estoque", 0))))],
            ["Retroativo financeiro definitivo", "Não calculado"],
            ["Retroativo (itens consumidos/ciclo)", moeda(valor_retroativo_consumo_itens_ciclo(res))],
            ["Execução atualizada por itens/ciclo", moeda(_m20_componentes_vta_documentos(res).get("execucao", 0))],
            ["Saldo Remanescente Atualizado", moeda(_m20_componentes_vta_documentos(res).get("remanescente", 0))],
            ["Base mensal por competência", "Não informada"],
        ]
    if eh_modo_reduzido_itens(res):
        return [
            ["Indicador", "Valor"],
            ["Modo de apuração", texto_seguro(res.get("modo_apuracao"), "Reduzido por Itens/Estoque")],
            ["Valor original", moeda(res.get("valor_original_contrato", 0))],
            ["Valor Total Atualizado do Contrato", moeda(_m20_componentes_vta_documentos(res).get("total", res.get("valor_atualizado_contrato", res.get("valor_global_estoque", 0))))],
            ["Retroativo financeiro definitivo", "Não calculado"],
            ["Retroativo estimado por itens/estoque", moeda(valor_retroativo_estimado_itens(res))],
            ["Execução estimada por itens/estoque", moeda(_m20_componentes_vta_documentos(res).get("execucao", 0))],
            ["Saldo remanescente atualizado", moeda(_m20_componentes_vta_documentos(res).get("remanescente", 0))],
            ["Aditivos/supressões registrados", moeda(_m20_componentes_vta_documentos(res).get("aditivos", 0))],
            ["Base mensal por competência", "Não informada"],
        ]
    return [
        ["Indicador", "Valor"],
        ["Valor original", moeda(res.get("valor_original_contrato", 0))],
        ["Valor pago efetivo", moeda(res.get("total_pago_faturado", 0))],
        ["Valor teórico calculado", moeda(res.get("total_devido_reajustado", 0))],
        ["Valor represado a pagar", moeda(res.get("valor_represado_a_pagar", 0))],
        ["Valor executado atualizado por ciclos", moeda(_m20_componentes_vta_documentos(res).get("execucao", 0))],
        ["Saldo remanescente atualizado", moeda(_m20_componentes_vta_documentos(res).get("remanescente", 0))],
        ["Valor total de aditivos/supressões", moeda(_m20_componentes_vta_documentos(res).get("aditivos", 0))],
        ["Valor Total Atualizado do Contrato", moeda(_m20_componentes_vta_documentos(res).get("total", res.get("valor_atualizado_contrato", res.get("valor_global_estoque", 0))))],
        ["Metodologia de corte", metodologia_corte_operacional_info(res)["titulo"].replace("Metodologia aplicada: ", "")],
    ]



def _formatar_data_corte_relatorio(valor):
    try:
        if valor is None:
            return "Não informada"
        texto = str(valor).strip()
        if texto == "" or texto.lower() in ["nan", "none", "nat"]:
            return "Não informada"
        data = pd.to_datetime(valor, dayfirst=True, errors="coerce")
        if pd.isna(data):
            return texto
        return data.strftime("%d/%m/%Y")
    except Exception:
        return str(valor or "Não informada")


def _config_corte_relatorio(res):
    if not isinstance(res, dict):
        return {}, False
    config = res.get("config_ciclo_em_execucao", {}) or {}
    aplicado = bool(res.get("corte_operacional_solicitado", False) or res.get("corte_operacional_aplicado", False))
    return config, aplicado


def texto_metodologia_corte_relatorio(res):
    config, aplicado = _config_corte_relatorio(res)
    if not aplicado:
        return (
            "Metodologia aplicada: corte padrão no início dos ciclos",
            "A aba CICLO_EM_EXECUCAO está ausente, vazia ou marcada como Não. "
            "O Valor Total Atualizado permanece calculado pela metodologia padrão: execução atualizada por ciclo + saldo remanescente atualizado."
        )

    ciclo = texto_seguro(config.get("ciclo", ""), "Não informado")
    data_corte = _formatar_data_corte_relatorio(config.get("data_corte", ""))
    fonte = texto_seguro(config.get("fonte", ""), "Não informada")
    rem_atualizado = config.get("valor_remanescente_atualizado_corte", config.get("remanescente_atualizado_corte", ""))
    rem_txt = moeda_ou_texto(rem_atualizado) if str(rem_atualizado).strip() else "Não informado"

    return (
        "Metodologia aplicada: corte operacional no ciclo em execução",
        f"Foi aplicado corte operacional no {ciclo}, com data de corte em {data_corte}. "
        f"A execução realizada foi apurada pela fonte preferencial: {fonte}. "
        f"O saldo remanescente atualizado considerado no corte foi: {rem_txt}."
    )


def render_metodologia_corte_relatorio(res):
    titulo, texto = texto_metodologia_corte_relatorio(res)
    _, aplicado = _config_corte_relatorio(res)
    if aplicado:
        html = (
            '<div style="background:#CCFBF1; border:1px solid #14B8A6; border-left:7px solid #0F766E; border-radius:12px; padding:14px 16px; margin:10px 0 16px 0; color:#134E4A;">'
            f'<div style="font-weight:900; margin-bottom:5px;">{titulo}</div>'
            f'<div style="font-size:0.94rem; line-height:1.45;">{texto}</div>'
            '</div>'
        )
    else:
        html = (
            '<div style="background:#EFF6FF; border:1px solid #93C5FD; border-left:7px solid #2563EB; border-radius:12px; padding:14px 16px; margin:10px 0 16px 0; color:#1E3A8A;">'
            f'<div style="font-weight:900; margin-bottom:5px;">{titulo}</div>'
            f'<div style="font-size:0.94rem; line-height:1.45;">{texto}</div>'
            '</div>'
        )
    st.markdown(html, unsafe_allow_html=True)

def aviso_modo_reduzido_html():
    return """
    <div class="modo-reduzido-box">
        <div class="modo-reduzido-titulo">Modo Reduzido por Itens/Estoque</div>
        <div class="modo-reduzido-texto">
            A análise foi processada sem base mensal por competência. O retroativo financeiro definitivo não é calculado neste modo.
            O valor exibido como retroativo estimado por itens/estoque possui natureza estimativa e deve ser validado antes de qualquer formalização de pagamento.
        </div>
    </div>
    """


def aviso_modo_consumo_html():
    return """
    <div style="background:#F6F3EE; border:1px solid #7A8F63; border-left:6px solid #4E6E58; border-radius:12px; padding:14px 16px; margin:10px 0 16px 0; color:#2F3E2F;">
        <div style="font-weight:800; margin-bottom:4px;">Modo Consumo por Itens/Ciclo</div>
        <div style="font-size:0.95rem; line-height:1.45;">A análise foi processada por consumo itemizado por ciclo, sem base mensal por competência. O retroativo financeiro mensal definitivo não é calculado neste modo.</div>
    </div>
    """



def metodologia_corte_operacional_info(res):
    """Retorna dados de metodologia do corte operacional para o Relatório Global."""
    if not isinstance(res, dict):
        return {
            "aplicado": False,
            "titulo": "Metodologia aplicada: corte padrão no início dos ciclos",
            "texto": "A aba CICLO_EM_EXECUCAO está ausente, vazia ou marcada como Não. O Valor Total Atualizado permanece calculado pela metodologia padrão.",
            "cor": "azul",
        }

    config = res.get("config_ciclo_em_execucao", {}) or {}
    solicitado = bool(res.get("corte_operacional_solicitado", False) or res.get("corte_operacional_aplicado", False) or config.get("aplicar", False))

    if not solicitado:
        return {
            "aplicado": False,
            "titulo": "Metodologia aplicada: corte padrão no início dos ciclos",
            "texto": "A aba CICLO_EM_EXECUCAO está ausente, vazia ou marcada como Não. O Valor Total Atualizado permanece calculado pela metodologia padrão: execução atualizada por ciclo + saldo remanescente atualizado.",
            "cor": "azul",
        }

    ciclo = texto_seguro(config.get("ciclo"), "não informado")
    data_corte = texto_seguro(config.get("data_corte"), "não informada")
    fonte = texto_seguro(config.get("fonte"), "não informada")
    rem_atualizado = config.get("valor_remanescente_atualizado_corte", "")
    rem_original = config.get("valor_remanescente_original_corte", "")

    detalhes = [
        f"Ciclo em execução: {ciclo}.",
        f"Data de corte operacional: {data_corte}.",
        f"Fonte preferencial da execução realizada: {fonte}.",
    ]

    if str(rem_original).strip():
        detalhes.append(f"Remanescente original informado no corte: {moeda_ou_texto(rem_original)}.")
    if str(rem_atualizado).strip():
        detalhes.append(f"Remanescente atualizado informado no corte: {moeda_ou_texto(rem_atualizado)}.")

    detalhes.append("Metodologia: execução realizada priorizando a base financeira até o corte + saldo remanescente informado no corte operacional.")

    return {
        "aplicado": True,
        "titulo": "Metodologia aplicada: corte operacional no ciclo em execução",
        "texto": " ".join(detalhes),
        "cor": "teal",
    }


def aviso_metodologia_corte_operacional_html(res):
    info = metodologia_corte_operacional_info(res)
    if info.get("aplicado"):
        return f"""
        <div style="background:#CCFBF1; border:1px solid #14B8A6; border-left:8px solid #0F766E; border-radius:14px; padding:16px 20px; margin:10px 0 16px 0; color:#134E4A;">
            <div style="font-weight:900; font-size:1.00rem; margin-bottom:6px;">{info["titulo"]}</div>
            <div style="font-size:0.94rem; line-height:1.45;">{info["texto"]}</div>
        </div>
        """

    return f"""
    <div style="background:#EAF2F8; border:1px solid #93C5FD; border-left:8px solid #2563EB; border-radius:14px; padding:16px 20px; margin:10px 0 16px 0; color:#1E3A8A;">
        <div style="font-weight:900; font-size:1.00rem; margin-bottom:6px;">{info["titulo"]}</div>
        <div style="font-size:0.94rem; line-height:1.45;">{info["texto"]}</div>
    </div>
    """


def formatar_data_corte_br_relatorio(valor):
    if valor is None:
        return "Não informada"
    texto = str(valor).strip()
    if texto.lower() in ["", "nan", "none", "nat", "<na>"]:
        return "Não informada"
    try:
        data = pd.to_datetime(valor, dayfirst=True, errors="coerce")
        if not pd.isna(data):
            return data.strftime("%d/%m/%Y")
    except Exception:
        pass
    return texto


def _config_corte_relatorio(res):
    if not isinstance(res, dict):
        return {}
    return res.get("config_ciclo_em_execucao", {}) or {}


def _corte_operacional_aplicado_relatorio(res):
    if not isinstance(res, dict):
        return False
    return bool(res.get("corte_operacional_aplicado", False) or res.get("corte_operacional_solicitado", False))


def metodologia_corte_operacional_texto(res):
    config = _config_corte_relatorio(res)
    if _corte_operacional_aplicado_relatorio(res):
        ciclo = texto_seguro(config.get("ciclo", ""), "Não informado")
        data_corte = formatar_data_corte_br_relatorio(config.get("data_corte", ""))
        fonte = texto_seguro(config.get("fonte", ""), "Não informada")
        rem_atualizado = config.get("valor_remanescente_atualizado_corte", config.get("remanescente_atualizado_corte", ""))
        rem_txt = moeda_ou_texto(rem_atualizado) if str(rem_atualizado).strip() else "não informado"
        return (
            "Metodologia aplicada: corte operacional no ciclo em execução. "
            f"Ciclo em execução: {ciclo}. Data de corte: {data_corte}. "
            f"Fonte da execução realizada: {fonte}. "
            f"Saldo remanescente atualizado informado no corte: {rem_txt}. "
            "Composição: execução atualizada até o corte + saldo remanescente informado no corte operacional."
        )
    return (
        "Metodologia aplicada: corte padrão no início dos ciclos. "
        "A aba CICLO_EM_EXECUCAO está ausente, vazia ou marcada como Não. "
        "O Valor Total Atualizado permanece calculado pela metodologia padrão: execução atualizada por ciclo + saldo remanescente atualizado."
    )


def aviso_metodologia_corte_relatorio_html(res):
    config = _config_corte_relatorio(res)
    if _corte_operacional_aplicado_relatorio(res):
        ciclo = texto_seguro(config.get("ciclo", ""), "Não informado")
        data_corte = formatar_data_corte_br_relatorio(config.get("data_corte", ""))
        fonte = texto_seguro(config.get("fonte", ""), "Não informada")
        rem_original = config.get("valor_remanescente_original_corte", config.get("remanescente_original_corte", ""))
        rem_atualizado = config.get("valor_remanescente_atualizado_corte", config.get("remanescente_atualizado_corte", ""))
        return (
            '<div style="background:#CCFBF1; border:1px solid #14B8A6; border-left:7px solid #0F766E; border-radius:12px; padding:14px 16px; margin:8px 0 16px 0; color:#134E4A;">'
            '<div style="font-weight:900; font-size:0.98rem; margin-bottom:6px;">Metodologia aplicada: corte operacional no ciclo em execução</div>'
            '<div style="font-size:0.92rem; line-height:1.45;">'
            f'<b>Ciclo em execução:</b> {ciclo}<br>'
            f'<b>Data de corte:</b> {data_corte}<br>'
            f'<b>Fonte da execução realizada:</b> {fonte}<br>'
            f'<b>Remanescente original no corte:</b> {moeda_ou_texto(rem_original) if str(rem_original).strip() else "Não informado"}<br>'
            f'<b>Remanescente atualizado no corte:</b> {moeda_ou_texto(rem_atualizado) if str(rem_atualizado).strip() else "Não informado"}'
            '</div>'
            '<div style="font-size:0.86rem; margin-top:8px; color:#0F766E;">Composição: execução atualizada até o corte + saldo remanescente informado no corte operacional.</div>'
            '</div>'
        )

    return (
        '<div style="background:#EFF6FF; border:1px solid #93C5FD; border-left:7px solid #2563EB; border-radius:12px; padding:13px 16px; margin:8px 0 16px 0; color:#1E3A8A;">'
        '<div style="font-weight:900; font-size:0.98rem; margin-bottom:5px;">Metodologia aplicada: corte padrão no início dos ciclos</div>'
        '<div style="font-size:0.92rem; line-height:1.45;">A aba CICLO_EM_EXECUCAO está ausente, vazia ou marcada como Não. O Valor Total Atualizado permanece calculado pela metodologia padrão.</div>'
        '</div>'
    )


# ============================================================
# Metodologia do corte operacional - relatório/PDF/apostila
# ============================================================
def _numero_br_relatorio(valor, padrao=0.0):
    try:
        if valor is None:
            return padrao
        if isinstance(valor, str):
            txt = valor.strip().replace("R$", "").replace(" ", "")
            if not txt:
                return padrao
            if "," in txt:
                txt = txt.replace(".", "").replace(",", ".")
            return float(txt)
        return float(valor)
    except Exception:
        return padrao


def _info_corte_operacional_relatorio(res):
    if not isinstance(res, dict):
        return {"ativo": False}
    cfg = res.get("config_ciclo_em_execucao") or {}
    ativo = bool(res.get("corte_operacional_aplicado") or res.get("corte_operacional_solicitado") or cfg.get("aplicar"))
    ciclo = texto_seguro(cfg.get("ciclo", ""), "Não informado")
    competencia = texto_seguro(cfg.get("competencia_corte") or cfg.get("data_corte") or "", "Não informada")
    fonte = texto_seguro(cfg.get("fonte", ""), "Não informada")
    c0_val = _numero_br_relatorio(cfg.get("valor_c0_manual", 0), 0.0)
    rem_atual = _numero_br_relatorio(cfg.get("valor_remanescente_atualizado_corte", 0), 0.0)
    rem_orig = _numero_br_relatorio(cfg.get("valor_remanescente_original_corte", 0), 0.0)
    rem_final = rem_atual if abs(rem_atual) > 0.004 else _numero_br_relatorio(res.get("remanescente_reajustado", 0), 0.0)

    exec_corte = 0.0
    df_exec = res.get("df_execucao_atualizada")
    if isinstance(df_exec, pd.DataFrame) and not df_exec.empty and "Ciclo" in df_exec.columns:
        ciclo_alvo = str(ciclo or "").upper().strip()
        if ciclo_alvo and ciclo_alvo != "NÃO INFORMADO":
            mask = df_exec["Ciclo"].astype(str).str.upper().str.strip().eq(ciclo_alvo)
            if mask.any() and "Valor executado atualizado" in df_exec.columns:
                exec_corte = df_exec.loc[mask, "Valor executado atualizado"].apply(_numero_br_relatorio).sum()
        if abs(exec_corte) <= 0.004 and "Valor executado atualizado" in df_exec.columns:
            exec_corte = df_exec["Valor executado atualizado"].apply(_numero_br_relatorio).sum()

    return {
        "ativo": ativo,
        "ciclo": ciclo,
        "competencia": competencia,
        "fonte": fonte,
        "c0_val": c0_val,
        "tem_c0": abs(c0_val) > 0.004,
        "remanescente_atualizado": rem_final,
        "remanescente_original": rem_orig,
        "execucao_corte": exec_corte,
        "valor_total": _numero_br_relatorio(res.get("valor_atualizado_contrato", res.get("valor_global_estoque", 0)), 0.0),
    }


def aviso_metodologia_corte_html(res):
    info = _info_corte_operacional_relatorio(res)
    if not info.get("ativo"):
        return """
        <div style="background:#EAF2F8; border:1px solid #93C5FD; border-left:6px solid #2563EB; border-radius:12px; padding:14px 16px; margin:10px 0 16px 0; color:#0B1F3A;">
            <div style="font-weight:800; margin-bottom:4px;">Metodologia aplicada: corte padrão no início dos ciclos</div>
            <div style="font-size:0.93rem; line-height:1.45;">O Valor Total Atualizado do Contrato foi composto pela execução atualizada por ciclo somada ao saldo remanescente atualizado do ciclo de referência.</div>
        </div>
        """
    c0 = f"<br><strong>C0 financeiro manual:</strong> {moeda(info['c0_val'])}" if info.get("tem_c0") else ""
    return f"""
    <div style="background:#CCFBF1; border:1px solid #14B8A6; border-left:6px solid #0F766E; border-radius:12px; padding:14px 16px; margin:10px 0 16px 0; color:#134E4A;">
        <div style="font-weight:800; margin-bottom:4px;">Metodologia aplicada: corte operacional no ciclo em execução</div>
        <div style="font-size:0.93rem; line-height:1.45;">
            <strong>Ciclo em execução:</strong> {info['ciclo']}<br>
            <strong>Competência de corte:</strong> {info['competencia']}<br>
            <strong>Fonte da execução realizada:</strong> {info['fonte']}{c0}<br>
            <strong>Execução atualizada do ciclo até o corte:</strong> {moeda(info['execucao_corte'])}<br>
            <strong>Remanescente atualizado informado:</strong> {moeda(info['remanescente_atualizado'])}<br>
            <strong>Valor Total Atualizado:</strong> {moeda(info['valor_total'])}
        </div>
    </div>
    """


def dados_metodologia_corte_tabela(res):
    info = _info_corte_operacional_relatorio(res)
    if not info.get("ativo"):
        return [
            ["Item", "Informação"],
            ["Metodologia aplicada", "Corte padrão no início dos ciclos"],
            ["Composição", "Execução atualizada por ciclo + saldo remanescente atualizado"],
            ["Valor Total Atualizado", moeda(info.get("valor_total", 0))],
        ]
    linhas = [
        ["Item", "Informação"],
        ["Metodologia aplicada", "Corte operacional no ciclo em execução"],
        ["Ciclo em execução", info.get("ciclo", "Não informado")],
        ["Competência de corte", info.get("competencia", "Não informada")],
        ["Fonte da execução realizada", info.get("fonte", "Não informada")],
        ["Execução atualizada do ciclo até o corte", moeda(info.get("execucao_corte", 0))],
        ["Remanescente atualizado informado", moeda(info.get("remanescente_atualizado", 0))],
        ["Valor Total Atualizado", moeda(info.get("valor_total", 0))],
    ]
    if info.get("tem_c0"):
        linhas.insert(5, ["C0 financeiro manual", moeda(info.get("c0_val", 0))])
    return linhas


def texto_metodologia_corte_minuta(res):
    info = _info_corte_operacional_relatorio(res)
    if not info.get("ativo"):
        return (
            "Metodologia de consolidação: foi adotado o corte padrão no início dos ciclos, "
            "com composição do Valor Total Atualizado pela execução atualizada por ciclo somada ao saldo remanescente atualizado."
        )
    partes = [
        f"Metodologia de consolidação: foi adotado corte operacional no ciclo em execução {info.get('ciclo', 'não informado')}, com competência de corte {info.get('competencia', 'não informada')}.",
        f"A execução atualizada do ciclo até o corte corresponde a {moeda(info.get('execucao_corte', 0))}.",
        f"O remanescente atualizado informado para a consolidação corresponde a {moeda(info.get('remanescente_atualizado', 0))}.",
    ]
    if info.get("tem_c0"):
        partes.insert(1, f"Foi utilizado C0 financeiro manual no valor de {moeda(info.get('c0_val', 0))}.")
    return " ".join(partes)

def formatar_data_br(valor):
    data = pd.to_datetime(valor, dayfirst=True, errors="coerce")
    if pd.isna(data):
        return "" if pd.isna(valor) else str(valor)
    return data.strftime("%d/%m/%Y")


def normalizar_status(status):
    texto = str(status or "").upper()
    if "PRECLUS" in texto:
        return "▲ PRECLUSO"
    if "RESSALVA" in texto:
        return "● ADMISSÍVEL COM RESSALVA"
    if "TEMPEST" in texto:
        return "■ TEMPESTIVO"
    if "ADIANT" in texto:
        return "▲ ADIANTADO"
    return texto or "NÃO INFORMADO"


def _status_relatorio(valor):
    """Status textual, sem emojis e sem caixa alta, para relatórios executivos."""
    texto = texto_seguro(valor, "")
    texto = re.sub(r"^[▲■●•\-\s]+", "", str(texto)).strip()
    if not texto:
        return "não informado"
    t = texto.upper()
    if "NEGOC" in t or "ACORDO" in t:
        return "aceito por negociação entre as partes"
    if "PRECLUS" in t:
        return "precluso"
    if "RESSALVA" in t:
        return "admissível com ressalva"
    if "TEMPEST" in t:
        return "tempestivo"
    if "ADIANT" in t:
        return "adiantado"
    return texto.lower()


def _ciclos_df_para_relatorio(adm, res):
    """Monta quadro de ciclos para relatório, incluindo C0 como ciclo-base."""
    df = None
    if isinstance(res, dict) and isinstance(res.get("df_ciclos"), pd.DataFrame) and not res.get("df_ciclos").empty:
        df = res.get("df_ciclos").copy()
    else:
        ciclos = []
        if adm:
            ciclos = adm.get("ciclos") or adm.get("detalhamento_ciclos") or []
        df = pd.DataFrame(ciclos)

    linhas = [{
        "Ciclo": "C0",
        "Data-base": "ciclo-base inicial",
        "Data do pedido": "não se aplica",
        "Classificação": "base sem reajuste",
        "Percentual aplicado": "0,00%",
        "Observação": "Período inicial do contrato, sem aplicação de reajuste.",
    }]

    # Fallback: se o arquivo foi enviado diretamente no módulo Valores e df_ciclos não veio completo,
    # monta ao menos a síntese dos ciclos a partir do delta/retroativo por ciclo.
    if (not isinstance(df, pd.DataFrame) or df.empty) and isinstance(res, dict):
        df_alt = res.get("df_delta_por_ciclo")
        if isinstance(df_alt, pd.DataFrame) and not df_alt.empty and "Ciclo" in df_alt.columns:
            df = df_alt.copy()

    if isinstance(df, pd.DataFrame) and not df.empty:
        ciclos_ja = {"C0"}
        for _, row in df.iterrows():
            ciclo = texto_seguro(row.get("Ciclo", row.get("ciclo", "")), "")
            ciclo_up = str(ciclo).strip().upper()
            if not ciclo_up or ciclo_up in ["C0", "TOTAL", "CICLO"] or ciclo_up in ciclos_ja:
                continue
            ciclos_ja.add(ciclo_up)
            acordo = texto_seguro(row.get("Acordo negocial", row.get("acordo_negocial", "")), "")
            situacao = row.get("Situação aplicada", row.get("situacao_aplicada", row.get("Situação", row.get("situacao", ""))))
            if str(acordo).strip().lower() in ["sim", "s", "true", "1"]:
                classificacao = "aceito por negociação entre as partes"
            else:
                classificacao = _status_relatorio(situacao)

            percentual_val = row.get("Percentual aplicado", row.get("percentual_aplicado", row.get("Variação", row.get("variacao", None))))
            if (percentual_val is None or str(percentual_val).strip() == "") and "Fator" in row.index:
                try:
                    percentual_val = float(row.get("Fator") or 1.0) - 1
                except Exception:
                    percentual_val = 0
            try:
                percentual_txt = percentual(float(percentual_val), 2)
            except Exception:
                percentual_txt = texto_seguro(percentual_val, "0,00%")

            obs = []
            if classificacao == "aceito por negociação entre as partes":
                obs.append("Tratamento negocial informado na análise.")
            justificativa = texto_seguro(row.get("Justificativa negocial", row.get("justificativa_negocial", "")), "")
            if justificativa:
                obs.append(justificativa)
            obs_extra = texto_seguro(row.get("Observação", row.get("Observacao", "")), "")
            if obs_extra and obs_extra not in obs:
                obs.append(obs_extra)

            linhas.append({
                "Ciclo": ciclo_up,
                "Data-base": formatar_data_br(row.get("Data-base", row.get("data_base", ""))),
                "Data do pedido": formatar_data_br(row.get("Data do pedido", row.get("data_pedido", ""))),
                "Classificação": classificacao,
                "Percentual aplicado": percentual_txt,
                "Observação": " ".join(obs) if obs else "",
            })

    return pd.DataFrame(linhas)


def texto_clausula_oito(adm):
    """Síntese enxuta, sem afirmar tempestividade quando houver tratamento negocial."""
    ciclos = []
    if adm:
        ciclos = adm.get("ciclos") or adm.get("detalhamento_ciclos") or []

    if not ciclos:
        return (
            "A análise deve ser lida a partir da Cláusula Oitava, separando a existência do ciclo, "
            "a admissibilidade do pedido e o início dos efeitos financeiros."
        )

    qtd = len(ciclos)
    plural = "pleitos" if qtd > 1 else "pleito"
    status_gerais = []
    tem_negocial = False
    for c in ciclos:
        status_gerais.append(str(c.get("situacao") or c.get("Situação") or c.get("status") or c.get("situacao_aplicada") or "").upper())
        if c.get("superacao_negocial") or str(c.get("Acordo negocial", "")).strip().lower() in ["sim", "s", "true", "1"]:
            tem_negocial = True
    texto_status = " ".join(status_gerais)

    partes = [
        f"A análise contempla {qtd} {plural} de reajuste, devendo distinguir ciclo, admissibilidade e efeitos financeiros."
    ]
    if "PRECLUS" in texto_status:
        partes.append("Há ciclo classificado como precluso, mantido para memória e rastreabilidade, sem gerar retroativo a pagar.")
    if tem_negocial or "NEGOC" in texto_status or "ACORDO" in texto_status:
        partes.append("Há ciclo aceito por negociação entre as partes, razão pela qual o relatório não deve descrevê-lo como simples pleito tempestivo.")
    if "RESSALVA" in texto_status:
        partes.append("Há pleito admissível com ressalva, com efeitos financeiros a serem conferidos conforme a cláusula contratual.")
    if not any(["PRECLUS" in texto_status, tem_negocial, "NEGOC" in texto_status, "ACORDO" in texto_status, "RESSALVA" in texto_status]):
        partes.append("Nos ciclos classificados como tempestivos, os efeitos financeiros devem observar a data juridicamente aplicável ao pedido e a cláusula de reajuste.")
    return " ".join(partes)


def texto_contexto_contrato(res):
    contexto = (res or {}).get("contexto_contratual_anterior", {}) or {}
    if not contexto or not contexto.get("contexto_informado"):
        return "Não houve contexto contratual anterior informado para esta análise."
    linhas = []
    if contexto.get("valor_formalizado_anterior"):
        linhas.append(f"Valor formalizado antes desta análise: {moeda(contexto.get('valor_formalizado_anterior', 0))}.")
    if contexto.get("ultimo_ciclo_concedido"):
        linhas.append(f"Último ciclo já concedido/formalizado: {contexto.get('ultimo_ciclo_concedido')}.")
    if texto_seguro(contexto.get("observacao_historico", ""), ""):
        linhas.append(f"Observação: {texto_seguro(contexto.get('observacao_historico'), 'Não')}.")
    eventos = contexto.get("eventos_historicos_anteriores", []) or []
    if eventos:
        linhas.append(f"Eventos históricos anteriores registrados: {len(eventos)}.")
    linhas.append("O contexto informado serve para memória e governança. Seus efeitos só impactam os valores quando estiverem refletidos na execução, nos itens ou no saldo remanescente do Arquivo de Coleta.")
    return "\n".join(linhas)


def texto_contexto_analise(adm, res):
    indice = res.get("indice", (adm or {}).get("indice", "Não informado"))
    fator = res.get("fator_acumulado", (adm or {}).get("fator_acumulado", (adm or {}).get("fator", 1.0)))
    df_ciclos = _ciclos_df_para_relatorio(adm, res)
    ciclos_validos = df_ciclos[df_ciclos["Ciclo"].astype(str).str.upper() != "C0"] if not df_ciclos.empty else pd.DataFrame()
    qtd_ciclos = len(ciclos_validos)
    classificacoes = []
    if not ciclos_validos.empty:
        for _, row in ciclos_validos.iterrows():
            ciclo = row.get("Ciclo", "")
            classificacao = row.get("Classificação", "")
            pct_apl = row.get("Percentual aplicado", "")
            trecho = f"{ciclo}: {classificacao}"
            if pct_apl:
                trecho += f", percentual aplicado {pct_apl}"
            classificacoes.append(trecho)
    resumo_ciclos = "; ".join(classificacoes) if classificacoes else "sem ciclos de reajuste informados"

    contexto_txt = texto_contexto_contrato(res)
    if eh_modo_consumo_itens_ciclo(res):
        modo_txt = (
            "A análise foi processada em Modo Consumo por Itens/Ciclo, sem base mensal por competência. "
            "Nesse cenário, o retroativo financeiro mensal definitivo não é calculado; apresenta-se o Retroativo (itens consumidos/ciclo), com base nos quantitativos consumidos informados pela fiscalização. "
        )
    elif eh_modo_reduzido_itens(res):
        modo_txt = (
            "A análise foi processada em Modo Reduzido por Itens/Estoque, sem base mensal por competência. "
            "Nesse cenário, o retroativo financeiro definitivo não é calculado; apresenta-se apenas o retroativo estimado por itens/estoque, com natureza estimativa. "
        )
    else:
        modo_txt = "O bloco financeiro serve para apurar o valor represado/retroativo. "
    return (
        "A análise consolida a admissibilidade dos pleitos de reajuste e a quantificação financeira apurada com base no Arquivo de Coleta. "
        "O C0 é tratado como ciclo-base inicial, sem aplicação de reajuste. "
        f"Foram considerados {qtd_ciclos} ciclo(s) de reajuste: {resumo_ciclos}. "
        f"Índice utilizado: {indice}. Fator acumulado considerado: {fator_fmt(fator)}. "
        "O Valor Total Atualizado do Contrato mantém a composição execução atualizada por ciclo + saldo remanescente atualizado + aditivos/supressões computáveis, quando aplicáveis. "
        f"{modo_txt}"
        "Aditivos e supressões computáveis são considerados no Valor Total Atualizado quando não estiverem refletidos na execução, no saldo remanescente ou no valor formalizado anterior. "
        f"{contexto_txt}"
    )


def gerar_texto_instrucao(adm, res):
    if eh_modo_consumo_itens_ciclo(res):
        df_ri = df_retroativo_consumo_itens_ciclo(res)
        tabela_ri = "Sem tabela detalhada disponível."
        if isinstance(df_ri, pd.DataFrame) and not df_ri.empty:
            tabela_ri = df_visual(
                df_ri,
                moeda_cols=["Valor original consumido", "Valor atualizado consumido", "Retroativo"],
                fator_cols=["Fator acumulado"],
            ).to_string(index=False)

        return f"""
RELATÓRIO EXECUTIVO — VALOR ATUALIZADO DO CONTRATO

1. Síntese dos ciclos

{_ciclos_df_para_relatorio(adm, res).to_string(index=False)}

2. Modo de apuração

A análise foi processada em Modo Consumo por Itens/Ciclo, sem base mensal por competência. A fiscalização informou os quantitativos consumidos por item e por ciclo, com premissa de equivalência entre consumo, medição/aprovação e faturamento devido.

Retroativo financeiro definitivo: Não calculado
Retroativo (itens consumidos/ciclo): {moeda(valor_retroativo_consumo_itens_ciclo(res))}
Execução atualizada por itens/ciclo: {moeda(res.get('valor_executado_atualizado', 0))}
Saldo Remanescente Atualizado: {moeda(res.get('remanescente_reajustado', 0))}

3. Retroativo por itens consumidos/ciclo

{tabela_ri}

4. Consolidação contratual

Valor original do contrato: {moeda(res.get('valor_original_contrato', 0))}
Valor Total Atualizado do Contrato: {moeda(res.get('valor_atualizado_contrato', res.get('valor_global_estoque', 0)))}
Aditivos/supressões registrados para controle: {moeda(res.get('total_aditivos_atualizados', 0))}

Data/hora de geração: {datetime.now(ZoneInfo('America/Sao_Paulo')).strftime('%d/%m/%Y %H:%M')}
""".strip()

    if eh_modo_reduzido_itens(res):
        df_ri = df_retroativo_estimado_itens(res)
        tabela_ri = "Sem tabela detalhada disponível."
        if isinstance(df_ri, pd.DataFrame) and not df_ri.empty:
            tabela_ri = df_visual(
                df_ri,
                moeda_cols=["Valor executado original", "Valor executado atualizado", "Retroativo estimado por itens/estoque"],
            ).to_string(index=False)

        return f"""
RELATÓRIO EXECUTIVO — VALOR ATUALIZADO DO CONTRATO

1. Síntese dos ciclos

{_ciclos_df_para_relatorio(adm, res).to_string(index=False)}

2. Modo de apuração

A análise foi processada em Modo Reduzido por Itens/Estoque, sem base mensal por competência. Por essa razão, o retroativo financeiro definitivo não foi calculado. O valor abaixo possui natureza estimativa e foi apurado a partir das informações de itens/remanescentes disponíveis.

Retroativo financeiro definitivo: Não calculado
Retroativo estimado por itens/estoque: {moeda(valor_retroativo_estimado_itens(res))}
Execução estimada por itens/estoque: {moeda(res.get('valor_executado_atualizado', 0))}
Saldo remanescente atualizado: {moeda(res.get('remanescente_reajustado', 0))}

3. Retroativo estimado por itens/estoque

{tabela_ri}

4. Consolidação contratual

Valor original do contrato: {moeda(res.get('valor_original_contrato', 0))}
Valor Total Atualizado do Contrato: {moeda(res.get('valor_atualizado_contrato', res.get('valor_global_estoque', 0)))}
Aditivos/supressões registrados para controle: {moeda(res.get('total_aditivos_atualizados', 0))}

Data/hora de geração: {datetime.now(ZoneInfo('America/Sao_Paulo')).strftime('%d/%m/%Y %H:%M')}
""".strip()

    return f"""
RELATÓRIO EXECUTIVO — VALOR ATUALIZADO DO CONTRATO

1. Síntese dos ciclos

{_ciclos_df_para_relatorio(adm, res).to_string(index=False)}

2. Síntese financeira

Valor original do contrato: {moeda(res.get('valor_original_contrato', 0))}
Valor financeiro pago até o mês mais recente: {moeda(res.get('total_pago_faturado', 0))}
Valor teórico calculado: {moeda(res.get('total_devido_reajustado', 0))}
Valor represado a pagar: {moeda(res.get('valor_represado_a_pagar', 0))}
Valor executado total atualizado, considerado até o início do ciclo atual: {moeda(res.get('valor_executado_atualizado', 0))}
Saldo remanescente atualizado: {moeda(res.get('remanescente_reajustado', 0))}
Aditivos da análise atual, para controle: {moeda(res.get('total_aditivos_atualizados', 0))}
Valor Total Atualizado do Contrato: {moeda(res.get('valor_atualizado_contrato', res.get('valor_global_estoque', 0)))}

Data/hora de geração: {datetime.now(ZoneInfo('America/Sao_Paulo')).strftime('%d/%m/%Y %H:%M')}
""".strip()

def df_visual(df, moeda_cols=None, fator_cols=None, pct_cols=None):
    if df is None or not isinstance(df, pd.DataFrame):
        return pd.DataFrame()
    visual = df.copy()
    for col in moeda_cols or []:
        if col in visual.columns:
            visual[col] = visual[col].apply(moeda_ou_texto)
    for col in fator_cols or []:
        if col in visual.columns:
            visual[col] = visual[col].apply(fator_fmt)
    for col in pct_cols or []:
        if col in visual.columns:
            visual[col] = visual[col].apply(percentual)
    for col in visual.columns:
        if "data" in str(col).lower():
            visual[col] = visual[col].apply(formatar_data_br)

    # Quadro Executivo costuma vir como Indicador/Valor. Formatar como moeda
    # apenas os indicadores financeiros, preservando textos e percentuais.
    if "Indicador" in visual.columns and "Valor" in visual.columns:
        termos_monetarios = (
            "valor", "saldo", "aditivo", "supress", "delta", "pago", "teórico", "teorico",
            "remanescente", "formalizado", "original", "atualizado", "represado"
        )
        def _formatar_valor_quadro(row):
            indicador = str(row.get("Indicador", "")).lower()
            valor = row.get("Valor", "")
            if any(t in indicador for t in termos_monetarios):
                return moeda_ou_texto(valor)
            return valor
        visual["Valor"] = visual.apply(_formatar_valor_quadro, axis=1)

    if "Situação" in visual.columns:
        visual["Situação"] = visual["Situação"].apply(normalizar_status)
    return visual



# ── Compatibilidade Matriz 2.0: componentes do VTA para documentos ──
def _m20_componentes_vta_documentos(res):
    out = {"execucao": 0.0, "remanescente": 0.0, "aditivos": 0.0, "total": 0.0}

    def _num_local(v, pad=0.0):
        try:
            if v is None:
                return pad
            if isinstance(v, (int, float)):
                return float(v)
            s = str(v).strip().replace("R$", "").replace(" ", "")
            if not s:
                return pad
            if "," in s:
                s = s.replace(".", "").replace(",", ".")
            return float(s)
        except Exception:
            return pad

    def _col(df, candidatos):
        try:
            cols = list(df.columns)
        except Exception:
            return None
        normal = {str(c).lower().strip(): c for c in cols}
        for cand in candidatos:
            c = normal.get(str(cand).lower().strip())
            if c is not None:
                return c
        for c in cols:
            cn = str(c).lower()
            if any(str(cand).lower() in cn for cand in candidatos):
                return c
        return None

    df = res.get("df_composicao_valor_total") if isinstance(res, dict) else None

    try:
        import pandas as _pd_m20_docs
        eh_df = isinstance(df, _pd_m20_docs.DataFrame) and not df.empty
    except Exception:
        eh_df = hasattr(df, "iterrows") and getattr(df, "empty", True) is False

    if eh_df:
        col_comp = _col(df, ["Componente", "Parcela", "Descrição", "Descricao", "Indicador"])
        col_origem = _col(df, ["Origem", "Fonte"])
        col_valor = _col(df, ["Valor", "valor"])

        if col_valor:
            for _, row in df.iterrows():
                comp = str(row.get(col_comp, "") if col_comp else "").strip()
                origem = str(row.get(col_origem, "") if col_origem else "").strip()
                val = _num_local(row.get(col_valor, 0.0), 0.0)

                comp_norm = comp.lower()
                origem_norm = origem.lower()

                if "total" in comp_norm and ("vta" in comp_norm or "atualizado" in comp_norm or comp_norm == "total"):
                    out["total"] = val
                elif "aditivo" in comp_norm or "supress" in comp_norm or "aditivo" in origem_norm:
                    out["aditivos"] += val
                elif "remanescente" in comp_norm:
                    out["remanescente"] += val
                elif comp_norm.startswith("c") or "financeiro" in origem_norm or "itens" in origem_norm or "execut" in comp_norm:
                    out["execucao"] += val

    if isinstance(res, dict):
        if abs(out["execucao"]) <= 0.004:
            out["execucao"] = _num_local(res.get("valor_executado_atualizado", res.get("total_devido_reajustado", 0)), 0.0)
        if abs(out["remanescente"]) <= 0.004:
            out["remanescente"] = _num_local(res.get("remanescente_reajustado", 0), 0.0)
        if abs(out["aditivos"]) <= 0.004:
            out["aditivos"] = _num_local(res.get("total_aditivos_atualizados", 0), 0.0)
        if abs(out["total"]) <= 0.004:
            out["total"] = _num_local(res.get("valor_atualizado_contrato", res.get("valor_global_estoque", 0)), 0.0)

    return out

# >>> M21_PDF_COMPOSICAO_UNIFICADA_V3
def _m21_pdf_unificado_eh_matriz21(res):
    if not isinstance(res, dict):
        return False
    origem = str(res.get("origem_coleta", "") or "").lower()
    versao = str(res.get("versao", "") or "").lower()
    modo = str(res.get("modo_detectado", "") or "").lower()
    return (
        "matriz_2_1" in origem
        or "matriz_2_1" in versao
        or "matriz 2.1" in modo
        or bool(res.get("memoria_vta_m21"))
    )


def _m21_pdf_append_intro_composicao_unificada(story, styles, res):
    if not _m21_pdf_unificado_eh_matriz21(res):
        return

    try:
        from reportlab.platypus import Paragraph, Spacer
        from reportlab.lib.units import cm
    except Exception:
        return

    try:
        style_normal = styles["Normal"]
    except Exception:
        return

    texto = (
        "Na Matriz 2.1, esta secao concentra a leitura especifica da composicao do "
        "Valor Total Atualizado do Contrato. O VTA e apresentado como memoria "
        "auditavel da composicao economica do contrato, separando a execucao "
        "atualizada por ciclos, o saldo remanescente informado apos a competencia "
        "de corte e os aditivos/supressoes computaveis. Assim, eventual indicacao "
        "de ausencia de dados na secao Financeiro por Ciclo nao significa ausencia "
        "de composicao do VTA; a memoria consolidada esta demonstrada nesta secao."
    )

    formula = (
        "<b>Formula operacional da Matriz 2.1:</b> VTA = execucao atualizada por ciclos "
        "+ saldo remanescente apos a ultima competencia financeira "
        "+ aditivos/supressoes computaveis."
    )

    story.append(Paragraph(texto, style_normal))
    story.append(Paragraph(formula, style_normal))
    story.append(Spacer(1, 0.10 * cm))
# <<< M21_PDF_COMPOSICAO_UNIFICADA_V3

def criar_pdf_relatorio(adm, res):
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, PageBreak
    except Exception as exc:
        raise RuntimeError("A biblioteca reportlab não está instalada. Inclua 'reportlab' no requirements.txt.") from exc

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.05 * cm,
        leftMargin=1.05 * cm,
        topMargin=1.0 * cm,
        bottomMargin=1.0 * cm,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Titulo", parent=styles["Title"], fontSize=13, leading=16, alignment=1, spaceAfter=6))
    styles.add(ParagraphStyle(name="Subtitulo", parent=styles["Heading2"], fontSize=10, leading=12, spaceBefore=6, spaceAfter=4))
    styles.add(ParagraphStyle(name="Texto", parent=styles["BodyText"], fontSize=8.2, leading=10.5, alignment=4))
    styles.add(ParagraphStyle(name="Celula", parent=styles["BodyText"], fontSize=6.8, leading=8.2))
    styles.add(ParagraphStyle(name="CelulaCab", parent=styles["BodyText"], fontSize=6.8, leading=8.2, fontName="Helvetica-Bold", textColor=colors.white))

    story = []
    story.append(Paragraph("Análise de Reajuste Contratual", styles["Titulo"]))
    story.append(Paragraph("Relatório Executivo", styles["Titulo"]))
    story.append(Paragraph(f"Gerado em: {datetime.now(ZoneInfo('America/Sao_Paulo')).strftime('%d/%m/%Y %H:%M')}", styles["Texto"]))
    story.append(Spacer(1, 6))

    origem = (adm or {}).get("origem") or (adm or {}).get("tipo") or "Não informado"
    indice = res.get("indice", (adm or {}).get("indice", "Não informado"))
    fator = res.get("fator_acumulado", (adm or {}).get("fator_acumulado", (adm or {}).get("fator", 1.0)))

    story.append(Paragraph("1. Identificação da Análise", styles["Subtitulo"]))
    dados_identificacao = []
    if str(origem).strip().lower() not in ["", "não informado", "nao informado"]:
        dados_identificacao.append(["Origem da análise", origem])
    dados_identificacao.extend([
        ["Índice aplicado", indice],
        ["Fator acumulado", fator_fmt(fator)],
        ["Data de geração", datetime.now(ZoneInfo("America/Sao_Paulo")).strftime("%d/%m/%Y %H:%M")],
    ])
    story.append(tabela_pdf(dados_identificacao, col_widths=[6 * cm, 11 * cm]))

    df_ciclos_rel = _ciclos_df_para_relatorio(adm, res)
    if isinstance(df_ciclos_rel, pd.DataFrame) and not df_ciclos_rel.empty:
        story.append(Paragraph("2. Síntese dos ciclos", styles["Subtitulo"]))
        story.append(tabela_dataframe_pdf(df_ciclos_rel, max_linhas=20))

    # _m20_comp_vta_docs removido: uso direto do helper

    story.append(Paragraph("3. Indicadores Executivos", styles["Subtitulo"]))
    story.append(tabela_pdf(indicadores_executivos_relatorio(res), header=True, col_widths=[8.5 * cm, 8.5 * cm]))
    info_metodologia = metodologia_corte_operacional_info(res)
    story.append(Paragraph("3.1. Metodologia de corte da apuração", styles["Subtitulo"]))
    story.append(Paragraph(f"{info_metodologia['titulo']}. {info_metodologia['texto']}", styles["Texto"]))

    if eh_modo_consumo_itens_ciclo(res):
        story.append(Paragraph("3.2. Modo Consumo por Itens/Ciclo", styles["Subtitulo"]))
        story.append(Paragraph(
            "A análise foi processada sem base mensal por competência. O retroativo financeiro definitivo por competência não é calculado neste modo. "
            "O Retroativo (itens consumidos/ciclo) foi apurado com base nos quantitativos consumidos por item e por ciclo informados pela fiscalização.",
            styles["Texto"],
        ))
        df_ri_pdf = df_retroativo_consumo_itens_ciclo(res)
        if isinstance(df_ri_pdf, pd.DataFrame) and not df_ri_pdf.empty:
            df_ri_pdf = df_visual(
                df_ri_pdf,
                moeda_cols=["Valor original consumido", "Valor atualizado consumido", "Retroativo"],
                fator_cols=["Fator acumulado"],
            )
            story.append(tabela_dataframe_pdf(df_ri_pdf, max_linhas=20))
    elif eh_modo_reduzido_itens(res):
        story.append(Paragraph("3.2. Modo Reduzido por Itens/Estoque", styles["Subtitulo"]))
        story.append(Paragraph(
            "A análise foi processada sem base mensal por competência. O retroativo financeiro definitivo não é calculado neste modo. "
            "O valor de retroativo estimado por itens/estoque possui natureza estimativa e deve ser validado antes de qualquer formalização de pagamento.",
            styles["Texto"],
        ))
        df_ri_pdf = df_retroativo_estimado_itens(res)
        if isinstance(df_ri_pdf, pd.DataFrame) and not df_ri_pdf.empty:
            df_ri_pdf = df_visual(
                df_ri_pdf,
                moeda_cols=["Valor executado original", "Valor executado atualizado", "Retroativo estimado por itens/estoque"],
            )
            story.append(tabela_dataframe_pdf(df_ri_pdf, max_linhas=20))
    story.append(Paragraph("4. Financeiro por Ciclo", styles["Subtitulo"]))
    df_fin = df_visual(
        res.get("df_financeiro_por_ciclo"),
        moeda_cols=["Valor pago efetivo", "Valor teórico calculado", "Valor pago/faturado", "Valor devido reajustado", "Delta do ciclo", "Delta acumulado"],
        fator_cols=["Fator aplicado ao retroativo", "Fator aplicado"],
    )
    keep_fin = [c for c in ["Ciclo", "Situação", "Tratamento financeiro", "Fator aplicado ao retroativo", "Fator aplicado", "Valor pago efetivo", "Valor teórico calculado", "Valor pago/faturado", "Valor devido reajustado", "Delta do ciclo"] if c in df_fin.columns]
    story.append(tabela_dataframe_pdf(df_fin[keep_fin] if keep_fin else df_fin, max_linhas=20))

    story.append(Paragraph("5. Composição do Valor Total Atualizado do Contrato", styles["Subtitulo"]))
    # >>> M21_PDF_COMPOSICAO_UNIFICADA_V3_CALL
    try:
        _m21_pdf_append_intro_composicao_unificada(story, styles, res)
    except Exception as _e_m21_pdf_unificado:
        try:
            story.append(Paragraph(f"Observacao Matriz 2.1: nao foi possivel inserir a introducao unificada da composicao do VTA ({_e_m21_pdf_unificado}).", styles["Normal"]))
        except Exception:
            pass
    # <<< M21_PDF_COMPOSICAO_UNIFICADA_V3_CALL
    df_comp = res.get("df_composicao_valor_total")
    if isinstance(df_comp, pd.DataFrame) and not df_comp.empty:
        df_comp_pdf = df_comp.copy()
        if "Valor" in df_comp_pdf.columns:
            df_comp_pdf["Valor"] = df_comp_pdf["Valor"].apply(moeda)
        keep_comp = [c for c in ["Componente", "Ciclo/Referência", "Valor", "Observação"] if c in df_comp_pdf.columns]
        story.append(tabela_dataframe_pdf(df_comp_pdf[keep_comp], max_linhas=20))
    else:
        story.append(tabela_pdf([
            ["Componente", "Valor"],
            ["Valor executado atualizado", moeda(_m20_componentes_vta_documentos(res).get("execucao", 0))],
            ["Saldo remanescente atualizado", moeda(_m20_componentes_vta_documentos(res).get("remanescente", 0))],
            ["Valor Total Atualizado do Contrato", moeda(_m20_componentes_vta_documentos(res).get("total", res.get("valor_atualizado_contrato", res.get("valor_global_estoque", 0))))],
        ], header=True, col_widths=[10 * cm, 7 * cm]))
        story.append(Paragraph(
            "Aditivos e supressões registrados são apresentados em seção própria para controle e não são somados como parcela autônoma ao Valor Total Atualizado quando já refletidos na execução ou no saldo remanescente.",
            styles["Texto"],
        ))

    df_ad = res.get("df_aditivos_executivo", res.get("df_aditivos"))
    story.append(Paragraph("6. Aditivos", styles["Subtitulo"]))

    if isinstance(df_ad, pd.DataFrame) and not df_ad.empty:
        df_adv = df_visual(
            df_ad,
            moeda_cols=[
                "Valor do aditivo na assinatura",
                "Valor do aditivo reajustado",
                "Valor original da alteração",
                "Valor atualizado da alteração",
            ],
            fator_cols=["Fator aplicado"],
        )

        keep_ad = [
            c for c in [
                "Aditivo",
                "Ciclo/Marco",
                "Tipo de alteração",
                "Tratamento do aditivo",
                "Quantidade de linhas",
                "Valor do aditivo na assinatura",
                "Fator aplicado",
                "Valor do aditivo reajustado",
            ]
            if c in df_adv.columns
        ]

        df_ad_visivel = df_adv[keep_ad].copy() if keep_ad else pd.DataFrame()

        if df_ad_visivel.empty or df_ad_visivel.replace("", pd.NA).dropna(how="all").empty:
            story.append(Paragraph("Não há aditivos a serem informados.", styles["Texto"]))
        else:
            story.append(tabela_dataframe_pdf(df_ad_visivel, max_linhas=12))
    else:
        story.append(Paragraph("Não há aditivos a serem informados.", styles["Texto"]))

    doc.build(story)
    pdf = buffer.getvalue()
    buffer.close()
    return pdf



# >>> PDF_WRAP_TABELAS_V1
def _pdf_wrap_celula_tabela(valor, header=False):
    # Converte textos longos em Paragraph para evitar extrapolação da célula.
    try:
        from reportlab.platypus import Paragraph
        from reportlab.lib.styles import ParagraphStyle
        from xml.sax.saxutils import escape
    except Exception:
        return valor

    if valor is None:
        return ""

    try:
        if hasattr(valor, "wrap") and hasattr(valor, "drawOn"):
            return valor
    except Exception:
        pass

    if isinstance(valor, (int, float)):
        return valor

    txt = str(valor)
    if not txt:
        return ""

    deve_wrap = header or len(txt) > 24 or "|" in txt or "\n" in txt or ";" in txt
    if not deve_wrap:
        return txt

    tamanho = 7
    leading = 8
    nome = "PDFTableCellHeaderWrap" if header else "PDFTableCellWrap"

    style = ParagraphStyle(
        name=nome,
        fontName="Helvetica-Bold" if header else "Helvetica",
        fontSize=tamanho,
        leading=leading,
        wordWrap="CJK",
        splitLongWords=True,
        spaceAfter=0,
        spaceBefore=0,
    )

    safe = escape(txt).replace("\n", "<br/>")
    return Paragraph(safe, style)


def _pdf_wrap_tabela_celulas(dados):
    # Aplica quebra automática às células textuais longas de uma tabela.
    try:
        linhas = list(dados or [])
    except Exception:
        return dados

    saida = []
    for i, linha in enumerate(linhas):
        try:
            celulas = list(linha)
        except Exception:
            saida.append(linha)
            continue

        header = i == 0
        saida.append([_pdf_wrap_celula_tabela(c, header=header) for c in celulas])

    return saida
# <<< PDF_WRAP_TABELAS_V1

def tabela_pdf(dados, header=False, col_widths=None):
    # >>> PDF_WRAP_TABELAS_V1_CALL_TABELA_PDF
    try:
        dados = _pdf_wrap_tabela_celulas(dados)
    except Exception:
        pass
    # <<< PDF_WRAP_TABELAS_V1_CALL_TABELA_PDF
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle

    table = Table(dados, colWidths=col_widths, hAlign="CENTER")
    estilo = [
        ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 7.4),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
    ]
    if header:
        estilo.extend([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ])
    table.setStyle(TableStyle(estilo))
    return table


def tabela_dataframe_pdf(df, max_linhas=12):
    from reportlab.platypus import Paragraph
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm

    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        return tabela_pdf([["Informação", "Sem dados disponíveis"]], col_widths=[5 * cm, 11 * cm])

    from reportlab.lib import colors

    styles = getSampleStyleSheet()
    cel = ParagraphStyle("Cel", parent=styles["BodyText"], fontSize=6.6, leading=8.0)
    cab = ParagraphStyle(
        "Cab",
        parent=styles["BodyText"],
        fontSize=6.6,
        leading=8.0,
        fontName="Helvetica-Bold",
        textColor=colors.white,
    )
    dados = [[Paragraph(str(c), cab) for c in df.columns]]
    for _, row in df.head(max_linhas).iterrows():
        linha = []
        for valor in row.tolist():
            linha.append(Paragraph(str(valor), cel))
        dados.append(linha)

    ncols = max(len(df.columns), 1)
    largura_total = 17.0 * cm
    col_widths = [largura_total / ncols] * ncols
    return tabela_pdf(dados, header=True, col_widths=col_widths)




# ============================================================
# Layout responsivo e Minuta de Apostilamento
# ============================================================

def aplicar_css_responsivo_relatorio():
    st.markdown(
        """
        <style>
        div[data-testid="stMetric"] {
            background: #FFFFFF;
            border: 1px solid #E5EAF0;
            border-radius: 12px;
            padding: 10px 12px;
            min-height: 82px;
        }
        div[data-testid="stMetricLabel"] p {
            color: #475569;
            font-size: clamp(0.74rem, 1.15vw, 0.90rem);
            line-height: 1.2;
            white-space: normal;
            word-break: normal;
        }
        div[data-testid="stMetricValue"] {
            font-size: clamp(1.02rem, 1.75vw, 1.50rem);
            line-height: 1.20;
            white-space: normal;
            overflow-wrap: anywhere;
        }
        div[data-testid="stMetricDelta"] {
            font-size: 0.78rem;
        }
        .telebras-kpi-destaque {
            background:#EAF2F8;
            border:1px solid #BFD7EA;
            border-radius:14px;
            padding:16px 20px;
            margin:10px 0 16px 0;
        }
        .telebras-kpi-destaque-label {
            font-size:0.92rem;
            color:#27496D;
            font-weight:600;
        }
        .telebras-kpi-destaque-valor {
            font-size:clamp(1.25rem, 2.45vw, 1.95rem);
            color:#0B1F3A;
            font-weight:800;
            line-height:1.25;
            word-break:break-word;
        }
        @media (max-width: 1200px) {
            div[data-testid="stMetricValue"] {
                font-size: 1.05rem;
            }
            div[data-testid="stMetric"] {
                padding: 8px 10px;
                min-height: 76px;
            }
        }

        .modo-reduzido-box {
            background: #F3E8FF;
            border: 1px solid #A855F7;
            border-left: 8px solid #7E22CE;
            border-radius: 14px;
            padding: 18px 22px;
            margin: 12px 0 18px 0;
        }
        .modo-reduzido-titulo {
            color: #581C87;
            font-weight: 900;
            font-size: 1.02rem;
            margin-bottom: 7px;
        }
        .modo-reduzido-texto {
            color: #3B0764;
            font-size: 0.95rem;
            line-height: 1.45;
        }
        .modo-reduzido-card {
            background: #F3E8FF;
            border: 1px solid #A855F7;
            border-radius: 14px;
            padding: 16px 18px;
            margin: 8px 0 16px 0;
        }
        .modo-reduzido-card-label {
            color: #581C87;
            font-weight: 800;
            font-size: 0.92rem;
            margin-bottom: 6px;
        }
        .modo-reduzido-card-valor {
            color: #0F172A;
            font-size: clamp(1.20rem, 2.10vw, 1.65rem);
            font-weight: 900;
            line-height: 1.2;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def _valor_resumo(res, *chaves, padrao=0.0):
    for chave in chaves:
        if isinstance(res, dict) and chave in res:
            return res.get(chave)
    return padrao


def _placeholder(valor=None):
    if valor is None:
        return "[campo a preencher]"
    texto = str(valor).strip()
    return texto if texto else "[campo a preencher]"


def _limpar_texto_formal(valor):
    """Remove emojis e marcadores visuais informais de textos destinados à minuta formal."""
    if valor is None:
        return ""
    texto = str(valor)
    substituicoes = {
        "❌": "",
        "✅": "",
        "⚠️": "",
        "⚠": "",
        "🟡": "",
        "🔴": "",
        "🟢": "",
        "🔵": "",
        "🛡️": "",
        "🛡": "",
        "📊": "",
        "📝": "",
        "⚖️": "",
        "⚖": "",
        "🔄": "",
        "📥": "",
        "🔍": "",
    }
    for antigo, novo in substituicoes.items():
        texto = texto.replace(antigo, novo)
    # Remove caracteres fora do plano multilíngue básico, onde ficam a maior parte dos emojis.
    texto = "".join(ch for ch in texto if ord(ch) <= 0xFFFF)
    return " ".join(texto.split())


def _adicionar_item_numerado(document, numero, texto):
    p = document.add_paragraph()
    p.add_run(f"{numero}. ").bold = True
    p.add_run(_limpar_texto_formal(texto))
    return p


def _adicionar_subitem(document, marcador, texto):
    from docx.shared import Cm
    p = document.add_paragraph(f"{marcador} {_limpar_texto_formal(texto)}")
    p.paragraph_format.left_indent = Cm(0.7)
    return p


def _romano(numero):
    mapa = [
        (10, "x"), (9, "ix"), (5, "v"), (4, "iv"), (1, "i"),
    ]
    n = int(numero)
    saida = ""
    for valor, simbolo in mapa:
        while n >= valor:
            saida += simbolo
            n -= valor
    return saida


def _ciclos_para_minuta(adm, res):
    ciclos = []
    if adm:
        ciclos = adm.get("ciclos") or adm.get("detalhamento_ciclos") or []
    if not ciclos and isinstance(res.get("df_ciclos"), pd.DataFrame) and not res.get("df_ciclos").empty:
        ciclos = res.get("df_ciclos").to_dict("records")
    if not ciclos and isinstance(res.get("df_delta_por_ciclo"), pd.DataFrame) and not res.get("df_delta_por_ciclo").empty:
        ciclos = res.get("df_delta_por_ciclo").to_dict("records")
    return ciclos or []


def _percentual_ciclo_minuta(ciclo):
    for chave in ["percentual_aplicado", "Percentual aplicado", "Variação", "variacao", "var"]:
        if chave in ciclo and ciclo.get(chave) not in [None, ""]:
            try:
                return percentual(float(ciclo.get(chave)), 2)
            except Exception:
                return str(ciclo.get(chave))
    return "[campo a preencher]"


def _efeito_ciclo_minuta(ciclo):
    for chave in [
        "data_inicio_efeito_financeiro",
        "inicio_financeiro_acordo",
        "financeiro_inicio",
        "Início financeiro",
        "Data de início dos efeitos financeiros",
        "data_pedido",
        "Data do pedido",
    ]:
        valor = ciclo.get(chave) if isinstance(ciclo, dict) else None
        if valor not in [None, ""]:
            return formatar_data_br(valor)
    return "[campo a preencher]"


def _nome_ciclo_minuta(ciclo, idx):
    return str(ciclo.get("ciclo") or ciclo.get("Ciclo") or f"C{idx}")


def _adicionar_paragrafo_justificado(document, texto):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = document.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _aplicar_estilo_docx(document):
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for nome in ["Title", "Heading 1", "Heading 2"]:
        if nome in styles:
            styles[nome].font.name = "Arial"
    styles["Title"].font.size = Pt(13)
    styles["Title"].font.bold = True


def _destacar_campos_preencher(document):
    """Destaca em amarelo os trechos [campo a preencher...] na minuta DOCX."""
    try:
        from copy import deepcopy
        from docx.enum.text import WD_COLOR_INDEX
        from docx.oxml import OxmlElement
        from docx.text.run import Run
    except Exception:
        return document

    padrao = re.compile(r"(\[campo a preencher[^\]]*\])", flags=re.IGNORECASE)
    for paragraph in document.paragraphs:
        runs_originais = list(paragraph.runs)
        for run in runs_originais:
            texto = run.text
            if not texto or not padrao.search(texto):
                continue
            partes = padrao.split(texto)
            run.text = partes[0]
            anterior = run._r
            for parte in partes[1:]:
                novo_r = OxmlElement('w:r')
                if run._r.rPr is not None:
                    novo_r.append(deepcopy(run._r.rPr))
                novo_t = OxmlElement('w:t')
                novo_t.text = parte
                novo_r.append(novo_t)
                anterior.addnext(novo_r)
                novo_run = Run(novo_r, paragraph)
                if padrao.fullmatch(parte):
                    novo_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                anterior = novo_r
    return document



# ============================================================
# Apostila em modelo tabelado executivo
# ============================================================

def _docx_add_titulo_secao(document, texto):
    from docx.shared import Pt
    p = document.add_paragraph()
    r = p.add_run(str(texto))
    r.bold = True
    r.font.size = Pt(10.5)
    return p


def _docx_add_texto(document, texto, justificar=True):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = document.add_paragraph(_limpar_texto_formal(str(texto)))
    if justificar:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _docx_set_cell_text(cell, texto, bold=False, align="left"):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(texto))
    r.bold = bool(bold)
    r.font.size = Pt(8.6)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _docx_shade_cell(cell, fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _docx_set_cell_width(cell, width_cm):
    try:
        from docx.shared import Cm
        cell.width = Cm(float(width_cm))
    except Exception:
        pass


def _docx_tabela_executiva(document, headers, rows, widths=None, total_last=False):
    """Tabela cinza corporativa, com cabeçalho discreto, similar ao modelo de apostila."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

    if rows is None:
        rows = []
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _docx_set_cell_text(hdr[i], h, bold=True, align="center")
        _docx_shade_cell(hdr[i], "D9D9D9")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths and i < len(widths):
            _docx_set_cell_width(hdr[i], widths[i])

    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        is_total = total_last and ridx == len(rows) - 1
        for cidx, valor in enumerate(row):
            align = "right" if cidx == len(row) - 1 else ("center" if cidx == 0 else "left")
            _docx_set_cell_text(cells[cidx], valor, bold=is_total, align=align)
            cells[cidx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if is_total:
                _docx_shade_cell(cells[cidx], "EAF2F8")
            if widths and cidx < len(widths):
                _docx_set_cell_width(cells[cidx], widths[cidx])
    document.add_paragraph("")
    return table


def _letra_ref(indice):
    letras = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    if indice < len(letras):
        return letras[indice]
    return f"A{indice - len(letras) + 1}"


def _fmt_pct_minuta(valor):
    try:
        v = float(valor)
        if abs(v) > 1.5:
            v = v / 100
        return percentual(v, 2)
    except Exception:
        txt = str(valor or "").strip()
        return txt if txt else "[campo a preencher]"


def _fmt_moeda_minuta(valor):
    try:
        return moeda(valor)
    except Exception:
        return texto_seguro(valor, "")


def _normalizar_nome_ciclo_doc(valor):
    txt = str(valor or "").strip().upper()
    if not txt:
        return ""
    if txt.startswith("C"):
        return txt
    if txt.isdigit():
        return f"C{txt}"
    return txt


def _df_seguro(res, *chaves):
    for chave in chaves:
        df = res.get(chave) if isinstance(res, dict) else None
        if isinstance(df, pd.DataFrame) and not df.empty:
            return df.copy()
    return pd.DataFrame()


def _linhas_quadro_reajustes(adm, res):
    ciclos = _ciclos_para_minuta(adm, res)
    linhas = []
    vistos = set()
    ref_idx = 0
    for idx, ciclo in enumerate(ciclos, start=1):
        if not isinstance(ciclo, dict):
            continue
        nome = _normalizar_nome_ciclo_doc(ciclo.get("ciclo") or ciclo.get("Ciclo") or f"C{idx}")
        if not nome or nome in ["C0", "TOTAL", "CICLO"] or nome in vistos:
            continue
        vistos.add(nome)
        pct = _percentual_ciclo_minuta(ciclo)
        efeito = _efeito_ciclo_minuta(ciclo)
        situacao_bruta = ciclo.get("situacao_aplicada") or ciclo.get("Situação aplicada") or ciclo.get("situacao") or ciclo.get("Situação") or ""
        acordo = texto_seguro(ciclo.get("Acordo negocial", ciclo.get("acordo_negocial", "")), "")
        if str(acordo).strip().lower() in ["sim", "s", "true", "1"]:
            situacao = "Admitido por negociação entre as partes"
        else:
            situacao = _status_relatorio(situacao_bruta).capitalize()
        linhas.append([_letra_ref(ref_idx), nome, pct, f"A partir de {efeito}", situacao])
        ref_idx += 1

    fator_acumulado = res.get("fator_acumulado", adm.get("fator_acumulado", adm.get("fator", 1.0))) if isinstance(res, dict) else adm.get("fator_acumulado", adm.get("fator", 1.0))
    try:
        pct_acum = percentual(float(fator_acumulado) - 1, 2)
    except Exception:
        pct_acum = "[campo a preencher]"
    if linhas:
        linhas.append([_letra_ref(ref_idx), "Acumulado", pct_acum, "Conforme composição dos ciclos", "Percentual acumulado apurado"])
    return linhas


def _coluna_existente(df, candidatos):
    if not isinstance(df, pd.DataFrame) or df.empty:
        return None
    mapa = {str(c).strip().lower(): c for c in df.columns}
    for cand in candidatos:
        c = mapa.get(str(cand).strip().lower())
        if c is not None:
            return c
    return None


def _linhas_quadro_financeiro(res):
    df = _df_seguro(res, "df_financeiro_por_ciclo", "df_delta_por_ciclo")
    if df.empty:
        return []
    col_ciclo = _coluna_existente(df, ["Ciclo", "ciclo"])
    col_pago = _coluna_existente(df, ["Valor pago efetivo", "Valor pago/faturado", "Valor pago", "Valor nominal pago"])
    col_teorico = _coluna_existente(df, ["Valor teórico calculado", "Valor devido reajustado", "Valor devido atualizado"])
    col_delta = _coluna_existente(df, ["Diferença/retroativo", "Delta do ciclo", "Delta acumulado", "Retroativo a pagar"])
    if not col_ciclo:
        return []
    linhas = []
    total_pago = total_teorico = total_delta = 0.0
    for _, row in df.iterrows():
        ciclo = _normalizar_nome_ciclo_doc(row.get(col_ciclo, ""))
        if not ciclo or ciclo in ["C0", "TOTAL", "CICLO"]:
            continue
        pago = _numero_br_relatorio(row.get(col_pago, 0), 0.0) if col_pago else 0.0
        teorico = _numero_br_relatorio(row.get(col_teorico, 0), 0.0) if col_teorico else 0.0
        delta = _numero_br_relatorio(row.get(col_delta, teorico - pago), teorico - pago) if col_delta else teorico - pago
        total_pago += pago
        total_teorico += teorico
        total_delta += delta
        linhas.append([ciclo, moeda(pago), moeda(teorico), moeda(delta)])
    if linhas:
        linhas.append(["Total", moeda(total_pago), moeda(total_teorico), moeda(total_delta)])
    return linhas


def _linhas_quadro_memoria_fiscal(res):
    """Monta memória fiscal evolutiva do Valor Total Atualizado.

    O Quadro 3 da Apostila deve contar a história econômico-financeira do contrato:
    valor original, execuções por ciclo, remanescentes intermediários, corte operacional,
    saldo remanescente atualizado e aditivos/supressões computáveis.

    O Quadro 4 permanece como composição sintética final.
    """
    if not isinstance(res, dict):
        return []

    # Se futuramente o módulo Valores fornecer uma memória fiscal analítica própria,
    # ela tem prioridade absoluta.
    for chave in ["df_memoria_fiscal_valor_total", "df_memoria_valor_total"]:
        df_mem = res.get(chave)
        if isinstance(df_mem, pd.DataFrame) and not df_mem.empty:
            col_desc = _coluna_existente(df_mem, ["Descrição", "Descricao", "Componente", "Parcela", "Indicador"])
            col_valor = _coluna_existente(df_mem, ["Valor", "valor"])
            if col_desc and col_valor:
                linhas_prontas = []
                for _, row in df_mem.iterrows():
                    desc = texto_seguro(row.get(col_desc, ""), "")
                    valor = _numero_br_relatorio(row.get(col_valor, 0), 0.0)
                    if desc:
                        linhas_prontas.append([_letra_ref(len(linhas_prontas)), desc, moeda(valor)])
                if linhas_prontas:
                    return linhas_prontas

    linhas = []

    def _adicionar(descricao, valor, permitir_zero=False):
        valor_num = _numero_br_relatorio(valor, 0.0)
        if permitir_zero or abs(valor_num) > 0.004:
            linhas.append([_letra_ref(len(linhas)), descricao, moeda(valor_num)])

    def _num(row, candidatos, padrao=0.0):
        if row is None:
            return padrao
        for col in candidatos:
            if col in row.index:
                return _numero_br_relatorio(row.get(col, padrao), padrao)
        return padrao

    def _txt(row, candidatos, padrao=""):
        if row is None:
            return padrao
        for col in candidatos:
            if col in row.index:
                return texto_seguro(row.get(col, ""), padrao)
        return padrao

    valor_original = _numero_br_relatorio(res.get("valor_original_contrato", 0), 0.0)
    if abs(valor_original) > 0.004:
        _adicionar("Valor original do contrato", valor_original)

    # Execução por ciclo.
    df_exec = _df_seguro(res, "df_execucao_atualizada")
    exec_rows = []

    if isinstance(df_exec, pd.DataFrame) and not df_exec.empty:
        col_ciclo = _coluna_existente(df_exec, ["Ciclo", "ciclo", "Ciclo/Referência"])
        if col_ciclo:
            df_tmp = df_exec.copy()
            df_tmp["_ordem_memoria"] = (
                df_tmp[col_ciclo]
                .astype(str)
                .str.upper()
                .str.extract(r"(\d+)")
                .fillna(999)
                .astype(int)
            )
            df_tmp = df_tmp.sort_values("_ordem_memoria", kind="stable")

            for _, row in df_tmp.iterrows():
                ciclo = _normalizar_nome_ciclo_doc(row.get(col_ciclo, ""))
                if not ciclo or ciclo in ["TOTAL", "CICLO"]:
                    continue

                valor_original_exec = _num(
                    row,
                    [
                        "Valor executado original",
                        "Valor original executado",
                        "Valor original/base",
                        "Valor base",
                        "Valor pago efetivo",
                        "Valor pago/faturado",
                    ],
                    0.0,
                )
                valor_atualizado_exec = _num(
                    row,
                    [
                        "Valor executado atualizado",
                        "Valor atualizado",
                        "Valor teórico calculado",
                        "Valor devido reajustado",
                    ],
                    valor_original_exec,
                )

                if abs(valor_atualizado_exec) <= 0.004 and abs(valor_original_exec) <= 0.004:
                    continue

                status = _txt(row, ["Status financeiro", "Situação", "Tratamento financeiro", "Observação"], "")
                status_norm = str(status).lower()

                if ciclo == "C0":
                    descricao = "C0 - execução sem reajuste"
                elif "preclus" in status_norm or "sem efeito" in status_norm or "sem_efeito" in status_norm:
                    descricao = f"{ciclo} - execução sem reajuste, em razão da preclusão/ausência de efeito financeiro"
                else:
                    descricao = f"{ciclo} - execução atualizada"

                exec_rows.append({
                    "ciclo": ciclo,
                    "valor_original": valor_original_exec,
                    "valor_atualizado": valor_atualizado_exec,
                    "descricao": descricao,
                })

    acumulado_original = 0.0

    for idx, item in enumerate(exec_rows):
        ciclo = item["ciclo"]
        valor_original_exec = _numero_br_relatorio(item["valor_original"], 0.0)
        valor_atualizado_exec = _numero_br_relatorio(item["valor_atualizado"], valor_original_exec)

        # Quando houver valor original/base, usa-o para memória de remanescente.
        # Se não houver, usa valor atualizado apenas como fallback documental.
        base_memoria = valor_original_exec if abs(valor_original_exec) > 0.004 else valor_atualizado_exec
        acumulado_original += base_memoria

        _adicionar(item["descricao"], valor_atualizado_exec)

        if abs(valor_original) > 0.004:
            rem_original = valor_original - acumulado_original
            if rem_original > 0.004:
                if idx + 1 < len(exec_rows):
                    prox_ciclo = exec_rows[idx + 1]["ciclo"]
                    desc_rem = f"Remanescente após {ciclo}, em base original, antes da apuração do {prox_ciclo}"
                else:
                    desc_rem = f"Remanescente após {ciclo}, em base original"
                _adicionar(desc_rem, rem_original)

    # Remanescente final: prioriza tabela de composição quando houver linha específica.
    rem_final = 0.0
    rem_final_desc = "Saldo remanescente atualizado"

    df_comp = _df_seguro(res, "df_composicao_valor_total")
    if isinstance(df_comp, pd.DataFrame) and not df_comp.empty:
        col_desc = _coluna_existente(df_comp, ["Componente", "Parcela", "Descrição", "Descricao", "Indicador"])
        col_valor = _coluna_existente(df_comp, ["Valor", "valor"])
        if col_desc and col_valor:
            for _, row in df_comp.iterrows():
                desc = texto_seguro(row.get(col_desc, ""), "")
                desc_norm = desc.lower()
                if "remanescente" in desc_norm and "total" not in desc_norm:
                    rem_final = _numero_br_relatorio(row.get(col_valor, 0), 0.0)
                    rem_final_desc = desc
                    break

    if abs(rem_final) <= 0.004:
        rem_final = _numero_br_relatorio(res.get("remanescente_reajustado", 0), 0.0)

    if abs(rem_final) > 0.004:
        config = res.get("config_ciclo_em_execucao", {}) or {}
        ciclo_corte = texto_seguro(config.get("ciclo", ""), "")
        data_corte = texto_seguro(config.get("data_corte") or config.get("competencia_corte") or "", "")
        if ciclo_corte or data_corte:
            complemento = []
            if ciclo_corte:
                complemento.append(str(ciclo_corte))
            if data_corte:
                complemento.append(str(data_corte))
            rem_final_desc = f"Saldo remanescente atualizado no corte operacional ({' - '.join(complemento)})"
        _adicionar(rem_final_desc, rem_final)

    # Aditivos computáveis: preferir dataframe executivo e respeitar tratamento.
    total_aditivos_computaveis = 0.0
    df_ad = _df_seguro(res, "df_aditivos_executivo", "df_aditivos")

    if isinstance(df_ad, pd.DataFrame) and not df_ad.empty:
        col_valor = _coluna_existente(df_ad, ["Valor do aditivo reajustado", "Valor atualizado da alteração", "Valor atualizado", "Valor"])
        col_trat = _coluna_existente(df_ad, ["Tratamento do aditivo", "Tratamento", "Computa no Valor Global"])
        if col_valor:
            for _, row in df_ad.iterrows():
                trat = texto_seguro(row.get(col_trat, ""), "") if col_trat else ""
                trat_norm = str(trat).lower()
                computa_txt = str(row.get("Computa no Valor Global", "")).strip().lower() if "Computa no Valor Global" in row.index else ""
                eh_informativo = any(t in trat_norm for t in ["informativo", "ja incorporado", "já incorporado", "ja incluido", "já incluído"])
                computa = (not eh_informativo) and computa_txt not in ["não", "nao", "n", "false"]
                if computa:
                    total_aditivos_computaveis += _numero_br_relatorio(row.get(col_valor, 0), 0.0)

    if abs(total_aditivos_computaveis) <= 0.004 and bool(res.get("aditivos_somados_ao_valor_total", False)):
        total_aditivos_computaveis = _numero_br_relatorio(res.get("total_aditivos_atualizados", 0), 0.0)

    if abs(total_aditivos_computaveis) > 0.004:
        _adicionar("Aditivos/supressões computáveis atualizados", total_aditivos_computaveis)

    total = _numero_br_relatorio(res.get("valor_atualizado_contrato", res.get("valor_global_estoque", 0)), 0.0)
    if abs(total) > 0.004:
        linhas.append(["Total", "Valor Total Atualizado do Contrato", moeda(total)])

    # Fallback de segurança: se a memória evolutiva não conseguiu ser formada,
    # usa a composição final para evitar quadro vazio.
    if len(linhas) <= 2 and isinstance(df_comp, pd.DataFrame) and not df_comp.empty:
        col_desc = _coluna_existente(df_comp, ["Componente", "Parcela", "Descrição", "Descricao", "Indicador"])
        col_valor = _coluna_existente(df_comp, ["Valor", "valor"])
        if col_desc and col_valor:
            linhas = []
            for _, row in df_comp.iterrows():
                desc = texto_seguro(row.get(col_desc, ""), "")
                valor = _numero_br_relatorio(row.get(col_valor, 0), 0.0)
                if desc and abs(valor) > 0.004:
                    linhas.append([_letra_ref(len(linhas)), desc, moeda(valor)])

    return linhas


def _linhas_quadro_composicao_sintetica(res):
    df = _df_seguro(res, "df_composicao_valor_total")
    linhas = []
    if not df.empty:
        col_desc = _coluna_existente(df, ["Componente", "Parcela", "Descrição", "Descricao"])
        col_valor = _coluna_existente(df, ["Valor", "valor"])
        if col_desc and col_valor:
            for _, row in df.iterrows():
                desc = texto_seguro(row.get(col_desc, ""), "")
                valor = _numero_br_relatorio(row.get(col_valor, 0), 0.0)
                if desc and abs(valor) > 0.004:
                    linhas.append([_letra_ref(len(linhas)), desc, moeda(valor)])
    if not linhas:
        execucao = _numero_br_relatorio(res.get("valor_executado_atualizado", res.get("total_devido_reajustado", 0)), 0.0)
        remanescente = _numero_br_relatorio(res.get("remanescente_reajustado", 0), 0.0)
        aditivos = _numero_br_relatorio(res.get("total_aditivos_atualizados", 0), 0.0)
        if abs(execucao) > 0.004:
            linhas.append(["A", "Execução atualizada", moeda(execucao)])
        if abs(remanescente) > 0.004:
            linhas.append(["B", "Saldo remanescente atualizado", moeda(remanescente)])
        if abs(aditivos) > 0.004:
            linhas.append(["C", "Aditivos/supressões registrados", moeda(aditivos)])
    total = _numero_br_relatorio(res.get("valor_atualizado_contrato", res.get("valor_global_estoque", 0)), 0.0)
    if abs(total) > 0.004:
        linhas.append(["Total", "Valor Total Atualizado", moeda(total)])
    return linhas


def _texto_formula_refs(linhas):
    refs = [str(l[0]) for l in linhas if l and str(l[0]).strip().upper() not in ["TOTAL", "P"]]
    refs = [r for r in refs if len(r) <= 2]
    if not refs:
        return ""
    return " + ".join(refs)


def gerar_minuta_apostilamento_docx(adm, res):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, Cm, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except Exception as exc:
        raise RuntimeError("A biblioteca python-docx não está instalada. Inclua 'python-docx' no requirements.txt.") from exc

    adm = adm or {}
    res = res or {}

    document = Document()
    _aplicar_estilo_docx(document)

    # ── Cabeçalho ──────────────────────────────────────────────────
    p_titulo = document.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_titulo.add_run("TERMO DE APOSTILA nº [campo a preencher]")
    r.bold = True
    r.font.size = Pt(13)

    document.add_paragraph("")

    contrato = _placeholder(adm.get("contrato") or res.get("contrato") or res.get("numero_contrato"))
    processo = _placeholder(adm.get("processo") or res.get("processo"))

    p_ref = document.add_paragraph()
    p_ref.add_run(f"PROCESSO Nº {processo}").bold = True
    p_ref2 = document.add_paragraph()
    p_ref2.add_run(f"CONTRATO Nº {contrato}").bold = True

    document.add_paragraph("")

    # ── Identificação ──────────────────────────────────────────────
    _adicionar_paragrafo_justificado(
        document,
        "A TELECOMUNICAÇÕES BRASILEIRAS S.A. - TELEBRAS, sociedade de economia mista, "
        "vinculada ao Ministério das Comunicações, com sede no SIG, Quadra 04, Bloco A, "
        "Salas 201 a 224, Edifício Capital Financial Center, CEP nº 70.610-440, inscrita no CNPJ sob "
        "o nº 00.336.701/0001-04, com seus atos constitutivos devidamente arquivados na Junta "
        "Comercial do Distrito Federal, sob o nº 7.665, em 20/02/1978, publicada no Diário Oficial da "
        "União de 13/03/1978, doravante denominada TELEBRAS, neste ato representada por seu "
        "[campo a preencher], matrícula [campo a preencher], e por seu [campo a preencher], "
        f"nos termos da Diretriz nº 229/2018, apostila o Contrato nº {contrato}, "
        "celebrado com a empresa [campo a preencher], doravante denominada CONTRATADA, nos "
        "termos do parágrafo 7º do art. 81 da Lei nº 13.303, de 30 de junho de 2016, legislação "
        "complementar, e:"
    )

    document.add_paragraph("")

    # ── CONSIDERANDO ──────────────────────────────────────────────
    p_cons = document.add_paragraph()
    p_cons.add_run("CONSIDERANDO:").bold = True

    indice = res.get("indice", adm.get("indice", "[campo a preencher]"))
    variacao_acumulada = res.get("variacao_acumulada", None)
    fator_acum = res.get("fator_acumulado", adm.get("fator_acumulado", 1.0))
    if variacao_acumulada is None:
        try: variacao_acumulada = float(fator_acum) - 1
        except Exception: variacao_acumulada = 0.0

    ciclos = _ciclos_para_minuta(adm, res)
    ciclos_validos = []
    ciclos_ja = set()
    for idx, ciclo in enumerate(ciclos, start=1):
        if not isinstance(ciclo, dict): continue
        nome = str(ciclo.get("ciclo") or ciclo.get("Ciclo") or f"C{idx}").strip().upper()
        if not nome or nome in ["C0", "TOTAL"] or nome in ciclos_ja: continue
        ciclos_ja.add(nome)
        if not (bool(ciclo.get("ciclo_ja_concedido", False)) or
                str(ciclo.get("Objeto da análise atual", "")).strip().lower() in ["não", "nao", "false"]):
            ciclos_validos.append(ciclo)

    considerandos_romanos = [
        (
            "I",
            f"A Cláusula Oitava do Contrato nº {contrato}, que prevê o reajuste contratual e "
            "disciplina o marco para concessão, os efeitos financeiros, os reajustes subsequentes e a regra de preclusão;"
        ),
        (
            "II",
            "A deliberação da Diretoria Executiva da Telebras que autorizou o processamento do "
            "presente reajuste contratual;"
        ),
        (
            "III",
            "O Despacho Saneador [campo a preencher], por meio do qual a Gerência de Compras e "
            "Contratos consolidou o histórico do pleito, a admissibilidade, a apuração dos índices, os "
            "efeitos financeiros e os documentos de suporte necessários à formalização do presente apostilamento;"
        ),
        (
            "IV",
            f"A memória de cálculo constante em [campo a preencher], que apurou o reajuste "
            f"acumulado de {percentual(variacao_acumulada, 2)}, assim como o detalhamento dos valores nominais pagos, os valores "
            "devidos, o saldo retroativo a pagar por ciclo e a composição do Valor Total Atualizado estimado do contrato;"
        ),
        (
            "V",
            "A manifestação da CONTRATADA constante de [campo a preencher], pela qual anuiu "
            "com os cálculos apresentados pela TELEBRAS;"
        ),
        (
            "VI",
            "As certidões de regularidade da CONTRATADA juntadas em [campo a preencher] e a "
            "adequação orçamentária registrada em [campo a preencher];",
        ),
    ]

    # Aditivos nos considerandos
    df_ad = res.get("df_aditivos_executivo", res.get("df_aditivos", pd.DataFrame())) if isinstance(res, dict) else pd.DataFrame()
    if isinstance(df_ad, pd.DataFrame) and not df_ad.empty:
        for i_ad, (_, ad) in enumerate(df_ad.head(5).iterrows(), start=7):
            num_romano = _romano(i_ad)
            ident = texto_seguro(ad.get("Aditivo", ad.get("Identificação", f"{i_ad}º Termo Aditivo")), f"{i_ad}º Termo Aditivo")
            ciclo_ref = texto_seguro(ad.get("Ciclo/Marco", ""), "[campo a preencher]")
            val_orig = ad.get("Valor do aditivo na assinatura", ad.get("Valor original da alteração", 0))
            val_atual = ad.get("Valor do aditivo reajustado", ad.get("Valor atualizado da alteração", 0))
            considerandos_romanos.append((
                num_romano,
                f"O {ident}, com valor de {moeda(val_orig)} na data de assinatura e valor atualizado de {moeda(val_atual)}, conforme o ciclo de referência {ciclo_ref};"
            ))
    else:
        considerandos_romanos.append(("VII", "Os aditivos e supressões contratuais, conforme instrumentos constantes do processo;"))

    for num_romano, texto in considerandos_romanos:
        _adicionar_paragrafo_justificado(document, f"{num_romano}. {_limpar_texto_formal(texto)}")

    document.add_paragraph("")

    # ── FORMALIZA-SE ──────────────────────────────────────────────
    p_form = document.add_paragraph()
    p_form.add_run("FORMALIZA-SE O PRESENTE TERMO DE APOSTILA:").bold = True

    document.add_paragraph("")

    # Item 1 — Ciclos (tabela executiva)
    _adicionar_paragrafo_justificado(
        document,
        f"1. Ao Contrato nº {contrato}, em razão da concessão dos reajustes contratuais apurados, nos seguintes termos:"
    )

    tab1 = document.add_table(rows=1, cols=4)
    tab1.style = "Table Grid"
    hdr1 = tab1.rows[0].cells
    for i, h in enumerate(["Ciclo", "Percentual aplicado", "Efeitos financeiros a partir de", "Situação"]):
        hdr1[i].text = h
        hdr1[i].paragraphs[0].runs[0].bold = True

    for ciclo in ciclos_validos:
        nome = str(ciclo.get("ciclo") or ciclo.get("Ciclo", "")).strip().upper()
        pct  = _percentual_ciclo_minuta(ciclo)
        efeito = _efeito_ciclo_minuta(ciclo)
        sit_raw = (ciclo.get("situacao_aplicada") or ciclo.get("Situação aplicada") or
                   ciclo.get("situacao") or ciclo.get("Situação") or "")
        # Limpar emojis e normalizar capitalização para documento formal
        import re as _re
        _sit_limpa = _re.sub(r"[^\w\s,()/-]", "", sit_raw).strip()
        if _sit_limpa.lower() in ("tempestivo", "admissivel com ressalva", "admissível com ressalva"):
            sit_label = _sit_limpa.capitalize() if _sit_limpa else "Tempestivo"
        elif "negoci" in _sit_limpa.lower():
            sit_label = "Negociado"
        else:
            sit_label = _sit_limpa.capitalize() if _sit_limpa else "[campo a preencher]"
        row1 = tab1.add_row().cells
        row1[0].text = nome
        row1[1].text = pct
        row1[2].text = efeito
        row1[3].text = sit_label

    row_acum = tab1.add_row().cells
    row_acum[0].text = "Percentual acumulado apurado"
    row_acum[0].paragraphs[0].runs[0].bold = True
    row_acum[1].text = percentual(variacao_acumulada, 2)
    row_acum[1].paragraphs[0].runs[0].bold = True
    row_acum[2].text = ""
    row_acum[3].text = ""

    document.add_paragraph("")

    # Item 2 — Tabela financeira por ciclo
    val_pago   = res.get("total_pago_faturado", 0)
    val_teorico = res.get("total_devido_reajustado", 0)
    val_retro  = res.get("valor_represado_a_pagar", res.get("delta_acumulado", 0))

    _adicionar_paragrafo_justificado(
        document,
        f"2. A apuração financeira consolidada indicou valor pago efetivo de {moeda(val_pago)} "
        f"e valor teórico calculado de {moeda(val_teorico)}, "
        f"resultando em valor retroativo a pagar de {moeda(val_retro)}, conforme quadro abaixo."
    )

    df_fin_ciclo = res.get("df_financeiro_por_ciclo")
    if isinstance(df_fin_ciclo, pd.DataFrame) and not df_fin_ciclo.empty:
        tab = document.add_table(rows=1, cols=4)
        tab.style = "Table Grid"
        hdr = tab.rows[0].cells
        for i, h in enumerate(["Ciclo", "Valor pago efetivo", "Valor teórico calculado", "Diferença/retroativo"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True

        col_ciclo  = next((c for c in df_fin_ciclo.columns if "ciclo" in str(c).lower()), None)
        col_pago   = next((c for c in df_fin_ciclo.columns if "pago" in str(c).lower() and "efetivo" in str(c).lower()), None)
        col_teo    = next((c for c in df_fin_ciclo.columns if "teórico" in str(c).lower() or "teorico" in str(c).lower()), None)
        col_delta  = next((c for c in df_fin_ciclo.columns if "delta" in str(c).lower() or "diferença" in str(c).lower()), None)

        for _, row_data in df_fin_ciclo.iterrows():
            row = tab.add_row().cells
            row[0].text = str(row_data.get(col_ciclo, "")) if col_ciclo else ""
            row[1].text = moeda(row_data.get(col_pago, 0))  if col_pago  else ""
            row[2].text = moeda(row_data.get(col_teo, 0))   if col_teo   else ""
            row[3].text = moeda(row_data.get(col_delta, 0)) if col_delta else ""

        # Linha total
        row_tot = tab.add_row().cells
        row_tot[0].text = "Total"
        row_tot[0].paragraphs[0].runs[0].bold = True
        row_tot[1].text = moeda(val_pago)
        row_tot[1].paragraphs[0].runs[0].bold = True
        row_tot[2].text = moeda(val_teorico)
        row_tot[2].paragraphs[0].runs[0].bold = True
        row_tot[3].text = moeda(val_retro)
        row_tot[3].paragraphs[0].runs[0].bold = True

    document.add_paragraph("")

    # Item 3 — Composição detalhada (tabela Ref/Descrição/Valor)
    df_comp = res.get("df_composicao_valor_total")
    val_total = res.get("valor_atualizado_contrato", res.get("valor_global_estoque", 0))

    _adicionar_paragrafo_justificado(
        document,
        "3. Para fins de consolidação contratual, a composição do Valor Total Atualizado do Contrato, "
        "considerando execução realizada por ciclo, saldo remanescente atualizado e aditivos, "
        "está demonstrada na tabela a seguir:"
    )

    if isinstance(df_comp, pd.DataFrame) and not df_comp.empty:
        tab3 = document.add_table(rows=1, cols=3)
        tab3.style = "Table Grid"
        hdr3 = tab3.rows[0].cells
        for i, h in enumerate(["Ref.", "Descrição", "Valor"]):
            hdr3[i].text = h
            hdr3[i].paragraphs[0].runs[0].bold = True
        letras = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        col_desc = next((c for c in df_comp.columns if "descri" in str(c).lower() or "parcela" in str(c).lower()), df_comp.columns[0])
        col_val  = next((c for c in df_comp.columns if "valor" in str(c).lower()), df_comp.columns[-1])
        for i, (_, r_data) in enumerate(df_comp.iterrows()):
            row3 = tab3.add_row().cells
            row3[0].text = letras[i] if i < len(letras) else str(i+1)
            row3[1].text = str(r_data.get(col_desc, ""))
            row3[2].text = moeda(r_data.get(col_val, 0))
        row_tot3 = tab3.add_row().cells
        row_tot3[0].text = ""
        row_tot3[1].text = "Valor Total Atualizado"
        row_tot3[1].paragraphs[0].runs[0].bold = True
        row_tot3[2].text = moeda(val_total)
        row_tot3[2].paragraphs[0].runs[0].bold = True
    else:
        _adicionar_paragrafo_justificado(
            document,
            f"O Valor Total Atualizado do Contrato, composto pela execução atualizada por ciclo e saldo remanescente, "
            f"corresponde a {moeda(val_total)}."
        )

    document.add_paragraph("")

    # Item 4 — Sumário sintético
    _adicionar_paragrafo_justificado(
        document,
        f"4. De forma sintética, o Valor Total Atualizado do Contrato de {moeda(val_total)} "
        "compreende a execução realizada em cada ciclo, atualizada pelo respectivo fator de reajuste, "
        "acrescida do saldo remanescente atualizado e dos aditivos/supressões, conforme apuração constante da memória de cálculo."
    )

    document.add_paragraph("")

    # Item 5 — Aditivos
    if isinstance(df_ad, pd.DataFrame) and not df_ad.empty:
        partes_aditivos = []
        for _, ad in df_ad.head(5).iterrows():
            ident = texto_seguro(ad.get("Aditivo", ad.get("Identificação", "")), "Aditivo")
            ciclo_ref = texto_seguro(ad.get("Ciclo/Marco", ""), "")
            val_orig = ad.get("Valor do aditivo na assinatura", ad.get("Valor original da alteração", 0))
            val_atual = ad.get("Valor do aditivo reajustado", ad.get("Valor atualizado da alteração", 0))
            partes_aditivos.append(
                f"o {ident}, atualizado a {percentual(variacao_acumulada,2)}, "
                f"gerou impacto de {moeda(val_atual)}"
                + (f", conforme o ciclo de referência {ciclo_ref}" if ciclo_ref else "")
            )
        texto_adi = "5. Quanto aos aditivos e supressões considerados, registra-se que " + "; e que ".join(partes_aditivos) + "."
    else:
        texto_adi = "5. Não foram identificados aditivos ou supressões específicos para esta análise."
    _adicionar_paragrafo_justificado(document, texto_adi)

    document.add_paragraph("")

    # Item 6 — Demais cláusulas
    _adicionar_paragrafo_justificado(
        document,
        "6. Permanecem inalteradas e em pleno vigor as demais cláusulas e condições do Contrato "
        "e de seus instrumentos posteriores não modificadas por este Termo de Apostila."
    )

    document.add_paragraph("")

    # Item 7 — Garantia
    _adicionar_paragrafo_justificado(
        document,
        "7. A CONTRATADA deverá atualizar a garantia contratual, prevista na Cláusula Décima do "
        "Contrato, no prazo contratualmente estabelecido, observado o novo valor após a formalização "
        "deste Termo de Apostila."
    )

    document.add_paragraph("")

    # Item 8 — Vinculação ao processo
    _adicionar_paragrafo_justificado(
        document,
        f"8. O presente apostilamento vincula-se, para todos os fins, aos documentos instruídos no "
        f"Processo {processo}."
    )

    document.add_paragraph("")
    p_data = document.add_paragraph("Brasília, [campo a preencher].")
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")

    # Assinaturas
    for cargo, diretoria in [
        ("[campo a preencher]", "Gerente\nGerência de Compras e Contratos"),
        ("[campo a preencher]", "Diretor\nDiretoria Técnico-operacional"),
        ("[campo a preencher]", "Presidente\nPresidência"),
    ]:
        p_ass = document.add_paragraph()
        p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_nome = p_ass.add_run(cargo)
        r_nome.bold = True
        p_cargo = document.add_paragraph()
        p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for linha in diretoria.split("\n"):
            p_cargo.add_run(linha + "\n")
        document.add_paragraph("")

    _destacar_campos_preencher(document)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

    titulo = document.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("MINUTA DE TERMO DE APOSTILAMENTO")
    run.bold = True
    run.font.size = Pt(13)

    contrato = _placeholder(adm.get("contrato") or adm.get("numero_contrato") or res.get("contrato") or res.get("numero_contrato"))
    p = document.add_paragraph(f"Contrato nº {contrato}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")

    _adicionar_paragrafo_justificado(
        document,
        "A TELECOMUNICAÇÕES BRASILEIRAS S.A. - TELEBRAS, sociedade de economia mista, vinculada ao Ministério das Comunicações, "
        "com sede no SIG, Quadra 04, Bloco A, Salas 201 a 224, Edifício Capital Financial Center, CEP nº 70.610-440, inscrita no CNPJ sob o "
        "n.º 00.336.701/0001-04, doravante denominada TELEBRAS, neste ato representada por [campo a preencher], Matrícula [campo a preencher], "
        "e por seu [campo a preencher], Matrícula [campo a preencher], nos termos da Diretriz nº 229/2018, apostila o Contrato nº [campo a preencher], "
        "celebrado com a empresa [campo a preencher], doravante denominada CONTRATADA, com fundamento no parágrafo 7º do art. 81 da Lei nº 13.303, "
        "de 30 de junho de 2016, na legislação aplicável, no Regulamento de Licitações e Contratos da Telebras e nos documentos constantes do processo."
    )

    document.add_paragraph("")
    document.add_paragraph("CONSIDERANDO:")

    indice = res.get("indice", adm.get("indice", "[campo a preencher]"))
    fator_acumulado = res.get("fator_acumulado", adm.get("fator_acumulado", adm.get("fator", 1.0)))
    variacao_acumulada = res.get("variacao_acumulada", None)
    if variacao_acumulada is None:
        try:
            variacao_acumulada = float(fator_acumulado) - 1
        except Exception:
            variacao_acumulada = 0.0

    considerandos = [
        "A Cláusula Oitava do Contrato nº [campo a preencher], que disciplina o reajuste contratual, os ciclos de apuração, a admissibilidade dos pedidos e os respectivos efeitos financeiros;",
        "A necessidade de distinguir o histórico já formalizado anteriormente do objeto da presente análise, evitando duplicidade de contagem ou sobreposição de efeitos financeiros;",
        "A memória de cálculo constante em [campo a preencher], que apurou os ciclos de reajuste, os percentuais aplicáveis e os efeitos financeiros correspondentes;",
        f"O índice contratual utilizado na análise, qual seja {indice}, e o percentual acumulado apurado de {percentual(variacao_acumulada, 2)};",
        "As informações encaminhadas pela área gestora/fiscal do contrato quanto à execução, ao saldo remanescente, aos itens contratuais e aos documentos de suporte da apuração;",
        "A adequação orçamentária constante em [campo a preencher], sem prejuízo das validações orçamentárias e financeiras cabíveis antes da execução da despesa;",
        "A necessidade de atualização da garantia contratual, quando aplicável, em razão da alteração do valor contratual atualizado;",
        "A manifestação de concordância da CONTRATADA constante em [campo a preencher], quando aplicável.",
    ]

    for item in considerandos:
        document.add_paragraph(_limpar_texto_formal(item), style="List Bullet")

    contexto = res.get("contexto_contratual_anterior", {}) or adm.get("contexto_contratual_anterior", {}) or {}
    ciclos = _ciclos_para_minuta(adm, res)

    ciclos_validos = []
    ciclos_historicos = []
    ciclos_ja = set()

    for idx, ciclo in enumerate(ciclos, start=1):
        if not isinstance(ciclo, dict):
            continue
        nome = str(ciclo.get("ciclo") or ciclo.get("Ciclo") or f"C{idx}").strip().upper()
        if not nome or nome in ["C0", "TOTAL", "CICLO"] or nome in ciclos_ja:
            continue
        ciclos_ja.add(nome)
        if bool(ciclo.get("ciclo_ja_concedido", False)) or str(ciclo.get("Objeto da análise atual", "")).strip().lower() in ["não", "nao", "false"]:
            ciclos_historicos.append(ciclo)
        else:
            ciclos_validos.append(ciclo)

    _adicionar_item_numerado(
        document,
        1,
        "O presente Termo de Apostilamento tem por objeto formalizar o reajuste contratual apurado no âmbito da análise atual, observada a separação entre histórico formalizado anteriormente e ciclos que compõem o objeto deste apostilamento."
    )

    _adicionar_subitem(
        document,
        "1.1.",
        "O C0 corresponde ao ciclo-base inicial do contrato e não recebe reajuste, servindo apenas como referência de partida para a apuração dos ciclos subsequentes."
    )

    _adicionar_subitem(
        document,
        "1.2.",
        "Os ciclos C1, C2, C3 e seguintes devem ser lidos como ciclos contratuais de reajuste, conforme a linha temporal de admissibilidade e efeitos financeiros prevista no contrato."
    )

    _adicionar_item_numerado(
        document,
        2,
        "Histórico formalizado anterior."
    )

    valor_formalizado = contexto.get("valor_formalizado_anterior", "")
    ultimo_ciclo = contexto.get("ultimo_ciclo_concedido", "")
    data_pedido_ultimo = contexto.get("data_pedido_ultimo_ciclo", "")
    obs_historico = contexto.get("observacao_historico", "")

    if valor_formalizado or ultimo_ciclo or data_pedido_ultimo or obs_historico or ciclos_historicos:
        if valor_formalizado:
            _adicionar_subitem(document, "2.1.", f"Valor contratual formalizado antes desta análise: {moeda(valor_formalizado)}.")
        if ultimo_ciclo and str(ultimo_ciclo).strip().upper() not in ["C0 / NENHUM", "C0/NENHUM", "C0", "NENHUM", ""]:
            _adicionar_subitem(document, "2.2.", f"Último ciclo já concedido/formalizado antes desta análise: {ultimo_ciclo}.")
        if data_pedido_ultimo:
            _adicionar_subitem(document, "2.3.", f"Data do pedido do último ciclo concedido/formalizado: {formatar_data_br(data_pedido_ultimo)}.")
        if obs_historico:
            _adicionar_subitem(document, "2.4.", f"Observação do histórico anterior: {obs_historico}.")
        if ciclos_historicos:
            for idx, ciclo in enumerate(ciclos_historicos, start=1):
                nome = _nome_ciclo_minuta(ciclo, idx)
                pct = _percentual_ciclo_minuta(ciclo)
                efeito = _efeito_ciclo_minuta(ciclo)
                _adicionar_subitem(
                    document,
                    f"2.{idx + 4}.",
                    f"{nome} consta como ciclo já concedido/formalizado anteriormente, preservado como histórico/âncora, com percentual de {pct} e efeitos financeiros a partir de {efeito}, sem compor o objeto novo desta análise."
                )
    else:
        _adicionar_subitem(
            document,
            "2.1.",
            "Não foi informado histórico formalizado anterior específico para esta análise, sem prejuízo da conferência dos instrumentos já constantes do processo."
        )

    _adicionar_item_numerado(
        document,
        3,
        "Objeto da análise atual: ciclos, admissibilidade, percentuais e efeitos financeiros."
    )

    subitem_3 = 1
    if ciclos_validos:
        for idx, ciclo in enumerate(ciclos_validos, start=1):
            nome = _nome_ciclo_minuta(ciclo, idx)
            data_base = formatar_data_br(ciclo.get("data_base", ciclo.get("Data-base", "")))
            data_pedido = formatar_data_br(ciclo.get("data_pedido", ciclo.get("Data do pedido", "")))
            efeito = _efeito_ciclo_minuta(ciclo)
            pct = _percentual_ciclo_minuta(ciclo)
            situacao_bruta = (
                ciclo.get("situacao_aplicada")
                or ciclo.get("Situação aplicada")
                or ciclo.get("situacao")
                or ciclo.get("Situação")
                or ""
            )
            situacao = _status_relatorio(situacao_bruta)
            _adicionar_subitem(
                document,
                f"3.{subitem_3}.",
                f"{nome}: data-base {data_base or '[campo a preencher]'}, pedido em {data_pedido or '[campo a preencher]'}, classificação {situacao}, percentual aplicado {pct} e efeitos financeiros a partir de {efeito}."
            )
            subitem_3 += 1
    else:
        _adicionar_subitem(
            document,
            f"3.{subitem_3}.",
            "[campo a preencher: informar os ciclos que compõem o objeto da análise atual, com data-base, data do pedido, classificação, percentual aplicado e início dos efeitos financeiros]."
        )
        subitem_3 += 1

    _adicionar_subitem(
        document,
        f"3.{subitem_3}.",
        f"O percentual acumulado considerado na presente análise corresponde a {percentual(variacao_acumulada, 2)}."
    )
    subitem_3 += 1

    if not (eh_modo_consumo_itens_ciclo(res) or eh_modo_reduzido_itens(res)):
        _adicionar_subitem(
            document,
            f"3.{subitem_3}.",
            f"O valor represado a pagar apurado corresponde a {moeda(res.get('valor_represado_a_pagar', res.get('delta_acumulado', 0)))}."
        )
        subitem_3 += 1


    _adicionar_item_numerado(
        document,
        4,
        "Consolidação do Valor Total Atualizado do Contrato."
    )

    _adicionar_subitem(
        document,
        "4.1.",
        "O Valor Total Atualizado do Contrato foi composto pela execução atualizada por ciclo somada ao saldo remanescente atualizado, preservada a metodologia de consolidação adotada na memória de cálculo."
    )

    _adicionar_subitem(
        document,
        "4.2.",
        f"Valor original do contrato: {moeda(res.get('valor_original_contrato', 0))}."
    )

    _adicionar_subitem(
        document,
        "4.3.",
        f"Execução atualizada por ciclo: {moeda(res.get('valor_executado_atualizado', res.get('total_devido_reajustado', 0)))}."
    )

    _adicionar_subitem(
        document,
        "4.4.",
        f"Saldo remanescente atualizado: {moeda(res.get('remanescente_reajustado', 0))}."
    )

    _adicionar_subitem(
        document,
        "4.5.",
        f"Valor Total Atualizado do Contrato: {moeda(res.get('valor_atualizado_contrato', res.get('valor_global_estoque', 0)))}."
    )

    _adicionar_item_numerado(
        document,
        5,
        "Aditivos e supressões."
    )

    df_ad = res.get("df_aditivos_executivo", res.get("df_aditivos", pd.DataFrame())) if isinstance(res, dict) else pd.DataFrame()
    proximo_subitem_aditivos = 2

    if isinstance(df_ad, pd.DataFrame) and not df_ad.empty:
        _adicionar_subitem(
            document,
            "5.1.",
            "Os aditivos e supressões registrados na análise foram considerados para fins de controle formal e governança do valor contratual, observada a classificação por ciclo/marco financeiro."
        )
        limite = min(len(df_ad), 10)
        for idx, (_, ad) in enumerate(df_ad.head(limite).iterrows(), start=2):
            identificacao = texto_seguro(ad.get("Aditivo", ad.get("Identificação", "Aditivo/Supressão")), "Aditivo/Supressão")
            ciclo = texto_seguro(ad.get("Ciclo/Marco", ""), "[campo a preencher]")
            tratamento = texto_seguro(ad.get("Tratamento do aditivo", ""), "[campo a preencher]")
            valor = ad.get("Valor do aditivo reajustado", ad.get("Valor atualizado da alteração", ad.get("Valor original da alteração", 0)))
            _adicionar_subitem(
                document,
                f"5.{idx}.",
                f"{identificacao}: ciclo/marco {ciclo}, tratamento {tratamento}, valor de referência {moeda(valor)}."
            )
            proximo_subitem_aditivos = idx + 1
    else:
        _adicionar_subitem(
            document,
            "5.1.",
            "Não foram identificados aditivos ou supressões específicos na base processada, sem prejuízo da conferência dos instrumentos já formalizados no processo."
        )
        proximo_subitem_aditivos = 2

    _adicionar_subitem(
        document,
        f"5.{proximo_subitem_aditivos}.",
        "Os aditivos e supressões não devem ser somados de forma autônoma ao Valor Total Atualizado quando seus efeitos já estiverem refletidos na execução atualizada ou no saldo remanescente."
    )


    _adicionar_item_numerado(
        document,
        6,
        "Garantia contratual."
    )

    _adicionar_subitem(
        document,
        "6.1.",
        "A CONTRATADA deverá atualizar a garantia contratual, quando exigida pelo contrato, de modo compatível com o Valor Total Atualizado do Contrato e com os instrumentos já formalizados."
    )


    _adicionar_item_numerado(
        document,
        7,
        "Ratificação."
    )

    _adicionar_subitem(
        document,
        "7.1.",
        "Permanecem inalteradas e em pleno vigor as demais cláusulas e condições do Contrato e de seus instrumentos posteriores não modificadas por este Termo de Apostilamento."
    )

    _adicionar_subitem(
        document,
        "7.2.",
        "O presente apostilamento vincula-se, para todos os fins, aos documentos [campo a preencher] instruídos no Processo [campo a preencher]."
    )


    document.add_paragraph("")
    document.add_paragraph("Brasília/DF, [Data].")
    document.add_paragraph("")
    document.add_paragraph("TELECOMUNICAÇÕES BRASILEIRAS S.A. - TELEBRAS")
    document.add_paragraph("Representante Legal 1")
    document.add_paragraph("")
    document.add_paragraph("TELECOMUNICAÇÕES BRASILEIRAS S.A. - TELEBRAS")
    document.add_paragraph("Representante Legal 2")

    _destacar_campos_preencher(document)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()



def gerar_despacho_saneador_docx(adm, res):
    """
    Despacho Saneador — template conforme padrão Telebras/GCC.
    Estrutura: 14 itens + 4 quadros.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except Exception as exc:
        raise RuntimeError("python-docx não instalado.") from exc

    adm    = adm or {}
    res    = res or {}
    params = res.get("params_v10") or res.get("params") or {}

    # ── Campos básicos ─────────────────────────────────────────────
    contrato      = _placeholder(adm.get("contrato") or res.get("contrato") or params.get("contrato"))
    processo      = _placeholder(adm.get("processo") or res.get("processo") or params.get("processo"))
    indice        = str(res.get("indice", adm.get("indice", "[campo a preencher]")))
    fator         = float(res.get("fator_acumulado", 1.0) or 1.0)
    variacao      = float(res.get("variacao_acumulada", fator - 1.0) or (fator - 1.0))
    val_retro     = float(res.get("valor_represado_a_pagar", res.get("delta_acumulado", 0)) or 0)
    val_total     = float(res.get("valor_atualizado_contrato", res.get("valor_global_estoque", 0)) or 0)
    val_pago      = float(res.get("valor_pago_efetivo", res.get("total_pago_faturado", 0)) or 0)
    val_teorico   = float(res.get("valor_teorico_calculado", res.get("total_devido_reajustado", 0)) or 0)
    val_original  = float(params.get("valor_original_do_contrato", res.get("valor_original_contrato", 0)) or 0)

    # Ciclos
    ciclos_lista = adm.get("ciclos", res.get("ciclos", []))
    n_ciclos = len(ciclos_lista) if ciclos_lista else 0
    if ciclos_lista:
        nomes = [str(c.get("ciclo", c) if isinstance(c, dict) else c).strip().upper() for c in ciclos_lista]
        ciclos_ref = ", ".join(nomes)
    else:
        ciclos_ref = "[campo a preencher]"

    document = Document()
    _aplicar_estilo_docx(document)

    p_titulo = document.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_titulo.add_run("DESPACHO SANEADOR")
    r.bold = True
    r.font.size = Pt(13)

    document.add_paragraph("")

    for linha, bold in [
        (f"PROCESSO Nº {processo}", True),
        (f"CONTRATO Nº {contrato}", True),
    ]:
        p = document.add_paragraph()
        p.add_run(linha).bold = bold

    document.add_paragraph("")

    # Item 1
    _adicionar_paragrafo_justificado(
        document,
        "1. Este despacho saneador consolida os elementos documentais, financeiros e formais "
        "necessários à instrução do Termo de Apostila destinado ao registro de reajuste contratual, "
        "com a finalidade de demonstrar a regularidade mínima da instrução antes da formalização."
    )
    document.add_paragraph("")

    # Item 2 — pleito e verificação de anualidade
    pedidos_txt = "[campo a preencher]"
    datas_pedido = []
    for c in ciclos_lista:
        if isinstance(c, dict):
            nome = str(c.get("ciclo", "")).strip().upper()
            data = str(c.get("data_pedido", "")).strip()
            if nome and data:
                datas_pedido.append(f"{nome} em {data}")
    if datas_pedido:
        pedidos_txt = ", ".join(datas_pedido)

    data_base = str((ciclos_lista[0].get("data_base", "") if ciclos_lista and isinstance(ciclos_lista[0], dict) else "")).strip() or "[campo a preencher]"

    _adicionar_paragrafo_justificado(
        document,
        f"2. A contratada apresentou pleito de reajuste por meio de [campo a preencher]. "
        f"Para fins de verificação da anualidade, foram consideradas a data da proposta em {data_base}, "
        f"o índice contratual {indice} e as datas de pedido registradas na coleta: {pedidos_txt}."
    )
    document.add_paragraph("")

    # Item 3 — acordos ou ressalvas da instrução
    _ciclo_concedido = next(
        (str(c.get("ciclo","")).strip().upper() for c in ciclos_lista
         if isinstance(c, dict) and
         str(c.get("situacao", c.get("situacao_aplicada",""))).lower() in
         ("admitido por negociação", "admitido por negociacao", "negociado")),
        None
    )
    if _ciclo_concedido:
        _txt3 = (
            f"Acordou-se na concessão do {_ciclo_concedido} mediante negociação entre as partes, "
            "tendo a contratada manifestado concordância com os percentuais propostos, "
            "conforme registrado em [campo a preencher]."
        )
    else:
        _txt3 = (
            "Os ciclos objeto desta análise foram admitidos nos termos da cláusula contratual de "
            "reajuste, sem ressalvas quanto à admissibilidade. Eventuais negociações ou "
            "ajustes de percentual deverão ser registrados neste item antes da formalização."
        )
    _adicionar_paragrafo_justificado(document, f"3. {_limpar_texto_formal(_txt3)}")
    document.add_paragraph("")

    # Item 4 + Quadro 1
    _adicionar_paragrafo_justificado(
        document,
        f"4. A análise de reajuste considerou {n_ciclos} ciclo(s), "
        f"com variação acumulada de {percentual(variacao, 2)}. "
        f"O valor original do contrato informado foi de {moeda(val_original)}."
        if val_original > 0 else
        f"4. A análise de reajuste considerou {n_ciclos} ciclo(s), "
        f"com variação acumulada de {percentual(variacao, 2)}."
    )
    document.add_paragraph("")

    _docx_add_titulo_secao(document, "Quadro 1 – Síntese dos ciclos de reajuste")
    if ciclos_lista:
        headers_q1 = ["Ciclo", "Data-base", "Data do pedido", "Início financeiro", "Fim financeiro", "Situação", "Percentual aplicado"]
        rows_q1 = []
        for c in ciclos_lista:
            if not isinstance(c, dict):
                continue
            sit = str(c.get("situacao", c.get("situacao_aplicada", "")) or "").strip()
            pct_val = float(c.get("percentual_aplicado", 0) or 0)
            rows_q1.append([
                str(c.get("ciclo", "")).strip().upper(),
                str(c.get("data_base", "")).strip(),
                str(c.get("data_pedido", "")).strip(),
                str(c.get("inicio_financeiro", "")).strip(),
                str(c.get("fim_financeiro", "")).strip(),
                sit.capitalize() if sit else "[campo a preencher]",
                percentual(pct_val / 100 if pct_val > 1 else pct_val, 2),
            ])
        _docx_tabela_executiva(document, headers_q1, rows_q1)
    else:
        document.add_paragraph("[Quadro 1 não disponível — sem dados de ciclos]")

    document.add_paragraph("")

    # Item 5 + Quadro 2
    _adicionar_paragrafo_justificado(
        document,
        f"5. A apuração financeira consolidada indicou valor pago efetivo de {moeda(val_pago)} "
        f"e valor teórico calculado de {moeda(val_teorico)}, "
        f"resultando em valor retroativo a pagar de {moeda(val_retro)}."
    )
    document.add_paragraph("")

    _docx_add_titulo_secao(document, "Quadro 2 – Apuração financeira por ciclo")
    df_fin = res.get("df_financeiro_por_ciclo")
    import pandas as pd
    if isinstance(df_fin, pd.DataFrame) and not df_fin.empty:
        headers_q2 = ["Ciclo", "Valor pago efetivo", "Valor teórico calculado", "Diferença/retroativo"]
        rows_q2 = []
        for _, row in df_fin.iterrows():
            ciclo_nome = str(row.get("Ciclo", "")).strip().upper()
            if ciclo_nome == "TOTAL":
                continue
            vpe  = float(row.get("Valor pago efetivo", 0) or 0)
            vtc  = float(row.get("Valor teórico calculado", 0) or 0)
            delta = float(row.get("Delta do ciclo", vtc - vpe) or 0)
            rows_q2.append([ciclo_nome, moeda(vpe), moeda(vtc), moeda(delta)])
        rows_q2.append(["Total", moeda(val_pago), moeda(val_teorico), moeda(val_retro)])
        _docx_tabela_executiva(document, headers_q2, rows_q2, total_last=True)
    else:
        document.add_paragraph("[Quadro 2 não disponível — sem dados financeiros por ciclo]")

    document.add_paragraph("")

    # Item 6 + Quadro 3
    _adicionar_paragrafo_justificado(
        document,
        "6. Para fins de consolidação contratual, [campo a preencher — premissa do corte operacional e memória fiscal adotada]."
    )
    document.add_paragraph("")

    _docx_add_titulo_secao(document, "Quadro 3 – Memória fiscal do Valor Total Atualizado Estimado")
    df_comp = res.get("df_composicao_valor_total")
    if isinstance(df_comp, pd.DataFrame) and not df_comp.empty:
        col_desc = next((c for c in ("Componente", "Parcela", "Descrição") if c in df_comp.columns), None)
        col_val  = next((c for c in ("Valor", "Valor R$", "Total R$") if c in df_comp.columns), None)
        if col_desc and col_val:
            headers_q3 = ["Descrição", "Valor"]
            rows_q3 = []
            for _, r3 in df_comp.iterrows():
                v = float(r3.get(col_val, 0) or 0)
                rows_q3.append([str(r3.get(col_desc, "")), moeda(v)])
            rows_q3.append(["Valor total do contrato estimado", moeda(val_total)])
            _docx_tabela_executiva(document, headers_q3, rows_q3, total_last=True)
        else:
            document.add_paragraph(f"Valor Total Atualizado Estimado: {moeda(val_total)}")
    else:
        document.add_paragraph(f"[Quadro 3 não disponível] — Valor Total Atualizado Estimado: {moeda(val_total)}")

    document.add_paragraph("")

    # Item 7 + Quadro 4 composição sintética
    _adicionar_paragrafo_justificado(
        document,
        "7. De forma didática, o Valor Total Atualizado Estimado do Contrato pode ser lido pela seguinte composição:"
    )
    document.add_paragraph("")

    df_exec = res.get("df_execucao_atualizada")
    if isinstance(df_exec, pd.DataFrame) and not df_exec.empty:
        headers_q4 = ["Parcela", "Valor"]
        rows_q4 = []
        col_parc = next((c for c in ("Parcela", "Componente", "Ciclo") if c in df_exec.columns), None)
        col_valu = next((c for c in ("Valor executado atualizado", "Valor teórico calculado", "Valor R$") if c in df_exec.columns), None)
        if col_parc and col_valu:
            for _, r4 in df_exec.iterrows():
                rows_q4.append([str(r4.get(col_parc, "")), moeda(float(r4.get(col_valu, 0) or 0))])
        rows_q4.append(["Valor Total Atualizado Estimado", moeda(val_total)])
        _docx_tabela_executiva(document, headers_q4, rows_q4, total_last=True)
    else:
        document.add_paragraph(f"Valor Total Atualizado Estimado: {moeda(val_total)}")

    document.add_paragraph("")

    # Item 8 — aditivos
    aditivos = res.get("aditivos") or res.get("df_aditivos_normalizado")
    if isinstance(aditivos, pd.DataFrame) and not aditivos.empty:
        linhas_adi = []
        for _, ra in aditivos.iterrows():
            ident = str(ra.get("identificacao", ra.get("Identificação", "")) or "").strip()
            tipo  = str(ra.get("tipo", ra.get("Tipo", "")) or "").strip()
            vt    = float(ra.get("valor_original", ra.get("Valor original", 0)) or 0)
            incorp = str(ra.get("incorporar", ra.get("Incorporar no Valor Total?", "")) or "").strip().lower()
            if ident or abs(vt) > 0.01:
                linhas_adi.append(f"{ident} ({tipo}), impacto de {moeda(vt)}{', incorporado ao valor total' if incorp in ('sim','s','true') else ''}")
        if linhas_adi:
            _adicionar_paragrafo_justificado(
                document,
                "8. Quanto aos aditivos e supressões, registra-se: " + "; e ".join(linhas_adi) + "."
            )
        else:
            _adicionar_paragrafo_justificado(document, "8. Não foram registrados aditivos ou supressões contratuais no período objeto desta análise.")
    else:
        _adicionar_paragrafo_justificado(document, "8. Não foram registrados aditivos ou supressões contratuais no período objeto desta análise.")
    document.add_paragraph("")

    # Itens 3 e 9-14 — redação padrão baseada em modelo institucional Telebras
    # Item 3 já foi inserido acima; aqui complementamos com itens 9-14
    itens_restantes = [
        ("9",
         f"Foi realizada a adequação orçamentária necessária ao prosseguimento da instrução, "
         f"no valor de {moeda(val_retro)}, conforme documento [campo a preencher]."),
        ("10",
         "As certidões de regularidade fiscal, trabalhista e previdenciária da contratada "
         "estão presentes no processo, em [campo a preencher]."),
        ("11",
         f"A contratada manifestou concordância com os valores apurados "
         f"conforme registrado em [campo a preencher]."),
        ("12",
         "A contratada foi informada da necessidade de apresentação do endosso complementar "
         "da garantia contratual, quando aplicável, observando-se o prazo e as condições "
         "previstos no contrato."),
        ("13",
         "Após atualizações e alinhamentos internos, alguns documentos anteriormente "
         "instruídos mostram-se desatualizados e devem ser desconsiderados: [campo a preencher]. "
         "Caso não haja documentos nessa situação, este item não se aplica."),
        ("14",
         "Diante do exposto, estando conferidos os elementos documentais, financeiros e "
         "formais acima indicados, e inexistindo pendência crítica impeditiva, a instrução "
         "poderá prosseguir para formalização do Termo de Apostila, observadas as alçadas "
         "competentes e os procedimentos internos aplicáveis."),
    ]
    for num, texto in itens_restantes:
        _adicionar_paragrafo_justificado(document, f"{num}. {_limpar_texto_formal(texto)}")
        document.add_paragraph("")

    # Quadro 4 — síntese dos principais valores
    _docx_add_titulo_secao(document, "Quadro 4 – Síntese dos principais valores")
    # Calcular valor original a partir de itens se não vier de params
    if not val_original > 0:
        df_itens_raw = res.get("df_itens")
        if isinstance(df_itens_raw, __import__('pandas').DataFrame) and not df_itens_raw.empty:
            _tot = 0.0
            for _, _ir in df_itens_raw.iterrows():
                _qtd = float(_ir.get("Quantidade contratada C0", _ir.get("qtd_c0", 0)) or 0)
                _vu  = float(_ir.get("Valor unitário original C0", _ir.get("vu_c0", 0)) or 0)
                _vt  = float(_ir.get("Valor total original C0", _ir.get("vt_c0", 0)) or 0)
                val_original += _vt if _vt > 0 else round(_qtd * _vu, 2)

    headers_sint = ["Descrição", "Valor"]
    rows_sint = [
        ("Valor original do contrato",      moeda(val_original) if val_original > 0 else "[campo a preencher]"),
        ("Variação acumulada do reajuste",   percentual(variacao, 2)),
        ("Valor retroativo/represado a pagar", moeda(val_retro)),
        ("Valor Total Atualizado Estimado do Contrato", moeda(val_total)),
        ("Adequação orçamentária registrada", "[campo a preencher]"),
    ]
    for c in ciclos_lista:
        if isinstance(c, dict):
            nome = str(c.get("ciclo", "")).strip().upper()
            retro_c = float(c.get("delta", c.get("delta_ciclo", c.get("retroativo", 0))) or 0)
            if nome and abs(retro_c) > 0.01:
                rows_sint.insert(-2, (f"Retroativo {nome}", moeda(retro_c)))
    _docx_tabela_executiva(document, headers_sint, rows_sint)

    document.add_paragraph("")
    document.add_paragraph("")

    for alinhamento, texto, negrito in [
        (WD_ALIGN_PARAGRAPH.CENTER, "Brasília, [campo a preencher]", False),
        (WD_ALIGN_PARAGRAPH.CENTER, "", False),
        (WD_ALIGN_PARAGRAPH.CENTER, "_" * 48, False),
        (WD_ALIGN_PARAGRAPH.CENTER, "[campo a preencher]", False),
        (WD_ALIGN_PARAGRAPH.CENTER, "Gerente de Compras e Contratos", True),
        (WD_ALIGN_PARAGRAPH.CENTER, "Telecomunicações Brasileiras S.A. — Telebras", False),
    ]:
        p = document.add_paragraph()
        p.alignment = alinhamento
        run = p.add_run(texto)
        run.bold = negrito

    _destacar_campos_preencher(document)

    buf = BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.getvalue()


def gerar_previsao_orcamentaria_docx(adm, res):
    """
    Memorando de Adequação Orçamentária.
    Lógica:
        - Retroativo: res["valor_represado_a_pagar"] — ano corrente
        - Média últimos 6 pagamentos: df_financeiro_por_ciclo ou df_financeiro_normalizado
        - Meses futuros: última_competência + 1 → vigência_final (params["vigencia_final"])
        - Diferença futura: (média × fator − média) × meses_futuros
        - Complementação: retroativo + diferença_futura
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except Exception as exc:
        raise RuntimeError("python-docx não instalado.") from exc

    from datetime import date, datetime
    import math

    adm   = adm or {}
    res   = res or {}
    params = res.get("params_v10") or res.get("params") or {}

    # ── Campos básicos ─────────────────────────────────────────────
    contrato   = _placeholder(adm.get("contrato") or res.get("contrato") or params.get("contrato"))
    processo   = _placeholder(adm.get("processo") or res.get("processo") or params.get("processo"))
    val_retro  = float(res.get("valor_represado_a_pagar", res.get("delta_acumulado", 0)) or 0)
    fator      = float(res.get("fator_acumulado", 1.0) or 1.0)
    ano_atual  = date.today().year

    # ── Ciclos referência ──────────────────────────────────────────
    ciclos_lista = adm.get("ciclos", res.get("ciclos", []))
    if ciclos_lista:
        nomes = [str(c.get("ciclo", c) if isinstance(c, dict) else c).strip().upper() for c in ciclos_lista]
        ciclos_ref = ", ".join(nomes[:-1]) + (" e " + nomes[-1] if len(nomes) > 1 else (nomes[0] if nomes else ""))
    else:
        ciclos_ref = "[campo a preencher]"

    # ── Vigência final ─────────────────────────────────────────────
    vig_final_raw = params.get("vigencia_final", "")
    vig_final = None
    if vig_final_raw:
        for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d/%m/%y"):
            try:
                vig_final = datetime.strptime(str(vig_final_raw).strip()[:10], fmt).date()
                break
            except Exception:
                pass
        if vig_final is None:
            try:
                import pandas as pd
                vig_final = pd.to_datetime(vig_final_raw).date()
            except Exception:
                pass

    # ── Última competência e média últimos 6 pagamentos ───────────
    media_sem_reajuste = 0.0
    ultima_competencia = None
    meses_futuros = 0

    _df_fin_a = res.get("df_financeiro_por_ciclo")
    _df_fin_b = res.get("df_financeiro_normalizado")
    import pandas as _pd_orc
    df_fin = _df_fin_a if isinstance(_df_fin_a, _pd_orc.DataFrame) and not _df_fin_a.empty else _df_fin_b
    if isinstance(df_fin, _pd_orc.DataFrame) and not df_fin.empty:
        # Pegar coluna de valor pago
        col_val = None
        for c in ("Valor pago efetivo", "Valor informado pelo fiscal", "Valor líquido considerado"):
            if c in df_fin.columns:
                col_val = c
                break
        col_comp = None
        for c in ("Competência", "competencia"):
            if c in df_fin.columns:
                col_comp = c
                break

        if col_val:
            import pandas as pd
            df_ok = df_fin[df_fin[col_val].apply(lambda x: float(x or 0) > 0.01)].copy()
            if len(df_ok) >= 1:
                ultimos6 = df_ok.tail(6)
                media_sem_reajuste = float(ultimos6[col_val].apply(lambda x: float(x or 0)).mean())
            if col_comp and not df_ok.empty:
                try:
                    ultima_competencia = pd.to_datetime(df_ok[col_comp].iloc[-1]).date()
                except Exception:
                    pass

    # Calcular meses futuros
    if ultima_competencia and vig_final:
        from datetime import date
        prox = date(ultima_competencia.year + (ultima_competencia.month // 12),
                    (ultima_competencia.month % 12) + 1, 1)
        meses_futuros = max(0, (vig_final.year - prox.year) * 12 + (vig_final.month - prox.month) + 1)
    else:
        meses_futuros = 0

    media_reajustada  = media_sem_reajuste * fator
    delta_mensal      = media_reajustada - media_sem_reajuste
    diferenca_futura  = round(delta_mensal * meses_futuros, 2)
    complementacao    = round(val_retro + diferenca_futura, 2)

    comp_ini_txt = ""
    if ultima_competencia:
        from datetime import date
        prox_m = (ultima_competencia.month % 12) + 1
        prox_a = ultima_competencia.year + (1 if ultima_competencia.month == 12 else 0)
        meses_pt = ["jan","fev","mar","abr","mai","jun","jul","ago","set","out","nov","dez"]
        comp_ini_txt = f"{meses_pt[prox_m-1]}/{prox_a}"
    vig_final_txt = vig_final.strftime("%d/%m/%Y") if vig_final else "[campo a preencher]"

    # ── Montar documento ───────────────────────────────────────────
    document = Document()
    _aplicar_estilo_docx(document)

    p_titulo = document.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_titulo.add_run("MEMORANDO DE ADEQUAÇÃO ORÇAMENTÁRIA")
    r.bold = True
    r.font.size = Pt(13)

    document.add_paragraph("")

    for linha, bold in [(f"PROCESSO Nº {processo}", True), (f"CONTRATO Nº {contrato}", True)]:
        p = document.add_paragraph()
        p.add_run(linha).bold = bold

    document.add_paragraph("")

    # Item 1 — tabela por exercício (retroativo = ano atual)
    _adicionar_paragrafo_justificado(
        document,
        f"1. Solicita-se adequação orçamentária para o Contrato {contrato}, "
        f"em razão dos reajustes {ciclos_ref}, conforme programação abaixo:"
    )
    document.add_paragraph("")

    tab1 = document.add_table(rows=1, cols=2)
    tab1.style = "Table Grid"
    for i, h in enumerate(["Exercício", "Valor"]):
        tab1.rows[0].cells[i].text = h
        tab1.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    row = tab1.add_row().cells
    row[0].text = str(ano_atual)
    row[1].text = moeda(complementacao)

    document.add_paragraph("")

    # Item 2 — justificativa
    _adicionar_paragrafo_justificado(
        document,
        "2. O ajuste é necessário para compatibilizar a programação orçamentária com o novo "
        "valor contratual estimado, considerando a alteração informada, a base de comparação "
        "adotada e os efeitos financeiros correspondentes."
    )

    document.add_paragraph("")

    # Item 3 — memória de cálculo
    _adicionar_paragrafo_justificado(
        document,
        "3. A memória de cálculo abaixo fundamenta tecnicamente esta solicitação em estrita "
        "observância aos princípios da motivação, legalidade e eficiência, garantindo a "
        "transparência, a impessoalidade e a devida publicidade dos critérios técnicos que "
        "balizam o presente ato administrativo."
    )
    document.add_paragraph("")

    # Tabela memória
    obs_dif = (
        f"Soma dos deltas mensais dos valores a serem pagos de {comp_ini_txt} "
        f"ao final do contrato ({vig_final_txt}), com base na média dos últimos 6 pagamentos "
        f"({moeda(media_sem_reajuste)}) reajustada pelo fator {fator:.4f} "
        f"({meses_futuros} meses × {moeda(delta_mensal)}/mês)."
        if comp_ini_txt else
        "Soma dos deltas mensais dos valores a serem pagos até o final do contrato."
    )

    tab2 = document.add_table(rows=1, cols=3)
    tab2.style = "Table Grid"
    for i, h in enumerate(["Descrição", "Valor", "Memória"]):
        tab2.rows[0].cells[i].text = h
        tab2.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    linhas_mem = [
        ("Diferença futura projetada", diferenca_futura, obs_dif),
        ("Retroativo apurado",         val_retro,
         "Valor já apurado a pagar em razão do reajuste represado."),
    ]
    for desc, val, obs in linhas_mem:
        r2 = tab2.add_row().cells
        r2[0].text = desc
        r2[1].text = moeda(val)
        r2[2].text = obs

    row_total = tab2.add_row().cells
    row_total[0].text = f"Total: {moeda(complementacao)}."
    row_total[0].paragraphs[0].runs[0].bold = True
    row_total[1].text = ""
    row_total[2].text = ""

    document.add_paragraph("")

    document.add_paragraph("")
    document.add_paragraph("")

    for alinhamento, texto, negrito in [
        (WD_ALIGN_PARAGRAPH.CENTER, "Brasília, [campo a preencher]", False),
        (WD_ALIGN_PARAGRAPH.CENTER, "", False),
        (WD_ALIGN_PARAGRAPH.CENTER, "_" * 48, False),
        (WD_ALIGN_PARAGRAPH.CENTER, "[campo a preencher]", False),
        (WD_ALIGN_PARAGRAPH.CENTER, "Gerente de Compras e Contratos", True),
        (WD_ALIGN_PARAGRAPH.CENTER, "Telecomunicações Brasileiras S.A. — Telebras", False),
    ]:
        p = document.add_paragraph()
        p.alignment = alinhamento
        run = p.add_run(texto)
        run.bold = negrito

    _destacar_campos_preencher(document)

    buf = BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.getvalue()


def gerar_publicacao_dou_docx(adm, res):
    """
    Gera o Extrato de Apostilamento para publicação no Diário Oficial da União (DOCX).
    Formato padronizado conforme padrão DOU.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except Exception as exc:
        raise RuntimeError("A biblioteca python-docx não está instalada.") from exc

    adm = adm or {}
    res = res or {}

    document = Document()
    _aplicar_estilo_docx(document)

    for linha, tam, negrito in [
        ("MINISTÉRIO DAS COMUNICAÇÕES",                         11, True),
        ("TELECOMUNICAÇÕES BRASILEIRAS S.A. — TELEBRAS",       10, True),
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(linha)
        run.bold = negrito
        run.font.size = Pt(tam)

    document.add_paragraph("")

    p_titulo = document.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_titulo.add_run("EXTRATO DE APOSTILAMENTO")
    r.bold = True
    r.font.size = Pt(11)

    document.add_paragraph("")

    contrato = _placeholder(adm.get("contrato") or res.get("contrato") or res.get("numero_contrato"))
    indice   = res.get("indice", adm.get("indice", "[campo a preencher]"))
    fator    = float(res.get("fator_acumulado", 1.0) or 1.0)
    variacao = float(res.get("variacao_acumulada", fator - 1.0) or (fator - 1.0))
    val_total= float(res.get("valor_atualizado_contrato", res.get("valor_global_estoque", 0)) or 0)

    from docx.shared import RGBColor
    campos_simples = [
        ("Espécie",              "Termo de Apostila"),
        ("Número do Contrato",   contrato),
        ("Contratante",          "Telecomunicações Brasileiras S.A. — Telebras, CNPJ 00.336.701/0001-04"),
        ("Contratada",           "[campo a preencher]"),
        ("CNPJ/CPF Contratada",  "[campo a preencher]"),
        ("Objeto do Contrato",   "[campo a preencher]"),
        ("Objeto do Apostilamento", "[campo a preencher]"),
        ("Data de Assinatura",   "[campo a preencher]"),
        ("Signatários",
         "[campo a preencher], pela Contratante, e [campo a preencher], pela Contratada"),
        ("Processo",             _placeholder(adm.get("processo") or res.get("processo"))),
    ]

    for label, valor in campos_simples:
        p = document.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(_limpar_texto_formal(valor))

    # Fundamento Legal com [confirmar] em laranja
    p_fl = document.add_paragraph()
    p_fl.add_run("Fundamento Legal: ").bold = True
    p_fl.add_run("§ 7º do art. 81 da Lei nº 13.303, de 30 de junho de 2016 ")
    _run_conf = p_fl.add_run("[confirmar]")
    _run_conf.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
    _run_conf.bold = True

    _destacar_campos_preencher(document)

    buf = BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── INICIO_UI ── não remover este comentário ─────────────────────────────────
render_marca_topo()
st.title("Relatório Global")

render_aviso_privacidade(tem_download=True)
adm = st.session_state.get("dados_admissibilidade")
res = st.session_state.get("resultado_valor_global")

if not res:
    st.warning(
        "Ainda não há dados processados para o Relatório Global. "
        "Acesse o módulo Valor Global, carregue o Arquivo de Coleta e processe a análise."
    )
    st.stop()

aplicar_css_responsivo_relatorio()

modo_reduzido = eh_modo_reduzido_itens(res)
modo_consumo = eh_modo_consumo_itens_ciclo(res)
retroativo_itens_valor = valor_retroativo_estimado_itens(res)

st.subheader("Resumo Executivo")
if modo_consumo:
    st.markdown(aviso_modo_consumo_html(), unsafe_allow_html=True)
elif modo_reduzido:
    st.markdown(aviso_modo_reduzido_html(), unsafe_allow_html=True)

col1, col2 = st.columns(2)
col1.metric("Índice", res.get("indice") or "Não informado no modelo")
col2.metric("Fator acumulado", fator_fmt(res.get("fator_acumulado", 1.0)))

st.markdown(
    f"""
    <div class="telebras-kpi-destaque">
        <div class="telebras-kpi-destaque-label">Valor Total Atualizado do Contrato</div>
        <div class="telebras-kpi-destaque-valor">{moeda(_m20_componentes_vta_documentos(res).get("total", res.get("valor_atualizado_contrato", res.get("valor_global_estoque", 0))))}</div>
    </div>
    """,
    unsafe_allow_html=True,
)

if modo_consumo:
    col3, col4 = st.columns(2)
    col3.metric("Retroativo financeiro definitivo", "Não calculado")
    col4.metric("Retroativo (itens consumidos/ciclo)", moeda(valor_retroativo_consumo_itens_ciclo(res)))
elif modo_reduzido:
    col3, col4 = st.columns(2)
    col3.metric("Retroativo financeiro definitivo", "Não calculado")
    with col4:
        st.markdown(
            f"""
            <div class="modo-reduzido-card">
                <div class="modo-reduzido-card-label">Retroativo estimado por itens/estoque</div>
                <div class="modo-reduzido-card-valor">{moeda(retroativo_itens_valor)}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    col5, col6 = st.columns(2)
    col5.metric("Base mensal por competência", "Não informada")
    col6.metric("Execução estimada por itens/estoque", moeda(_m20_componentes_vta_documentos(res).get("execucao", 0)))

    col7, col8 = st.columns(2)
    col7.metric("Saldo remanescente atualizado", moeda(_m20_componentes_vta_documentos(res).get("remanescente", 0)))
    col8.metric("Aditivos/supressões registrados", moeda(_m20_componentes_vta_documentos(res).get("aditivos", 0)))
else:
    col3, col4 = st.columns(2)
    col3.metric("Valor represado a pagar", moeda(res.get("valor_represado_a_pagar", 0)))
    col4.metric("Aditivos/supressões registrados", moeda(_m20_componentes_vta_documentos(res).get("aditivos", 0)))

    col5, col6 = st.columns(2)
    col5.metric("Valor pago efetivo", moeda(res.get("total_pago_faturado", 0)))
    col6.metric("Valor teórico calculado", moeda(res.get("total_devido_reajustado", 0)))

st.divider()

tab2, tab3, tab4 = st.tabs(["Tabelas", "PDF", "Minuta de Apostilamento"])

with tab2:
    st.markdown("### Metodologia de corte da apuração")
    st.markdown(aviso_metodologia_corte_html(res), unsafe_allow_html=True)

    st.markdown("### Quadro Executivo")
    st.dataframe(
        df_visual(res.get("df_comparativo"), moeda_cols=["Valor", "Antes do Reajuste", "Após Reajuste", "Diferença"]),
        use_container_width=True,
        hide_index=True,
    )

    if modo_reduzido:
        st.markdown("### Retroativo estimado por itens/estoque")
        st.markdown(aviso_modo_reduzido_html(), unsafe_allow_html=True)
        st.metric("Retroativo estimado por itens/estoque", moeda(retroativo_itens_valor))
        df_ri = df_retroativo_estimado_itens(res)
        if isinstance(df_ri, pd.DataFrame) and not df_ri.empty:
            st.dataframe(
                df_visual(
                    df_ri,
                    moeda_cols=["Valor executado original", "Valor executado atualizado", "Retroativo estimado por itens/estoque"],
                ),
                use_container_width=True,
                hide_index=True,
            )
        else:
            st.info("Tabela detalhada do retroativo estimado por itens/estoque não disponível nesta sessão.")

    st.markdown("### Composição do Valor Total Atualizado do Contrato")
    st.caption("Composição considerada: execução atualizada por ciclo + saldo remanescente atualizado + aditivos/supressões computáveis, quando aplicáveis.")
    df_comp_valor = res.get("df_composicao_valor_total")
    if isinstance(df_comp_valor, pd.DataFrame) and not df_comp_valor.empty:
        st.dataframe(
            df_visual(df_comp_valor, moeda_cols=["Valor"]),
            use_container_width=True,
            hide_index=True,
        )
    else:
        st.info("Composição do valor total não disponível nesta sessão.")

    st.markdown("### Financeiro por Ciclo")
    st.dataframe(
        df_visual(
            res.get("df_financeiro_por_ciclo"),
            moeda_cols=["Valor pago efetivo", "Valor teórico calculado", "Valor pago/faturado", "Valor devido reajustado", "Delta do ciclo", "Delta acumulado"],
            fator_cols=["Fator aplicado ao retroativo", "Fator aplicado"],
        ),
        use_container_width=True,
        hide_index=True,
    )

    df_ad = res.get("df_aditivos_executivo", res.get("df_aditivos"))
    if isinstance(df_ad, pd.DataFrame) and not df_ad.empty:
        st.markdown("### Aditivos e Supressões")
        st.caption("Quadro de controle formal. Aditivos computáveis integram o Valor Total Atualizado; aditivos informativos permanecem apenas como memória quando já incorporados à execução, ao saldo remanescente ou ao valor formalizado anterior.")
        keep_ad = [c for c in ["Aditivo", "Ciclo/Marco", "Tratamento do aditivo", "Quantidade de linhas", "Valor do aditivo na assinatura", "Fator aplicado", "Valor do aditivo reajustado"] if c in df_ad.columns]
        st.dataframe(
            df_visual(df_ad[keep_ad].copy() if keep_ad else df_ad, moeda_cols=["Valor do aditivo na assinatura", "Valor do aditivo reajustado", "Valor original da alteração", "Valor atualizado da alteração"], fator_cols=["Fator aplicado"]),
            use_container_width=True,
            hide_index=True,
        )

with tab3:
    st.markdown("### Baixar Relatório Executivo em PDF")
    try:
        pdf_bytes = criar_pdf_relatorio(adm, res)
        st.session_state["arquivo_relatorio_executivo_pdf"] = pdf_bytes
        st.download_button(
            label="Baixar Relatório Executivo em PDF",
            data=pdf_bytes,
            file_name="Relatorio_Executivo_Analise_Reajuste.pdf",
            mime="application/pdf",
            type="primary",
        )
    except Exception as exc:
        st.error(f"Não foi possível gerar o PDF: {exc}")
        st.caption("Verifique se a biblioteca reportlab foi incluída no requirements.txt.")


with tab4:
    st.markdown("### Gerar Minuta de Termo de Apostilamento")
    if modo_reduzido:
        st.markdown(
            """
            <div class="modo-reduzido-box">
                <div class="modo-reduzido-titulo">Atenção: minuta em Modo Reduzido</div>
                <div class="modo-reduzido-texto">
                    A análise foi processada sem base mensal por competência. Recomenda-se validar a base financeira antes de usar a minuta para formalização de pagamento ou apostilamento.
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    st.info(
        "A minuta é gerada em DOCX editável. Os dados disponíveis no sistema são preenchidos automaticamente; "
        "as informações ainda não cadastradas permanecem como [campo a preencher]."
    )
    try:
        docx_bytes = gerar_minuta_apostilamento_docx(adm, res)
        st.session_state["arquivo_minuta_apostilamento_docx"] = docx_bytes
        st.download_button(
            label="Baixar Minuta de Apostilamento em DOCX",
            data=docx_bytes,
            file_name="minuta_termo_apostilamento.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=False,
        )
    except Exception as exc:
        st.error(f"Não foi possível gerar a minuta de apostilamento: {exc}")
        st.caption("Verifique se a biblioteca python-docx foi incluída no requirements.txt.")
