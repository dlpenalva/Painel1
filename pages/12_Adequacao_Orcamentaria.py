from datetime import datetime, date
from io import BytesIO
from zoneinfo import ZoneInfo
from html import escape
import re

import pandas as pd
import streamlit as st

from _ui_utils import render_marca_topo, render_aviso_privacidade, render_cabecalho_pagina
# Etapa 4: toda a matematica da Adequacao Orcamentaria vive no motor de dominio
# (_adequacao_orcamentaria), reproduzindo o golden normativo com paridade Excel.
from _adequacao_orcamentaria import (
    _round2,
    calcular_adequacao_orcamentaria,
    Pedido,
    OverrideMes,
)

st.set_page_config(page_icon="assets/cl8us_favicon_512.png", page_title="TLB · cl8us - Adequação Orçamentária", layout="wide")


def moeda(valor, com_prefixo=True):
    try:
        valor = round(float(valor or 0), 2)
    except Exception:
        valor = 0.0
    texto = f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"R$ {texto}" if com_prefixo else texto


def parse_moeda_br(valor):
    if valor is None:
        return 0.0
    if isinstance(valor, (int, float)):
        try:
            if pd.isna(valor):
                return 0.0
        except Exception:
            pass
        return float(valor)
    texto = str(valor).strip()
    if not texto:
        return 0.0
    texto = texto.replace("R$", "").replace(" ", "").replace("\xa0", "").replace("%", "")
    if "," in texto:
        texto = texto.replace(".", "").replace(",", ".")
    try:
        return float(texto)
    except Exception:
        return 0.0


def pct(valor):
    try:
        n = float(valor or 0)
    except Exception:
        n = 0.0
    if abs(n) < 1:
        n *= 100
    return f"{n:.2f}%".replace(".", ",")


def texto_seguro(valor, padrao="[campo a preencher]"):
    if valor is None:
        return padrao
    try:
        if pd.isna(valor):
            return padrao
    except Exception:
        pass
    texto = str(valor).strip()
    if not texto or texto.lower() in ["nan", "none", "null", "nat", "<na>"]:
        return padrao
    return texto


def data_hora_brasilia():
    return datetime.now(ZoneInfo("America/Sao_Paulo")).strftime("%d/%m/%Y %H:%M")


def normalizar_texto(valor):
    if valor is None:
        return ""
    texto = str(valor).strip().lower()
    mapa = str.maketrans("áàâãäéèêëíìîïóòôõöúùûüç", "aaaaaeeeeiiiiooooouuuuc")
    texto = texto.translate(mapa)
    texto = re.sub(r"[^a-z0-9]+", "_", texto)
    return texto.strip("_")


def localizar_coluna(df, opcoes):
    if not isinstance(df, pd.DataFrame) or df.empty:
        return None
    mapa = {normalizar_texto(c): c for c in df.columns}
    for opcao in opcoes:
        alvo = normalizar_texto(opcao)
        if alvo in mapa:
            return mapa[alvo]
    for col_norm, col_original in mapa.items():
        for opcao in opcoes:
            alvo = normalizar_texto(opcao)
            if alvo and alvo in col_norm:
                return col_original
    return None


def input_moeda(label, valor_padrao, key, help=None):
    txt = st.text_input(label, value=moeda(valor_padrao, com_prefixo=False), key=key, help=help)
    valor = parse_moeda_br(txt)
    st.caption(moeda(valor))
    return valor


def render_card_valor(label, valor, nota="", destaque=False, formato="moeda"):
    bg = "#EAF2F8" if destaque else "#FFFFFF"
    border = "#9EC5E8" if destaque else "#E5EAF0"
    cor_valor = "#0B1F3A" if destaque else "#0F172A"
    fs = "1.65rem" if destaque else "1.12rem"
    nota_html = f"<div style='color:#64748B; font-size:0.82rem; margin-top:5px;'>{escape(str(nota))}</div>" if nota else ""
    if formato == "inteiro":
        try:
            valor_fmt = f"{int(round(float(valor or 0)))}"
        except Exception:
            valor_fmt = "0"
    elif formato == "texto":
        valor_fmt = escape(str(valor))
    else:
        valor_fmt = moeda(valor)
    st.markdown(
        f"""<div style="background:{bg}; border:1px solid {border}; border-radius:14px; padding:14px 16px; min-height:96px;">
            <div style="color:#475569; font-size:0.84rem; font-weight:700; margin-bottom:7px;">{escape(str(label))}</div>
            <div style="color:{cor_valor}; font-size:{fs}; font-weight:900; line-height:1.2; overflow-wrap:anywhere;">{valor_fmt}</div>
            {nota_html}</div>""",
        unsafe_allow_html=True,
    )


MESES_PT = {
    "jan": 1, "janeiro": 1, "fev": 2, "fevereiro": 2, "mar": 3, "marco": 3, "março": 3,
    "abr": 4, "abril": 4, "mai": 5, "maio": 5, "jun": 6, "junho": 6,
    "jul": 7, "julho": 7, "ago": 8, "agosto": 8, "set": 9, "setembro": 9,
    "out": 10, "outubro": 10, "nov": 11, "novembro": 11, "dez": 12, "dezembro": 12,
}


def periodo_para_label(periodo):
    if periodo is None or pd.isna(periodo):
        return ""
    p = pd.Period(periodo, freq="M")
    nomes = ["jan", "fev", "mar", "abr", "mai", "jun", "jul", "ago", "set", "out", "nov", "dez"]
    return f"{nomes[p.month - 1]}/{str(p.year)[-2:]}"


def normalizar_competencia(valor):
    if valor is None:
        return None
    try:
        if pd.isna(valor):
            return None
    except Exception:
        pass
    if isinstance(valor, pd.Period):
        return valor.asfreq("M")
    if isinstance(valor, (datetime, date, pd.Timestamp)):
        data = pd.to_datetime(valor, errors="coerce")
        if pd.notna(data):
            return data.to_period("M")
    texto = str(valor).strip().lower()
    if not texto or texto in ["nan", "none", "nat", "total"]:
        return None
    texto = texto.replace(".", "/").replace("-", "/")
    m = re.search(r"([a-zçãéíóú]+)\s*/\s*(\d{2,4})", texto, flags=re.IGNORECASE)
    if m:
        mes_txt = normalizar_texto(m.group(1))
        ano = int(m.group(2))
        if ano < 100:
            ano += 2000
        mes = MESES_PT.get(mes_txt)
        if mes:
            return pd.Period(f"{ano}-{mes:02d}", freq="M")
    m = re.search(r"(\d{1,2})\s*/\s*(\d{2,4})", texto)
    if m:
        mes = int(m.group(1))
        ano = int(m.group(2))
        if 1 <= mes <= 12:
            if ano < 100:
                ano += 2000
            return pd.Period(f"{ano}-{mes:02d}", freq="M")
    data = pd.to_datetime(valor, dayfirst=True, errors="coerce")
    if pd.notna(data):
        return data.to_period("M")
    return None


def periodo_de_data_final(valor):
    data = pd.to_datetime(valor, dayfirst=True, errors="coerce")
    if pd.isna(data):
        return None
    return data.to_period("M")


def extrair_contexto_valores():
    res = st.session_state.get("resultado_valor_global", {}) or {}
    modo = res.get("modo_apuracao", "Completo") if isinstance(res, dict) else "Completo"
    if modo == "Reduzido por Itens/Estoque":
        valor_represado = parse_moeda_br(res.get("valor_retroativo_estimado_itens_estoque", 0))
    else:
        valor_represado = parse_moeda_br(res.get("valor_represado_a_pagar", res.get("delta_total", 0)))
    variacao = parse_moeda_br(res.get("variacao_acumulada", res.get("fator_acumulado", 1) - 1 if res.get("fator_acumulado") else 0))
    indice = texto_seguro(res.get("indice", ""), "não informado")
    quantidade_ciclos = texto_seguro(res.get("quantidade_ciclos", ""), "[campo a preencher]")
    return {
        "disponivel": bool(res),
        "resultado": res,
        "valor_represado": valor_represado,
        "valor_retroativo_estimado_itens_estoque": parse_moeda_br(res.get("valor_retroativo_estimado_itens_estoque", 0)),
        "variacao": variacao,
        "indice": indice,
        "quantidade_ciclos": quantidade_ciclos,
    }


def financeiro_mensal_consolidado(resultado):
    if not isinstance(resultado, dict):
        return pd.DataFrame(columns=["Competência", "Valor pago/medido"]), "resultado_valor_global.df_financeiro_mensal indisponível"
    df = resultado.get("df_financeiro_mensal")
    origem = "resultado_valor_global.df_financeiro_mensal"
    if not isinstance(df, pd.DataFrame) or df.empty:
        return pd.DataFrame(columns=["Competência", "Valor pago/medido"]), origem + " indisponível"
    col_comp = localizar_coluna(df, ["Competência", "Competencia", "Mês/Ano", "Mes/Ano", "Mês", "Mes"])
    col_valor = localizar_coluna(df, ["Valor bruto medido/aprovado por competência", "Valor bruto medido", "Valor medido/aprovado", "Valor pago/faturado", "Valor bruto faturado", "Valor faturado", "Valor pago", "Valor medido", "Valor"])
    if col_comp is None or col_valor is None:
        return pd.DataFrame(columns=["Competência", "Valor pago/medido"]), origem + " sem colunas reconhecidas"
    temp = df[[col_comp, col_valor]].copy()
    temp["_periodo"] = temp[col_comp].apply(normalizar_competencia)
    temp["_valor"] = temp[col_valor].apply(parse_moeda_br)
    temp = temp[temp["_periodo"].notna()].copy()
    temp = temp[temp["_valor"].abs() > 0.004].copy()
    if temp.empty:
        return pd.DataFrame(columns=["Competência", "Valor pago/medido"]), origem
    mensal = (temp.groupby("_periodo", as_index=False)["_valor"].sum()
              .sort_values("_periodo").reset_index(drop=True))
    mensal["Competência"] = mensal["_periodo"].apply(periodo_para_label)
    mensal["Valor pago/medido"] = mensal["_valor"].round(2)
    return mensal[["_periodo", "Competência", "Valor pago/medido"]], origem


def ultimos_meses_para_media(mensal, n=6):
    if not isinstance(mensal, pd.DataFrame) or mensal.empty:
        return pd.DataFrame(columns=["_periodo", "Competência", "Valor pago/medido"])
    return mensal.sort_values("_periodo").tail(n).reset_index(drop=True)


def gerar_periodos_projecao(ultima_competencia, data_final_vigencia):
    if ultima_competencia is None:
        return []
    fim = periodo_de_data_final(data_final_vigencia)
    if fim is None:
        return []
    inicio = pd.Period(ultima_competencia, freq="M") + 1
    if fim < inicio:
        return []
    return list(pd.period_range(inicio, fim, freq="M"))


def carregar_itens_pc_da_sessao(resultado):
    """Reutiliza os Pedidos de Compra ja estruturados (itens_PC) lidos da Coleta,
    sem redigitacao de NUMERO_PC/DATA_PC/VALOR_PC. Procura os registros no
    resultado e no diagnostico da Coleta ja disponivel em sessao."""
    fontes = []
    if isinstance(resultado, dict):
        fontes.append(resultado)
    diag = st.session_state.get("diagnostico_coleta_v2")
    if isinstance(diag, dict):
        fontes.append(diag)
    for fonte in fontes:
        ipc = fonte.get("itens_pc_v10")
        if isinstance(ipc, dict) and isinstance(ipc.get("itens"), list) and ipc["itens"]:
            return ipc["itens"]
        if isinstance(fonte.get("itens_pc"), list) and fonte["itens_pc"]:
            return fonte["itens_pc"]
    return []


def montar_base_editor(periodos, media_mensal):
    return pd.DataFrame([{
        "Competência": periodo_para_label(p),
        "Base automática pela média": moeda(media_mensal, com_prefixo=False),
        "Valor informado pelo fiscal": "",
        "Premissa do valor informado": "Valor sem reajuste",
        "Observação": "",
    } for p in periodos])


def calcular_projecao(df_editor, media_mensal, fator_reajuste):
    linhas = []
    if not isinstance(df_editor, pd.DataFrame) or df_editor.empty:
        return pd.DataFrame(columns=["Competência", "Origem", "Premissa usada", "Valor base considerado", "Valor reajustado estimado", "Diferença futura a adequar", "Observação"])
    for _, row in df_editor.iterrows():
        competencia = texto_seguro(row.get("Competência"), "")
        informado_raw = row.get("Valor informado pelo fiscal", "")
        informado = parse_moeda_br(informado_raw)
        premissa = texto_seguro(row.get("Premissa do valor informado"), "Valor sem reajuste")
        observacao = texto_seguro(row.get("Observação"), "")
        # Base G, valor reajustado H = ROUND(G*fator,2) e diferenca I = ROUND(H-G,2)
        # seguem o golden; o arredondamento e o do motor (_round2, paridade Excel).
        if abs(informado) > 0.004:
            origem = "Valor informado pelo fiscal"
            if premissa == "Valor já reajustado":
                base_considerada = informado / fator_reajuste if fator_reajuste else informado
                premissa_usada = "Valor já reajustado"
            else:
                base_considerada = informado
                premissa_usada = "Valor sem reajuste"
        else:
            origem = "Média dos últimos 6 meses"
            base_considerada = media_mensal
            premissa_usada = "Média sem reajuste"
        valor_reajustado = _round2(base_considerada * fator_reajuste)
        diferenca = _round2(valor_reajustado - base_considerada)
        linhas.append({
            "Competência": competencia, "Origem": origem, "Premissa usada": premissa_usada,
            "Valor base considerado": _round2(base_considerada),
            "Valor reajustado estimado": valor_reajustado,
            "Diferença futura a adequar": diferenca,
            "Observação": observacao,
        })
    return pd.DataFrame(linhas)


def cronograma_por_exercicio(df_projecao, retroativo):
    # Programacao por exercicio (golden): diferencas futuras de cada ano + o
    # retroativo somente no PRIMEIRO exercicio da projecao (nao no ano corrente).
    linhas = {}
    anos = []
    if isinstance(df_projecao, pd.DataFrame) and not df_projecao.empty:
        for _, row in df_projecao.iterrows():
            periodo = normalizar_competencia(row.get("Competência"))
            if periodo is None:
                continue
            ano = int(periodo.year)
            anos.append(ano)
            linhas[ano] = linhas.get(ano, 0.0) + parse_moeda_br(row.get("Diferença futura a adequar", 0))
    if anos:
        primeiro_exercicio = min(anos)
        linhas[primeiro_exercicio] = linhas.get(primeiro_exercicio, 0.0) + float(retroativo or 0)
    return pd.DataFrame([
        {"Exercício": str(ano), "Valor": _round2(valor)}
        for ano, valor in sorted(linhas.items()) if abs(valor) > 0.004
    ])


def gerar_xlsx_projecao(df_ultimos, df_projecao, resumo):
    output = BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        workbook = writer.book
        fmt_title = workbook.add_format({"bold": True, "font_size": 14, "font_color": "#0B1F3A"})
        fmt_header = workbook.add_format({"bold": True, "font_color": "white", "bg_color": "#1F4E78", "border": 1, "align": "center", "valign": "vcenter", "text_wrap": True})
        fmt_money = workbook.add_format({"num_format": 'R$ #,##0.00', "border": 1})
        fmt_text = workbook.add_format({"border": 1})
        fmt_note = workbook.add_format({"italic": True, "font_color": "#64748B"})
        fmt_total = workbook.add_format({"num_format": 'R$ #,##0.00', "border": 1, "bold": True, "bg_color": "#EAF2F8"})
        ws = workbook.add_worksheet("RESUMO")
        writer.sheets["RESUMO"] = ws
        ws.write(0, 0, "Adequação Orçamentária — Delta do Reajuste", fmt_title)
        ws.write(1, 0, "Complementação necessária = retroativo apurado + diferença futura projetada.", fmt_note)
        ws.write(3, 0, "Indicador", fmt_header)
        ws.write(3, 1, "Valor", fmt_header)
        for r, (label, valor) in enumerate(resumo, start=4):
            ws.write(r, 0, label, fmt_text)
            if isinstance(valor, (int, float)):
                fmt = fmt_total if "Complementação" in label else fmt_money
                ws.write_number(r, 1, float(valor), fmt)
            else:
                ws.write(r, 1, valor, fmt_text)
        ws.set_column("A:A", 42)
        ws.set_column("B:B", 26)
        df_ult = df_ultimos[["Competência", "Valor pago/medido"]].copy() if isinstance(df_ultimos, pd.DataFrame) and not df_ultimos.empty else pd.DataFrame(columns=["Competência", "Valor pago/medido"])
        df_ult.to_excel(writer, sheet_name="MEDIA", index=False)
        ws_m = writer.sheets["MEDIA"]
        for c, col in enumerate(df_ult.columns):
            ws_m.write(0, c, col, fmt_header)
        ws_m.set_column("A:A", 18)
        ws_m.set_column("B:B", 22, fmt_money)
        df_proj = df_projecao.copy() if isinstance(df_projecao, pd.DataFrame) else pd.DataFrame()
        df_proj.to_excel(writer, sheet_name="PROJECAO", index=False)
        ws_p = writer.sheets["PROJECAO"]
        for c, col in enumerate(df_proj.columns):
            ws_p.write(0, c, col, fmt_header)
        for c, col in enumerate(df_proj.columns):
            if col in ["Valor base considerado", "Valor reajustado estimado", "Diferença futura a adequar"]:
                ws_p.set_column(c, c, max(18, min(32, len(str(col)) + 4)), fmt_money)
                for r_idx, valor in enumerate(df_proj[col], start=1):
                    ws_p.write_number(r_idx, c, parse_moeda_br(valor), fmt_money)
            else:
                ws_p.set_column(c, c, max(16, min(32, len(str(col)) + 4)))
    output.seek(0)
    return output.getvalue()


def montar_texto_memorando(dados):
    contrato = texto_seguro(dados.get("contrato"))
    ciclos = texto_seguro(dados.get("ciclos_reajuste"))
    total = float(dados.get("complementacao", 0) or 0)
    retroativo = float(dados.get("retroativo", 0) or 0)
    diferenca_futura = float(dados.get("diferenca_futura", 0) or 0)
    periodo = texto_seguro(dados.get("periodo_projecao"), "[campo a preencher]")
    cronograma = dados.get("cronograma", pd.DataFrame())
    linhas = [
        "MEMORANDO", "",
        f"Assunto: Solicitação de adequação orçamentária — Contrato {contrato}.", "",
        f"1. Solicita-se adequação orçamentária para o Contrato {contrato}, em razão dos reajustes {ciclos}:", "",
        "Exercício\tValor",
    ]
    if isinstance(cronograma, pd.DataFrame) and not cronograma.empty:
        for _, row in cronograma.iterrows():
            linhas.append(f"{texto_seguro(row.get('Exercício'), '[campo a preencher]')}\t{moeda(row.get('Valor', 0))}")
    linhas.append(f"Total\t{moeda(total)}")
    linhas.extend([
        "",
        "2. A solicitação tem por finalidade compatibilizar a programação orçamentária com o impacto financeiro dos reajustes.",
        "",
        f"Diferença futura projetada\t{moeda(diferenca_futura)}\tSoma dos deltas mensais de {periodo} até o final da vigência.",
        f"Retroativo apurado\t{moeda(retroativo)}\tValor já apurado a pagar em razão do reajuste represado.",
        f"Total\t{moeda(total)}\tComplementação orçamentária necessária.",
        "",
        f"3. Recomenda-se a adequação orçamentária no valor total de {moeda(total)}.",
    ])
    return "\n".join(linhas)


def gerar_docx_memorando(dados):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except Exception as exc:
        raise RuntimeError("A biblioteca python-docx não está instalada.") from exc
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    contrato = texto_seguro(dados.get("contrato"))
    ciclos_ref = texto_seguro(dados.get("ciclos_reajuste"))
    total = float(dados.get("complementacao", 0) or 0)
    retroativo = float(dados.get("retroativo", 0) or 0)
    diferenca_futura = float(dados.get("diferenca_futura", 0) or 0)
    cronograma = dados.get("cronograma", pd.DataFrame())
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("MEMORANDO")
    r.bold = True
    r.font.size = Pt(13)
    p = document.add_paragraph()
    p.add_run("Assunto: ").bold = True
    p.add_run(f"Solicitação de adequação orçamentária — Contrato {contrato}.")
    document.add_paragraph(
        f"1. Solicita-se adequação orçamentária para o Contrato {contrato}, em razão dos reajustes {ciclos_ref}."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if isinstance(cronograma, pd.DataFrame) and not cronograma.empty:
        tabela = document.add_table(rows=1, cols=2)
        tabela.style = "Table Grid"
        hdr = tabela.rows[0].cells
        hdr[0].text = "Exercício"
        hdr[1].text = "Valor"
        for _, item in cronograma.iterrows():
            row = tabela.add_row().cells
            row[0].text = texto_seguro(item.get("Exercício"), "[campo a preencher]")
            row[1].text = moeda(item.get("Valor", 0))
        row = tabela.add_row().cells
        row[0].text = "Total"
        row[1].text = moeda(total)
    document.add_paragraph(
        f"2. Retroativo apurado: {moeda(retroativo)}. Diferença futura projetada: {moeda(diferenca_futura)}. "
        f"Complementação necessária: {moeda(total)}."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.add_paragraph(f"Gerado em: {data_hora_brasilia()}.")
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()




# >>> METODOLOGIA_VALOR_IMPORTADO_CORTE_OPERACIONAL
def _metodologia_valor_importado_html(resultado_valor_global):
    """Box informativo sobre a metodologia do Valor Total Atualizado importado do módulo Valores.
    Não altera cálculo. Apenas documenta a origem/metodologia usada.
    """
    if not isinstance(resultado_valor_global, dict) or not resultado_valor_global:
        return """
        <div style="background:#F8FAFC; border:1px solid #CBD5E1; border-left:6px solid #64748B; border-radius:12px; padding:13px 15px; margin:10px 0 16px 0; color:#334155;">
            <div style="font-weight:800; margin-bottom:5px;">Metodologia do Valor Total Atualizado importado</div>
            <div style="font-size:0.92rem; line-height:1.45;">Não há resultado do módulo Valores carregado nesta sessão. Os campos desta página devem ser conferidos/preenchidos manualmente.</div>
        </div>
        """

    cfg = resultado_valor_global.get("config_ciclo_em_execucao", {}) or {}
    corte = bool(resultado_valor_global.get("corte_operacional_aplicado") or resultado_valor_global.get("corte_operacional_solicitado") or cfg.get("aplicar"))

    try:
        valor_total = moeda(resultado_valor_global.get("valor_atualizado_contrato", resultado_valor_global.get("valor_global_estoque", 0)))
    except Exception:
        valor_total = str(resultado_valor_global.get("valor_atualizado_contrato", "Não informado"))
    try:
        execucao = moeda(resultado_valor_global.get("valor_executado_atualizado", 0))
    except Exception:
        execucao = str(resultado_valor_global.get("valor_executado_atualizado", "Não informado"))
    try:
        remanescente = moeda(resultado_valor_global.get("remanescente_reajustado", 0))
    except Exception:
        remanescente = str(resultado_valor_global.get("remanescente_reajustado", "Não informado"))

    valor_c0 = cfg.get("valor_c0_manual", 0)
    try:
        c0_manual = moeda(valor_c0) if parse_moeda_br(valor_c0) > 0 else "Não informado"
    except Exception:
        c0_manual = "Não informado"

    ciclo = str(cfg.get("ciclo") or resultado_valor_global.get("ciclo_ultimo_remanescente") or "Não informado")
    competencia = str(cfg.get("competencia_corte") or cfg.get("data_corte") or "Não informado")
    fonte = str(cfg.get("fonte") or "Base financeira preferencial, quando disponível")

    if corte:
        linhas = [
            "<b>Metodologia:</b> corte operacional no ciclo em execução.",
            f"<b>Ciclo em execução:</b> {ciclo}.",
            f"<b>Competência de corte:</b> {competencia}.",
            f"<b>Fonte da execução:</b> {fonte}.",
            f"<b>C0 financeiro manual:</b> {c0_manual}.",
            f"<b>Execução atualizada considerada:</b> {execucao}.",
            f"<b>Saldo remanescente atualizado considerado:</b> {remanescente}.",
            f"<b>Valor Total Atualizado importado:</b> {valor_total}.",
        ]
        return f"""
        <div style="background:#ECFEFF; border:1px solid #14B8A6; border-left:6px solid #0F766E; border-radius:12px; padding:13px 15px; margin:10px 0 16px 0; color:#134E4A;">
            <div style="font-weight:900; margin-bottom:6px;">Metodologia do Valor Total Atualizado importado</div>
            <div style="font-size:0.92rem; line-height:1.48;">{'<br>'.join(linhas)}</div>
        </div>
        """

    return f"""
    <div style="background:#F8FAFC; border:1px solid #CBD5E1; border-left:6px solid #64748B; border-radius:12px; padding:13px 15px; margin:10px 0 16px 0; color:#334155;">
        <div style="font-weight:900; margin-bottom:6px;">Metodologia do Valor Total Atualizado importado</div>
        <div style="font-size:0.92rem; line-height:1.48;">
            <b>Metodologia:</b> corte padrão, sem corte operacional específico no ciclo em execução.<br>
            <b>Composição:</b> execução atualizada por ciclo + saldo remanescente atualizado.<br>
            <b>Execução atualizada considerada:</b> {execucao}.<br>
            <b>Saldo remanescente atualizado considerado:</b> {remanescente}.<br>
            <b>Valor Total Atualizado importado:</b> {valor_total}.
        </div>
    </div>
    """
# <<< METODOLOGIA_VALOR_IMPORTADO_CORTE_OPERACIONAL

render_cabecalho_pagina(
    "Adequação Orçamentária",
    "Estimativa simplificada do delta orçamentário do reajuste: retroativo apurado + diferença futura projetada.",
)
if st.button("← Voltar para Central", key="voltar_central_adequacao"):
    st.switch_page("pages/03_Valor_Global.py")

ctx = extrair_contexto_valores()
resultado = ctx["resultado"]
st.markdown(_metodologia_valor_importado_html(resultado), unsafe_allow_html=True)
modo_apuracao = resultado.get("modo_apuracao", "Completo") if isinstance(resultado, dict) else "Completo"
modo_reduzido_estoque = modo_apuracao == "Reduzido por Itens/Estoque"
modo_consumo_itens_ciclo = modo_apuracao == "Consumo por Itens/Ciclo"

if modo_consumo_itens_ciclo:
    st.markdown(
        """<div style="background:#F6F3EE; border:1px solid #7A8F63; border-left:6px solid #4E6E58;
                border-radius:12px; padding:14px 16px; margin:10px 0 16px 0; color:#2F3E2F;">
            <div style="font-weight:800; margin-bottom:4px;">Modo Consumo por Itens/Ciclo</div>
            <div style="font-size:0.95rem;">A base mensal por competência não foi informada. A adequação utilizará o Retroativo (itens consumidos/ciclo) e deverá ser tratada como estimativa orçamentária apoiada na validação fiscal.</div>
        </div>""",
        unsafe_allow_html=True,
    )
elif modo_reduzido_estoque:
    st.markdown(
        """<div style="background:#F3E8FF; border:1px solid #A855F7; border-left:6px solid #7E22CE;
                border-radius:12px; padding:14px 16px; margin:10px 0 16px 0; color:#581C87;">
            <div style="font-weight:800; margin-bottom:4px;">Modo Reduzido por Itens/Estoque</div>
            <div style="font-size:0.95rem;">A base mensal por competência não foi informada. A adequação será tratada como estimativa.</div>
        </div>""",
        unsafe_allow_html=True,
    )

mensal, origem_financeira = financeiro_mensal_consolidado(resultado)
ultimos_6 = ultimos_meses_para_media(mensal, 6)
media_6 = float(ultimos_6["Valor pago/medido"].mean()) if not ultimos_6.empty else 0.0
ultima_comp = ultimos_6["_periodo"].iloc[-1] if not ultimos_6.empty else None
ultima_comp_txt = periodo_para_label(ultima_comp) if ultima_comp is not None else "[campo a preencher]"

st.subheader("Parâmetros principais")
col1, col2, col3 = st.columns(3)
with col1:
    label_retroativo = "Retroativo (itens consumidos/ciclo)" if modo_consumo_itens_ciclo else ("Retroativo estimado por itens/estoque" if modo_reduzido_estoque else "Retroativo apurado")
    retroativo = input_moeda(label_retroativo, ctx["valor_represado"], "adequacao_v2_retroativo")
with col2:
    percentual_txt = st.text_input("Percentual de reajuste aplicado", value=pct(ctx["variacao"]), key="adequacao_v2_percentual")
    percentual_reajuste = parse_moeda_br(percentual_txt) / 100
    fator_reajuste = 1 + percentual_reajuste
    st.caption(f"Fator usado: {fator_reajuste:.6f}".replace(".", ","))
with col3:
    data_final_vigencia = st.text_input("Data final da vigência contratual", value=st.session_state.get("adequacao_v2_data_final_vigencia", ""),
        placeholder="Ex.: 30/11/2026", key="adequacao_v2_data_final_vigencia")

st.subheader("Origem histórica")
origem_hist = st.radio(
    "Base histórica para a projeção da adequação",
    ["Financeiro", "Pedidos de compra"], horizontal=True, key="adequacao_v2_origem")

# A matematica vive no motor (_adequacao_orcamentaria); aqui apenas escolhemos a
# referencia mensal (media_ref) conforme a origem. O restante do fluxo (projecao,
# programacao, exports) e comum as duas origens.
media_ref = media_6
origem_hist_rotulo = "Financeiro mensal"
janela_meses_pc = None

if origem_hist == "Pedidos de compra":
    origem_hist_rotulo = "Pedidos de compra"
    registros_pc = carregar_itens_pc_da_sessao(resultado)
    ultima_comp_data = ultima_comp.to_timestamp().date() if ultima_comp is not None else None
    if not registros_pc:
        st.warning("NÃO HÁ PEDIDOS DE COMPRA DISPONÍVEIS PARA ESTA ADEQUAÇÃO. "
                   "Utilize a origem Financeiro ou carregue uma Coleta com itens_PC.")
        media_ref = 0.0
    elif ultima_comp_data is None:
        st.warning("Base financeira ausente: informe/di sponibilize a última competência "
                   "para delimitar a janela histórica dos Pedidos de Compra.")
        media_ref = 0.0
    else:
        janela_meses_pc = st.slider("Janela histórica dos pedidos (meses)", 1, 60,
            value=int(st.session_state.get("adequacao_v2_janela", 39)), key="adequacao_v2_janela")
        ids_disponiveis = [
            str(r.get("numero_pc") or r.get("NUMERO_PC") or r.get("id") or "")
            for r in registros_pc]
        exclusoes = st.multiselect(
            "Excluir Pedidos de Compra desta adequação (opcional)",
            options=[i for i in ids_disponiveis if i], key="adequacao_v2_exclusoes_pc")
        peds_pc = pedidos_de_itens_pc(registros_pc, exclusoes=exclusoes)
        base_pc = media_pedidos_compra(peds_pc, ultima_comp_data, janela_meses_pc)
        media_ref = base_pc["media_mensal"]
        st.subheader("Resumo histórico dos Pedidos de Compra")
        cA, cB, cC, cD = st.columns(4)
        with cA:
            render_card_valor(
                "Período histórico",
                f"{base_pc['inicio_janela'].strftime('%d/%m/%Y')} a "
                f"{base_pc['fim_janela'].strftime('%d/%m/%Y')}", formato="texto")
        with cB:
            render_card_valor("Meses da janela", janela_meses_pc, formato="inteiro")
        with cC:
            render_card_valor("PCs considerados", base_pc["pedidos_considerados"], formato="inteiro")
        with cD:
            render_card_valor("Média mensal (PCs)", media_ref, destaque=True)
        cE, cF, cG = st.columns(3)
        with cE:
            render_card_valor("Meses com PCs", base_pc["meses_com_pedido"], formato="inteiro")
        with cF:
            render_card_valor("Meses sem PCs", base_pc["meses_sem_pedido"], formato="inteiro")
        with cG:
            render_card_valor("Total histórico", base_pc["total_historico"])
        if base_pc["pedidos_considerados"] == 0:
            st.info("0 PCs considerados na janela escolhida. A janela permanece soberana "
                    "(não é movida para encaixar pedidos).")
        with st.expander("Ver Pedidos de Compra e situação na janela", expanded=False):
            cl = classificar_pedidos(peds_pc, ultima_comp_data, janela_meses_pc)
            df_pc = pd.DataFrame([{
                "PC": x["identificacao"],
                "Data": x["data"].strftime("%d/%m/%Y") if x["data"] else "",
                "Valor": moeda(x["valor"]), "Situação": x["situacao"]} for x in cl["pedidos"]])
            st.dataframe(df_pc, use_container_width=True, hide_index=True)
            st.caption("Apenas PCs 'Considerado' entram na média; 'Fora da janela' e "
                       "'Excluído' não participam do cálculo.")

if origem_hist == "Financeiro":
    st.subheader("Base automática pela média financeira")
    col_m1, col_m2, col_m3, col_m4 = st.columns(4)
    with col_m1:
        render_card_valor("Média dos últimos 6 meses", media_ref)
    with col_m2:
        render_card_valor("Última competência financeira", ultima_comp_txt, formato="texto", destaque=True)
    with col_m3:
        render_card_valor("Média mensal reajustada", media_ref * fator_reajuste)
    with col_m4:
        render_card_valor("Delta mensal estimado", (media_ref * fator_reajuste) - media_ref)

    with st.expander("Ver competências usadas no cálculo da média", expanded=False):
        st.caption(f"Base de execução mensal detectada: {origem_financeira}")
        if ultimos_6.empty:
            st.warning("Não foram localizadas competências financeiras válidas para cálculo da média.")
        else:
            df_ultimos_vis = ultimos_6[["Competência", "Valor pago/medido"]].copy()
            df_ultimos_vis["Valor pago/medido"] = df_ultimos_vis["Valor pago/medido"].apply(moeda)
            st.dataframe(df_ultimos_vis, use_container_width=True, hide_index=True)

periodos = gerar_periodos_projecao(ultima_comp, data_final_vigencia)
periodo_inicio_txt = periodo_para_label(periodos[0]) if periodos else "[campo a preencher]"
periodo_fim_txt = periodo_para_label(periodos[-1]) if periodos else "[campo a preencher]"
periodo_projecao_txt = f"{periodo_inicio_txt} a {periodo_fim_txt}" if periodos else "[campo a preencher]"

st.subheader("Projeção futura")
if not data_final_vigencia.strip():
    st.warning("Informe a data final da vigência contratual para calcular a diferença futura projetada.")
elif not periodos:
    st.info("Não há competências futuras a projetar com os dados informados.")

base_editor = montar_base_editor(periodos, media_ref)
editor_key = f"adequacao_v2_editor_{origem_hist}_{ultima_comp_txt}_{data_final_vigencia}_{round(media_ref, 2)}_{round(fator_reajuste, 6)}"

with st.expander("Ajustar projeção por competência, se necessário", expanded=False):
    df_editor = st.data_editor(
        base_editor, hide_index=True, use_container_width=True, num_rows="fixed", key=editor_key,
        column_config={
            "Competência": st.column_config.TextColumn("Competência", disabled=True),
            "Base automática pela média": st.column_config.TextColumn("Base automática pela média", disabled=True),
            "Valor informado pelo fiscal": st.column_config.TextColumn("Valor informado pelo fiscal"),
            "Premissa do valor informado": st.column_config.SelectboxColumn("Premissa do valor informado",
                options=["Valor sem reajuste", "Valor já reajustado"], required=True),
            "Observação": st.column_config.TextColumn("Observação"),
        },
    )

df_projecao = calcular_projecao(df_editor, media_ref, fator_reajuste)
if modo_reduzido_estoque and ultimos_6.empty:
    rem_original = parse_moeda_br(resultado.get("remanescente_original", 0)) if isinstance(resultado, dict) else 0.0
    rem_atualizado = parse_moeda_br(resultado.get("remanescente_reajustado", 0)) if isinstance(resultado, dict) else 0.0
    diferenca_estoque = round(max(rem_atualizado - rem_original, 0.0), 2)
    df_projecao = pd.DataFrame([{
        "Competência": "Estimativa por saldo remanescente", "Origem": "Modo reduzido por itens/estoque",
        "Premissa usada": "Saldo remanescente informado",
        "Valor base considerado": round(rem_original, 2), "Valor reajustado estimado": round(rem_atualizado, 2),
        "Diferença futura a adequar": diferenca_estoque,
        "Observação": "Estimativa sem base mensal; validar antes de formalizar pagamento.",
    }])

diferenca_futura = float(df_projecao["Diferença futura a adequar"].sum()) if not df_projecao.empty else 0.0
complementacao = _round2(float(retroativo or 0) + diferenca_futura)
referencia_reajustada = _round2(media_ref * fator_reajuste)
qtd_meses = 0 if modo_reduzido_estoque and ultimos_6.empty else len(df_projecao)
cronograma = cronograma_por_exercicio(df_projecao, retroativo)

# --- Referência mensal (comum as duas origens; calculada pelo motor) ---
st.subheader("Referência mensal")
col_r1, col_r2, col_r3 = st.columns(3)
with col_r1:
    render_card_valor("Média histórica", media_ref)
with col_r2:
    render_card_valor("Percentual aplicado", pct(percentual_reajuste), formato="texto")
with col_r3:
    render_card_valor("Referência mensal reajustada", referencia_reajustada, destaque=True)

# --- Detalhe da projeção mês a mês (sem tabela vazia) ---
if not str(data_final_vigencia).strip():
    st.info("Informe a data final da vigência para calcular a projeção futura.")
elif df_projecao.empty:
    st.info("Não há competências futuras a projetar com os dados informados.")
else:
    with st.expander(f"Ver projeção mês a mês ({qtd_meses} meses · {periodo_projecao_txt})", expanded=False):
        df_proj_vis = df_projecao.copy()
        for col in ["Valor base considerado", "Valor reajustado estimado", "Diferença futura a adequar"]:
            if col in df_proj_vis.columns:
                df_proj_vis[col] = df_proj_vis[col].apply(moeda)
        st.dataframe(df_proj_vis, use_container_width=True, hide_index=True)

# --- Bloco 6: Resultado da adequação ---
st.subheader("Resultado da adequação")
col_a1, col_a2, col_a3, col_a4 = st.columns(4)
with col_a1:
    render_card_valor("Referência mensal reajustada", referencia_reajustada)
with col_a2:
    render_card_valor("Diferença futura projetada", diferenca_futura, nota=f"{qtd_meses} meses")
with col_a3:
    render_card_valor("Retroativo já apurado", retroativo)
with col_a4:
    render_card_valor("COMPLEMENTAÇÃO NECESSÁRIA", complementacao, destaque=True)
st.caption("Complementação necessária = retroativo oficial + diferença futura projetada.")

# --- Bloco 7: Programação por exercício ---
st.subheader("Programação por exercício")
if isinstance(cronograma, pd.DataFrame) and not cronograma.empty:
    total_cron = float(pd.to_numeric(cronograma["Valor"], errors="coerce").sum())
    df_cron = cronograma.copy()
    df_cron["Valor"] = df_cron["Valor"].apply(moeda)
    df_cron = pd.concat(
        [df_cron, pd.DataFrame([{"Exercício": "TOTAL", "Valor": moeda(total_cron)}])],
        ignore_index=True)
    st.dataframe(df_cron, use_container_width=True, hide_index=True)
    if abs(total_cron - complementacao) < 0.01:
        st.caption("A soma dos exercícios confere com a complementação necessária.")
    else:
        st.warning("A soma dos exercícios não confere com a complementação — verificar.")
else:
    st.info("A programação por exercício depende de projeção futura calculada.")

# --- Bloco 8: Arquivos da adequação (ação secundária) ---
resumo_xlsx = [
    ("Origem histórica", origem_hist_rotulo),
    ("Retroativo apurado", retroativo),
    ("Média mensal histórica", media_ref),
    ("Percentual de reajuste", pct(percentual_reajuste)),
    ("Referência mensal reajustada", referencia_reajustada),
    ("Quantidade de meses projetados", str(qtd_meses)),
    ("Diferença futura projetada", diferenca_futura),
    ("Complementação necessária", complementacao),
]
if origem_hist == "Pedidos de compra" and janela_meses_pc is not None:
    _base_pc_exp = media_pedidos_compra(
        pedidos_de_itens_pc(carregar_itens_pc_da_sessao(resultado),
                            exclusoes=st.session_state.get("adequacao_v2_exclusoes_pc", [])),
        ultima_comp.to_timestamp().date(), janela_meses_pc)
    resumo_xlsx[3:3] = [
        ("Janela histórica (meses)", str(janela_meses_pc)),
        ("PCs considerados", str(_base_pc_exp["pedidos_considerados"])),
        ("Meses com PCs", str(_base_pc_exp["meses_com_pedido"])),
        ("Meses sem PCs", str(_base_pc_exp["meses_sem_pedido"])),
        ("Total histórico dos PCs", _base_pc_exp["total_historico"]),
    ]
st.subheader("Arquivos da adequação")
with st.expander("Baixar planilha de validação da adequação (XLSX)", expanded=False):
    xlsx_bytes = gerar_xlsx_projecao(ultimos_6, df_projecao, resumo_xlsx)
    st.download_button("Baixar XLSX", data=xlsx_bytes,
        file_name="adequacao_orcamentaria.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=False)
