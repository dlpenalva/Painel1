from io import BytesIO
from decimal import Decimal, ROUND_HALF_UP
from datetime import datetime
from zoneinfo import ZoneInfo
import re

import pandas as pd
import streamlit as st

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_COLOR_INDEX
    DOCX_OK = True
except Exception:
    DOCX_OK = False

try:
    from _ui_utils import render_marca_topo, render_aviso_privacidade
except Exception:
    def render_marca_topo():
        st.markdown("### TLB · cl8us")
        st.caption("apoio à gestão de contratos")
    def render_aviso_privacidade(tem_upload=False, tem_download=False):
        return

st.set_page_config(page_title="TLB · cl8us - DOU", layout="wide")

DOU_VERSAO = "20260516_botao_laranja_dados_auto"


# ============================================================
# Utilitários
# ============================================================

def texto_seguro(valor, padrao="[preencher campo]"):
    if valor is None:
        return padrao
    try:
        if pd.isna(valor):
            return padrao
    except Exception:
        pass
    texto = str(valor).strip()
    if not texto or texto.lower() in ["nan", "none", "null", "nat", "<na>"]:
        return padrao
    return texto


def moeda(valor):
    try:
        valor = Decimal(str(valor or 0)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    except Exception:
        valor = Decimal("0.00")
    sinal = "-" if valor < 0 else ""
    valor_abs = abs(valor)
    inteiro = int(valor_abs)
    centavos = int((valor_abs - inteiro) * 100)
    inteiro_fmt = f"{inteiro:,}".replace(",", ".")
    return f"{sinal}R$ {inteiro_fmt},{centavos:02d}"


def percentual(valor, casas=2):
    try:
        v = float(valor or 0)
    except Exception:
        v = 0.0
    if abs(v) <= 1:
        v *= 100
    return f"{v:.{casas}f}%".replace(".", ",")


def agora_brasilia():
    return datetime.now(ZoneInfo("America/Sao_Paulo")).strftime("%d/%m/%Y %H:%M")


UNIDADES = ["", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito", "nove"]
DEZ_A_DEZENOVE = ["dez", "onze", "doze", "treze", "quatorze", "quinze", "dezesseis", "dezessete", "dezoito", "dezenove"]
DEZENAS = ["", "", "vinte", "trinta", "quarenta", "cinquenta", "sessenta", "setenta", "oitenta", "noventa"]
CENTENAS = ["", "cento", "duzentos", "trezentos", "quatrocentos", "quinhentos", "seiscentos", "setecentos", "oitocentos", "novecentos"]


def _ate_999(n):
    n = int(n)
    if n == 0:
        return ""
    if n == 100:
        return "cem"
    partes = []
    c = n // 100
    resto = n % 100
    if c:
        partes.append(CENTENAS[c])
    if resto:
        if resto < 10:
            partes.append(UNIDADES[resto])
        elif resto < 20:
            partes.append(DEZ_A_DEZENOVE[resto - 10])
        else:
            d = resto // 10
            u = resto % 10
            if u:
                partes.append(f"{DEZENAS[d]} e {UNIDADES[u]}")
            else:
                partes.append(DEZENAS[d])
    return " e ".join([p for p in partes if p])


def inteiro_por_extenso(n):
    n = int(n)
    if n == 0:
        return "zero"

    escalas = [
        (1_000_000_000, "bilhão", "bilhões"),
        (1_000_000, "milhão", "milhões"),
        (1_000, "mil", "mil"),
        (1, "", ""),
    ]
    partes = []
    resto = n
    for base, singular, plural in escalas:
        grupo = resto // base
        resto = resto % base
        if grupo == 0:
            continue
        if base == 1_000:
            texto = "mil" if grupo == 1 else f"{_ate_999(grupo)} mil"
        elif base == 1:
            texto = _ate_999(grupo)
        else:
            texto = f"{_ate_999(grupo)} {singular if grupo == 1 else plural}"
        partes.append(texto)

    if len(partes) == 1:
        return partes[0]
    return ", ".join(partes[:-1]) + " e " + partes[-1]


def valor_por_extenso(valor):
    try:
        d = Decimal(str(valor or 0)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    except Exception:
        d = Decimal("0.00")
    inteiro = int(abs(d))
    centavos = int((abs(d) - inteiro) * 100)
    reais = "real" if inteiro == 1 else "reais"
    texto = f"{inteiro_por_extenso(inteiro)} {reais}"
    if centavos:
        cent = "centavo" if centavos == 1 else "centavos"
        texto += f" e {inteiro_por_extenso(centavos)} {cent}"
    if d < 0:
        texto = "menos " + texto
    return texto


def _coluna(df, candidatos):
    if not isinstance(df, pd.DataFrame) or df.empty:
        return None
    mapa = {str(c).strip().lower(): c for c in df.columns}
    for cand in candidatos:
        chave = str(cand).strip().lower()
        if chave in mapa:
            return mapa[chave]
    for c in df.columns:
        nome = str(c).strip().lower()
        for cand in candidatos:
            if str(cand).strip().lower() in nome:
                return c
    return None


def obter_df_ciclos(res, adm):
    if isinstance(res, dict):
        for chave in ["df_ciclos", "df_ciclos_considerados", "df_ciclos_apurados"]:
            df = res.get(chave)
            if isinstance(df, pd.DataFrame) and not df.empty:
                return df.copy()
    if isinstance(adm, dict):
        ciclos = adm.get("ciclos") or adm.get("detalhamento_ciclos") or []
        if ciclos:
            return pd.DataFrame(ciclos)
    return pd.DataFrame()


def ciclos_reajuste_texto(res, adm):
    df = obter_df_ciclos(res, adm)
    if not isinstance(df, pd.DataFrame) or df.empty:
        return "[preencher campo]"

    col_ciclo = _coluna(df, ["Ciclo"])
    col_pct = _coluna(df, ["Percentual aplicado", "Variação", "Variacao", "percentual_aplicado", "variacao"])
    col_fator_acum = _coluna(df, ["Fator acumulado", "fator_acumulado"])

    linhas = []
    fator_acum_final = None
    for _, row in df.iterrows():
        ciclo = texto_seguro(row.get(col_ciclo, "") if col_ciclo else "", "")
        if not ciclo or str(ciclo).upper() == "C0":
            continue
        pct_val = row.get(col_pct, 0) if col_pct else 0
        linhas.append(f"{ciclo}: {percentual(pct_val)}")
        if col_fator_acum:
            try:
                f = float(row.get(col_fator_acum, 0) or 0)
                if f > 0:
                    fator_acum_final = f
            except Exception:
                pass

    acumulado = None
    if isinstance(res, dict):
        for chave in ["variacao_acumulada", "percentual_acumulado"]:
            try:
                v = float(res.get(chave, 0) or 0)
                if abs(v) > 0.000001:
                    acumulado = v
                    break
            except Exception:
                pass
        if acumulado is None:
            try:
                f = float(res.get("fator_acumulado", 0) or 0)
                if f > 0:
                    acumulado = f - 1
            except Exception:
                pass
    if acumulado is None and fator_acum_final:
        acumulado = fator_acum_final - 1

    texto = "; ".join(linhas) if linhas else "[preencher campo]"
    if acumulado is not None:
        texto += f". Percentual acumulado: {percentual(acumulado)}"
    else:
        texto += ". Percentual acumulado: [preencher campo]"
    return texto


def valor_total_atualizado(res):
    if not isinstance(res, dict):
        return 0.0
    for chave in ["valor_atualizado_contrato", "valor_global_estoque", "valor_global_contrato"]:
        try:
            v = float(res.get(chave, 0) or 0)
            if abs(v) > 0.004:
                return v
        except Exception:
            pass
    return 0.0


def extrair_apostilado(res, adm):
    candidatos = []
    if isinstance(res, dict):
        contexto = res.get("contexto_contratual_anterior", {}) or {}
        candidatos.extend([
            res.get("contratada"), res.get("apostilado"), res.get("fornecedor"),
            contexto.get("contratada"), contexto.get("apostilado"), contexto.get("fornecedor"),
        ])
    if isinstance(adm, dict):
        candidatos.extend([adm.get("contratada"), adm.get("apostilado"), adm.get("fornecedor")])
    for valor in candidatos:
        texto = texto_seguro(valor, "")
        if texto:
            return texto
    return "[preencher campo]"


def montar_texto_dou(campos, auto):
    return f"""Assunto: EXTRATO DE PUBLICAÇÃO TERMO DE APOSTILA

Referência(s): {campos['referencias']}

Contrato: {campos['contrato']}
Instrumento assinado: {campos['instrumento']}
Processo n.º: {campos['processo']}
Data de Assinatura: {campos['data_assinatura']}

Apostilado: {auto['apostilado']}
CNPJ/MF: {campos['cnpj']}

Objeto: {campos['objeto']}
Reajuste Contratual: {auto['reajuste']}
Valor Total Atualizado: {auto['valor_total']} ({auto['valor_total_extenso']}).

Fundamentação Legal: Art. 81, § 7º da Lei nº 13.303/2016.
Signatários: {campos['signatarios']}"""


def _add_texto_com_destaque(paragraph, texto, bold=False):
    """Adiciona texto ao parágrafo, destacando placeholders entre colchetes em amarelo."""
    partes = re.split(r"(\[[^\]]+\])", str(texto))
    for parte in partes:
        if parte == "":
            continue
        run = paragraph.add_run(parte)
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(11)
        if re.fullmatch(r"\[[^\]]+\]", parte):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def gerar_docx_dou(texto):
    if not DOCX_OK:
        raise RuntimeError("python-docx não está disponível neste ambiente.")
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    for idx, bloco in enumerate(str(texto).split("\n\n")):
        bloco = bloco.strip()
        if not bloco:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        if idx == 0 and bloco.upper().startswith("ASSUNTO:"):
            _add_texto_com_destaque(p, bloco, bold=True)
        else:
            # Preserva linhas internas do bloco, sem quebrar a estrutura do texto.
            linhas = bloco.split("\n")
            for i, linha in enumerate(linhas):
                if i > 0:
                    p.add_run().add_break()
                _add_texto_com_destaque(p, linha, bold=False)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================
# Interface
# ============================================================

render_marca_topo()
st.title("DOU")
st.caption("Geração simples de DOCX pré-preenchido para extrato de publicação de Termo de Apostila.")
render_aviso_privacidade(tem_download=True)

st.markdown(
    """
    <style>
    .dou-auto-grid {
        display: grid;
        grid-template-columns: repeat(3, minmax(0, 1fr));
        gap: 12px;
        margin: 8px 0 18px 0;
    }
    .dou-auto-card {
        border: 1px solid #D6E0EA;
        border-radius: 12px;
        padding: 12px 14px;
        background: #F8FAFC;
        min-height: 92px;
    }
    .dou-auto-label {
        color: #475569;
        font-size: 0.86rem;
        font-weight: 700;
        margin-bottom: 5px;
    }
    .dou-auto-value {
        color: #0F172A;
        font-size: 0.98rem;
        line-height: 1.35;
        overflow-wrap: anywhere;
    }
    div.stDownloadButton > button {
        background-color: #F59E0B !important;
        color: white !important;
        border: 1px solid #D97706 !important;
        border-radius: 10px !important;
        font-weight: 700 !important;
    }
    div.stDownloadButton > button:hover {
        background-color: #D97706 !important;
        color: white !important;
        border: 1px solid #B45309 !important;
    }
    @media (max-width: 900px) {
        .dou-auto-grid { grid-template-columns: 1fr; }
    }
    </style>
    """,
    unsafe_allow_html=True,
)

res = st.session_state.get("resultado_valor_global", {}) or {}
adm = st.session_state.get("dados_admissibilidade", {}) or {}

if not isinstance(res, dict) or not res:
    st.info("Para pré-preencher os campos automáticos, processe primeiro uma análise na página Valores. Ainda assim, você pode gerar o DOCX com campos a preencher.")

auto = {
    "apostilado": extrair_apostilado(res, adm),
    "reajuste": ciclos_reajuste_texto(res, adm),
}
valor_total = valor_total_atualizado(res)
auto["valor_total"] = moeda(valor_total) if valor_total else "[preencher campo]"
auto["valor_total_extenso"] = valor_por_extenso(valor_total) if valor_total else "[preencher campo]"

st.markdown("### Campos manuais")
col1, col2 = st.columns(2)
with col1:
    referencias = st.text_input("Referência(s)", value="[preencher campo]")
    contrato = st.text_input("Contrato", value="[preencher campo]")
    instrumento = st.text_input("Instrumento assinado", value="[N. do Termo de Apostila]")
with col2:
    processo = st.text_input("Processo n.º", value="[preencher campo]")
    data_assinatura = st.text_input("Data de Assinatura", value="[preencher campo]")
    cnpj = st.text_input("CNPJ/MF", value="[preencher campo]")

objeto = st.text_area("Objeto", value="[preencher campo]", height=80)
signatarios = st.text_input("Signatários", value="[preencher campo]")

st.markdown("### Dados automáticos")
st.markdown(
    f"""
    <div class="dou-auto-grid">
      <div class="dou-auto-card">
        <div class="dou-auto-label">Apostilado</div>
        <div class="dou-auto-value">{auto['apostilado']}</div>
      </div>
      <div class="dou-auto-card">
        <div class="dou-auto-label">Reajuste Contratual</div>
        <div class="dou-auto-value">{auto['reajuste']}</div>
      </div>
      <div class="dou-auto-card">
        <div class="dou-auto-label">Valor Total Atualizado</div>
        <div class="dou-auto-value">{auto['valor_total']}<br><span style="color:#64748B; font-size:0.86rem;">{auto['valor_total_extenso']}</span></div>
      </div>
    </div>
    """,
    unsafe_allow_html=True,
)

campos = {
    "referencias": texto_seguro(referencias),
    "contrato": texto_seguro(contrato),
    "instrumento": texto_seguro(instrumento),
    "processo": texto_seguro(processo),
    "data_assinatura": texto_seguro(data_assinatura),
    "cnpj": texto_seguro(cnpj),
    "objeto": texto_seguro(objeto),
    "signatarios": texto_seguro(signatarios),
}

texto_final = montar_texto_dou(campos, auto)

with st.expander("Prévia do texto que será gerado no DOCX", expanded=True):
    st.text_area("Texto do extrato", value=texto_final, height=360, key="previa_dou", disabled=False)
    texto_para_docx = st.session_state.get("previa_dou", texto_final)

try:
    arquivo_docx = gerar_docx_dou(texto_para_docx)
    st.download_button(
        "Baixar DOCX pré-preenchido",
        data=arquivo_docx,
        file_name="DOU_Extrato_Termo_de_Apostila.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
except Exception as exc:
    st.error(f"Não foi possível gerar o DOCX: {exc}")
