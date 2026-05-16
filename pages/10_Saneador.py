from datetime import datetime
from io import BytesIO
from zoneinfo import ZoneInfo
import re

import pandas as pd
import streamlit as st

st.set_page_config(page_title="TLB · cl8us - Saneador", layout="wide")

from _ui_utils import render_marca_topo, render_aviso_privacidade


COLUNAS_INFOS = [
    "Grupo", "Informação necessária", "Documento/fonte mínima",
    "Valor / informação levantada", "Link SIGA (opcional)", "Status", "Observação",
]


def moeda(valor):
    try:
        valor = round(float(valor or 0), 2)
    except Exception:
        valor = 0.0
    return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def pct(valor):
    try:
        valor = float(valor or 0)
    except Exception:
        valor = 0.0
    if abs(valor) <= 1:
        valor *= 100
    return f"{valor:.2f}%".replace(".", ",")


def texto_seguro(valor, padrao="[campo a preencher]"):
    if valor is None:
        return padrao
    try:
        if pd.isna(valor):
            return padrao
    except Exception:
        pass
    texto = str(valor).strip()
    if not texto or texto.lower() in ["nan", "none", "null", "nat", "<na>"]:
        return padrao
    return texto


def limpar_texto_saneador(valor):
    texto = "" if valor is None else str(valor)
    substituicoes = {
        "✅": "", "❌": "", "⚠️": "", "⚠": "", "🟡": "", "🔴": "", "🟢": "", "🔵": "",
        "🛡️": "", "🛡": "", "📊": "", "📝": "", "⚖️": "", "⚖": "", "🔄": "", "📥": "", "🔍": "",
        "▲": "", "■": "", "●": "", "•": "",
    }
    for antigo, novo in substituicoes.items():
        texto = texto.replace(antigo, novo)
    texto = "".join(ch for ch in texto if ord(ch) <= 0xFFFF)
    trocas_status = {
        "TEMPESTIVO": "tempestivo", "PRECLUSO": "precluso",
        "ADMISSÍVEL": "admissível", "ADMISSIVEL": "admissível",
        "RESSALVA": "ressalva", "ADIANTADO": "adiantado",
        "CICLO NEGATIVO": "ciclo negativo", "APLICADO 0,00%": "aplicado 0,00%",
        "CICLO ADMITIDO POR NEGOCIAÇÃO ENTRE AS PARTES": "ciclo admitido por negociação entre as partes",
    }
    for antigo, novo in trocas_status.items():
        texto = texto.replace(antigo, novo)
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r" *\n *", "\n", texto)
    return texto.strip()


def normalizar_df_infos(df):
    if not isinstance(df, pd.DataFrame) or df.empty:
        return pd.DataFrame(columns=COLUNAS_INFOS)
    df = df.copy()
    for col in COLUNAS_INFOS:
        if col not in df.columns:
            df[col] = ""
    return df[COLUNAS_INFOS].fillna("")


def _normalizar_chave(texto):
    import unicodedata
    texto = "" if texto is None else str(texto).strip().lower()
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(ch for ch in texto if not unicodedata.combining(ch))
    texto = re.sub(r"\s+", " ", texto)
    return texto


def info_por_alias(df, aliases, padrao="[campo a preencher]", incluir_link=True):
    df = normalizar_df_infos(df)
    if df.empty:
        return padrao
    aliases_norm = [_normalizar_chave(a) for a in aliases]
    chaves = df["Informação necessária"].astype(str).map(_normalizar_chave)
    for alvo in aliases_norm:
        mask = chaves == alvo
        if not mask.any():
            continue
        row = df.loc[mask].iloc[0]
        valor = texto_seguro(row.get("Valor / informação levantada", ""), "")
        link = texto_seguro(row.get("Link SIGA (opcional)", ""), "") if incluir_link else ""
        if valor and link and link not in valor:
            return f"{valor} ({link})"
        if valor:
            return valor
        if link:
            return link
    return padrao


def resumo_fontes(df_infos, resultado_vg, resultado_garantia, df_checklist, eventos_aditivos):
    linhas = [
        {"Fonte": "Infos Prévias", "Status": "Disponível" if isinstance(df_infos, pd.DataFrame) and not df_infos.empty else "Não disponível"},
        {"Fonte": "Valor Global", "Status": "Disponível" if isinstance(resultado_vg, dict) and bool(resultado_vg) else "Não disponível"},
        {"Fonte": "Garantia", "Status": "Disponível" if isinstance(resultado_garantia, dict) and bool(resultado_garantia) else "Não disponível"},
        {"Fonte": "Checklist Processual", "Status": "Disponível" if isinstance(df_checklist, pd.DataFrame) and not df_checklist.empty else "Não disponível"},
        {"Fonte": "Aditivos: 25%", "Status": "Disponível" if isinstance(eventos_aditivos, pd.DataFrame) and not eventos_aditivos.empty else "Não disponível"},
    ]
    return pd.DataFrame(linhas)


def _valor_numerico_positivo(valor):
    try:
        n = float(valor or 0)
    except Exception:
        return None
    if abs(n) > 0.004:
        return n
    return None


def extrair_delta_por_ciclo(resultado):
    df = resultado.get("df_delta_por_ciclo") if isinstance(resultado, dict) else None
    if not isinstance(df, pd.DataFrame) or df.empty:
        df = resultado.get("df_financeiro_por_ciclo") if isinstance(resultado, dict) else None
    if not isinstance(df, pd.DataFrame) or df.empty:
        return ""
    partes = []
    for _, row in df.head(8).iterrows():
        ciclo = texto_seguro(row.get("Ciclo", ""), "")
        pago = row.get("Valor pago efetivo", row.get("Valor pago", None))
        teorico = row.get("Valor teórico calculado", row.get("Valor devido", None))
        delta = row.get("Delta do ciclo", row.get("Valor represado", None))
        trecho = []
        if ciclo:
            trecho.append(str(ciclo))
        if pago is not None:
            trecho.append(f"valor pago efetivo {moeda(pago)}")
        if teorico is not None:
            trecho.append(f"valor teórico calculado {moeda(teorico)}")
        if delta is not None:
            trecho.append(f"diferença/retroativo {moeda(delta)}")
        if trecho:
            partes.append(", ".join(trecho))
    if not partes:
        return ""
    return " Por ciclo, a apuração registrou: " + "; ".join(partes) + "."


def extrair_valor_aditivos(resultado_vg, eventos_aditivos):
    if isinstance(resultado_vg, dict):
        for chave in ["total_aditivos_atualizados", "total_aditivos_quantitativos_atualizados",
                      "valor_aditivos_atualizados", "valor_atualizado_aditivos"]:
            valor = _valor_numerico_positivo(resultado_vg.get(chave))
            if valor is not None:
                return moeda(valor)
        for chave_df in ["df_aditivos_executivo", "df_aditivos"]:
            df = resultado_vg.get(chave_df)
            if isinstance(df, pd.DataFrame) and not df.empty:
                candidatos = [c for c in df.columns if "valor" in c.lower() and ("atual" in c.lower() or "reaj" in c.lower())]
                if not candidatos:
                    candidatos = [c for c in df.columns if "valor" in c.lower()]
                for col in candidatos[:1]:
                    serie = pd.to_numeric(df[col], errors="coerce").fillna(0)
                    total = float(serie.sum())
                    if abs(total) > 0.004:
                        return moeda(total)
    if isinstance(eventos_aditivos, pd.DataFrame) and not eventos_aditivos.empty:
        candidatos = [c for c in eventos_aditivos.columns if "valor" in c.lower()]
        for col in candidatos[:1]:
            serie = pd.to_numeric(eventos_aditivos[col], errors="coerce").fillna(0)
            total = float(serie.sum())
            if abs(total) > 0.004:
                return moeda(total)
    return "[campo a preencher]"


def montar_saneador_integrado(df_infos, resultado_vg, resultado_garantia, df_checklist, eventos_aditivos, complemento_manual=""):
    df_infos = normalizar_df_infos(df_infos)

    numero_contrato = info_por_alias(df_infos, ["Número do contrato", "Contrato", "Nº do contrato"])
    contratada = info_por_alias(df_infos, ["Contratada", "Nome da contratada"])
    objeto = info_por_alias(df_infos, ["Objeto do contrato", "Objeto"])
    data_final = info_por_alias(df_infos, ["Data final vigente", "Data final da vigência", "Vigência final"])
    data_proposta = info_por_alias(df_infos, ["Data da proposta", "Data-base da proposta"])
    indice_info = info_por_alias(df_infos, ["Índice contratual", "Índice de reajuste"])
    data_pedido = info_por_alias(df_infos, ["Data do pedido da empresa", "Data do pedido", "Data do pleito"])
    documento_pleito = info_por_alias(df_infos, ["Ofício de Pleito da Empresa", "Documento do pleito", "Pleito da empresa"])
    houve_reajuste = info_por_alias(df_infos, ["Já houve reajuste anterior?", "Reajuste anterior"])
    percentual_anterior = info_por_alias(df_infos, ["Percentual já concedido", "Percentual anterior"])
    data_efeitos_anterior = info_por_alias(df_infos, ["Data de efeitos financeiros anterior", "Efeitos financeiros anteriores"])
    ultimo_termo = info_por_alias(df_infos, ["Último Termo de Apostila", "Último termo"])
    adequacao_doc = info_por_alias(df_infos, ["Documento da adequação orçamentária", "Adequação orçamentária"])
    adequacao_valor = info_por_alias(df_infos, ["Valor da adequação orçamentária", "Valor adequado"], padrao="")
    certidoes = info_por_alias(df_infos, ["Certidões de regularidade", "Certidões", "Regularidade fiscal"])
    concordancia = info_por_alias(df_infos, ["Concordância da contratada", "Manifestação da contratada", "Aceite da contratada"])
    documentos_desatualizados = info_por_alias(df_infos, ["Documentos desatualizados", "Documentos a desconsiderar"])
    documentacao_apoio = info_por_alias(df_infos, ["Documentação de apoio", "Documentos de apoio"])

    resultado_adequacao = st.session_state.get("resultado_adequacao_orcamentaria")
    if isinstance(resultado_adequacao, dict) and resultado_adequacao:
        for chave in ["complementacao_necessaria", "complementacao", "valor_adequacao", "valor_total"]:
            valor = _valor_numerico_positivo(resultado_adequacao.get(chave))
            if valor is not None:
                adequacao_valor = moeda(valor)
                break
        if adequacao_doc == "[campo a preencher]":
            adequacao_doc = texto_seguro(resultado_adequacao.get("documento"), "[campo a preencher]")

    if not adequacao_valor:
        adequacao_valor = "[campo a preencher]"

    vg_disponivel = isinstance(resultado_vg, dict) and bool(resultado_vg)
    modo_apuracao = texto_seguro(resultado_vg.get("modo_apuracao") if vg_disponivel else None, "Completo")
    modo_reduzido_estoque = modo_apuracao == "Reduzido por Itens/Estoque"
    modo_consumo_itens_ciclo = modo_apuracao == "Consumo por Itens/Ciclo"

    indice_resultado = texto_seguro(resultado_vg.get("indice") if vg_disponivel else None, indice_info)
    qtd_ciclos = texto_seguro(resultado_vg.get("quantidade_ciclos") if vg_disponivel else None)
    valor_original = moeda(resultado_vg.get("valor_original_contrato", 0)) if vg_disponivel else "[campo a preencher]"
    valor_pago = moeda(resultado_vg.get("valor_pago_efetivo", resultado_vg.get("total_pago_faturado", 0))) if vg_disponivel else "[campo a preencher]"
    valor_teorico = moeda(resultado_vg.get("valor_teorico_calculado", resultado_vg.get("total_devido_reajustado", 0))) if vg_disponivel else "[campo a preencher]"
    valor_represado = moeda(resultado_vg.get("valor_represado_a_pagar", 0)) if vg_disponivel else "[campo a preencher]"
    remanescente = moeda(resultado_vg.get("remanescente_reajustado", 0)) if vg_disponivel else "[campo a preencher]"
    valor_executado = moeda(resultado_vg.get("valor_executado_atualizado", 0)) if vg_disponivel else "[campo a preencher]"
    valor_total = moeda(resultado_vg.get("valor_atualizado_contrato", 0)) if vg_disponivel else "[campo a preencher]"
    ciclo_rem = texto_seguro(resultado_vg.get("ciclo_ultimo_remanescente") if vg_disponivel else None)
    variacao = pct(resultado_vg.get("variacao_acumulada", 0)) if vg_disponivel else "[campo a preencher]"
    delta_ciclos = extrair_delta_por_ciclo(resultado_vg if vg_disponivel else {})
    aditivos_atualizados = extrair_valor_aditivos(resultado_vg if vg_disponivel else {}, eventos_aditivos)

    contrato_txt = "contrato"
    if numero_contrato and numero_contrato != "[campo a preencher]":
        contrato_txt = f"Contrato {numero_contrato}"
    if contratada and contratada != "[campo a preencher]":
        contrato_txt += f", celebrado com {contratada}"

    titulo = "Despacho Saneador para Formalização de Termo de Apostila"
    assunto = f"Assunto: Saneamento da instrução — Contrato {numero_contrato}."

    paragrafos = [titulo, assunto]

    itens = [
        (f"1. Realiza-se o saneamento processual relativo ao {contrato_txt}, cujo objeto é {objeto}. "
         f"A vigência final informada é {data_final}."),
        (f"2. A contratada apresentou pleito de reajuste por meio de {documento_pleito}, com data de pedido registrada em {data_pedido}. "
         f"Para fins de verificação da anualidade, foi considerada a data da proposta em {data_proposta}, bem como o índice contratual {indice_resultado}."),
        (f"3. Quanto ao histórico anterior, foi informado que {houve_reajuste} quanto à existência de reajuste anterior, "
         f"com percentual já concedido de {percentual_anterior}, efeitos financeiros anteriores em {data_efeitos_anterior} "
         f"e último Termo de Apostila identificado como {ultimo_termo}."),
        (f"4. A análise de reajuste considerou {qtd_ciclos} ciclo(s), com variação acumulada de {variacao}. "
         f"O valor original do contrato informado foi de {valor_original}."),
        ("5. A área gestora encaminhou a execução sob a forma de itens consumidos por ciclo, sem apresentar base financeira mensal por competência. "
         f"Com base na premissa fiscal de equivalência entre consumo, medição/aprovação e faturamento devido, foi apurado Retroativo (itens consumidos/ciclo) de {moeda(resultado_vg.get('valor_retroativo_consumo_itens_ciclo', 0))}.{delta_ciclos}"
         if modo_consumo_itens_ciclo else
         (f"5. A apuração financeira consolidada indicou valor pago efetivo de {valor_pago} e valor teórico calculado de {valor_teorico}, "
          f"resultando em valor represado a pagar de {valor_represado}.{delta_ciclos}"
          if not modo_reduzido_estoque else
          "5. A área gestora encaminhou as informações exclusivamente sob a forma de itens e saldos remanescentes, sem apresentar a base de execução mensal por competência. "
          "Por essa razão, a presente análise foi processada em modo reduzido, com natureza estimativa, não substituindo a validação financeira prévia à formalização de pagamento.")),
        (f"6. Para fins de consolidação contratual, o Valor Total Atualizado do Contrato foi apurado em {valor_total}, composto pela "
         f"execução atualizada por ciclo, no montante de {valor_executado}, somada ao saldo remanescente atualizado de {remanescente}, "
         f"correspondente ao ciclo/referência {ciclo_rem}."),
        (f"7. Quanto aos aditivos e supressões, registra-se o valor atualizado consolidado de {aditivos_atualizados}."),
        (f"8. Foi realizada a adequação orçamentária necessária ao prosseguimento da instrução, no valor de {adequacao_valor}, "
         f"conforme documento {adequacao_doc}."),
        (f"9. As certidões de regularidade foram juntadas aos autos, conforme documento(s) {certidoes}."),
        (f"10. A contratada manifestou concordância com os valores propostos, conforme registrado em {concordancia}."),
        ("11. A contratada deverá ser informada quanto à necessidade de apresentação do endosso da garantia contratual, quando aplicável, "
         "observando-se o prazo e as condições previstos no contrato."),
        (f"12. Após atualizações e alinhamentos internos, alguns documentos instruídos mostram-se desatualizados, devendo ser desconsiderados: {documentos_desatualizados}."),
    ]

    complemento_manual = (complemento_manual or "").strip()
    if complemento_manual:
        itens.append(f"13. Registro complementar: {complemento_manual}")
        numero_final = 14
    else:
        numero_final = 13

    itens.append(
        f"{numero_final}. Após a conferência dos documentos acima indicados e saneadas eventuais pendências remanescentes, "
        "a instrução poderá prosseguir para formalização do Termo de Apostila, observadas as alçadas competentes e os procedimentos internos aplicáveis."
    )
    itens.append(f"{numero_final + 1}. Documentação de apoio: {documentacao_apoio}.")

    paragrafos.extend(itens)
    texto_final = "\n\n".join(paragrafos)
    return limpar_texto_saneador(texto_final)


def gerar_docx_texto(texto):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    def aplicar_sombreamento(paragraph, fill="FFF2CC"):
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        pPr.append(shd)

    def add_texto_com_placeholder(paragraph, conteudo):
        partes = re.split(r"(\[campo a preencher[^\]]*\])", str(conteudo), flags=re.IGNORECASE)
        for parte in partes:
            if not parte:
                continue
            run = paragraph.add_run(parte)
            if re.fullmatch(r"\[campo a preencher[^\]]*\]", parte, flags=re.IGNORECASE):
                run.bold = True
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    paragrafos = [p.strip() for p in str(texto).split("\n\n") if p.strip()]
    corpo = []
    if paragrafos:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(paragrafos[0])
        run.bold = True
        run.font.size = Pt(14)
        corpo = paragrafos[1:]

    for par in corpo:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if par.startswith("Assunto:"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run("Assunto:")
            run.bold = True
            add_texto_com_placeholder(p, par[len("Assunto:"):])
            continue
        if "Após atualizações e alinhamentos internos" in par:
            aplicar_sombreamento(p, "EAF2F8")
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            add_texto_com_placeholder(p, par)
            for run in p.runs:
                run.font.color.rgb = RGBColor(18, 59, 99)
            continue
        add_texto_com_placeholder(p, par)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


render_marca_topo()
st.title("Saneador")
st.caption("Minuta narrativa integrada que consolida Infos Prévias, análises e resultados antes da assinatura do Termo de Apostila.")
render_aviso_privacidade(tem_download=True)
resultado_vg_ui = st.session_state.get("resultado_valor_global", {}) or {}
if isinstance(resultado_vg_ui, dict) and resultado_vg_ui.get("modo_apuracao") == "Consumo por Itens/Ciclo":
    st.markdown("""<div style="background:#F6F3EE; border:1px solid #7A8F63; border-left:6px solid #4E6E58; border-radius:12px; padding:14px 16px; margin:10px 0 16px 0; color:#2F3E2F;">
        <div style="font-weight:800; margin-bottom:4px;">Modo Consumo por Itens/Ciclo</div>
        <div style="font-size:0.95rem; line-height:1.45;">O saneador será redigido com base em consumo itemizado por ciclo e saldo remanescente atualizado, sem tratar o resultado como retroativo financeiro mensal definitivo.</div>
    </div>""", unsafe_allow_html=True)

df_infos = normalizar_df_infos(st.session_state.get("infos_previas_df"))
resultado_vg = st.session_state.get("resultado_valor_global")
resultado_garantia = st.session_state.get("resultado_garantia")
df_checklist = st.session_state.get("checklist_processual")
eventos_aditivos = st.session_state.get("avaliacao_aditivos_eventos")

st.info("O Saneador gera uma minuta formal editável. Revise o texto antes de baixar o DOCX.")

with st.expander("Complemento manual opcional", expanded=False):
    complemento_manual = st.text_area(
        "Usar apenas para ressalvas ou informações que ainda não estejam nos módulos da ferramenta.",
        value=st.session_state.get("saneador_complemento_manual", ""),
        height=100,
    )
    st.session_state["saneador_complemento_manual"] = complemento_manual

texto_gerado = montar_saneador_integrado(
    df_infos,
    resultado_vg if isinstance(resultado_vg, dict) else {},
    resultado_garantia if isinstance(resultado_garantia, dict) else {},
    df_checklist if isinstance(df_checklist, pd.DataFrame) else pd.DataFrame(),
    eventos_aditivos if isinstance(eventos_aditivos, pd.DataFrame) else pd.DataFrame(),
    complemento_manual=complemento_manual,
)

if st.button("Gerar/atualizar texto integrado", type="primary"):
    st.session_state["saneador_texto"] = texto_gerado
    st.rerun()

if "saneador_texto" not in st.session_state:
    st.session_state["saneador_texto"] = texto_gerado

texto_editado = st.text_area(
    "Texto do Saneador",
    value=st.session_state["saneador_texto"],
    height=680,
    help="Revise o texto antes de baixar. Campos sem informação aparecerão como [campo a preencher].",
)

st.session_state["saneador_texto"] = texto_editado

docx_bytes = gerar_docx_texto(texto_editado)
st.session_state["arquivo_saneador_docx"] = docx_bytes

st.download_button(
    "Baixar Saneador em DOCX", data=docx_bytes,
    file_name="Saneador_Instrucao_Processual.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    type="primary", use_container_width=True,
)

with st.expander("Base de Infos Prévias utilizada", expanded=False):
    st.dataframe(df_infos, use_container_width=True, hide_index=True)
