# -*- coding: utf-8 -*-
"""Vinculo XLSX -> leitura Python -> documento das TRES referencias do VTA.

Garante que Saneador, Apostila e Sumario Executivo consomem a posicao mais
recente disponivel (posicao atual / ultima abertura / integral reajustado) COM
classificacao explicita OFICIAL / REFERENCIA AUDITAVEL / COMPARATIVO, sem nunca
promover valor sombra a oficial (VTA_FINAL/B26 preservado).
"""
from __future__ import annotations

import datetime as _dt
import sys
from io import BytesIO
from pathlib import Path

import openpyxl
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "templates" / "COLETA_REAJUSTE_OFICIAL.xlsx"
sys.path.insert(0, str(ROOT / "tests"))

from test_sumario_executivo import leitura_simples_financeiro  # noqa: E402

from _leitor_masterfile_v10 import _ler_referencias_vta  # noqa: E402
from _templates_documentos import (  # noqa: E402
    gerar_despacho_saneador,
    gerar_termo_apostila,
)
from _sumario_executivo import (  # noqa: E402
    gerar_sumario_executivo_pdf,
    montar_dados_sumario_executivo,
)

REF_FIXTURE = {
    "disponivel": True,
    "forma1_posicao_atual": 127_975_842.65,
    "forma2_ultima_abertura": 137_375_560.29,
    "forma3_integral_reajustado": 145_000_000.00,
    "reconciliacao_valor": -9_399_717.64,
    "reconciliacao_status": "REVISE",
    "forma1_situacao": "DISPONIVEL PARA CONFERENCIA",
    "forma2_situacao": "ADOTADO - ABERTURA DO CICLO VIGENTE C1",
    "ciclo_ultima_abertura": 1,
    "fallback_razao": None,
    "ciclo_vigente": "C1",
    "data_posicao_atual": _dt.date(2027, 6, 15),
    "posicao_atual_disponivel": True,
}


def _texto_docx(b: bytes) -> str:
    doc = Document(BytesIO(b))
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cel in row.cells:
                partes.append(cel.text)
    return "\n".join(partes)


# --------------------------------------------------------------------------- #
# 1. Fonte no XLSX -> leitura Python (fail-closed sem recalculo)                #
# --------------------------------------------------------------------------- #
def test_leitor_referencias_fail_closed_no_template():
    wb = openpyxl.load_workbook(TEMPLATE, data_only=True)
    ref = _ler_referencias_vta(wb)
    assert ref["disponivel"] is True
    assert ref["forma1_posicao_atual"] is None
    assert ref["posicao_atual_disponivel"] is False
    assert ref["origem_leitura"] == "defined_names"
    assert ref["fontes"]["forma1"] == "AUDITORIA_SITUACAO_ATUAL_CONTRATO"


def test_leitor_referencias_le_da_tabela1_sintetica():
    wb = openpyxl.Workbook()
    res = wb.active
    res.title = "RESULTADOS"
    res["B10"], res["B11"], res["B12"], res["B13"] = 100.0, 200.0, 300.0, -100.0
    res["H10"], res["H11"], res["H13"] = "OK1", "OK2", "RECONCILIADO"
    mem = wb.create_sheet("MEMORIA_RESULTADOS")
    mem["W46"], mem["W47"] = 1, "abertura vigente completa"
    ctl = wb.create_sheet("CONTROLE")
    ctl["B2"] = "C1"
    ce = wb.create_sheet("CICLO_EM_EXECUCAO")
    ce["D5"] = _dt.date(2027, 6, 15)
    ce["A9"] = 999.0
    ref = _ler_referencias_vta(wb)
    assert ref["origem_leitura"] == "legacy_coordinates"
    assert ref["forma1_posicao_atual"] == 100.0
    assert ref["forma2_ultima_abertura"] == 200.0
    assert ref["forma3_integral_reajustado"] == 300.0
    assert ref["reconciliacao_valor"] == -100.0
    assert ref["reconciliacao_status"] == "RECONCILIADO"
    assert ref["ciclo_ultima_abertura"] == 1
    assert ref["ciclo_vigente"] == "C1"
    assert ref["posicao_atual_disponivel"] is True
    assert ref["data_posicao_atual"] == _dt.date(2027, 6, 15)


# --------------------------------------------------------------------------- #
# 2. Leitura Python -> documento (DOCX): classificacao e valores               #
# --------------------------------------------------------------------------- #
def test_saneador_nao_repete_referencias_alternativas_do_vta():
    leit = leitura_simples_financeiro()
    leit["referencias_vta"] = dict(REF_FIXTURE)
    texto = _texto_docx(gerar_despacho_saneador(leit))
    for termo in (
        "OFICIAL", "REFERÊNCIA AUDITÁVEL", "COMPARATIVO",
        "posição atual do contrato", "última posição de abertura",
        "integralmente reajustado", "Cadeia homologada do VTA",
    ):
        assert termo not in texto
    assert "Valor Total Atualizado do Contrato" in texto


def test_saneador_posicao_indisponivel_nao_vira_bloco_autonomo():
    leit = leitura_simples_financeiro()
    ref = dict(REF_FIXTURE)
    ref.update(forma1_posicao_atual=None, data_posicao_atual=None,
               posicao_atual_disponivel=False,
               forma1_situacao="INDISPONIVEL - POSICAO ATUAL NAO INFORMADA")
    leit["referencias_vta"] = ref
    texto = _texto_docx(gerar_despacho_saneador(leit))
    assert "posição atual" not in texto.lower()
    assert "REFERÊNCIA AUDITÁVEL" not in texto


# --------------------------------------------------------------------------- #
# 3. Leitura Python -> documento (PDF): bloco presente                          #
# --------------------------------------------------------------------------- #
def test_pdf_referencias_presentes():
    leit = leitura_simples_financeiro()
    leit["referencias_vta"] = dict(REF_FIXTURE)
    dados = montar_dados_sumario_executivo(leit)
    assert dados["referencias_vta"]["forma2_ultima_abertura"] == 137_375_560.29
    b = gerar_sumario_executivo_pdf(dados)
    assert b[:4] == b"%PDF" and len(b) > 3000


# --------------------------------------------------------------------------- #
# 4. VTA oficial nunca substituido pelas referencias auditaveis                 #
# --------------------------------------------------------------------------- #
def test_termo_preserva_vta_oficial_sem_referencias_alternativas():
    leit = leitura_simples_financeiro()
    leit["referencias_vta"] = dict(REF_FIXTURE)
    texto = _texto_docx(gerar_termo_apostila(leit))
    assert "Valor Total Atualizado do Contrato" in texto
    assert "Referências auditáveis" not in texto
    assert "REFERÊNCIA AUDITÁVEL" not in texto
    assert "COMPARATIVO" not in texto
    assert "integralmente reajustado" not in texto
