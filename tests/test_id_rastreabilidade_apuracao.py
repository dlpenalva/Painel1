"""Contratos da Prioridade 5: ID curto e rastreabilidade documental."""
from __future__ import annotations

import ast
import re
from io import BytesIO
from pathlib import Path

import openpyxl
import pandas as pd
from docx import Document
from _estado_apuracao_upload import (
    assinatura_conteudo_upload,
    id_apuracao_de_assinatura,
    invalidar_estado_caso,
)
from _seguranca_xlsx import opcoes_excel_writer_seguro
from _sumario_executivo import _rodape_factory
from _templates_documentos import (
    gerar_despacho_saneador,
    gerar_modelo_branco_despacho,
    gerar_modelo_branco_termo,
    gerar_termo_apostila,
)
from test_sumario_executivo import leitura_simples_financeiro


ROOT = Path(__file__).resolve().parents[1]
PAGINA = (ROOT / "pages" / "03_Valor_Global.py").read_text(encoding="utf-8")
ASSINATURA = assinatura_conteudo_upload(b"coleta-prioridade-5")
ID_APURACAO = id_apuracao_de_assinatura(ASSINATURA)


def _leitura_identificada() -> dict:
    leitura = leitura_simples_financeiro()
    leitura["assinatura_origem_xlsx"] = ASSINATURA
    leitura["id_apuracao"] = ID_APURACAO
    return leitura


def _texto_docx_com_rodape(conteudo: bytes) -> str:
    doc = Document(BytesIO(conteudo))
    partes = [p.text for p in doc.paragraphs]
    for secao in doc.sections:
        partes.extend(p.text for p in secao.footer.paragraphs)
    return "\n".join(partes)


def _gerador_planilha_executiva():
    arvore = ast.parse(PAGINA, filename=str(ROOT / "pages" / "03_Valor_Global.py"))
    nomes = {
        "numero_ciclo",
        "numero_seguro",
        "limpar_nan_inf_df",
        "normalizar_ciclo",
        "normalizar_texto",
        "numero_br",
        "texto_seguro",
        "formatar_data_br",
        "gerar_planilha_executiva",
    }
    funcoes = [
        no for no in arvore.body
        if isinstance(no, ast.FunctionDef) and no.name in nomes
    ]
    modulo = ast.Module(body=funcoes, type_ignores=[])
    ast.fix_missing_locations(modulo)
    espaco = {
        "BytesIO": BytesIO,
        "pd": pd,
        "re": re,
        "opcoes_excel_writer_seguro": opcoes_excel_writer_seguro,
    }
    exec(compile(modulo, str(ROOT / "pages" / "03_Valor_Global.py"), "exec"), espaco)
    return espaco["gerar_planilha_executiva"]


def test_id_tem_formato_canonico_e_valor_conhecido() -> None:
    assinatura = "0123456789abcdef" + "0" * 48
    assert id_apuracao_de_assinatura(assinatura) == "APUR-0123456789ABCDEF"


def test_mesmos_bytes_mesmo_id_e_bytes_diferentes_ids_diferentes() -> None:
    assert id_apuracao_de_assinatura(assinatura_conteudo_upload(b"A")) == id_apuracao_de_assinatura(assinatura_conteudo_upload(b"A"))
    assert id_apuracao_de_assinatura(assinatura_conteudo_upload(b"A")) != id_apuracao_de_assinatura(assinatura_conteudo_upload(b"B"))


def test_id_exige_sha_completo_e_nao_faz_hash_do_hash() -> None:
    assert ID_APURACAO == f"APUR-{ASSINATURA[:16].upper()}"
    for invalida in ("", ASSINATURA[:16], "g" * 64):
        try:
            id_apuracao_de_assinatura(invalida)
        except ValueError:
            pass
        else:
            raise AssertionError("assinatura incompleta/invalida foi aceita")


def test_a_para_a_preserva_id_e_a_para_b_remove_resultado_antigo() -> None:
    estado = {
        "assinatura_upload_docs": ASSINATURA,
        "assinatura_processada_upload_docs": ASSINATURA,
        "resultado_valor_global": {"id_apuracao": ID_APURACAO},
    }
    assert invalidar_estado_caso(estado, ASSINATURA) is False
    assert estado["resultado_valor_global"]["id_apuracao"] == ID_APURACAO
    assert invalidar_estado_caso(estado, assinatura_conteudo_upload(b"B")) is True
    assert "resultado_valor_global" not in estado


def test_id_so_entra_no_resultado_depois_do_processamento_bem_sucedido() -> None:
    bloco = PAGINA[PAGINA.index('if st.button("Processar"'):PAGINA.index('if st.session_state.get("assinatura_processada_upload_docs")')]
    atribuicao = 'resultado_processado["id_apuracao"] = id_apuracao_de_assinatura(assinatura_upload)'
    assert bloco.index("processar_coleta_oficial_runtime(conteudo_upload)") < bloco.index(atribuicao)
    assert bloco.index(atribuicao) < bloco.index('st.session_state["resultado_valor_global"] = resultado_processado')


def test_resultado_preserva_sha_completo_sem_expor_na_interface() -> None:
    assert 'resultado_processado["assinatura_origem_xlsx"] = assinatura_upload' in PAGINA
    fim = PAGINA.rindex("render_documentos_funcionais_upload(resultado)")
    inicio = PAGINA.rfind("if resultado:", 0, fim)
    trecho_ui = PAGINA[inicio:fim]
    assert 'st.caption(f"ID da apuração: {resultado[\'id_apuracao\']}")' in trecho_ui
    assert "assinatura_origem_xlsx" not in trecho_ui


def test_tres_documentos_principais_recebem_o_mesmo_id() -> None:
    leitura = _leitura_identificada()
    textos_canvas: list[str] = []

    class CanvasFalso:
        def saveState(self): pass
        def setStrokeColor(self, _cor): pass
        def setLineWidth(self, _largura): pass
        def line(self, *_args): pass
        def setFont(self, *_args): pass
        def setFillColor(self, _cor): pass
        def drawString(self, _x, _y, texto): textos_canvas.append(texto)
        def drawRightString(self, _x, _y, texto): textos_canvas.append(texto)
        def restoreState(self): pass

    class DocumentoFalso:
        leftMargin = 10
        rightMargin = 10
        pagesize = (600, 800)

    _rodape_factory({"id_apuracao": ID_APURACAO})(CanvasFalso(), DocumentoFalso())
    textos = [
        "\n".join(textos_canvas),
        _texto_docx_com_rodape(gerar_despacho_saneador(leitura)),
        _texto_docx_com_rodape(gerar_termo_apostila(leitura)),
    ]
    assert all(texto.count(ID_APURACAO) >= 1 for texto in textos)
    assert all(ASSINATURA not in texto for texto in textos)


def test_modelos_em_branco_nao_recebem_id_nem_placeholder() -> None:
    for conteudo in (gerar_modelo_branco_despacho(), gerar_modelo_branco_termo()):
        texto = _texto_docx_com_rodape(conteudo)
        assert "APUR-" not in texto
        assert "ID da apuração" not in texto


def test_planilha_executiva_usa_id_canonico_em_todas_as_abas() -> None:
    conteudo = _gerador_planilha_executiva()({"id_apuracao": ID_APURACAO})
    wb = openpyxl.load_workbook(BytesIO(conteudo), data_only=False)
    assert wb.sheetnames
    for ws in wb.worksheets:
        assert ID_APURACAO in (ws.oddFooter.left.text or "")
    assert ASSINATURA not in conteudo.decode("latin1", errors="ignore")
    assert "hashlib" not in PAGINA


def test_quatro_destinos_nao_derivam_id_proprio() -> None:
    fontes = [
        (ROOT / "_sumario_executivo.py").read_text(encoding="utf-8"),
        (ROOT / "_templates_documentos.py").read_text(encoding="utf-8"),
        PAGINA,
    ]
    assert all("APUR-" not in fonte for fonte in fontes)
    assert all("sha256(" not in fonte for fonte in fontes)
