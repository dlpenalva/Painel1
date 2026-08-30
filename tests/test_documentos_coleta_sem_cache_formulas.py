"""Fluxo documental ponta a ponta a partir de uma Coleta sem cache de formulas.

Cenario reproduzido (caso real relatado): ciclo vigente C1, percentual de
3,08%, efeitos financeiros a partir de 01/04/2026, situacao TEMPESTIVO e
nenhum aditivo.

A Coleta e gerada pelo mesmo caminho da Calculadora
(``gerar_coleta_oficial_preenchida``), que escreve o arquivo por openpyxl. Um
arquivo assim NAO possui valor calculado gravado nas formulas: openpyxl com
``data_only=True`` devolve ``None`` para todas elas. Esse e exatamente o estado
do arquivo que o fiscal reenvia quando nao o abriu no Excel.

O que estes testes fixam:

* dados que existem na Coleta como entrada literal (percentual do ciclo) devem
  chegar ao documento, mesmo quando a celula-formula que os consolida
  (parametros!FATOR_ACUMULADO) esta sem valor gravado;
* dados que realmente nao existem na Coleta continuam ausentes, sem virar zero
  nem valor inventado;
* campos de preenchimento manual continuam como marcador;
* a classificacao juridica das tres referencias do VTA nao se altera;
* o inicio dos efeitos financeiros permanece apurado como DATA EXATA no dado
  canonico e e apresentado nos documentos pela COMPETENCIA (mm/aaaa) dela
  derivada — o efeito alcanca a competencia inteira, nunca uma fracao de mes.
"""

from __future__ import annotations

import io
from datetime import date

import docx
import pytest
from openpyxl import load_workbook

from _coleta_oficial import gerar_coleta_oficial_preenchida
from _coleta_reajuste_documentos import processar_coleta_oficial_runtime
from _sumario_executivo import montar_dados_sumario_executivo
from _templates_documentos import gerar_despacho_saneador, gerar_termo_apostila


CICLO = "C1"
PERCENTUAL_C1 = 0.0307853139440224
# Marco temporal do cenario: unica fonte tanto da Coleta gerada quanto do
# valor esperado nos documentos. Nada de data fixa repetida a mao.
FINANCEIRO_INICIO_C1 = date(2026, 4, 1)


def _data_br(valor: date) -> str:
    return valor.strftime("%d/%m/%Y")


def _competencia(valor: date) -> str:
    return valor.strftime("%m/%Y")


def _dados_cenario() -> dict:
    return {
        "origem": "Cenario C1 sem cache de formulas",
        "indice": "IST (Anatel)",
        "data_base_original": "01/02/2025",
        "data_corte": date(2027, 1, 31),
        "ciclos": [{
            "ciclo": "C1",
            "data_inicio": date(2026, 2, 1),
            "data_fim": date(2027, 1, 31),
            "data_pedido": date(2026, 2, 1),
            "financeiro_inicio": FINANCEIRO_INICIO_C1,
            "percentual_aplicado": PERCENTUAL_C1,
            "situacao": "✅ TEMPESTIVO",
            "objeto_analise_atual": True,
        }],
    }


def _texto_docx(conteudo: bytes) -> str:
    documento = docx.Document(io.BytesIO(conteudo))
    partes = [p.text for p in documento.paragraphs]
    for tabela in documento.tables:
        for linha in tabela.rows:
            partes.append(" | ".join(c.text.strip() for c in linha.cells))
    return "\n".join(partes)


def _linha_do_quadro(texto: str, ciclo: str) -> str:
    """Linha da tabela de ciclos (docx serializado com ' | ') do ciclo dado."""
    linhas = [
        linha for linha in texto.splitlines()
        if ciclo in [celula.strip() for celula in linha.split("|")]
    ]
    assert len(linhas) == 1, (
        f"esperada exatamente 1 linha de quadro para {ciclo}; achadas: {linhas}"
    )
    return linhas[0]


@pytest.fixture(scope="module")
def coleta_sem_cache() -> bytes:
    return gerar_coleta_oficial_preenchida(_dados_cenario())


@pytest.fixture(scope="module")
def payload(coleta_sem_cache: bytes) -> dict:
    resultado, _diagnostico = processar_coleta_oficial_runtime(coleta_sem_cache)
    return resultado


@pytest.fixture(scope="module")
def texto_saneador(payload: dict) -> str:
    return _texto_docx(gerar_despacho_saneador(payload))


@pytest.fixture(scope="module")
def texto_apostila(payload: dict) -> str:
    return _texto_docx(gerar_termo_apostila(payload))


# --- pre-condicao: o arquivo realmente nao tem valor calculado nas formulas ---

def test_coleta_gerada_nao_possui_cache_de_formulas(coleta_sem_cache: bytes) -> None:
    com_formula = load_workbook(io.BytesIO(coleta_sem_cache), data_only=False)
    com_valor = load_workbook(io.BytesIO(coleta_sem_cache), data_only=True)
    formulas = 0
    for aba in com_formula.sheetnames:
        origem, destino = com_formula[aba], com_valor[aba]
        for linha in origem.iter_rows():
            for celula in linha:
                if isinstance(celula.value, str) and celula.value.startswith("="):
                    formulas += 1
                    assert destino[celula.coordinate].value is None, (
                        f"{aba}!{celula.coordinate} nao deveria ter valor gravado"
                    )
    assert formulas > 0, "cenario invalido: a Coleta precisa conter formulas"


def test_percentual_do_ciclo_e_literal_e_fator_acumulado_e_formula(
    coleta_sem_cache: bytes,
) -> None:
    """Fixa a assimetria que origina o defeito: E literal, F formula."""
    com_formula = load_workbook(io.BytesIO(coleta_sem_cache), data_only=False)
    parametros = com_formula["parametros"]
    assert parametros["E3"].value == pytest.approx(PERCENTUAL_C1)
    assert str(parametros["F3"].value or "").startswith("=")


# --- dado disponivel na Coleta chega ao documento ---

def test_payload_recompoe_o_percentual_acumulado(payload: dict) -> None:
    assert payload.get("variacao_acumulada") == pytest.approx(PERCENTUAL_C1)
    assert payload.get("fator_acumulado") == pytest.approx(1.0 + PERCENTUAL_C1)


def test_saneador_traz_a_variacao_acumulada(texto_saneador: str) -> None:
    assert "3,08%" in texto_saneador
    assert "[PREENCHER: Variacao acumulada]" not in texto_saneador


def test_apostila_traz_o_percentual_acumulado(texto_apostila: str) -> None:
    assert "percentual acumulado apurado de 3,08%" in texto_apostila
    assert "[PREENCHER: Percentual acumulado apurado]" not in texto_apostila


@pytest.fixture(scope="module")
def ciclo_apurado(payload: dict) -> dict:
    """Ciclo como a consolidacao canonica o entrega aos dois geradores."""
    dados = montar_dados_sumario_executivo(payload, None)
    ciclos = [c for c in (dados.get("ciclos") or []) if c.get("ciclo") == CICLO]
    assert len(ciclos) == 1, f"cenario invalido: {CICLO} nao foi consolidado"
    return ciclos[0]


def test_inicio_dos_efeitos_permanece_data_exata_no_dado_apurado(
    ciclo_apurado: dict,
) -> None:
    """A conversao para competencia e de apresentacao: a data nao se perde."""
    assert ciclo_apurado["inicio_efeito_financeiro"] == _data_br(
        FINANCEIRO_INICIO_C1
    )


def test_ciclo_percentual_efeitos_e_situacao_no_quadro_de_ciclos(
    ciclo_apurado: dict, texto_saneador: str, texto_apostila: str
) -> None:
    # Esperado derivado do proprio marco do cenario: data exata apurada ->
    # competencia declarada. Trocar o marco muda o esperado junto.
    data_exata = _data_br(FINANCEIRO_INICIO_C1)
    assert ciclo_apurado["inicio_efeito_financeiro"] == data_exata
    competencia = _competencia(FINANCEIRO_INICIO_C1)
    for texto in (texto_saneador, texto_apostila):
        linha = _linha_do_quadro(texto, CICLO)
        assert "3,08%" in linha
        assert f"A partir de {competencia}" in linha
        assert "TEMPESTIVO" in linha
        # O quadro declara a competencia, nunca o dia do inicio do efeito.
        assert data_exata not in linha


# --- ausencias legitimas permanecem ausentes ---

def test_ausencia_de_aditivos_gera_a_redacao_correta(texto_apostila: str) -> None:
    assert (
        "Não foram identificados aditivos ou supressões específicos na base processada"
        in texto_apostila
    )


def test_campos_manuais_permanecem_como_marcador(
    texto_saneador: str, texto_apostila: str
) -> None:
    assert "[PREENCHER: Numero do contrato]" in texto_saneador
    assert "[PREENCHER: Numero do contrato]" in texto_apostila
    assert "[PREENCHER: Numero do processo de instrucao]" in texto_apostila


def test_posicao_atual_nao_e_repetida_no_saneador_enxuto(
    texto_saneador: str,
) -> None:
    assert "Posição atual" not in texto_saneador


# --- classificacao juridica das referencias do VTA ---

def test_referencias_alternativas_do_vta_nao_sao_renderizadas_no_saneador(
    texto_saneador: str,
) -> None:
    for termo in ("REFERÊNCIA AUDITÁVEL", "COMPARATIVO", "cadeia homologada"):
        assert termo not in texto_saneador
    assert "Valor Total Atualizado do Contrato" in texto_saneador


def test_documentos_nao_expoem_jargao_tecnico(
    texto_saneador: str, texto_apostila: str
) -> None:
    for texto in (texto_saneador, texto_apostila):
        for jargao in ("parametros!", "CONTROLE!B", "RESULTADOS!B", "fator_acumulado",
                       "variacao_acumulada", "data_only"):
            assert jargao not in texto


# --- os dois documentos consomem o mesmo payload consolidado ---

def test_saneador_e_apostila_consomem_o_mesmo_payload(payload: dict) -> None:
    primeiro = gerar_despacho_saneador(payload)
    segundo = gerar_termo_apostila(payload)
    assert primeiro and segundo
    for texto in (_texto_docx(primeiro), _texto_docx(segundo)):
        assert "3,08%" in texto
        assert "IST (Anatel)" in texto


def test_sumario_executivo_nao_regride(payload: dict) -> None:
    from _sumario_executivo import gerar_sumario_executivo

    pdf = gerar_sumario_executivo(payload)
    assert pdf[:4] == b"%PDF"
