"""Protecao dos dois ajustes documentais desta entrega.

1A  Data do pedido no Quadro 1 do Despacho Saneador, vinda da apuracao
    (parametros!DATA_PEDIDO) e nao de redigitacao manual.
1B  Declaracao expressa das competencias que nao produzem efeitos
    financeiros, no Saneador e no Termo de Apostila, a partir da MESMA
    fonte temporal canonica ja consolidada pela apuracao.
"""
from __future__ import annotations

import sys
from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import Workbook

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from _gerador_masterfile import _registrar_datas_pedido  # noqa: E402
from _leitor_masterfile_v10 import _ler_parametros_v10  # noqa: E402
from _sumario_executivo import montar_dados_sumario_executivo  # noqa: E402
from _templates_documentos import (  # noqa: E402
    gerar_despacho_saneador,
    gerar_modelo_branco_despacho,
    gerar_modelo_branco_termo,
    gerar_termo_apostila,
)

from test_sumario_executivo import (  # noqa: E402
    leitura_multiciclo_pc,
    leitura_simples_financeiro,
)
from test_templates_documentos import (  # noqa: E402
    CAMPOS_SANEADOR,
    CAMPOS_TERMO,
)

FRASE_INICIO = "Em razão da data do pedido, os efeitos financeiros do reajuste"
FRASE_BRANCO = "Havendo competências não alcançadas pelos efeitos financeiros"


def _texto(docx_bytes: bytes) -> str:
    doc = Document(BytesIO(docx_bytes))
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                partes.append(cell.text)
    return "\n".join(partes)


def _paragrafos(docx_bytes: bytes) -> list[str]:
    return [p.text for p in Document(BytesIO(docx_bytes)).paragraphs if p.text.strip()]


def _quadro1(docx_bytes: bytes):
    return Document(BytesIO(docx_bytes)).tables[0]


def _com_datas_pedido(leitura: dict, datas: dict[str, date]) -> dict:
    por_ciclo = leitura["parametros_v10"]["por_ciclo"]
    for nome, valor in datas.items():
        por_ciclo[nome]["data_pedido"] = valor
    return leitura


# ---------------------------------------------------------------------------
# 1A - data do pedido
# ---------------------------------------------------------------------------

def test_data_pedido_do_ciclo_aparece_no_quadro1():
    leitura = _com_datas_pedido(leitura_simples_financeiro(), {"C1": date(2025, 3, 20)})
    tabela = _quadro1(gerar_despacho_saneador(leitura, campos_manuais=CAMPOS_SANEADOR))
    linha = next(r for r in tabela.rows if r.cells[0].text == "C1")
    assert linha.cells[2].text == "20/03/2025"


def test_cada_ciclo_recebe_a_sua_propria_data_de_pedido():
    leitura = _com_datas_pedido(
        leitura_multiciclo_pc(),
        {"C1": date(2024, 4, 15), "C2": date(2025, 7, 10)},
    )
    tabela = _quadro1(gerar_despacho_saneador(leitura, campos_manuais=CAMPOS_SANEADOR))
    datas = {r.cells[0].text: r.cells[2].text for r in tabela.rows[1:]}
    assert datas["C1"] == "15/04/2024"
    assert datas["C2"] == "10/07/2025"


def test_ausencia_real_de_data_mantem_politica_documental():
    """Sem data na apuracao: marcador de preenchimento, jamais data inventada."""
    tabela = _quadro1(
        gerar_despacho_saneador(
            leitura_simples_financeiro(), campos_manuais=CAMPOS_SANEADOR
        )
    )
    linha = next(r for r in tabela.rows if r.cells[0].text == "C1")
    assert linha.cells[2].text == "[PREENCHER: Data do pedido]"


def test_identificacao_explicita_tem_precedencia_sobre_a_apuracao():
    leitura = _com_datas_pedido(leitura_simples_financeiro(), {"C1": date(2025, 3, 20)})
    dados = montar_dados_sumario_executivo(
        leitura, identificacao={"datas_pedido": {"C1": date(2025, 5, 2)}}
    )
    c1 = next(c for c in dados["ciclos"] if c["ciclo"] == "C1")
    assert c1["data_pedido"] == "02/05/2025"


def test_modelo_em_branco_nao_carrega_data_real():
    texto = _texto(gerar_modelo_branco_despacho())
    assert "[PREENCHER: Data do pedido]" in texto
    assert "20/03/2025" not in texto


def test_quadro1_preserva_as_demais_colunas():
    leitura = _com_datas_pedido(leitura_multiciclo_pc(), {"C2": date(2025, 7, 10)})
    tabela = _quadro1(gerar_despacho_saneador(leitura, campos_manuais=CAMPOS_SANEADOR))
    assert [c.text for c in tabela.rows[0].cells] == [
        "Ciclo", "Data-base", "Data do pedido", "Situação",
        "Efeito financeiro", "Percentual",
    ]
    linha = next(r for r in tabela.rows if r.cells[0].text == "C2")
    assert linha.cells[1].text == "01/05/2025"
    assert linha.cells[4].text == "A partir de 01/08/2025"


def test_gerador_grava_e_leitor_recupera_data_pedido():
    wb = Workbook()
    ws = wb.active
    ws.title = "parametros"
    cabecalhos = [
        "COMPUTAR_NESTA_APURACAO", "CICLO", "DATA_INICIO", "DATA_FIM",
        "PERCENTUAL_DO_CICLO", "FATOR_ACUMULADO", "SITUACAO",
    ]
    for coluna, titulo in enumerate(cabecalhos, start=1):
        ws.cell(1, coluna).value = titulo
    for linha, nome in enumerate(("C0", "C1", "C2", "C3", "C4"), start=2):
        ws.cell(linha, 2).value = nome

    _registrar_datas_pedido(wb, {"C1": {"data_pedido": date(2025, 3, 20)}})

    assert ws["U1"].value == "DATA_PEDIDO"
    lido = _ler_parametros_v10(wb)
    assert lido["por_ciclo"]["C1"]["data_pedido"] == date(2025, 3, 20)
    assert lido["por_ciclo"]["C2"]["data_pedido"] is None


def test_arquivo_sem_a_coluna_nao_inventa_data():
    wb = Workbook()
    ws = wb.active
    ws.title = "parametros"
    ws["A1"], ws["B1"] = "COMPUTAR_NESTA_APURACAO", "CICLO"
    ws["B2"] = "C1"
    lido = _ler_parametros_v10(wb)
    assert lido["por_ciclo"]["C1"]["data_pedido"] is None


# ---------------------------------------------------------------------------
# 1B - competencias sem efeitos financeiros
# ---------------------------------------------------------------------------

def test_perda_de_duas_competencias_e_declarada_no_saneador():
    # C1 abre em 01/02/2025 e o efeito financeiro so comeca em 01/04/2025.
    texto = _texto(
        gerar_despacho_saneador(
            leitura_simples_financeiro(), campos_manuais=CAMPOS_SANEADOR
        )
    )
    assert (
        "Em razão da data do pedido, os efeitos financeiros do reajuste deste "
        "ciclo iniciam-se em 01/04/2025, não alcançando as competências de "
        "fevereiro e março de 2025." in texto
    )


def test_perda_de_uma_unica_competencia_usa_singular():
    leitura = leitura_simples_financeiro()
    leitura["parametros_v10"]["por_ciclo"]["C1"]["inicio_efeito_financeiro"] = (
        date(2025, 3, 1)
    )
    texto = _texto(gerar_despacho_saneador(leitura, campos_manuais=CAMPOS_SANEADOR))
    assert "não alcançando a competência de fevereiro de 2025." in texto
    assert "não alcançando as competências" not in texto


def test_ciclo_tempestivo_sem_perda_nao_recebe_texto():
    leitura = leitura_simples_financeiro()
    leitura["parametros_v10"]["por_ciclo"]["C1"]["inicio_efeito_financeiro"] = (
        date(2025, 2, 1)
    )
    assert FRASE_INICIO not in _texto(
        gerar_despacho_saneador(leitura, campos_manuais=CAMPOS_SANEADOR)
    )


def test_somente_o_ciclo_afetado_recebe_o_texto():
    # C1 abre e produz efeito em 01/05/2024 (sem perda);
    # C2 abre em 01/05/2025 e so produz efeito em 01/08/2025 (perde 3 meses).
    texto = _texto(
        gerar_despacho_saneador(leitura_multiciclo_pc(), campos_manuais=CAMPOS_SANEADOR)
    )
    assert (
        "os efeitos financeiros do reajuste do ciclo C2 iniciam-se em "
        "01/08/2025, não alcançando as competências de maio, junho e julho "
        "de 2025." in texto
    )
    assert "do ciclo C1 iniciam-se" not in texto
    assert texto.count(FRASE_INICIO) == 1


def test_ciclo_precluso_preserva_semantica_anterior():
    leitura = leitura_simples_financeiro()
    leitura["parametros_v10"]["por_ciclo"]["C1"]["situacao"] = "PRECLUSO"
    texto = _texto(gerar_despacho_saneador(leitura, campos_manuais=CAMPOS_SANEADOR))
    assert FRASE_INICIO not in texto
    assert "Sem efeitos financeiros" in texto


def test_sem_inicio_de_efeito_registrado_nada_e_declarado():
    leitura = leitura_simples_financeiro()
    leitura["parametros_v10"]["por_ciclo"]["C1"]["inicio_efeito_financeiro"] = None
    assert FRASE_INICIO not in _texto(
        gerar_despacho_saneador(leitura, campos_manuais=CAMPOS_SANEADOR)
    )


def test_termo_de_apostila_declara_a_mesma_perda():
    frase_saneador = [
        p for p in _paragrafos(
            gerar_despacho_saneador(
                leitura_multiciclo_pc(), campos_manuais=CAMPOS_SANEADOR
            )
        )
        if p.startswith(FRASE_INICIO)
    ]
    frase_termo = [
        p for p in _paragrafos(
            gerar_termo_apostila(leitura_multiciclo_pc(), campos_manuais=CAMPOS_TERMO)
        )
        if p.startswith(FRASE_INICIO)
    ]
    assert frase_saneador == frase_termo
    assert len(frase_termo) == 1


def test_modelos_em_branco_usam_texto_generico():
    for gerar in (gerar_modelo_branco_despacho, gerar_modelo_branco_termo):
        texto = _texto(gerar())
        assert FRASE_BRANCO in texto
        assert FRASE_INICIO not in texto
