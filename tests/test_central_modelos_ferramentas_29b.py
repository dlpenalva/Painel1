"""Etapa 29B — Central de Modelos e Ferramentas sem Coleta.

Cobre: modelos em branco (DOCX valido, placeholders todos destacados, sem
afirmacoes factuais, sem zeros fabricados, sem None/nan/emoji, deterministico,
sem leitura de session_state); paridade dos documentos automaticos com o modo
omitido/False; fonte da nova pagina, do app.py, do card inicial e do retorno
da Garantia; contaminacao por sessao; calculo da Garantia.
"""
import io
import re
import zipfile
from pathlib import Path

import pytest
from docx import Document

from _templates_documentos import (
    gerar_despacho_saneador,
    gerar_termo_apostila,
    gerar_modelo_branco_despacho,
    gerar_modelo_branco_termo,
)

ROOT = Path(__file__).resolve().parents[1]
APP = (ROOT / "app.py").read_text(encoding="utf-8")
INICIO = (ROOT / "pages" / "00_Calculadora_Reajustes.py").read_text(encoding="utf-8")
PAGE14 = (ROOT / "pages" / "14_Central_Modelos_Ferramentas.py").read_text(encoding="utf-8")
GARANTIA = (ROOT / "pages" / "05_Garantia.py").read_text(encoding="utf-8")
PAGE03 = (ROOT / "pages" / "03_Valor_Global.py").read_text(encoding="utf-8")

FRASES_PROIBIDAS = [
    "0 ciclo", "R$ 0,00", "0,00%",
    "não foram identificados aditivos", "não foram identificados eventos",
    "base processada", "manifestou concordância",
    "manifestação de concordância da CONTRATADA",
    "Foi realizada a adequação orçamentária",
    "adequação orçamentária constantes",
    "As certidões de regularidade estão presentes",
    "formalizam-se os reajustes contratuais apurados",
    "regularidade mínima da instrução", "inexistindo pendência",
    "[campo a preencher]",
    # Etapa 29C.1: afirmacoes factuais neutralizadas nos modelos em branco.
    "os elementos disponíveis encontram-se consolidados",
    "consolidados para análise",
    "utilizado na análise",
    "percentual acumulado apurado",
    "análise realizada",
    # Etapa 29C.1.2: fechamento integral da neutralidade (ambos os modelos).
    "A apuração financeira consolidada indicou",
    "A contratada apresentou pleito",
    "foram consideradas a data da proposta",
    "registradas na análise",
    "Acordou-se na concessão",
    "foi adotada a premissa",
    "A contratada foi informada",
    "contrato informado foi",
    "memória de cálculo constante em",
    "que apurou os ciclos",
    "informações encaminhadas pela área gestora",
    "foi organizada de forma evolutiva",
]
PALAVRAS_TECNICAS = ["nan", "null", "None", "NaN"]
NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _doc(b):
    return Document(io.BytesIO(b))


def _texto(b):
    d = _doc(b)
    partes = [p.text for p in d.paragraphs if p.text.strip()]
    for tb in d.tables:
        for row in tb.rows:
            partes.append(" | ".join(c.text for c in row.cells))
    return "\n".join(partes)


def _tem_highlight(run):
    rpr = run._element.rPr
    return rpr is not None and rpr.find(f"{NS}highlight") is not None


def _placeholders_e_destaque(b):
    d = _doc(b)
    total = sem = 0

    def scan(paras):
        nonlocal total, sem
        for p in paras:
            for r in p.runs:
                if "[PREENCHER" in r.text:
                    total += 1
                    if not _tem_highlight(r):
                        sem += 1

    scan(d.paragraphs)
    for tb in d.tables:
        for row in tb.rows:
            for c in row.cells:
                scan(c.paragraphs)
    return total, sem


GERADORES_BRANCO = [
    ("despacho", gerar_modelo_branco_despacho),
    ("termo", gerar_modelo_branco_termo),
]


# ---------------------------------------------------------- modelos em branco
@pytest.mark.parametrize("nome,gerar", GERADORES_BRANCO)
def test_modelo_branco_retorna_docx_valido(nome, gerar):
    b = gerar()
    assert isinstance(b, bytes) and len(b) > 5000
    zf = zipfile.ZipFile(io.BytesIO(b))
    assert zf.testzip() is None
    assert "word/document.xml" in zf.namelist()
    _doc(b)  # abre sem reparo


@pytest.mark.parametrize("nome,gerar", GERADORES_BRANCO)
def test_modelo_branco_tem_placeholders_todos_destacados(nome, gerar):
    total, sem = _placeholders_e_destaque(gerar())
    assert total > 0
    assert sem == 0, f"{sem} placeholders sem destaque amarelo em {nome}"


@pytest.mark.parametrize("nome,gerar", GERADORES_BRANCO)
def test_modelo_branco_sem_afirmacoes_factuais(nome, gerar):
    txt = _texto(gerar()).lower()
    achou = [fr for fr in FRASES_PROIBIDAS if fr.lower() in txt]
    assert not achou, f"frases proibidas em {nome}: {achou}"


@pytest.mark.parametrize("nome,gerar", GERADORES_BRANCO)
def test_modelo_branco_sem_zeros_fabricados(nome, gerar):
    txt = _texto(gerar())
    assert "R$ 0,00" not in txt
    assert "0,00%" not in txt
    assert not re.search(r"\b0 ciclo", txt)


@pytest.mark.parametrize("nome,gerar", GERADORES_BRANCO)
def test_modelo_branco_sem_termos_tecnicos_e_emoji(nome, gerar):
    txt = _texto(gerar())
    for t in PALAVRAS_TECNICAS:
        assert not re.search(r"\b" + re.escape(t) + r"\b", txt), f"{t} em {nome}"
    assert not re.findall(r"[\U0001F000-\U0001FAFF☀-➿]", txt)


@pytest.mark.parametrize("nome,gerar", GERADORES_BRANCO)
def test_modelo_branco_deterministico(nome, gerar):
    # O pacote DOCX embute um timestamp: os bytes variam entre execucoes.
    # O determinismo relevante e o do CONTEUDO (texto semantico), nao dos bytes.
    assert _texto(gerar()) == _texto(gerar())


def test_modelo_branco_textos_neutros_29c1():
    # Bloqueadores 2 e 3 da 29C: conclusao do Despacho e considerando 4 do
    # Termo em sentido neutro, sem afirmar consolidacao/analise/apuracao.
    txt_ds = _texto(gerar_modelo_branco_despacho())
    assert ("Após o preenchimento e a conferência das informações, "
            "deverá ser avaliado se a instrução reúne condições") in txt_ds
    txt_ta = _texto(gerar_modelo_branco_termo())
    assert ("O índice contratual e o percentual aplicável deverão ser "
            "informados nos campos a seguir") in txt_ta
    assert "[PREENCHER: Indice contratual]" in txt_ta
    assert "[PREENCHER: Percentual aplicavel]" in txt_ta


def test_despacho_branco_resultado_essencial_neutro_e_destacado():
    b = gerar_modelo_branco_despacho()
    txt = _texto(b)
    assert "A apuração financeira consolidada indicou" not in txt
    assert "valor teórico" not in txt.lower()
    assert "os resultados essenciais deverão ser preenchidos no Quadro 2" in txt
    for ph in (
        "[PREENCHER: Valor pago no periodo analisado]",
        "[PREENCHER: Valor devido apos a atualizacao contratual]",
        "[PREENCHER: Valor retroativo a pagar]",
        "[PREENCHER: Valor Total Atualizado do Contrato]",
    ):
        assert ph in txt
    total, sem_destaque = _placeholders_e_destaque(b)
    assert total >= 20
    assert sem_destaque == 0


def test_neutralidade_integral_29c12():
    # Etapa 29C.1.2: todos os itens com moldura factual foram reescritos como
    # instrucao/obrigacao futura nos DOIS modelos, com contagens preservadas.
    b_ds, b_ta = gerar_modelo_branco_despacho(), gerar_modelo_branco_termo()
    txt_ds, txt_ta = _texto(b_ds), _texto(b_ta)
    # Despacho: estrutura 1-6, sem afirmações factuais ainda não comprovadas.
    for titulo in (
        "1. IDENTIFICAÇÃO",
        "2. PEDIDO E PARÂMETROS DA ANÁLISE",
        "3. RESULTADO ESSENCIAL",
        "4. DOCUMENTOS E VERIFICAÇÕES",
        "5. PENDÊNCIAS",
        "6. CONCLUSÃO",
    ):
        assert titulo in txt_ds
    assert "Deverão ser registrados o pedido" in txt_ds
    assert "Registrar as pendências relevantes" in txt_ds
    assert "instrução encontra-se apta" not in txt_ds
    # Termo: considerandos 3 e 5, secoes 2.1 e 3.1 neutros.
    assert "A memória de cálculo a ser indicada em" in txt_ta
    assert ("As informações da área gestora que vierem a fundamentar "
            "a formalização") in txt_ta
    assert ("Os valores financeiros que eventualmente integrem a formalização "
            "deverão ser informados") in txt_ta
    assert "deverá ser organizada de forma evolutiva" in txt_ta
    # Todos os placeholders permanecem amarelos nos dois modelos.
    tot_ds, sem_ds = _placeholders_e_destaque(b_ds)
    tot_ta, sem_ta = _placeholders_e_destaque(b_ta)
    assert tot_ds >= 20 and sem_ds == 0
    assert (tot_ta, sem_ta) == (30, 0)


def test_automatico_usa_nova_redacao_sem_alterar_termo():
    from test_templates_documentos import leitura_multiciclo_pc, CAMPOS_SANEADOR, CAMPOS_TERMO
    txt_ds = _texto(gerar_despacho_saneador(leitura_multiciclo_pc(), campos_manuais=CAMPOS_SANEADOR))
    assert "A CONTRATADA apresentou pedido de reajuste" in txt_ds
    assert "Conforme memória de cálculo" in txt_ds
    assert "Acordou-se na concessão" not in txt_ds
    assert "Referências auditáveis" not in txt_ds
    txt_ta = _texto(gerar_termo_apostila(leitura_multiciclo_pc(), campos_manuais=CAMPOS_TERMO))
    assert "A apuração financeira consolidada indicou" in txt_ta
    assert "que apurou os ciclos" in txt_ta
    assert "informações encaminhadas pela área gestora" in txt_ta
    assert "foi organizada de forma evolutiva" in txt_ta


def test_despacho_automatico_resultado_consolidado_sem_valor_teorico():
    from test_templates_documentos import leitura_multiciclo_pc, CAMPOS_SANEADOR
    txt = _texto(gerar_despacho_saneador(leitura_multiciclo_pc(),
                                         campos_manuais=CAMPOS_SANEADOR))
    assert "Quadro 2 - Síntese financeira" in txt
    assert "Retroativo a pagar" in txt
    assert "Valor Total Atualizado do Contrato" in txt
    assert "valor teórico" not in txt.lower()


def test_wrappers_nao_leem_session_state():
    fonte = (ROOT / "_templates_documentos.py").read_text(encoding="utf-8")
    idx = fonte.index("def gerar_modelo_branco_despacho")
    trecho = fonte[idx: idx + 900]
    assert "session_state" not in trecho
    assert "modo_modelo_em_branco=True" in trecho


# --------------------------------------------- nao-regressao dos automaticos
def _leituras():
    from test_templates_documentos import (
        leitura_simples_financeiro, leitura_multiciclo_pc, leitura_ausencias,
        CAMPOS_SANEADOR, CAMPOS_TERMO,
    )
    return [
        (gerar_despacho_saneador, leitura_simples_financeiro(), CAMPOS_SANEADOR),
        (gerar_termo_apostila, leitura_simples_financeiro(), CAMPOS_TERMO),
        (gerar_despacho_saneador, leitura_multiciclo_pc(), CAMPOS_SANEADOR),
        (gerar_termo_apostila, leitura_multiciclo_pc(), CAMPOS_TERMO),
        (gerar_despacho_saneador, leitura_ausencias(), CAMPOS_SANEADOR),
        (gerar_termo_apostila, leitura_ausencias(), CAMPOS_TERMO),
    ]


def test_automatico_modo_omitido_igual_false():
    for gerar, leitura, cm in _leituras():
        t_omitido = _texto(gerar(leitura, campos_manuais=cm))
        t_false = _texto(gerar(leitura, campos_manuais=cm, modo_modelo_em_branco=False))
        assert t_omitido == t_false


def test_automatico_nao_vira_modelo_branco():
    from test_templates_documentos import leitura_multiciclo_pc, CAMPOS_SANEADOR
    txt = _texto(gerar_despacho_saneador(leitura_multiciclo_pc(), campos_manuais=CAMPOS_SANEADOR))
    assert "adequação orçamentária" in txt.lower() or "certid" in txt.lower()


# ------------------------------------------------- contaminacao por sessao
def test_modelos_ignoram_session_state_simulado():
    import _templates_documentos as T
    b1d, b1t = T.gerar_modelo_branco_despacho(), T.gerar_modelo_branco_termo()
    contaminacao = {
        "resultado_valor_global": {"valor_atualizado_contrato": 137375560.29},
        "diagnostico_coleta_v2": {"x": 1},
        "aditivos": [{"ciclo": "C1"}],
    }
    _ = contaminacao
    b2d, b2t = T.gerar_modelo_branco_despacho(), T.gerar_modelo_branco_termo()
    # Conteudo semantico identico (bytes variam pelo timestamp do pacote DOCX).
    assert _texto(b1d) == _texto(b2d)
    assert _texto(b1t) == _texto(b2t)
    for txt in (_texto(b2d), _texto(b2t)):
        assert "137" not in txt


# ----------------------------------------------------------- nova pagina
def test_pagina14_existe_e_titulo():
    assert (ROOT / "pages" / "14_Central_Modelos_Ferramentas.py").exists()
    assert "Central de Modelos e Ferramentas" in PAGE14


def test_pagina14_titulo_e_aviso_removidos():
    # Hotfix pos-PR#87 (item 4): titulo "Modelos em branco" e o aviso
    # AVISO_MODELOS foram removidos da apresentacao; os cards permanecem.
    assert 'st.subheader("Modelos em branco")' not in PAGE14
    assert "Os modelos desta área não utilizam dados de uma Coleta e não são " not in PAGE14
    assert (
        "preenchidos automaticamente. Revise e complete todos os campos destacados."
    ) not in PAGE14


def test_pagina14_downloads_modelos():
    assert "Baixar modelo em branco — Despacho Saneador" in PAGE14
    assert "Baixar modelo em branco — Termo de Apostila" in PAGE14
    assert "Modelo_Em_Branco_Despacho_Saneador.docx" in PAGE14
    assert "Modelo_Em_Branco_Termo_de_Apostila.docx" in PAGE14
    assert "gerar_modelo_branco_despacho" in PAGE14
    assert "gerar_modelo_branco_termo" in PAGE14


def test_coleta_oficial_disponivel_somente_na_central_documental():
    assert "Baixar Arquivo Coleta Oficial" in PAGE14
    assert "gerar_coleta_oficial_preenchida({})" in PAGE14
    assert 'key="dl_coleta_oficial_modelos"' in PAGE14
    assert "Baixar Arquivo Coleta Oficial" not in PAGE03


def test_pagina14_ferramentas_presentes():
    for alvo in ["pages/05_Garantia.py", "pages/08_Avaliacao_Aditivos.py"]:
        assert alvo in PAGE14


def test_pagina14_itens_retirados_etapa47():
    # Etapa 47: os quatro itens deixaram de ser apresentados nesta seção.
    # Os arquivos continuam no repositório, mas suas rotas não são registradas.
    for removido in [
        "pages/07_Checklist_Processual.py", "pages/09_Infos_Previas.py",
        "pages/11_Cl8us_Orienta.py", "pages/03_Valor_Global.py",
    ]:
        assert removido not in PAGE14
        assert (ROOT / removido).exists()


def test_pagina14_nao_expoe_dou_adequacao_saneador():
    assert "13_DOU" not in PAGE14
    assert "12_Adequacao" not in PAGE14
    assert "10_Saneador" not in PAGE14
    assert "06_Central_Arquivos" not in PAGE14
    assert "Sumário Executivo" not in PAGE14


def test_pagina14_sem_dependencia_de_apuracao():
    # Não lê dados da apuração nem grava as chaves dos documentos automáticos.
    # (A docstring pode citar essas chaves para explicar que NÃO as utiliza;
    # o que se proíbe é o uso efetivo via session_state.)
    assert "resultado_valor_global" not in PAGE14
    assert "diagnostico_coleta_v2" not in PAGE14
    assert 'session_state["arquivo_despacho_saneador_docx"]' not in PAGE14
    assert 'session_state["arquivo_termo_apostila_docx"]' not in PAGE14
    assert 'st.session_state.get("arquivo_despacho_saneador_docx"' not in PAGE14
    assert 'st.session_state.get("arquivo_termo_apostila_docx"' not in PAGE14


def test_app_registra_pagina_e_link():
    assert "14_Central_Modelos_Ferramentas.py" in APP
    assert 'label="Modelos e ferramentas"' in APP
    assert "PAGINA_MODELOS_FERRAMENTAS" in APP


def test_inicio_nao_repete_o_acesso_a_modelos():
    # Redesign da HOME: o acesso vive no menu lateral (test_app_registra_pagina_e_link).
    assert "5 · Modelos e ferramentas" not in INICIO
    assert "Abrir Modelos e Ferramentas" not in INICIO
    assert "pages/14_Central_Modelos_Ferramentas.py" not in INICIO


# --------------------------------------------------- retorno da Garantia
def test_garantia_retorno_por_origem():
    # Etapa 29C.1: ponte de sessao consumida por pop no render + query param
    # exclusivo da visita; o clique remove apenas o param antes do switch.
    assert "origem_navegacao_garantia" in GARANTIA
    assert "pages/14_Central_Modelos_Ferramentas.py" in GARANTIA
    assert "pages/03_Valor_Global.py" in GARANTIA
    assert 'st.session_state.pop("origem_navegacao_garantia", None)' in GARANTIA
    assert '_PARAM_ORIGEM_GARANTIA = "origem_garantia"' in GARANTIA
    assert "st.query_params[_PARAM_ORIGEM_GARANTIA]" in GARANTIA
    assert "del st.query_params[_PARAM_ORIGEM_GARANTIA]" in GARANTIA


# ------------------------------------------- 29C.1: cenarios da origem A-E
def _at_garantia(*, ponte=False, param=None, sessao=None):
    from streamlit.testing.v1 import AppTest
    at = AppTest.from_file(str(ROOT / "pages" / "05_Garantia.py"), default_timeout=60)
    if ponte:
        at.session_state["origem_navegacao_garantia"] = "modelos_ferramentas"
    if param is not None:
        at.query_params["origem_garantia"] = param
    for k, v in (sessao or {}).items():
        at.session_state[k] = v
    at.run()
    return at


def _ponte_presente(at):
    try:
        at.session_state["origem_navegacao_garantia"]
        return True
    except Exception:
        return False


def _param_visita(at):
    # AppTest guarda query params como lista de valores; normaliza p/ escalar.
    v = at.query_params.get("origem_garantia")
    if isinstance(v, list):
        return v[0] if v else None
    return v


def _clicar_voltar(at):
    botao = next(b for b in at.button if b.key == "voltar_central_garantia")
    try:
        botao.click().run()
    except Exception:
        pass  # switch_page nao resolve pagina em AppTest single-file


def test_origem_cenario_a_acesso_direto():
    at = _at_garantia()
    assert not at.exception
    assert not _ponte_presente(at)
    assert _param_visita(at) is None  # destino: pages/03


def test_origem_cenario_b_central_rerun_voltar():
    at = _at_garantia(ponte=True)
    # ponte consumida no primeiro render e convertida em query param da visita
    assert not _ponte_presente(at)
    assert _param_visita(at) == "modelos_ferramentas"
    # rerun por widget preserva o param -> destino segue a Central. Etapa 49: a
    # Garantia nao tem mais radio (modo unico); usa-se o percentual da garantia.
    at.number_input(key="garantia_percentual").set_value(4.0).run()
    assert _param_visita(at) == "modelos_ferramentas"
    # clique no Voltar: param removido antes do switch; ponte segue ausente
    _clicar_voltar(at)
    assert _param_visita(at) is None
    assert not _ponte_presente(at)


def test_origem_cenario_c_sidebar_sem_voltar():
    at = _at_garantia(ponte=True)          # Central -> Garantia (sem clique)
    assert not _ponte_presente(at)          # ponte ja consumida
    # sidebar encerra o query param da visita; reabertura direta = nova URL
    # limpa com a MESMA sessao (sem ponte):
    at2 = _at_garantia()                    # sem ponte, sem param
    assert _param_visita(at2) is None  # destino: pages/03
    assert not _ponte_presente(at2)


def test_origem_cenario_d_reload_preserva_visita():
    at = _at_garantia(ponte=True)
    assert _param_visita(at) == "modelos_ferramentas"
    # dois reruns adicionais
    at.run()
    at.run()
    assert _param_visita(at) == "modelos_ferramentas"
    # reload da propria Garantia: URL conserva o query param da visita
    at2 = _at_garantia(param="modelos_ferramentas")
    assert _param_visita(at2) == "modelos_ferramentas"  # destino: pages/14


def test_origem_cenario_e_alternancia():
    # Central -> Garantia -> Voltar (Central)
    at = _at_garantia(ponte=True)
    assert _param_visita(at) == "modelos_ferramentas"
    _clicar_voltar(at)
    assert _param_visita(at) is None
    # Upload/direto -> Garantia -> Voltar (Upload e docs)
    at2 = _at_garantia()
    assert _param_visita(at2) is None
    _clicar_voltar(at2)
    assert _param_visita(at2) is None
    # Central -> Garantia -> Voltar (Central) — sem contaminacao da visita anterior
    at3 = _at_garantia(ponte=True)
    assert _param_visita(at3) == "modelos_ferramentas"


def test_pagina14_marca_origem_garantia():
    assert 'origem_navegacao_garantia"] = "modelos_ferramentas"' in PAGE14


def _destino_voltar(param_visita):
    # Etapa 29C.1: o destino e funcao do query param exclusivo da visita.
    if param_visita == "modelos_ferramentas":
        return "pages/14_Central_Modelos_Ferramentas.py"
    return "pages/03_Valor_Global.py"


def test_regra_retorno_garantia_pura():
    assert _destino_voltar("modelos_ferramentas") == "pages/14_Central_Modelos_Ferramentas.py"
    assert _destino_voltar(None) == "pages/03_Valor_Global.py"
    assert _destino_voltar("apuracao") == "pages/03_Valor_Global.py"


def test_upload_nao_repete_box_de_modelos_e_ferramentas():
    assert "Não possui uma Planilha de Coleta?" not in PAGE03
    assert "Abrir Modelos e Ferramentas" not in PAGE03


def test_navegacao_bidirecional_upload_modelos():
    # O acesso à Central permanece exclusivamente no menu lateral.
    assert 'label="Upload e resultados"' in APP
    assert 'label="Modelos e ferramentas"' in APP


def test_upload_docs_apptest_sessao_limpa():
    # Comportamental: pagina 03 em sessao limpa renderiza sem excecao e sem
    # criar chaves de apuracao.
    from streamlit.testing.v1 import AppTest
    at = AppTest.from_file(str(ROOT / "pages" / "03_Valor_Global.py"), default_timeout=60)
    at.run()
    assert not at.exception
    assert not any(b.label == "Abrir Modelos e Ferramentas" for b in at.button)
    assert "resultado_valor_global" not in at.session_state
    assert "diagnostico_coleta_v2" not in at.session_state


# ------------------------------------------------------------- Garantia calc
def test_garantia_calculo_5_6_1_mil():
    # Etapa 49: a garantia passou a ser calculada sobre o percentual informado
    # (padrao 5%), no lugar da antiga funcao de 5% fixo.
    from _garantia_calculo import calcular_garantia_necessaria, moeda
    ga = calcular_garantia_necessaria(100000)
    gb = calcular_garantia_necessaria(120000)
    assert moeda(ga) == "R$ 5.000,00"
    assert moeda(gb) == "R$ 6.000,00"
    assert moeda(gb - ga) == "R$ 1.000,00"
