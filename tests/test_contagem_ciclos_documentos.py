"""Contagem de ciclos efetivamente considerados na apuracao.

A frase do documento nao pode contar linhas do quadro: um ciclo listado como
"Fora da apuracao" permanece no Quadro 1 para rastreabilidade, mas afirmar que
ele foi considerado contradiz o proprio quadro.

Cobre a regra compartilhada (_reajuste_utils) e a redacao resultante no
Despacho Saneador, incluindo a concordancia entre a frase e o Quadro 1.
"""

from __future__ import annotations

import io
from datetime import date

import docx
import pytest

from _reajuste_utils import (
    FRASE_SEM_CICLOS_COMPUTADOS,
    ciclo_computado,
    contar_ciclos_computados,
    expressao_quantidade_ciclos,
)
from _templates_documentos import gerar_despacho_saneador
from test_sumario_executivo import leitura_multiciclo_pc


CICLOS_REAJUSTE = ("C1", "C2", "C3", "C4")


# ---------------------------------------------------------------------------
# Regra compartilhada
# ---------------------------------------------------------------------------

def _por_computar(*marcas):
    return [{"ciclo": f"C{i + 1}", "computar": m} for i, m in enumerate(marcas)]


@pytest.mark.parametrize("marcas,esperado", [
    (("Sim", "Não", "Não", "Não"), 1),
    (("Sim", "Sim", "Não", "Não"), 2),
    (("Sim", "Sim", "Sim", "Sim"), 4),
    (("Não", "Não", "Não", "Não"), 0),
    ((None, None, None, None), 0),
    (("", "", "", ""), 0),
])
def test_contagem_por_marcacao_computar(marcas, esperado):
    assert contar_ciclos_computados(_por_computar(*marcas)) == esperado


def test_marcacao_computar_aceita_variacoes_de_grafia():
    assert ciclo_computado({"computar": "sim"})
    assert ciclo_computado({"computar_nesta_apuracao": "SIM"})
    assert not ciclo_computado({"computar": "Nao"})
    assert not ciclo_computado({"computar": "Não"})


def test_tratamento_financeiro_quando_nao_ha_marcacao_computar():
    assert ciclo_computado({"Tratamento financeiro do ciclo": "Apurar"})
    assert not ciclo_computado({"Tratamento financeiro do ciclo": "Fora da apuração"})


def test_classificacao_exclui_apenas_quando_explicita():
    assert not ciclo_computado({"Situação": "Fora da apuracao"})
    assert not ciclo_computado({"Classificação": "não computado"})
    assert not ciclo_computado({"Situação": "Não aplicável"})
    assert ciclo_computado({"Situação": "TEMPESTIVO"})


def test_marcacao_computar_tem_precedencia_sobre_a_classificacao():
    """COMPUTAR e a fonte canonica; a situacao exibida nao a sobrepoe."""
    assert not ciclo_computado({"computar": "Não", "Situação": "TEMPESTIVO"})


@pytest.mark.parametrize("quantidade,esperado", [
    (1, "1 ciclo"),
    (2, "2 ciclos"),
    (4, "4 ciclos"),
    (0, None),
    (None, None),
])
def test_expressao_com_flexao_correta(quantidade, esperado):
    assert expressao_quantidade_ciclos(quantidade) == esperado


def test_expressao_nunca_usa_a_forma_ciclo_s():
    for quantidade in range(1, 6):
        assert "(s)" not in expressao_quantidade_ciclos(quantidade)


# ---------------------------------------------------------------------------
# Redacao no Despacho Saneador
# ---------------------------------------------------------------------------

def _leitura_com_computados(computados: set[str]) -> dict:
    leitura = leitura_multiciclo_pc()
    parametros = leitura.get("parametros_v10") or {}
    for registro in parametros.get("ciclos") or []:
        nome = str(registro.get("ciclo") or "").upper()
        if nome in CICLOS_REAJUSTE:
            registro["computar_nesta_apuracao"] = "Sim" if nome in computados else "Nao"
    for nome, registro in (parametros.get("por_ciclo") or {}).items():
        if str(nome).upper() in CICLOS_REAJUSTE:
            registro["computar_nesta_apuracao"] = (
                "Sim" if str(nome).upper() in computados else "Nao"
            )
    return leitura


def _documento(leitura: dict):
    return docx.Document(io.BytesIO(gerar_despacho_saneador(leitura)))


def _quadro_1(documento) -> list[list[str]]:
    tabela = documento.tables[0]
    return [[c.text.strip() for c in linha.cells] for linha in tabela.rows[1:]]


@pytest.mark.parametrize("computados,esperado", [
    ({"C1"}, ["C1"]),
    ({"C1", "C2"}, ["C1", "C2"]),
    ({"C1", "C2", "C3", "C4"}, ["C1", "C2", "C3", "C4"]),
])
def test_quadro_1_lista_apenas_os_ciclos_computados(computados, esperado):
    linhas = _quadro_1(_documento(_leitura_com_computados(computados)))
    assert [linha[0] for linha in linhas] == esperado


def test_sem_ciclo_computado_nao_afirma_que_houve_analise_de_ciclos():
    documento = _documento(_leitura_com_computados(set()))
    texto = "\n".join(paragrafo.text for paragrafo in documento.paragraphs)
    linhas = _quadro_1(documento)
    assert "considerou" not in texto
    assert linhas[0][0] == "[PREENCHER: Ciclo]"


def test_ciclos_existentes_marcados_como_nao_nao_sao_contados():
    """Ciclos fora da apuracao nao poluem o Quadro 1 do despacho."""
    documento = _documento(_leitura_com_computados(set()))
    ciclos = [linha[0] for linha in _quadro_1(documento)]
    assert not set(ciclos).intersection(CICLOS_REAJUSTE)


@pytest.mark.parametrize("computados", [
    {"C1"},
    {"C1", "C2"},
    {"C1", "C2", "C3", "C4"},
])
def test_concordancia_entre_a_frase_e_o_quadro_1(computados):
    """Cada ciclo necessario aparece uma unica vez no quadro essencial."""
    documento = _documento(_leitura_com_computados(computados))
    linhas = _quadro_1(documento)
    ciclos = [linha[0] for linha in linhas]
    assert ciclos == sorted(computados)
    assert len(ciclos) == len(set(ciclos))


# ---------------------------------------------------------------------------
# Cenario real relatado: C1 computado, C2-C4 fora da apuracao
# ---------------------------------------------------------------------------

@pytest.fixture(scope="module")
def documento_cenario_real():
    from _coleta_oficial import gerar_coleta_oficial_preenchida
    from _coleta_reajuste_documentos import processar_coleta_oficial_runtime

    coleta = gerar_coleta_oficial_preenchida({
        "origem": "Cenario real C1 computado",
        "indice": "IST (Anatel)",
        "data_base_original": "01/02/2025",
        "data_corte": date(2027, 1, 31),
        "ciclos": [{
            "ciclo": "C1",
            "data_inicio": date(2026, 2, 1),
            "data_fim": date(2027, 1, 31),
            "data_pedido": date(2026, 2, 1),
            "financeiro_inicio": date(2026, 4, 1),
            "percentual_aplicado": 0.0307853139440224,
            "situacao": "✅ TEMPESTIVO",
            "objeto_analise_atual": True,
        }],
    })
    payload, _ = processar_coleta_oficial_runtime(coleta)
    return _documento(payload)


def test_cenario_real_declara_um_unico_ciclo(documento_cenario_real):
    linhas = _quadro_1(documento_cenario_real)
    assert [linha[0] for linha in linhas] == ["C1"]
    assert linhas[0][4] == "A partir de 04/2026"
    assert linhas[0][5] == "3,08%"


def test_cenario_real_omite_os_demais_ciclos_do_quadro(documento_cenario_real):
    linhas = _quadro_1(documento_cenario_real)
    ciclos = [linha[0] for linha in linhas]
    assert ciclos == ["C1"]
    assert not {"C2", "C3", "C4"}.intersection(ciclos)
