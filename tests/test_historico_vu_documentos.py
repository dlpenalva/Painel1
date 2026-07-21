"""Etapa 7: historico de Valores Unitarios por ciclo no Saneador e na Apostila.

C0 aparece sempre; os demais ciclos vao ate o ultimo efetivamente analisado
(computar == "Sim"). Ciclos futuros presentes na planilha nao entram. Nao ha
zeros artificiais. A estrutura canonica e unica (uma linha por item) e alimenta
os dois documentos pela mesma funcao de render.
"""
import sys
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from _reconciliacao_xls_python import campos_nao_confiaveis_para_documentos
from _sumario_executivo import _montar_secao_historico_vu, _ultimo_ciclo_analisado
from _templates_documentos import (
    _secao_valores_unitarios_por_ciclo,
    gerar_despacho_saneador,
    gerar_termo_apostila,
)
from test_sumario_executivo import leitura_multiciclo_pc


def _ciclos(*analisados):
    """ciclos com C0 base e C1..C4 marcados (Sim) conforme informado."""
    saida = [{"ciclo": "C0", "computar": "Não", "eh_base": True}]
    for i in range(1, 5):
        nome = f"C{i}"
        saida.append({"ciclo": nome, "computar": "Sim" if nome in analisados else "Não"})
    return saida


def _historico(*n_itens):
    itens = []
    for k in range(n_itens[0] if n_itens else 2):
        itens.append({
            "item": f"ITEM_{k+1}",
            "descricao": f"Descricao {k+1}",
            "vu_original": 10.0 + k,
            "vu_ciclos": {
                "VU_C0": 10.0 + k, "VU_C1": 11.0 + k, "VU_C2": 12.0 + k,
                "VU_C3": 13.0 + k, "VU_C4": 14.0 + k,
            },
        })
    return {"itens": itens}


class TestUltimoCicloETruncamento(unittest.TestCase):
    def test_a_somente_c1(self):
        sec = _montar_secao_historico_vu(_historico(), _ciclos("C1"))
        self.assertEqual(sec["ciclos"], ["C0", "C1"])
        self.assertEqual(sec["ultimo_ciclo"], "C1")

    def test_b_ate_c2(self):
        sec = _montar_secao_historico_vu(_historico(), _ciclos("C1", "C2"))
        self.assertEqual(sec["ciclos"], ["C0", "C1", "C2"])

    def test_c_ate_c4(self):
        sec = _montar_secao_historico_vu(_historico(), _ciclos("C1", "C2", "C3", "C4"))
        self.assertEqual(sec["ciclos"], ["C0", "C1", "C2", "C3", "C4"])

    def test_d_ciclos_posteriores_nao_aparecem(self):
        sec = _montar_secao_historico_vu(_historico(), _ciclos("C1", "C2"))
        chaves_vu = set(sec["itens"][0]["vus"].keys())
        self.assertNotIn("C3", chaves_vu)
        self.assertNotIn("C4", chaves_vu)
        self.assertEqual(chaves_vu, {"C0", "C1", "C2"})

    def test_so_c0_quando_nada_analisado(self):
        sec = _montar_secao_historico_vu(_historico(), _ciclos())
        self.assertEqual(sec["ciclos"], ["C0"])
        self.assertEqual(_ultimo_ciclo_analisado(_ciclos()), 0)


class TestConteudoItens(unittest.TestCase):
    def test_g_uma_linha_por_item(self):
        sec = _montar_secao_historico_vu(_historico(3), _ciclos("C1"))
        self.assertEqual(len(sec["itens"]), 3)
        self.assertEqual([i["item"] for i in sec["itens"]], ["ITEM_1", "ITEM_2", "ITEM_3"])

    def test_nao_inventa_zeros(self):
        hist = {"itens": [{
            "item": "X", "descricao": "d", "vu_original": 10.0,
            "vu_ciclos": {"VU_C0": 10.0, "VU_C1": None, "VU_C2": None},
        }]}
        sec = _montar_secao_historico_vu(hist, _ciclos("C1"))
        self.assertEqual(sec["itens"][0]["vus"]["C0"], 10.0)
        self.assertIsNone(sec["itens"][0]["vus"]["C1"])  # ausente != zero


class TestRenderDocx(unittest.TestCase):
    def _render(self, sec):
        doc = Document()
        _secao_valores_unitarios_por_ciclo(doc, {"historico_vu": sec})
        return doc

    def test_e_mesma_estrutura_para_os_dois_documentos(self):
        # A funcao e a mesma usada por Saneador e Apostila: render deterministico.
        sec = _montar_secao_historico_vu(_historico(2), _ciclos("C1", "C2"))
        t1 = self._render(sec).tables[0]
        t2 = self._render(sec).tables[0]
        self.assertEqual(len(t1.columns), len(t2.columns))
        self.assertEqual(len(t1.rows), len(t2.rows))
        # cabecalho: Item, Descricao, VU C0, VU C1, VU C2
        cab = [c.text for c in t1.rows[0].cells]
        self.assertEqual(cab, ["Item", "Descrição", "VU C0", "VU C1", "VU C2"])

    def test_f_formatacao_monetaria_reais_duas_casas(self):
        sec = _montar_secao_historico_vu(_historico(1), _ciclos("C1"))
        tabela = self._render(sec).tables[0]
        valor_c0 = tabela.rows[1].cells[2].text  # primeira linha de dados, coluna VU C0
        self.assertTrue(valor_c0.startswith("R$"))
        self.assertRegex(valor_c0, r",\d{2}$")

    def test_h_sem_itens_nao_renderiza_tabela(self):
        doc = Document()
        _secao_valores_unitarios_por_ciclo(doc, {"historico_vu": {"itens": [], "ciclos": []}})
        self.assertEqual(len(doc.tables), 0)


class TestInteracaoEtapa5(unittest.TestCase):
    def test_i_divergencia_independente_nao_remove_historico_vu(self):
        # Divergencia em remanescente/retroativo NAO contamina VU (dimensao propria).
        nc = campos_nao_confiaveis_para_documentos(
            {"divergencias_relevantes": [
                {"campo": "REM_ATUALIZADO_OFICIAL"}, {"campo": "RETRO_OFICIAL"},
            ]}
        )
        # nenhum campo VU no conjunto nao-confiavel
        self.assertFalse(any("VU" in c for c in nc))
        # a secao historico_vu independe de nc e permanece completa
        sec = _montar_secao_historico_vu(_historico(2), _ciclos("C1", "C2"))
        self.assertTrue(sec["disponivel"])
        self.assertEqual(len(sec["itens"]), 2)


class TestDocumentoFinalEndToEnd(unittest.TestCase):
    """Prova de que a tabela de VU chega ao DOCX final gerado pelos builders
    reais (nao apenas pela funcao de secao). Cenario: analise ate C2."""

    def _leitura_c2(self):
        leit = leitura_multiciclo_pc()  # C1+C2 computar=Sim; C3/C4 nao
        leit["historico_vu"] = {"itens": [
            {"item": "ITEM_X", "descricao": "Servico X", "vu_original": 100.00,
             "vu_ciclos": {"VU_C0": 100.00, "VU_C1": 103.10, "VU_C2": 106.00,
                           "VU_C3": 999.99, "VU_C4": 888.88}},
            {"item": "ITEM_Y", "descricao": "Servico Y", "vu_original": 50.00,
             "vu_ciclos": {"VU_C0": 50.00, "VU_C1": 51.55, "VU_C2": 53.00,
                           "VU_C3": 777.77, "VU_C4": 666.66}},
        ]}
        return leit

    def _tabela_vu(self, docx_bytes):
        from io import BytesIO
        doc = Document(BytesIO(docx_bytes))
        for t in doc.tables:
            cab = [c.text for c in t.rows[0].cells]
            if cab[:1] == ["Item"] and "VU C0" in cab:
                linhas = [[c.text for c in r.cells] for r in t.rows[1:]]
                return cab, linhas
        return None, None

    def _assert_documento(self, docx_bytes):
        cab, linhas = self._tabela_vu(docx_bytes)
        self.assertIsNotNone(cab, "tabela de VU ausente no DOCX final")
        # C0..C2 presentes; C3/C4 ausentes (analise termina em C2)
        self.assertEqual(cab, ["Item", "Descrição", "VU C0", "VU C1", "VU C2"])
        self.assertNotIn("VU C3", cab)
        self.assertNotIn("VU C4", cab)
        # valores vindos de historico_vu, formatados R$ com duas casas
        self.assertEqual(linhas[0][0], "ITEM_X")
        self.assertEqual(linhas[0][2:], ["R$ 100,00", "R$ 103,10", "R$ 106,00"])
        self.assertEqual(linhas[1][0], "ITEM_Y")
        self.assertEqual(linhas[1][2:], ["R$ 50,00", "R$ 51,55", "R$ 53,00"])
        # valores futuros nao vazam para o documento
        texto = "\n".join("|".join(l) for l in linhas)
        self.assertNotIn("999,99", texto)
        self.assertNotIn("888,88", texto)

    def test_saneador_final_contem_vu(self):
        self._assert_documento(gerar_despacho_saneador(self._leitura_c2()))

    def test_apostila_final_contem_vu(self):
        self._assert_documento(gerar_termo_apostila(self._leitura_c2()))


if __name__ == "__main__":
    unittest.main()
