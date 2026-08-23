"""R2A — VTA documental alinhado ao VTA canonico do resultado_consolidado.

Protege exatamente o contrato desta correcao:

  - quando `resultado_consolidado` existe, o VTA dos 3 documentos (Sumario,
    Despacho, Termo) e IDENTICO ao VTA do painel (resultado_consolidado.vta),
    nunca a "PREVIA" reconstruida do XLS;
  - o fallback de origem do VTA (posicao atual / ultima posicao disponivel /
    indisponivel) e uma decisao exclusiva do resultado_consolidado — os
    documentos apenas repetem o numero, nunca recalculam;
  - ausencia (None) permanece ausencia — nunca vira R$ 0,00 nem cai para XLS;
  - retroativo reconhecido e valor pago efetivo continuam grandezas distintas
    do VTA, intocadas por esta correcao (R2D/R2E: nao confundir retroativo
    com o delta XLS x Python do VTA).
"""
from __future__ import annotations

import sys
import zipfile
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from _coleta_reajuste_documentos import processar_coleta_oficial_runtime  # noqa: E402
from _sumario_executivo import gerar_sumario_executivo, montar_dados_sumario_executivo  # noqa: E402
from _templates_documentos import gerar_despacho_saneador, gerar_termo_apostila  # noqa: E402
from test_sumario_executivo import leitura_multiciclo_pc  # noqa: E402
from test_templates_documentos import CAMPOS_SANEADOR, CAMPOS_TERMO  # noqa: E402

CASO_PC = r"C:\_DesktopReal\Coleta_Reajuste_C1_IST_20ago.xlsx"
CASO_FINANCEIRO = r"C:\Users\danie\OneDrive\Desktop\Coleta_Reajuste_C3_14ago-TARDE_2.xlsx"
_ARQUIVOS_REAIS_AUSENTES = not (Path(CASO_PC).exists() and Path(CASO_FINANCEIRO).exists())


def _texto_docx(conteudo: bytes) -> str:
    doc = Document(BytesIO(conteudo))
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            partes += [c.text for c in r.cells]
    for s in doc.sections:
        partes += [p.text for p in s.footer.paragraphs]
    return "\n".join(partes)


def _docx_valido(conteudo: bytes) -> bool:
    try:
        with zipfile.ZipFile(BytesIO(conteudo)) as z:
            return (
                conteudo[:2] == b"PK"
                and "word/document.xml" in z.namelist()
                and z.testzip() is None
            )
    except Exception:
        return False


def _pdf_valido(conteudo: bytes) -> bool:
    return conteudo.startswith(b"%PDF-") and conteudo.rstrip().endswith(b"%%EOF")


def _texto_pdf(conteudo: bytes) -> str:
    fitz = pytest.importorskip("fitz")
    doc = fitz.open(stream=conteudo, filetype="pdf")
    return "\n".join(p.get_text() for p in doc)


def _formatar_moeda_simples(valor: float) -> str:
    from _sumario_executivo import formatar_moeda
    return formatar_moeda(valor)


@pytest.mark.skipif(_ARQUIVOS_REAIS_AUSENTES, reason="arquivos reais dos casos PC/Financeiro nao disponiveis neste ambiente")
class TestCasosReaisParidadePainelDocumentos:
    """Item 11 da tarefa: casos reais, painel = Sumario = Despacho = Termo."""

    def _rodar(self, caminho, campos_saneador, campos_termo):
        conteudo = Path(caminho).read_bytes()
        resultado, _diagnostico = processar_coleta_oficial_runtime(conteudo)
        consolidado = resultado["resultado_consolidado"]
        dados = montar_dados_sumario_executivo(resultado)
        despacho = _texto_docx(gerar_despacho_saneador(resultado, campos_manuais=campos_saneador))
        termo = _texto_docx(gerar_termo_apostila(resultado, campos_manuais=campos_termo))
        sumario = _texto_pdf(gerar_sumario_executivo(resultado))
        return consolidado, dados, despacho, termo, sumario

    def test_caso_pc_real_vta_identico_em_todas_as_camadas(self):
        consolidado, dados, despacho, termo, sumario = self._rodar(
            CASO_PC, CAMPOS_SANEADOR, CAMPOS_TERMO
        )
        vta_painel = consolidado["vta"]
        assert vta_painel is not None
        assert dados["sintese"]["vta"] == vta_painel
        assert dados["sintese"].get("vta_previa") is None
        vta_fmt = _formatar_moeda_simples(vta_painel)
        for nome, texto in (("Despacho", despacho), ("Termo", termo), ("Sumario", sumario)):
            assert vta_fmt in texto, f"{nome} nao exibe o VTA canonico {vta_fmt}"
            assert "PRÉVIA" not in texto, f"{nome} ainda exibe rotulo PREVIA para o VTA"
        # Retroativo reconhecido permanece intocado por esta correcao.
        assert consolidado["retroativo_reconhecido"] == dados["sintese"]["retroativo_total"]

    def test_caso_financeiro_real_vta_identico_em_todas_as_camadas(self):
        consolidado, dados, despacho, termo, sumario = self._rodar(
            CASO_FINANCEIRO, CAMPOS_SANEADOR, CAMPOS_TERMO
        )
        vta_painel = consolidado["vta"]
        assert vta_painel is not None
        assert dados["sintese"]["vta"] == vta_painel
        assert dados["sintese"].get("vta_previa") is None
        vta_fmt = _formatar_moeda_simples(vta_painel)
        for nome, texto in (("Despacho", despacho), ("Termo", termo), ("Sumario", sumario)):
            assert vta_fmt in texto, f"{nome} nao exibe o VTA canonico {vta_fmt}"
            assert "PRÉVIA" not in texto, f"{nome} ainda exibe rotulo PREVIA para o VTA"
        # R2D/R2E: retroativo reconhecido (24.678,92) e o pago efetivo
        # (7.300.890,27) sao grandezas distintas do delta XLS x Python do VTA
        # (388.960,65) — nao confundir. Confirma que o retroativo permanece
        # separado e nao foi alterado por esta correcao.
        assert round(consolidado["retroativo_reconhecido"], 2) == 24678.92
        assert round(dados["sintese"]["retroativo_total"], 2) == 24678.92
        assert round(consolidado["retroativo_reconhecido"], 2) != round(vta_painel, 2)

    def test_documentos_reais_validos_e_layout_integro(self):
        from _estado_apuracao_upload import (
            assinatura_conteudo_upload,
            id_apuracao_de_assinatura,
        )

        conteudo = Path(CASO_FINANCEIRO).read_bytes()
        resultado, _diagnostico = processar_coleta_oficial_runtime(conteudo)
        # id_apuracao e atribuido pela pagina (pages/03_Valor_Global.py), nao
        # por processar_coleta_oficial_runtime — reproduz aqui o mesmo passo
        # para poder verificar APUR/R1A no rodape (R1A).
        assinatura = assinatura_conteudo_upload(conteudo)
        resultado["id_apuracao"] = id_apuracao_de_assinatura(assinatura)
        d_desp = gerar_despacho_saneador(resultado, campos_manuais=CAMPOS_SANEADOR)
        d_termo = gerar_termo_apostila(resultado, campos_manuais=CAMPOS_TERMO)
        d_sum = gerar_sumario_executivo(resultado)
        assert _docx_valido(d_desp)
        assert _docx_valido(d_termo)
        assert _pdf_valido(d_sum)
        # APUR e R1A preservados (rodape com ID da apuracao + Gerado em).
        for texto in (_texto_docx(d_desp), _texto_docx(d_termo)):
            assert "ID da apuração: APUR-" in texto
            assert "Gerado em" in texto


class TestFallbackUltimaPosicaoDisponivel:
    """Item 8 da tarefa: o documento NAO reconstroi a escolha de origem do VTA."""

    def test_vta_por_ultima_posicao_propaga_identico_aos_documentos(self):
        leitura = leitura_multiciclo_pc()
        vta_ultima_posicao = 999888.77
        leitura["resultado_consolidado"] = {
            "vta": vta_ultima_posicao,
            "vta_origem": "ultima_posicao_disponivel",
            "vta_usa_ultima_posicao": True,
        }
        dados = montar_dados_sumario_executivo(leitura)
        assert dados["sintese"]["vta"] == vta_ultima_posicao
        assert dados["sintese"].get("vta_previa") is None
        texto = _texto_docx(
            gerar_despacho_saneador(leitura, campos_manuais=CAMPOS_SANEADOR)
        )
        assert _formatar_moeda_simples(vta_ultima_posicao) in texto
        assert "PRÉVIA" not in texto


class TestAusenciaNaoViraZeroNemXls:
    """Item 9 da tarefa."""

    def test_vta_none_no_consolidado_permanece_none_nos_documentos(self):
        leitura = leitura_multiciclo_pc()
        leitura["resultado_consolidado"] = {
            "vta": None,
            "vta_origem": "indisponivel",
            "vta_usa_ultima_posicao": False,
        }
        dados = montar_dados_sumario_executivo(leitura)
        assert dados["sintese"]["vta"] is None
        # Nao deve ter caido para o XLS (vta_previa).
        assert dados["sintese"].get("vta_previa") is None


class TestComportamentoLegadoPreservado:
    """Chamadas sem resultado_consolidado continuam com a politica da PREVIA
    (Etapa 26H) — R2A nao pode quebrar geradores isolados/testes existentes."""

    def test_sem_resultado_consolidado_mantem_comportamento_legado(self):
        leitura = leitura_multiciclo_pc()
        assert "resultado_consolidado" not in leitura
        dados = montar_dados_sumario_executivo(leitura)
        assert "vta" in dados["sintese"]
