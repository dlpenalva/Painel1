"""RESULTADOS-BASELINE (PR 0) — fotografo das grandezas, em quatro camadas.

O QUE ESTE MODULO E
-------------------
Um LEITOR. Ele nao calcula, nao arredonda, nao interpreta e nao decide nada:
so recolhe o que as cadeias de producao ja produziram e organiza num dicionario
serializavel, comparavel ao centavo. Se um valor sair diferente depois da
refatoracao da aba RESULTADOS, foi a producao que mudou — nunca este arquivo.

AS QUATRO CAMADAS (e por que sao separadas)
-------------------------------------------
1. `fotografar_contrato_xls` — ESTRUTURA da aba RESULTADOS: formulas de
   A1:J87, intervalos nomeados e seus destinos, visibilidade, e o conjunto de
   coordenadas que o runtime realmente le. Nao depende de Excel nem de cache.
   E o que prova, no PR 1, que nenhuma formula/ancora mudou sem intencao.

2. `fotografar_valores_xls` — os VALORES calculados pelo Excel, lidos do cache
   de formulas (`data_only=True`). Um XLS montado por openpyxl nao tem cache:
   nesses casos a camada devolve `cache_ausente=True` e valores nulos, que e o
   estado real e fail-closed do produto — nao uma falha do baseline.

3. `fotografar_web` — a cadeia PYTHON completa, a partir de
   `processar_coleta_oficial_runtime` (o mesmo entry point do upload real).
   Congela metodo, status, ciclos, situacao, variacao, percentual, fatores,
   efeito financeiro, retroativo, remanescente, composicao, referencias
   auditaveis e mensagens.

4. `fotografar_documentos` — o conteudo negocial dos documentos que consomem o
   resultado consolidado.

CONVERGENCIA, NAO COPIA
-----------------------
A prova de que XLS e web concordam NAO e "web == celula da RESULTADOS". Ela ja
existe em producao: `resultado["reconciliacao_xls_python"]` compara, campo a
campo e com tolerancia propria, o numero do XLS contra o numero que o motor
Python calculou por conta propria. O baseline fotografa esse bloco inteiro
(`convergencia_xls_python`), preservando as DUAS colunas — `xls` e `python` — e
o status. Se um dia a web passar a copiar o XLS em vez de calcular, as colunas
continuariam iguais mas a origem mudaria; e por isso que se fotografa o bloco
inteiro, e nao apenas a igualdade.

PRECISAO
--------
Nenhum valor e arredondado aqui. Monetarios chegam ja na precisao economica da
grandeza (o centavo, quando e o caso) e percentuais/fatores chegam na precisao
matematica que a producao calculou; ambos sao gravados como vieram.
"""
from __future__ import annotations

import io
import math
import re
from datetime import date, datetime
from decimal import Decimal
from typing import Any

from openpyxl import load_workbook

ABA = "RESULTADOS"
ABA_MEMORIA = "MEMORIA_RESULTADOS"

# Ultima linha/coluna do leiaute atual da aba executiva (A1:J87).
ULTIMA_LINHA = 87
ULTIMA_COLUNA = 10

# Coordenadas da aba RESULTADOS efetivamente lidas pelo RUNTIME de producao.
# Fora desta lista, a aba e apresentacao — e esse e o achado que sustenta a
# refatoracao do PR 1.
COORDENADAS_LIDAS_EM_RUNTIME = {
    "A1": "_coleta_reajuste.py — titulo literal, gate de integridade da aba",
    "B3": "_coleta_reajuste.py — status_resultados['geral'] (STATUS_RESULTADOS)",
    "B10": "_leitor_masterfile_v10.py — referencias_vta.forma1_posicao_atual",
    "B11": "_leitor_masterfile_v10.py — referencias_vta.forma2_ultima_abertura",
    "B12": "_leitor_masterfile_v10.py — referencias_vta.forma3_integral_reajustado",
    "B13": "_leitor_masterfile_v10.py — referencias_vta.reconciliacao_valor",
    "H10": "_leitor_masterfile_v10.py — referencias_vta.forma1_situacao",
    "H11": "_leitor_masterfile_v10.py — referencias_vta.forma2_situacao",
    "H13": "_leitor_masterfile_v10.py — referencias_vta.reconciliacao_status",
}

# Ancoradas na aba RESULTADOS (as demais apontam para MEMORIA_RESULTADOS).
# B36/B38 nao tem consumidor em producao nem em teste: ficam fotografadas
# justamente para que o PR 1 decida conscientemente o que fazer com elas.
COORDENADAS_ANCORADAS_POR_NOME = ("B3", "B12", "B36", "B38")

# O contrato da aba NAO se esgota no que o Python le. As linhas 43-50 (bloco
# "5. AJUSTES MANUAIS") sao ENTRADAS do usuario, e formulas de outras abas do
# proprio workbook as consomem para compor VTA e retroativo. Mover ou renumerar
# essas linhas quebra o calculo dentro do Excel, sem que nenhum teste de Python
# perceba. `mapear_referencias_de_outras_abas` levanta esse acoplamento a
# partir do arquivo, em vez de confiar nesta lista.
_RE_REFERENCIA_A_ABA = re.compile(r"(?<!MEMORIA_)RESULTADOS!\$?([A-Z]{1,2})\$?(\d{1,4})")


# --------------------------------------------------------------------------- #
# Normalizacao — o baseline precisa comparar ao centavo e sobreviver a JSON.
# --------------------------------------------------------------------------- #
def normalizar(valor: Any) -> Any:
    """Converte para uma forma estavel, comparavel e serializavel.

    Datas viram ISO (`YYYY-MM-DD`) para o snapshot nao depender de locale.
    NaN vira None: pandas usa NaN para ausencia, e ausencia nunca e um numero.
    Nada e arredondado — a precisao entregue pela producao e a que se congela.
    """
    if valor is None:
        return None
    if isinstance(valor, bool):
        return valor
    if isinstance(valor, int):
        return valor
    if isinstance(valor, float):
        return None if math.isnan(valor) else valor
    if isinstance(valor, Decimal):
        return float(valor)
    if isinstance(valor, datetime):
        return (valor.date().isoformat() if valor.time() == datetime.min.time()
                else valor.isoformat())
    if isinstance(valor, date):
        return valor.isoformat()
    if isinstance(valor, dict):
        return {str(chave): normalizar(item) for chave, item in valor.items()}
    if isinstance(valor, (list, tuple)):
        return [normalizar(item) for item in valor]
    if isinstance(valor, set):
        return [normalizar(item) for item in sorted(valor, key=str)]
    texto = str(valor).strip()
    return texto or None


def _num(valor: Any) -> float | int | None:
    valor = normalizar(valor)
    if isinstance(valor, bool):
        return None
    return valor if isinstance(valor, (int, float)) else None


# --------------------------------------------------------------------------- #
# CAMADA 1 — contrato estrutural da aba RESULTADOS.
# --------------------------------------------------------------------------- #
def mapear_referencias_de_outras_abas(wb) -> dict[str, list[str]]:
    """Celulas da aba RESULTADOS que formulas de OUTRAS abas consultam.

    Acoplamento XLS -> XLS, invisivel para qualquer teste de Python. Hoje sao
    as entradas dos AJUSTES MANUAIS (linhas 43-50), lidas por
    MEMORIA_RESULTADOS para compor VTA e retroativo.
    """
    mapa: dict[str, set[str]] = {}
    for nome in wb.sheetnames:
        if nome == ABA:
            continue
        for linha in wb[nome].iter_rows():
            for celula in linha:
                valor = celula.value
                if not isinstance(valor, str) or "RESULTADOS!" not in valor:
                    continue
                for achado in _RE_REFERENCIA_A_ABA.finditer(valor):
                    coordenada = f"{achado.group(1)}{achado.group(2)}"
                    mapa.setdefault(coordenada, set()).add(nome)
    return {coordenada: sorted(abas) for coordenada, abas in sorted(mapa.items())}


def fotografar_contrato_xls(conteudo: bytes) -> dict[str, Any]:
    """Estrutura da aba RESULTADOS: formulas, ancoras e visibilidade."""
    wb = load_workbook(io.BytesIO(conteudo), data_only=False)
    try:
        if ABA not in wb.sheetnames:
            return {"aba_presente": False}
        ws = wb[ABA]
        formulas: dict[str, Any] = {}
        for linha in range(1, ULTIMA_LINHA + 1):
            for coluna in range(1, ULTIMA_COLUNA + 1):
                celula = ws.cell(row=linha, column=coluna)
                if celula.value is not None:
                    formulas[celula.coordinate] = normalizar(celula.value)
        nomes = {
            nome: normalizar(definicao.value)
            for nome, definicao in wb.defined_names.items()
        }
        return {
            "aba_presente": True,
            "visibilidade": ws.sheet_state,
            "e_ultima_aba": wb.sheetnames[-1] == ABA,
            "titulo_a1": normalizar(ws["A1"].value),
            "dimensoes": ws.dimensions,
            "quantidade_celulas_preenchidas": len(formulas),
            "formulas": formulas,
            "nomes_definidos": nomes,
            "nomes_ancorados_na_aba": sorted(
                nome for nome, destino in nomes.items()
                if isinstance(destino, str) and destino.startswith(f"{ABA}!")
            ),
            "coordenadas_lidas_em_runtime": dict(COORDENADAS_LIDAS_EM_RUNTIME),
            "coordenadas_lidas_por_outras_abas":
                mapear_referencias_de_outras_abas(wb),
        }
    finally:
        wb.close()


# --------------------------------------------------------------------------- #
# CAMADA 2 — valores calculados pelo Excel (cache de formulas).
# --------------------------------------------------------------------------- #
def fotografar_valores_xls(conteudo: bytes) -> dict[str, Any]:
    """Valores do cache do Excel. Sem recalculo, tudo nulo — e correto."""
    wb = load_workbook(io.BytesIO(conteudo), data_only=True)
    try:
        if ABA not in wb.sheetnames:
            return {"aba_presente": False, "cache_ausente": True}
        ws = wb[ABA]
        coordenadas = sorted(
            set(COORDENADAS_LIDAS_EM_RUNTIME) | set(COORDENADAS_ANCORADAS_POR_NOME)
        )
        valores = {coord: normalizar(ws[coord].value) for coord in coordenadas}
        nomes: dict[str, Any] = {}
        for nome, definicao in wb.defined_names.items():
            try:
                aba, referencia = list(definicao.destinations)[0]
                nomes[nome] = normalizar(wb[aba][referencia.replace("$", "")].value)
            except Exception:
                nomes[nome] = None
        # O cache e avaliado por B3 (STATUS_RESULTADOS): e a unica celula da
        # aba que o produto exige calculada para considerar o arquivo conferido.
        return {
            "aba_presente": True,
            "cache_ausente": valores.get("B3") in (None, ""),
            "coordenadas": valores,
            "nomes_definidos": nomes,
        }
    finally:
        wb.close()


# --------------------------------------------------------------------------- #
# CAMADA 3 — a cadeia web (Python), do upload ao consolidado.
# --------------------------------------------------------------------------- #
def _ciclos_do_dataframe(df: Any) -> list[dict[str, Any]]:
    """Uma linha por ciclo, com as grandezas exigidas pelo item 5 do PR 0."""
    if df is None or not hasattr(df, "iterrows"):
        return []
    campos = {
        "ciclo": "Ciclo",
        "data_base": "Data-base",
        "intervalo_do_indice": "Intervalo do índice",
        "janela_de_admissibilidade": "Janela de admissibilidade",
        "data_do_pedido": "Data do pedido",
        "situacao": "Situação",
        "efeito_financeiro": "Tratamento financeiro do ciclo",
        "variacao_apurada": "Variação",
        "fator_proprio": "Fator",
        "fator_acumulado": "Fator acumulado",
        "fator_acumulado_efetivo": "Fator acumulado efetivo",
        "fator_ciclo_efetivo": "Fator ciclo efetivo",
    }
    saida: list[dict[str, Any]] = []
    for _, linha in df.iterrows():
        registro = {
            chave: (normalizar(linha[coluna]) if coluna in df.columns else None)
            for chave, coluna in campos.items()
        }
        # "Percentual efetivamente aplicado" nao e coluna do DataFrame: e o
        # fator do ciclo DEPOIS do tratamento (negativo aplicado/neutralizado)
        # menos 1. Derivacao de leitura, nao regra de negocio nova.
        fator_efetivo = _num(registro.get("fator_ciclo_efetivo"))
        registro["percentual_efetivamente_aplicado"] = (
            fator_efetivo - 1.0 if fator_efetivo is not None else None
        )
        saida.append(registro)
    return saida


def _retroativo_por_ciclo(memoria: dict[str, Any] | None) -> list[dict[str, Any]]:
    saida: list[dict[str, Any]] = []
    for ciclo in ((memoria or {}).get("ciclos") or []):
        if not isinstance(ciclo, dict):
            continue
        retroativo = ciclo.get("retroativo") or {}
        residuais = ciclo.get("residuais") or {}
        saida.append({
            "ciclo": normalizar(ciclo.get("ciclo")),
            "data_inicio": normalizar(ciclo.get("data_inicio")),
            "data_fim": normalizar(ciclo.get("data_fim")),
            "fator_acumulado": normalizar(ciclo.get("fator_acumulado")),
            "retroativo": {
                fonte: normalizar(retroativo.get(fonte))
                for fonte in ("financeiro", "pc", "consumidos")
            },
            "residuais": {
                "quantidade": normalizar(residuais.get("quantidade")),
                "valor_original": normalizar(residuais.get("valor_original")),
                "valor_atualizado": normalizar(residuais.get("valor_atualizado")),
                "itens": normalizar(residuais.get("itens")),
            },
        })
    return saida


def fotografar_web(resultado: dict[str, Any],
                   diagnostico: dict[str, Any]) -> dict[str, Any]:
    """Fotografia da cadeia Python — o que a web mostra e por que mostra."""
    consolidado = resultado.get("resultado_consolidado") or {}
    metadados = diagnostico.get("metadados") or {}
    status_xls = metadados.get("status_resultados") or {}
    convergencia = resultado.get("reconciliacao_xls_python") or {}
    return {
        # --- identificacao da apuracao --------------------------------------
        "metodo": normalizar(consolidado.get("metodo")),
        "metodo_do_retroativo": normalizar(status_xls.get("metodo_retroativo")),
        "indice": normalizar(metadados.get("indice")),
        "ciclo_vigente": normalizar(consolidado.get("ciclo_vigente")),
        "data_corte": normalizar(metadados.get("data_corte")),
        "ciclos_considerados": normalizar(metadados.get("ciclos_em_analise")),
        "quantidade_ciclos": normalizar(resultado.get("quantidade_ciclos")),

        # --- status global ---------------------------------------------------
        "status_apuracao": normalizar(consolidado.get("status_apuracao")),
        "status_confiabilidade": normalizar(consolidado.get("status_confiabilidade")),
        "mensagem_status": normalizar(consolidado.get("mensagem_status")),
        "formalizacao": normalizar(consolidado.get("formalizacao")),

        # --- por ciclo: situacao, variacao, percentual, fatores, efeito ------
        "ciclos": _ciclos_do_dataframe(resultado.get("df_ciclos")),
        "retroativo_por_ciclo": _retroativo_por_ciclo(
            resultado.get("memoria_por_ciclo")
        ),

        # --- grandezas monetarias canonicas ----------------------------------
        "vta_oficial": normalizar(consolidado.get("vta")),
        "vta_origem": normalizar(consolidado.get("vta_origem")),
        "vta_usa_ultima_posicao": normalizar(consolidado.get("vta_usa_ultima_posicao")),
        "retroativo_total": normalizar(consolidado.get("retroativo_reconhecido")),
        "retroativo_potencial": normalizar(consolidado.get("retroativo_potencial")),
        "valor_atualizado_em_analise": normalizar(
            consolidado.get("valor_atualizado_em_analise")
        ),
        "valor_atualizado_contrato": normalizar(
            resultado.get("valor_atualizado_contrato")
        ),
        "execucao_atualizada_do_ciclo": normalizar(
            resultado.get("valor_executado_atualizado")
        ),
        "remanescente_original": normalizar(resultado.get("remanescente_original")),
        "remanescente_atualizado": normalizar(resultado.get("remanescente_reajustado")),
        "quantidade_remanescente": normalizar(
            (status_xls.get("valores") or {}).get("quantidade_remanescente")
        ),
        "fator_acumulado_global": normalizar(resultado.get("fator_acumulado")),
        "variacao_acumulada_global": normalizar(resultado.get("variacao_acumulada")),

        # --- composicao e conferencia do VTA ---------------------------------
        "composicao_vta": normalizar(consolidado.get("composicao_vta")),
        "medidas_pc_aplicaveis": normalizar(consolidado.get("medidas_pc_aplicaveis")),
        "totais_canonicos_pc": normalizar(resultado.get("totais_canonicos_pc")),
        "fora_do_corte": normalizar(consolidado.get("fora_do_corte")),

        # --- referencias para auditoria (nunca sao o VTA oficial) ------------
        "referencias_auditaveis": normalizar(consolidado.get("referencias_auditaveis")),
        "referencias_vta_xls": normalizar(resultado.get("referencias_vta")),

        # --- convergencia natural XLS x Python -------------------------------
        "convergencia_xls_python": {
            "disponivel": normalizar(convergencia.get("disponivel")),
            "sem_cache": normalizar(convergencia.get("sem_cache")),
            "tolerancia": normalizar(convergencia.get("tolerancia")),
            "status_geral": normalizar(convergencia.get("status_geral")),
            "campos": normalizar(convergencia.get("campos")),
            "divergencias_relevantes": normalizar(
                convergencia.get("divergencias_relevantes")
            ),
        },

        # --- mensagens de bloqueio / estimativa / revisao --------------------
        "bloqueios": normalizar(consolidado.get("bloqueios")),
        "ressalvas": normalizar(consolidado.get("ressalvas")),
        "informacoes": normalizar(consolidado.get("informacoes")),
        "campos_nao_confiaveis": normalizar(consolidado.get("campos_nao_confiaveis")),
        "avisos_da_leitura": normalizar(diagnostico.get("avisos")),
        "pendencias": normalizar(diagnostico.get("pendencias")),
        "status_resultados_xls": normalizar(status_xls),
    }


# --------------------------------------------------------------------------- #
# CAMADA 4 — documentos que consomem o resultado consolidado.
# --------------------------------------------------------------------------- #
def _texto_docx(conteudo: bytes) -> list[str]:
    from docx import Document

    documento = Document(io.BytesIO(conteudo))
    linhas = [p.text.strip() for p in documento.paragraphs if p.text.strip()]
    for tabela in documento.tables:
        for linha in tabela.rows:
            celulas = [celula.text.strip() for celula in linha.cells]
            if any(celulas):
                linhas.append(" | ".join(celulas))
    return linhas


_MARCADORES_NEGOCIAIS = (
    "R$", "%", "VTA", "Ciclo", "CICLO", "C0", "C1", "C2", "C3", "C4",
    "retroativ", "Retroativ", "RETROATIV", "remanescent", "Remanescent",
    "REMANESCENT", "TEMPESTIV", "PRECLUS", "VALIDADO", "REVISE", "ESTIMADO",
)


def _valores_negociais(linhas: list[str]) -> list[str]:
    """So o conteudo negocial: linhas com valor, percentual, ciclo ou situacao.

    Snapshot de bytes inteiros seria fragil (fonte, espacamento, data de
    geracao, ID de apuracao). O que precisa sobreviver a refatoracao e o NUMERO
    e a DECISAO — e so isso que entra no baseline.
    """
    return [linha for linha in linhas
            if any(marcador in linha for marcador in _MARCADORES_NEGOCIAIS)]


def _sumario_negocial(dados: dict[str, Any]) -> dict[str, Any]:
    """Do Sumario Executivo interessa a sintese e a composicao, nao o PDF."""
    return {
        "sintese": normalizar(dados.get("sintese")),
        "composicao_vta": normalizar(dados.get("composicao_vta")),
        "ciclos": normalizar(dados.get("ciclos")),
        "observacoes": normalizar(dados.get("observacoes")),
        "status_vta": normalizar(dados.get("status_vta")),
    }


def _ciclos_canonicos(resultado: dict[str, Any]) -> list[dict[str, Any]]:
    """Ciclos na forma que a comunicacao a contratada consome.

    ARMADILHA CONHECIDA: o `parametros_v10` canonico NAO esta em
    `resultado["parametros_v10"]`; ele vive em
    `objeto_processo.dados_operacionais.parametros_v10.por_ciclo`. Ler do lugar
    errado devolve dicionario vazio e a comunicacao sai com o texto-modelo,
    passando falso-verde no baseline.
    """
    objeto = resultado.get("objeto_processo") or {}
    operacionais = objeto.get("dados_operacionais") or {}
    por_ciclo = (operacionais.get("parametros_v10") or {}).get("por_ciclo") or {}
    return [dict(registro, ciclo=nome) for nome, registro in por_ciclo.items()]


def _comunicacao_negocial(rascunho: Any) -> dict[str, Any]:
    """`gerar_rascunho_email_contratada` devolve a tupla (assunto, corpo)."""
    if isinstance(rascunho, tuple) and len(rascunho) == 2:
        assunto, corpo = rascunho
    elif isinstance(rascunho, dict):
        assunto, corpo = rascunho.get("assunto"), rascunho.get("corpo")
    else:
        assunto, corpo = None, rascunho
    linhas = [linha.strip() for linha in str(corpo or "").splitlines() if linha.strip()]
    return {
        "assunto": normalizar(assunto),
        "linhas_negociais": _valores_negociais(linhas),
    }


def _dou_negocial(resultado: dict[str, Any]) -> dict[str, Any]:
    """Do DOU interessam os campos AUTOMATICOS — os que derivam da apuracao.

    O extrato do DOU nao tem gerador em modulo proprio: as funcoes vivem em
    `pages/13_DOU.py`. A pagina importa em modo bare (o Streamlit apenas avisa
    da falta de ScriptRunContext), entao da para chamar as funcoes puras sem
    subir um AppTest. Os campos manuais (contrato, processo, signatarios) nao
    entram: nao derivam do resultado e nada tem a ver com a RESULTADOS.
    """
    import importlib.util
    from pathlib import Path

    caminho = Path(__file__).resolve().parents[1] / "pages" / "13_DOU.py"
    especificacao = importlib.util.spec_from_file_location("_dou_baseline", caminho)
    modulo = importlib.util.module_from_spec(especificacao)
    especificacao.loader.exec_module(modulo)
    return {
        "valor_total_atualizado": normalizar(modulo.valor_total_atualizado(resultado)),
        "reajuste_contratual": normalizar(modulo.ciclos_reajuste_texto(resultado, {})),
    }


def fotografar_documentos(resultado: dict[str, Any]) -> dict[str, Any]:
    """Conteudo negocial dos documentos que consomem o resultado consolidado.

    Cada gerador e chamado como a pagina chama. Falha de um gerador e
    REGISTRADA, nunca engolida: um documento que deixa de ser gerado e uma
    mudanca de comportamento tao relevante quanto um numero diferente.
    """
    from _email_contratada import gerar_rascunho_email_contratada
    from _sumario_executivo import montar_dados_sumario_executivo
    from _templates_documentos import gerar_despacho_saneador, gerar_termo_apostila

    saida: dict[str, Any] = {}

    def _registrar(nome: str, funcao) -> None:
        try:
            saida[nome] = funcao()
        except Exception as excecao:
            saida[nome] = {"erro": f"{type(excecao).__name__}: {excecao}"}

    _registrar("termo_apostila", lambda: {
        "linhas_negociais": _valores_negociais(
            _texto_docx(gerar_termo_apostila(resultado))
        ),
    })
    _registrar("despacho_saneador", lambda: {
        "linhas_negociais": _valores_negociais(
            _texto_docx(gerar_despacho_saneador(resultado))
        ),
    })
    _registrar("sumario_executivo",
               lambda: _sumario_negocial(montar_dados_sumario_executivo(resultado)))
    # A comunicacao a contratada e PRE-apuracao: nao consome o consolidado, e
    # sim os ciclos e o acumulado canonico — exatamente como a pagina a chama.
    # Entra no baseline porque publica percentual e situacao por ciclo, que sao
    # as mesmas grandezas que a RESULTADOS apresenta.
    _registrar("comunicacao_contratada", lambda: _comunicacao_negocial(
        gerar_rascunho_email_contratada(
            _ciclos_canonicos(resultado),
            indice=resultado.get("indice"),
            fator_acumulado=resultado.get("fator_acumulado"),
        )
    ))
    _registrar("dou", lambda: _dou_negocial(resultado))
    return saida


# --------------------------------------------------------------------------- #
# Fotografia completa de um cenario.
# --------------------------------------------------------------------------- #
def fotografar_cenario(conteudo: bytes, *,
                       com_documentos: bool = True) -> dict[str, Any]:
    """Roda as cadeias de producao sobre `conteudo` e devolve a fotografia."""
    from _coleta_reajuste_documentos import processar_coleta_oficial_runtime

    resultado, diagnostico = processar_coleta_oficial_runtime(conteudo)
    fotografia: dict[str, Any] = {
        "contrato_xls": fotografar_contrato_xls(conteudo),
        "valores_xls": fotografar_valores_xls(conteudo),
        "web": fotografar_web(resultado, diagnostico),
    }
    if com_documentos:
        fotografia["documentos"] = fotografar_documentos(resultado)
    return fotografia
