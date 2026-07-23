"""§10.4 — Nenhum ARQUIVO entregue pode conter emoji/pictograma.

Cobre os seis documentos:
  - Sumario Executivo (PDF), Despacho Saneador (DOCX), Termo de Apostila (DOCX)
    sao gerados e tem o texto extraido e verificado.
  - Adequacao Orcamentaria (XLSX), Garantia (PDF) e DOU (DOCX) sao construidos
    dentro de paginas Streamlit; verifica-se, por inspecao de codigo, que nao ha
    emoji embutido no fluxo de escrita do arquivo (emojis so na UI web).
"""
from __future__ import annotations

import re
import sys
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from _sanitizacao_documental import contem_emoji, remover_emojis_leve  # noqa: E402
from _sumario_executivo import (  # noqa: E402
    montar_dados_sumario_executivo,
    gerar_sumario_executivo_pdf,
)
from _templates_documentos import gerar_despacho_saneador, gerar_termo_apostila  # noqa: E402
from test_sumario_executivo import leitura_simples_financeiro  # noqa: E402


def _leitura_com_emoji():
    """Injeta emoji nas situacoes de ciclo para exercitar a sanitizacao."""
    leit = leitura_simples_financeiro()
    for nome, reg in leit["parametros_v10"]["por_ciclo"].items():
        reg["situacao"] = f"🟢 {reg.get('situacao') or 'Computado'} ⚠️"
    return leit


def _texto_docx(b: bytes) -> str:
    doc = Document(BytesIO(b))
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                partes.append(cell.text)
    return "\n".join(partes)


def test_apostila_docx_sem_emoji_mesmo_com_situacao_emojizada():
    texto = _texto_docx(gerar_termo_apostila(_leitura_com_emoji()))
    assert not contem_emoji(texto)


def test_saneador_docx_sem_emoji_mesmo_com_situacao_emojizada():
    texto = _texto_docx(gerar_despacho_saneador(_leitura_com_emoji()))
    assert not contem_emoji(texto)


def test_sumario_dados_situacao_sanitizada():
    dados = montar_dados_sumario_executivo(_leitura_com_emoji())
    for c in dados["ciclos"]:
        assert not contem_emoji(c.get("situacao") or "")


def test_sumario_executivo_pdf_sem_emoji():
    fitz = pytest.importorskip("fitz")
    dados = montar_dados_sumario_executivo(_leitura_com_emoji())
    pdf = gerar_sumario_executivo_pdf(dados)
    assert pdf[:4] == b"%PDF"
    doc = fitz.open(stream=pdf, filetype="pdf")
    texto = "\n".join(page.get_text() for page in doc)
    assert not contem_emoji(texto)


def test_dou_builder_sanitiza_entrada():
    """gerar_docx_dou remove emoji do texto do usuario antes de escrever."""
    fonte = (ROOT / "pages" / "13_DOU.py").read_text(encoding="utf-8")
    assert "remover_emojis_leve" in fonte
    assert remover_emojis_leve("📢 Extrato — R$ 1,00 ✅") == " Extrato — R$ 1,00 "


# Faixa de pictogramas usada tambem em outros modulos (espelha o sanitizador).
_EMOJI = re.compile(
    "[\U0001F300-\U0001FAFF\U00002600-\U000027BF\U00002B00-\U00002BFF"
    "\U0001F1E6-\U0001F1FF\U0000FE0F\U00002190-\U000021FF"
    "\U00002300-\U000023FF\U000025A0-\U000025FF]"
)
# Linhas de UI web (emojis permitidos): rotulos/markdown/alertas Streamlit.
_UI_WEB = re.compile(r"st\.(button|markdown|title|header|subheader|caption|"
                     r"info|success|warning|error|write|metric|expander|tabs|radio)")


@pytest.mark.parametrize("pagina", [
    "pages/12_Adequacao_Orcamentaria.py",
    "pages/05_Garantia.py",
    "pages/13_DOU.py",
])
def test_paginas_documentais_emoji_apenas_na_ui_web(pagina):
    """Qualquer emoji nessas paginas deve estar em linha de UI web (permitido),
    nunca em string escrita diretamente no arquivo gerado."""
    for i, linha in enumerate((ROOT / pagina).read_text(encoding="utf-8").splitlines(), 1):
        if _EMOJI.search(linha):
            assert _UI_WEB.search(linha), f"{pagina}:{i} emoji fora de UI web: {linha.strip()!r}"
