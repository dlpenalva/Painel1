"""Testes dos modelos canonicos de Termo de Apostila (§6/§10.2) e
Despacho Saneador (§7/§10.3).
"""
from __future__ import annotations

import sys
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from _templates_documentos import (  # noqa: E402
    diagnosticar_campos_manuais,
    gerar_despacho_saneador,
    gerar_termo_apostila,
    _extrair_dados,
    _composicao_didatica_vta,
    _fmt_pct_doc,
    _ta_considerandos,
    _vta_texto_doc,
    formatar_moeda,
)
from _sanitizacao_documental import contem_emoji  # noqa: E402
from test_sumario_executivo import (  # noqa: E402
    leitura_ausencias,
    leitura_multiciclo_pc,
    leitura_simples_financeiro,
)

# ---------------------------------------------------------------------------
# Campos manuais
# ---------------------------------------------------------------------------

CAMPOS_TERMO = {
    "contrato": "TLB-CTR-2025/00001",
    "empresa_contratada": "Empresa XPTO S.A., CNPJ 00.000.000/0001-00",
    "representante_telebras_1_nome": "Fulano de Tal",
    "representante_telebras_1_matricula": "12345",
    "representante_telebras_2_cargo": "Diretor Financeiro",
    "representante_telebras_2_matricula": "67890",
    "solicitacao_data": "10/10/2025",
    "solicitacao_ref": "TLB-AUT-2025/00100",
    "memoria_calculo_ref": "TLB-AUT-2026/00700",
    "concordancia_ref": "TLB-AUT-2026/00500",
    "regularidade_ref": "TLB-AUT-2026/00400",
    "adequacao_orcamentaria_ref": "TLB-DES-2026/00300",
    "processo_ref": "TLB-PRO-2026/01100",
    "valor_original_contrato": 1000000.0,
    "local_data": "20/07/2026",
}

CAMPOS_SANEADOR = {
    "contrato": "TLB-CTR-2025/00001",
    "empresa_contratada": "Empresa XPTO S.A.",
    "objeto_contrato": "prestação de serviços especializados",
    "vigencia_ate": "31/12/2027",
    "tipo_atualizacao": "reajuste",
    "processo_pleito": "TLB-AUT-2025/00100",
    "referencia_analise": "TLB-AUT-2026/00200",
    "memoria_calculo_ref": "TLB-AUT-2026/00200",
    "adequacao_orcamentaria_ref": "TLB-DES-2026/00300",
    "adequacao_orcamentaria_valor": 123456.78,
    "regularidade_ref": "TLB-AUT-2026/00400",
    "regularidade_situacao": "documentação apresentada para conferência",
    "concordancia_ref": "TLB-AUT-2026/00500",
    "concordancia_situacao": "manifestação juntada ao processo",
    "garantia_situacao": "verificação documental pendente",
}

# Vocabulario de implementacao proibido nos documentos (§5).
TERMOS_TECNICOS_PROIBIDOS = [
    "Base executada do financeiro", "G não exclui base", "EFEITO_FINANCEIRO",
    "RESULTADOS!B23", "RESULTADOS!B26", "coluna G", "EFEITO_FINANCEIRO_PC",
    "QTD_REM_AJUSTADA", "fonte dupla", "sheet12.xml", "motor Python",
    "XLS × Python", "delta financeiro", "valor nominal", "base financeira",
    "parcelas_computadas", "vta_sombra", "fonte_parcela",
]


def _texto_docx(docx_bytes: bytes) -> str:
    doc = Document(BytesIO(docx_bytes))
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                partes.append(cell.text)
    return "\n".join(partes)


def _titulos_quadros(docx_bytes: bytes) -> list[str]:
    doc = Document(BytesIO(docx_bytes))
    return [" | ".join(c.text for c in t.rows[0].cells) for t in doc.tables]


# ---------------------------------------------------------------------------
# Validade basica
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("gerador,leitura", [
    (gerar_despacho_saneador, leitura_simples_financeiro),
    (gerar_termo_apostila, leitura_simples_financeiro),
    (gerar_despacho_saneador, leitura_multiciclo_pc),
    (gerar_termo_apostila, leitura_multiciclo_pc),
    (gerar_despacho_saneador, leitura_ausencias),
    (gerar_termo_apostila, leitura_ausencias),
])
def test_docx_valido(gerador, leitura):
    b = gerador(leitura())
    assert isinstance(b, bytes) and b[:2] == b"PK" and len(b) > 100
    Document(BytesIO(b))  # abre sem erro


# ---------------------------------------------------------------------------
# §10.2 — APOSTILA
# ---------------------------------------------------------------------------

def test_apostila_titulo_exato():
    texto = _texto_docx(gerar_termo_apostila(leitura_simples_financeiro(), campos_manuais=CAMPOS_TERMO))
    assert "MINUTA DE TERMO DE APOSTILAMENTO" in texto


def test_apostila_qualificacao_canonica():
    texto = _texto_docx(gerar_termo_apostila(leitura_multiciclo_pc(), campos_manuais=CAMPOS_TERMO))
    assert "sociedade de economia mista" in texto
    assert "00.336.701/0001-04" in texto
    assert "SIG, Quadra 04, Bloco A" in texto
    assert "70.610-440" in texto
    assert "parágrafo 7º do art. 81 da Lei nº 13.303" in texto
    assert "Diretriz nº 229/2018" in texto
    # Nao pode conter a qualificacao antiga divergente
    assert "33.200.056/0001-41" not in texto
    assert "empresa pública federal" not in texto
    assert "SAS Quadra 05" not in texto


def test_apostila_nove_considerandos_na_ordem_aprovada():
    doc = Document(BytesIO(gerar_termo_apostila(
        leitura_multiciclo_pc(), campos_manuais=CAMPOS_TERMO
    )))
    textos = [p.text for p in doc.paragraphs]
    inicio = textos.index("CONSIDERANDO:") + 1
    considerandos = textos[inicio:inicio + 9]
    assert len(considerandos) == 9
    chaves_ordenadas = (
        "Cláusula Oitava",
        "deliberação da Diretoria Executiva",
        "solicitação da CONTRATADA",
        "histórico já formalizado",
        "informações encaminhadas pela área gestora",
        "memória de cálculo",
        "índice contratual",
        "concordância da CONTRATADA",
        "certidões de regularidade",
    )
    for numero, (paragrafo, chave) in enumerate(
        zip(considerandos, chaves_ordenadas), start=1
    ):
        assert paragrafo.startswith(f"{numero}. ")
        assert chave in paragrafo
    texto = "\n".join(considerandos)
    assert "Ata da 1869ª Reunião Ordinária, de 13 de janeiro de 2026" in texto
    assert "10/10/2025" in texto
    assert "TLB-AUT-2025/00100" in texto


def test_apostila_estrutura_final_1_a_8_sem_duplicidades():
    texto = _texto_docx(gerar_termo_apostila(leitura_multiciclo_pc(), campos_manuais=CAMPOS_TERMO))
    assert "FORMALIZA-SE O PRESENTE TERMO DE APOSTILA:" in texto
    assert "1. Dos reajustes concedidos" in texto
    assert "2. Da apuração financeira do retroativo" in texto
    assert "3. Da composição do Valor Total Atualizado" in texto
    assert "4. Dos valores unitários" in texto
    assert "5. Dos aditivos e supressões considerados" in texto
    assert "6. Permanecem inalteradas e em pleno vigor" in texto
    assert "7. A CONTRATADA deverá atualizar a garantia contratual" in texto
    assert "8. O presente apostilamento vincula-se" in texto
    assert "Da composição sintética do Valor Total Atualizado" not in texto
    assert "4-A." not in texto
    assert "Referências auditáveis do Valor Total Atualizado" not in texto


def test_apostila_quadros_sem_composicao_duplicada():
    quadros = _titulos_quadros(gerar_termo_apostila(leitura_multiciclo_pc(), campos_manuais=CAMPOS_TERMO))
    assert "Ref. | Ciclo | Percentual aplicado | Efeitos financeiros | Situação" in quadros  # Q1
    assert "Ciclo | Valor pago efetivo | Valor devido após o reajuste | Diferença/retroativo" in quadros  # Q2
    assert "Ref. | Descrição | Valor" in quadros  # Q3
    assert "Ref. | Parcela | Valor" not in quadros
    assert quadros.count("Ref. | Descrição | Valor") == 1


def test_apostila_vu_ate_ultimo_ciclo_sem_futuro():
    b = gerar_termo_apostila(leitura_simples_financeiro(), campos_manuais=CAMPOS_TERMO)
    quadros = _titulos_quadros(b)
    # Etapa 26D: cabecalhos VU_Cn (estrutura canonica do quadro historico).
    vu = [q for q in quadros if "VU_C0" in q]
    assert vu, "tabela de VU ausente"
    assert "VU_C1" in vu[0]
    assert "Descrição" not in vu[0]
    assert "VU_C2" not in vu[0]  # ciclo futuro nao entra
    assert "VU_C3" not in vu[0]


def test_apostila_duas_assinaturas_telebras_sem_contratada():
    texto = _texto_docx(gerar_termo_apostila(leitura_simples_financeiro(), campos_manuais=CAMPOS_TERMO))
    assert texto.count("TELECOMUNICAÇÕES BRASILEIRAS S.A. - TELEBRAS") >= 3  # qualificacao + 2 assinaturas
    linhas = texto.splitlines()
    assert not any(l.strip() == "CONTRATADA" for l in linhas)


def test_apostila_composicao_vta_unica_preserva_componentes_canonicos():
    leitura = leitura_multiciclo_pc()
    leitura["composicao_vta"] = {
        "disponivel": True,
        "metodo": "pc",
        "total_execucao_atualizada": 13_973_327.58,
        "saldo_remanescente": {"valor_atualizado": 123_402_232.71},
        "vta_composicao": 137_375_560.29,
    }
    dados = _extrair_dados(leitura, None)
    componentes = _composicao_didatica_vta(dados)
    doc = Document(BytesIO(gerar_termo_apostila(
        leitura, campos_manuais=CAMPOS_TERMO
    )))
    tabela = next(
        t for t in doc.tables
        if [c.text for c in t.rows[0].cells] == ["Ref.", "Descrição", "Valor"]
    )
    linhas = [[c.text for c in row.cells] for row in tabela.rows[1:]]
    assert [linha[1] for linha in linhas[:-1]] == [d for d, _ in componentes]
    assert [linha[2] for linha in linhas[:-1]] == [
        formatar_moeda(valor) if valor is not None else ""
        for _, valor in componentes
    ]
    assert linhas[-1][1] == "Valor Total Atualizado do Contrato"
    assert linhas[-1][2] == _vta_texto_doc(dados)


def test_apostila_terminologia_exclusiva_e_seguranca_da_tempestividade():
    for bytes_docx in (
        gerar_termo_apostila(
            leitura_multiciclo_pc(), campos_manuais=CAMPOS_TERMO
        ),
        gerar_termo_apostila({}, {}, {}, modo_modelo_em_branco=True),
    ):
        texto = _texto_docx(bytes_docx)
        assert "fiscal" not in texto.lower()
        assert "valor teórico" not in texto.lower()
        assert "valor devido após o reajuste" in texto

    doc_neutro = Document()
    _ta_considerandos(doc_neutro, {
        "_modo_branco": False,
        "ciclos_computados": [{
            "situacao": "PRECLUSO", "data_pedido": "10/10/2025"
        }],
        "identificacao": {},
        "var_acumulada": None,
    }, CAMPOS_TERMO)
    texto_neutro = "\n".join(p.text for p in doc_neutro.paragraphs)
    assert "solicitação tempestiva" not in texto_neutro.lower()

    doc_tempestivo = Document()
    _ta_considerandos(doc_tempestivo, {
        "_modo_branco": False,
        "ciclos_computados": [{
            "situacao": "TEMPESTIVO*", "data_pedido": "15/10/2025"
        }],
        "identificacao": {},
        "var_acumulada": None,
    }, CAMPOS_TERMO)
    texto_tempestivo = "\n".join(p.text for p in doc_tempestivo.paragraphs)
    assert "solicitação tempestiva da CONTRATADA, de 15/10/2025" in texto_tempestivo


def test_apostila_espaco_visual_apos_capitulo_5():
    doc = Document(BytesIO(gerar_termo_apostila(
        leitura_multiciclo_pc(), campos_manuais=CAMPOS_TERMO
    )))
    textos = [p.text for p in doc.paragraphs]
    indice_52 = next(i for i, texto in enumerate(textos) if texto.startswith("5.2."))
    assert textos[indice_52 + 1] == ""
    assert textos[indice_52 + 2].startswith("6. Permanecem inalteradas")


def test_apostila_sem_termos_tecnicos_e_sem_emoji():
    for leit in (leitura_simples_financeiro, leitura_multiciclo_pc):
        texto = _texto_docx(gerar_termo_apostila(leit(), campos_manuais=CAMPOS_TERMO))
        assert not contem_emoji(texto)
        for termo in TERMOS_TECNICOS_PROIBIDOS:
            assert termo not in texto, f"termo tecnico proibido: {termo}"


# ---------------------------------------------------------------------------
# §10.3 — SANEADOR
# ---------------------------------------------------------------------------

def test_saneador_assunto_e_identificacao():
    texto = _texto_docx(gerar_despacho_saneador(
        leitura_simples_financeiro(), campos_manuais=CAMPOS_SANEADOR
    ))
    assert "DESPACHO SANEADOR" in texto
    assert "Saneamento para formalização de reajuste — TLB-CTR-2025/00001" in texto
    assert "Referência(s): TLB-AUT-2025/00100" in texto
    assert "celebrado com Empresa XPTO S.A." in texto
    assert "cujo objeto é prestação de serviços especializados" in texto
    assert "com vigência até 31/12/2027" in texto


def test_saneador_estrutura_final_1_a_6():
    doc = Document(BytesIO(gerar_despacho_saneador(
        leitura_multiciclo_pc(), campos_manuais=CAMPOS_SANEADOR
    )))
    titulos = [p.text for p in doc.paragraphs if p.text[:2] in {f"{n}." for n in range(1, 10)}]
    assert titulos == [
        "1. IDENTIFICAÇÃO",
        "2. PEDIDO E PARÂMETROS DA ANÁLISE",
        "3. RESULTADO ESSENCIAL",
        "4. DOCUMENTOS E VERIFICAÇÕES",
        "5. PENDÊNCIAS",
        "6. CONCLUSÃO",
    ]


def test_saneador_tres_quadros_enxutos_e_blocos_removidos():
    b = gerar_despacho_saneador(
        leitura_multiciclo_pc(), campos_manuais=CAMPOS_SANEADOR
    )
    texto = _texto_docx(b)
    quadros = _titulos_quadros(b)
    assert len(quadros) == 3
    assert quadros == [
        "Ciclo | Data-base | Data do pedido | Situação | Efeito financeiro | Percentual",
        "Resultado | Valor",
        "Documento ou verificação | Referência ou situação",
    ]
    assert "Quadro 1 - Síntese da análise" in texto
    assert "Quadro 2 - Síntese financeira" in texto
    assert "Quadro 3 - Documentos e verificações" in texto
    for removido in (
        "valor teórico",
        "Referências auditáveis",
        "De forma didática",
        "Memória fiscal do Valor Total Atualizado",
        "Valores Unitários por Ciclo",
        "Saldo remanescente",
    ):
        assert removido.lower() not in texto.lower()


def test_saneador_inicio_financeiro_usa_efeito_real():
    b = gerar_despacho_saneador(leitura_multiciclo_pc(), campos_manuais=CAMPOS_SANEADOR)
    tabela = Document(BytesIO(b)).tables[0]  # Quadro 1
    linha_c2 = next(r for r in tabela.rows if r.cells[0].text == "C2")
    assert linha_c2.cells[1].text == "01/05/2025"   # Data-base
    assert linha_c2.cells[4].text == "A partir de 01/08/2025"


def test_saneador_sem_tabela_de_valores_unitarios():
    texto = _texto_docx(gerar_despacho_saneador(leitura_simples_financeiro(), campos_manuais=CAMPOS_SANEADOR))
    assert "VU C0" not in texto
    assert "Valores Unitários" not in texto


def test_saneador_itens_administrativos_presentes():
    texto = _texto_docx(gerar_despacho_saneador(leitura_multiciclo_pc(), campos_manuais=CAMPOS_SANEADOR))
    for item in (
        "Memória de cálculo",
        "Adequação orçamentária",
        "Regularidade da contratada",
        "Concordância da contratada",
        "Garantia contratual",
    ):
        assert item in texto


def test_saneador_resultado_consolidado_sem_detalhamento_por_ciclo():
    texto = _texto_docx(gerar_despacho_saneador(
        leitura_multiciclo_pc(), campos_manuais=CAMPOS_SANEADOR
    ))
    for rotulo in (
        "Valor pago no período analisado",
        "Valor devido após o reajuste",
        "Retroativo a pagar",
        "Valor Total Atualizado do Contrato",
    ):
        assert rotulo in texto
    assert "Apuração financeira por ciclo" not in texto


def test_saneador_conclusao_neutra_sem_habilitacao_nova():
    texto = _texto_docx(gerar_despacho_saneador(
        leitura_multiciclo_pc(), campos_manuais=CAMPOS_SANEADOR
    ))
    assert "deverá ser avaliado o prosseguimento da instrução" in texto
    assert "instrução encontra-se apta" not in texto


def test_saneador_conclusao_impeditiva_com_pendencia_suportada():
    cm = dict(CAMPOS_SANEADOR, pendencia_critica=True)
    texto = _texto_docx(gerar_despacho_saneador(leitura_multiciclo_pc(), campos_manuais=cm))
    assert "A instrução deverá ser complementada quanto às pendências acima" in texto
    assert "instrução encontra-se apta" not in texto


def test_saneador_documentos_desatualizados_ficam_nas_pendencias():
    cm = dict(CAMPOS_SANEADOR, docs_desatualizados=["SEI 999/2026", "SEI 888/2026"])
    texto = _texto_docx(gerar_despacho_saneador(leitura_multiciclo_pc(), campos_manuais=cm))
    assert "Documentos desatualizados: SEI 999/2026, SEI 888/2026" in texto
    assert texto.count("SEI 999/2026") == 1


def test_saneador_identificacao_externa_confiavel_e_automatica():
    identificacao = {
        "contrato": "TLB-CTR-2026/00077",
        "empresa_contratada": "Contratada Canônica S.A.",
        "objeto_contrato": "serviço continuado",
        "vigencia_ate": "30/06/2028",
        "tipo_atualizacao": "repactuação",
    }
    texto = _texto_docx(gerar_despacho_saneador(
        leitura_simples_financeiro(), identificacao=identificacao,
        campos_manuais={"processo_pleito": "TLB-AUT-2026/00999"},
    ))
    for valor in identificacao.values():
        assert valor in texto
    for ausente in (
        "[PREENCHER: Numero do contrato]",
        "[PREENCHER: Nome da empresa contratada]",
        "[PREENCHER: Objeto resumido do contrato]",
        "[PREENCHER: Data final da vigencia contratual]",
    ):
        assert ausente not in texto


def test_saneador_sem_termos_tecnicos_e_sem_emoji():
    for leit in (leitura_simples_financeiro, leitura_multiciclo_pc):
        texto = _texto_docx(gerar_despacho_saneador(leit(), campos_manuais=CAMPOS_SANEADOR))
        assert not contem_emoji(texto)
        for termo in TERMOS_TECNICOS_PROIBIDOS:
            assert termo not in texto, f"termo tecnico proibido: {termo}"


# ---------------------------------------------------------------------------
# Robustez comum
# ---------------------------------------------------------------------------

def test_ausencias_nao_viram_zero():
    for gerador in (gerar_termo_apostila, gerar_despacho_saneador):
        texto = _texto_docx(gerador(leitura_ausencias(), campos_manuais={}))
        assert "R$ 0,00" not in texto or "[PREENCHER:" in texto


def test_campos_ausentes_geram_preencher():
    for gerador in (gerar_termo_apostila, gerar_despacho_saneador):
        texto = _texto_docx(gerador(leitura_simples_financeiro(), campos_manuais={}))
        assert "[PREENCHER:" in texto


def test_diagnostico_pendencias():
    pend = diagnosticar_campos_manuais(leitura_simples_financeiro(), campos_manuais=None)
    assert isinstance(pend, list) and pend
    for item in pend:
        assert {"campo", "descricao", "documento"} <= set(item)
    campos = dict(CAMPOS_TERMO, **CAMPOS_SANEADOR)
    pend2 = [p["campo"] for p in diagnosticar_campos_manuais(leitura_simples_financeiro(), campos_manuais=campos)]
    for chave in campos:
        assert chave not in pend2


def test_sem_dados_padtec_hardcoded():
    for gerador in (gerar_termo_apostila, gerar_despacho_saneador):
        texto = _texto_docx(gerador(leitura_multiciclo_pc(), campos_manuais={}))
        for proibida in ("TLB-CTR-2022/00067", "PADTEC", "158.292.598"):
            assert proibida not in texto


def test_fmt_pct_doc():
    assert _fmt_pct_doc(0.04) == "4,00%"
    assert _fmt_pct_doc(0.106231) == "10,62%"
    assert _fmt_pct_doc(-0.02) == "-2,00%"
