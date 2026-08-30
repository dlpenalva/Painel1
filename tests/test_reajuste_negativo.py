from __future__ import annotations

import io
import math
from copy import deepcopy
from datetime import date
from pathlib import Path

import pytest
from docx import Document
from openpyxl import load_workbook

from _coleta_oficial import gerar_coleta_oficial_preenchida
from _email_contratada import gerar_rascunho_email_contratada
from _leitor_masterfile_v10 import ler_masterfile_v10
from _motor_composicao_vta import _execucao_por_ciclo
from _objeto_processo_reajuste import _fatores_acumulados_por_percentual
from _reajuste_utils import (
    APLICAR_VARIACAO_NEGATIVA,
    NEUTRALIZAR_VARIACAO_NEGATIVA,
    resolver_tratamento_variacao_negativa,
    situacao_com_tratamento_variacao_negativa,
)
from _sumario_executivo import (
    _observacoes_reajuste_negativo,
    gerar_sumario_executivo,
    montar_dados_sumario_executivo,
)
from _templates_documentos import _extrair_dados
from _templates_documentos import gerar_despacho_saneador, gerar_termo_apostila
from test_sumario_executivo import leitura_simples_financeiro


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "templates" / "COLETA_REAJUSTE_OFICIAL.xlsx"
PAGINA_SIMPLES = ROOT / "pages" / "01_Calculo_Simples.py"
PAGINA_MULTIPLA = ROOT / "pages" / "02_Calculo_Represados.py"


def _wrapper_app(pagina: Path, anos_negativos: set[int]) -> str:
    return f'''
import runpy
import sys
from pathlib import Path

import pandas as pd
import streamlit as st

sys.path.insert(0, {str(ROOT)!r})
import _indice_utils

ANOS_NEGATIVOS = {sorted(anos_negativos)!r}

def _ist_deterministico(data_inicio):
    inicio = pd.Timestamp(data_inicio).replace(day=1)
    fim = inicio + pd.DateOffset(months=12)
    variacao = -0.02 if inicio.year in ANOS_NEGATIVOS else 0.05
    chamadas = list(st.session_state.get("_fake_ist_calls", []))
    chamadas.append(inicio.strftime("%Y-%m-%d"))
    st.session_state["_fake_ist_calls"] = chamadas
    return {{
        "variacao": variacao,
        "var": variacao,
        "i_ini": 100.0,
        "i_fim": 100.0 * (1.0 + variacao),
        "d_ini": inicio,
        "d_fim": fim,
        "p_ini": inicio,
        "p_fim": fim,
        "metodo": "IST determinístico de teste",
        "serie": "IST-TESTE",
        "dados": pd.DataFrame({{
            "data": [inicio, fim],
            "indice": [100.0, 100.0 * (1.0 + variacao)],
        }}),
    }}

_ist_original = _indice_utils.calcular_ist_numero_indice
try:
    _indice_utils.calcular_ist_numero_indice = _ist_deterministico
    runpy.run_path({str(pagina)!r}, run_name="__main__")
finally:
    _indice_utils.calcular_ist_numero_indice = _ist_original
'''


def _ciclo_payload(percentual: float, situacao: str) -> dict:
    return {
        "ciclo": "C1",
        "data_base": "01/02/2024",
        "periodo_inicio": "01/02/2024",
        "periodo_fim": "31/01/2025",
        "financeiro_inicio": "01/02/2025",
        "percentual_indice": -0.02,
        "percentual_aplicado": percentual,
        "variacao": percentual,
        "fator": 1.0 + percentual,
        "fator_acumulado": 1.0 + percentual,
        "situacao": situacao,
        "situacao_aplicada": situacao,
        "ciclo_negativo": True,
        "tratamento_ciclo_negativo": (
            APLICAR_VARIACAO_NEGATIVA
            if percentual < 0 else NEUTRALIZAR_VARIACAO_NEGATIVA
        ),
    }


def _dados_calculadora(percentual: float, situacao: str) -> dict:
    return {
        "origem": "Reajuste Simples",
        "tipo": "Simples",
        "indice": "IST (Anatel)",
        "data_base_original": "01/02/2023",
        "fator": 1.0 + percentual,
        "fator_acumulado": 1.0 + percentual,
        "variacao": percentual,
        "variacao_acumulada": percentual,
        "ciclos": [_ciclo_payload(percentual, situacao)],
    }


def _registro_c1(conteudo: bytes) -> dict:
    leitura = ler_masterfile_v10(conteudo)
    return leitura["parametros_v10"]["por_ciclo"]["C1"]


def _leitura_com_tratamento(situacao: str, percentual: float) -> dict:
    leitura = deepcopy(leitura_simples_financeiro())
    c1 = leitura["parametros_v10"]["por_ciclo"]["C1"]
    c1["percentual_reajuste"] = percentual
    c1["fator_acumulado"] = 1.0 + percentual
    c1["situacao"] = situacao
    return leitura


def _texto_docx(conteudo: bytes) -> str:
    doc = Document(io.BytesIO(conteudo))
    partes = [p.text for p in doc.paragraphs]
    for tabela in doc.tables:
        for linha in tabela.rows:
            partes.extend(celula.text for celula in linha.cells)
    return "\n".join(partes)


def _texto_pdf(conteudo: bytes) -> str:
    fitz = pytest.importorskip("fitz")
    doc = fitz.open(stream=conteudo, filetype="pdf")
    try:
        return "\n".join(pagina.get_text() for pagina in doc)
    finally:
        doc.close()


def _contrato_workbook(conteudo: bytes) -> dict:
    wb = load_workbook(io.BytesIO(conteudo), data_only=False)
    try:
        return {
            "abas": tuple(wb.sheetnames),
            "dimensoes": {
                nome: (wb[nome].max_row, wb[nome].max_column)
                for nome in wb.sheetnames
            },
            "cabecalhos": {
                nome: tuple(wb[nome].cell(1, c).value for c in range(1, wb[nome].max_column + 1))
                for nome in wb.sheetnames
            },
            "formulas": {
                nome: {
                    celula.coordinate: celula.value
                    for linha in wb[nome].iter_rows()
                    for celula in linha
                    if isinstance(celula.value, str) and celula.value.startswith("=")
                }
                for nome in wb.sheetnames
            },
            "validacoes": {
                nome: tuple(
                    (str(dv.sqref), dv.type, dv.formula1, dv.formula2)
                    for dv in wb[nome].data_validations.dataValidation
                )
                for nome in wb.sheetnames
            },
            "mesclas": {
                nome: tuple(sorted(str(r) for r in wb[nome].merged_cells.ranges))
                for nome in wb.sheetnames
            },
        }
    finally:
        wb.close()


@pytest.mark.parametrize("percentual", [0.0375, 0.0])
def test_resultado_nao_negativo_preserva_payload(percentual):
    resolvido = resolver_tratamento_variacao_negativa(percentual)
    assert resolvido == {
        "ciclo_negativo": False,
        "pendente": False,
        "percentual_indice": percentual,
        "percentual_aplicado": percentual,
        "fator": 1.0 + percentual,
        "tratamento": None,
    }


def test_resultado_negativo_sem_decisao_fica_fail_closed():
    resolvido = resolver_tratamento_variacao_negativa(-0.02)
    assert resolvido["pendente"] is True
    assert resolvido["percentual_indice"] == -0.02
    assert resolvido["percentual_aplicado"] is None
    assert resolvido["fator"] is None
    assert resolvido["tratamento"] is None


@pytest.mark.parametrize(
    ("tratamento", "percentual", "fator", "marcador"),
    (
        (
            APLICAR_VARIACAO_NEGATIVA,
            -0.02,
            0.98,
            "TEMPESTIVO — VARIAÇÃO NEGATIVA",
        ),
        (
            NEUTRALIZAR_VARIACAO_NEGATIVA,
            0.0,
            1.0,
            "TEMPESTIVO — VARIAÇÃO NEGATIVA NEUTRALIZADA EM 0,00%",
        ),
    ),
)
def test_tratamentos_negativos_e_marcadores_exatos(
    tratamento, percentual, fator, marcador
):
    resolvido = resolver_tratamento_variacao_negativa(-0.02, tratamento)
    assert resolvido["pendente"] is False
    assert resolvido["percentual_indice"] == -0.02
    assert resolvido["percentual_aplicado"] == percentual
    assert resolvido["fator"] == fator
    assert resolvido["tratamento"] == tratamento
    assert situacao_com_tratamento_variacao_negativa(
        "TEMPESTIVO", tratamento
    ) == marcador


def test_classificacao_base_e_tempestivo_asterisco_sao_preservados():
    assert situacao_com_tratamento_variacao_negativa(
        "ADIANTADO", APLICAR_VARIACAO_NEGATIVA
    ) == "ADIANTADO — VARIAÇÃO NEGATIVA"
    assert situacao_com_tratamento_variacao_negativa(
        "TEMPESTIVO*", NEUTRALIZAR_VARIACAO_NEGATIVA
    ) == "TEMPESTIVO* — VARIAÇÃO NEGATIVA NEUTRALIZADA EM 0,00%"


def test_meses_negativos_isolados_nao_abrem_decisao_se_final_for_positivo():
    taxas = (0.03, -0.01, 0.02)
    percentual_final = math.prod(1.0 + taxa for taxa in taxas) - 1.0
    assert percentual_final > 0
    resolvido = resolver_tratamento_variacao_negativa(percentual_final)
    assert resolvido["pendente"] is False
    assert resolvido["percentual_aplicado"] == percentual_final


def test_calculadora_simples_exige_decisao_e_invalida_estado_no_novo_resultado():
    from streamlit.testing.v1 import AppTest

    at = AppTest.from_string(
        _wrapper_app(PAGINA_SIMPLES, {2023, 2024}), default_timeout=180
    ).run()
    at.date_input[0].set_value(date(2023, 2, 1))
    at.run()
    at.date_input[1].set_value(date(2024, 2, 1))
    at.run()
    next(b for b in at.button if "Processar" in str(b.label)).click()
    at.run()

    escolha = at.radio(key="tratamento_variacao_negativa_simples_c1")
    assert escolha.value is None
    assert escolha.options == [
        "Aplicar a variação apurada (-2,00%)",
        "Neutralizar a variação e aplicar 0,00%",
    ]
    assert any("Decisão pendente" in str(info.value) for info in at.info)
    with pytest.raises(KeyError):
        at.session_state["dados_admissibilidade"]

    escolha.set_value(APLICAR_VARIACAO_NEGATIVA)
    at.run()
    ciclo = at.session_state["dados_admissibilidade"]["ciclos"][0]
    assert ciclo["percentual_indice"] == -0.02
    assert ciclo["percentual_aplicado"] == -0.02
    assert ciclo["fator"] == 0.98
    assert ciclo["situacao_aplicada"] == "TEMPESTIVO — VARIAÇÃO NEGATIVA"
    download_aplicar = at.get("download_button")[0].proto
    identidade_download_aplicar = download_aplicar.id
    url_download_aplicar = download_aplicar.url

    at.radio(key="tratamento_variacao_negativa_simples_c1").set_value(
        NEUTRALIZAR_VARIACAO_NEGATIVA
    )
    at.run()
    ciclo = at.session_state["dados_admissibilidade"]["ciclos"][0]
    assert ciclo["percentual_indice"] == -0.02
    assert ciclo["percentual_aplicado"] == 0.0
    assert ciclo["fator"] == 1.0
    download_neutralizar = at.get("download_button")[0].proto
    assert download_neutralizar.id != identidade_download_aplicar
    assert download_neutralizar.url != url_download_aplicar

    at.date_input[0].set_value(date(2024, 2, 1))
    at.run()
    next(b for b in at.button if "Processar" in str(b.label)).click()
    at.run()
    assert at.radio(key="tratamento_variacao_negativa_simples_c1").value is None


def test_multiciclo_para_em_c2_e_recompõe_cadeia_apos_decisao():
    from streamlit.testing.v1 import AppTest

    at = AppTest.from_string(
        _wrapper_app(PAGINA_MULTIPLA, {2023}), default_timeout=180
    ).run()
    at.selectbox(key="rep_ciclo_final_analise").select("C3")
    at.run()
    next(b for b in at.button if "Processar" in str(b.label)).click()
    at.run()

    assert at.session_state["_fake_ist_calls"] == ["2022-10-01", "2023-10-01"]
    assert at.radio(key="tratamento_variacao_negativa_multiplos_c2").value is None
    assert any("ciclos seguintes" in str(info.value) for info in at.info)
    with pytest.raises(KeyError):
        at.session_state["dados_admissibilidade"]

    at.radio(key="tratamento_variacao_negativa_multiplos_c2").set_value(
        APLICAR_VARIACAO_NEGATIVA
    )
    at.run()
    adm = at.session_state["dados_admissibilidade"]
    ciclos = {c["ciclo"]: c for c in adm["ciclos"]}
    assert ciclos["C2"]["fator"] == 0.98
    assert ciclos["C2"]["fator_acumulado"] == pytest.approx(1.05 * 0.98)
    assert ciclos["C3"]["fator_acumulado"] == pytest.approx(1.05 * 0.98 * 1.05)
    download_aplicar = at.get("download_button")[0].proto
    identidade_download_aplicar = download_aplicar.id
    url_download_aplicar = download_aplicar.url

    at.radio(key="tratamento_variacao_negativa_multiplos_c2").set_value(
        NEUTRALIZAR_VARIACAO_NEGATIVA
    )
    at.run()
    ciclos = {
        c["ciclo"]: c
        for c in at.session_state["dados_admissibilidade"]["ciclos"]
    }
    assert ciclos["C2"]["fator"] == 1.0
    assert ciclos["C2"]["fator_acumulado"] == pytest.approx(1.05)
    assert ciclos["C3"]["fator_acumulado"] == pytest.approx(1.05 * 1.05)
    download_neutralizar = at.get("download_button")[0].proto
    assert download_neutralizar.id != identidade_download_aplicar
    assert download_neutralizar.url != url_download_aplicar


@pytest.mark.parametrize(
    ("percentual", "situacao"),
    (
        (-0.02, "TEMPESTIVO — VARIAÇÃO NEGATIVA"),
        (0.0, "TEMPESTIVO — VARIAÇÃO NEGATIVA NEUTRALIZADA EM 0,00%"),
    ),
)
def test_xls_e_leitor_preservam_percentual_e_situacao(percentual, situacao):
    conteudo = gerar_coleta_oficial_preenchida(
        _dados_calculadora(percentual, situacao)
    )
    reg = _registro_c1(conteudo)
    assert reg["percentual_reajuste"] == percentual
    assert reg["situacao"] == situacao


def test_xls_nao_altera_geometria_formulas_validacoes_cabecalhos_ou_template():
    template_antes = TEMPLATE.read_bytes()
    contrato_base = _contrato_workbook(
        gerar_coleta_oficial_preenchida(
            _dados_calculadora(0.03, "TEMPESTIVO")
        )
    )
    for percentual, situacao in (
        (-0.02, "TEMPESTIVO — VARIAÇÃO NEGATIVA"),
        (0.0, "TEMPESTIVO — VARIAÇÃO NEGATIVA NEUTRALIZADA EM 0,00%"),
    ):
        saida = gerar_coleta_oficial_preenchida(
            _dados_calculadora(percentual, situacao)
        )
        assert _contrato_workbook(saida) == contrato_base
    assert TEMPLATE.read_bytes() == template_antes


def test_leitor_preserva_legado_zero_sem_reinterpretar():
    situacao_legada = "✅ TEMPESTIVO | 🔻 CICLO NEGATIVO (APLICADO 0,00%)"
    reg = _registro_c1(
        gerar_coleta_oficial_preenchida(
            _dados_calculadora(0.0, situacao_legada)
        )
    )
    assert reg["percentual_reajuste"] == 0.0
    assert reg["situacao"] == situacao_legada


def test_fator_negativo_flui_para_objeto_e_vta_sem_regra_duplicada():
    fatores = _fatores_acumulados_por_percentual([
        {"ciclo": "C1", "indice_percentual": 0.05},
        {"ciclo": "C2", "indice_percentual": -0.02},
        {"ciclo": "C3", "indice_percentual": 0.03},
    ])
    assert fatores["C2"] == pytest.approx(1.05 * 0.98)
    assert fatores["C3"] == pytest.approx(1.05 * 0.98 * 1.03)

    leitura = {"reconciliacao": {"registros": [{
        "ciclo": "C2",
        "valor_computado": 100.0,
        "fonte_principal": "teste",
    }]}}
    linhas = _execucao_por_ciclo(
        leitura, {"C2": {"fator_acumulado": 0.98}}, []
    )
    assert linhas[0]["valor_atualizado"] == 98.0


@pytest.mark.parametrize(
    ("situacao", "percentual"),
    (
        ("TEMPESTIVO — VARIAÇÃO NEGATIVA", -0.02),
        ("TEMPESTIVO — VARIAÇÃO NEGATIVA NEUTRALIZADA EM 0,00%", 0.0),
    ),
)
def test_percentual_e_situacao_chegam_aos_dados_dos_documentos(
    situacao, percentual
):
    leitura = _leitura_com_tratamento(situacao, percentual)
    dados = _extrair_dados(leitura, None)
    ciclo = next(c for c in dados["ciclos"] if c["ciclo"] == "C1")
    assert ciclo["percentual_reajuste"] == percentual
    assert ciclo["situacao"] == situacao

    saneador = _texto_docx(gerar_despacho_saneador(leitura))
    termo = _texto_docx(gerar_termo_apostila(leitura))
    percentual_txt = "-2,00%" if percentual < 0 else "0,00%"
    for texto in (saneador, termo):
        assert percentual_txt in texto
        assert situacao in texto
        assert "Manual de Licitações & Contratos" not in texto


def test_sumario_sem_negativo_nao_recebe_observacao():
    leitura = leitura_simples_financeiro()
    dados = montar_dados_sumario_executivo(leitura)
    assert dados["observacoes"]["reajuste_negativo"] == []
    texto = _texto_pdf(gerar_sumario_executivo(leitura))
    assert "Observação sobre o reajuste negativo" not in texto
    assert "Manual de Licitações & Contratos" not in texto


def test_sumario_aplicado_usa_texto_prudente_do_tcu():
    observacoes = _observacoes_reajuste_negativo([{
        "ciclo": "C1",
        "situacao": "TEMPESTIVO — VARIAÇÃO NEGATIVA",
    }])
    texto = " ".join(observacoes)
    assert "Manual de Licitações & Contratos" in texto
    assert "referência interpretativa" in texto
    assert "foi aplicada conforme a opção registrada" in texto
    assert "TCU determinou" not in texto

    pdf = _texto_pdf(gerar_sumario_executivo(_leitura_com_tratamento(
        "TEMPESTIVO — VARIAÇÃO NEGATIVA", -0.02
    )))
    assert "Manual de Licitações & Contratos" in pdf
    assert "foi aplicada conforme a opção registrada" in pdf


def test_sumario_neutralizado_e_factual_sem_atribuir_decisao_ao_tcu():
    observacoes = _observacoes_reajuste_negativo([{
        "ciclo": "C1",
        "situacao": "TEMPESTIVO — VARIAÇÃO NEGATIVA NEUTRALIZADA EM 0,00%",
    }])
    texto = " ".join(observacoes)
    assert texto.startswith("Observação sobre o reajuste negativo:")
    assert "neutralizada para 0,00%" in texto
    assert "TCU" not in texto

    pdf = _texto_pdf(gerar_sumario_executivo(_leitura_com_tratamento(
        "TEMPESTIVO — VARIAÇÃO NEGATIVA NEUTRALIZADA EM 0,00%", 0.0
    )))
    assert "neutralizada para 0,00%" in pdf
    assert "Manual de Licitações & Contratos" not in pdf


def test_sumario_multiciclo_identifica_cada_tratamento_sem_repetir_tcu():
    observacoes = _observacoes_reajuste_negativo([
        {"ciclo": "C1", "situacao": "TEMPESTIVO — VARIAÇÃO NEGATIVA"},
        {
            "ciclo": "C2",
            "situacao": "ADIANTADO — VARIAÇÃO NEGATIVA NEUTRALIZADA EM 0,00%",
        },
    ])
    texto = " ".join(observacoes)
    assert "C1:" in texto and "C2:" in texto
    assert texto.count("Manual de Licitações & Contratos") == 1


def test_email_contratada_nao_recebe_fundamentacao_do_tcu():
    ciclo = _ciclo_payload(-0.02, "TEMPESTIVO — VARIAÇÃO NEGATIVA")
    _assunto, corpo = gerar_rascunho_email_contratada(
        [ciclo], indice="IST", fator_acumulado=0.98
    )
    assert "Manual de Licitações & Contratos" not in corpo
    assert "TCU" not in corpo


def test_precedencias_permanecem_antes_da_nova_escolha():
    simples = PAGINA_SIMPLES.read_text(encoding="utf-8")
    multipla = PAGINA_MULTIPLA.read_text(encoding="utf-8")
    assert simples.index("if sem_pedido:", simples.index("percentual_indice =")) < simples.index(
        "elif ciclo_negativo:", simples.index("percentual_indice =")
    )
    assert simples.index('elif "PRECLUSO" in status_ped.upper():') < simples.index(
        "elif ciclo_negativo:", simples.index("percentual_indice =")
    )
    assert multipla.index(
        'if situacao_limpa == "PRECLUSO" and superacao_negocial:'
    ) < multipla.index("elif ciclo_negativo:", multipla.index("percentual_indice ="))
    assert multipla.index(
        'elif situacao_limpa == "PRECLUSO":', multipla.index("percentual_indice =")
    ) < multipla.index("elif ciclo_negativo:", multipla.index("percentual_indice ="))
