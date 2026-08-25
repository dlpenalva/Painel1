"""Ciclo em que a CONTRATADA nao apresentou pedido.

Ate aqui a unica forma de registrar esse fato era digitar uma data ficticia,
o que produzia documento falso ("a CONTRATADA apresentou pedido em ..."). O
estado passa a ser declarado no proprio rotulo de situacao, com
parametros!DATA_PEDIDO (U) vazia — sem coluna nova, sem template alterado e
sem data inventada.

Triestado protegido por estes testes:
  1. U preenchida .................. houve pedido, com data real;
  2. U vazia + marcador em G ....... nao houve pedido;
  3. U vazia sem marcador .......... nao informado / arquivo legado.
"""
from __future__ import annotations

import ast
import re
import sys
from datetime import date
from io import BytesIO
from pathlib import Path

import openpyxl
import pytest
from docx import Document
from openpyxl import Workbook

RAIZ = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(RAIZ))

from _gerador_masterfile import _registrar_datas_pedido  # noqa: E402
from _leitor_masterfile_v10 import _ler_parametros_v10  # noqa: E402
from _reajuste_utils import (  # noqa: E402
    MARCADOR_SEM_PEDIDO,
    SITUACAO_SEM_PEDIDO,
    classificar_pedido_por_data_exata,
    tem_sem_pedido,
)
from _sumario_executivo import (  # noqa: E402
    NAO_HOUVE_PEDIDO,
    NAO_INFORMADO,
    montar_dados_sumario_executivo,
)
from _templates_documentos import (  # noqa: E402
    gerar_despacho_saneador,
    gerar_termo_apostila,
)

from test_sumario_executivo import (  # noqa: E402
    leitura_multiciclo_pc,
    leitura_simples_financeiro,
)
from test_templates_documentos import (  # noqa: E402
    CAMPOS_SANEADOR,
    CAMPOS_TERMO,
)

PAGINA_SIMPLES = RAIZ / "pages" / "01_Calculo_Simples.py"
PAGINA_MULTIPLA = RAIZ / "pages" / "02_Calculo_Represados.py"

CABECALHOS_PARAMETROS = (
    "COMPUTAR_NESTA_APURACAO", "CICLO", "DATA_INICIO", "DATA_FIM",
    "PERCENTUAL_DO_CICLO", "FATOR_ACUMULADO", "SITUACAO",
)

PRECLUSO = "❌ PRECLUSO"
TEMPESTIVO = "✅ TEMPESTIVO"


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def _texto(docx_bytes: bytes) -> str:
    doc = Document(BytesIO(docx_bytes))
    partes = [p.text for p in doc.paragraphs]
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                partes.append(celula.text)
    return "\n".join(partes)


def _quadro1(docx_bytes: bytes):
    return Document(BytesIO(docx_bytes)).tables[0]


def _com_situacao(leitura: dict, situacoes: dict) -> dict:
    por_ciclo = leitura["parametros_v10"]["por_ciclo"]
    for nome, valor in situacoes.items():
        por_ciclo[nome]["situacao"] = valor
    return leitura


def _com_datas_pedido(leitura: dict, datas: dict) -> dict:
    por_ciclo = leitura["parametros_v10"]["por_ciclo"]
    for nome, valor in datas.items():
        por_ciclo[nome]["data_pedido"] = valor
    return leitura


def _planilha_parametros(situacoes: dict) -> Workbook:
    wb = Workbook()
    ws = wb.active
    ws.title = "parametros"
    for coluna, titulo in enumerate(CABECALHOS_PARAMETROS, start=1):
        ws.cell(1, coluna).value = titulo
    for linha, nome in enumerate(("C0", "C1", "C2", "C3", "C4"), start=2):
        ws.cell(linha, 2).value = nome
        if nome in situacoes:
            ws.cell(linha, 7).value = situacoes[nome]
    return wb


def _abrir_pagina(caminho, timeout=300):
    from streamlit.testing.v1 import AppTest

    at = AppTest.from_file(str(caminho), default_timeout=timeout)
    at.run()
    assert not at.exception
    return at


def _preparar_multiplos(at, data_lateral, ciclo_final="C2"):
    for campo in at.date_input:
        if "Data-base de referência" in str(campo.label):
            campo.set_value(data_lateral)
            break
    at.run()
    at.selectbox(key="rep_ciclo_final_analise").select(ciclo_final)
    at.run()
    assert not at.exception
    return at


def _relatorio(at) -> str:
    return next(i.value for i in at.info if "Resultado:" in i.value)


def _tem_acordo_negocial(at) -> bool:
    return any("Acordo negocial" in str(e.label) for e in at.expander)


# ---------------------------------------------------------------------------
# 1-3  regra semantica e vocabulario do classificador
# ---------------------------------------------------------------------------

def test_marcador_e_situacao_sao_definidos_em_um_unico_lugar():
    assert MARCADOR_SEM_PEDIDO == "SEM PEDIDO NESTE CICLO"
    # O marcador NAO pode conter palavra neutralizada no caminho documental.
    from _objeto_processo_reajuste import _termo_neutro_objeto
    assert _termo_neutro_objeto(SITUACAO_SEM_PEDIDO) == SITUACAO_SEM_PEDIDO
    assert SITUACAO_SEM_PEDIDO == PRECLUSO + " | " + MARCADOR_SEM_PEDIDO
    # PRECLUSO permanece como PREFIXO: ha consumidor documental que decide a
    # tempestividade por startswith("TEMPESTIVO").
    assert SITUACAO_SEM_PEDIDO.split("|")[0].strip().endswith("PRECLUSO")
    # Nenhum caractere de injecao de formula na abertura da celula do XLS.
    assert not SITUACAO_SEM_PEDIDO.startswith(("=", "+", "-", "@"))


@pytest.mark.parametrize(
    ("valor", "esperado"),
    [
        (SITUACAO_SEM_PEDIDO, True),
        ("❌ precluso | sem pedido neste ciclo", True),
        ("Sem Pedido Neste Ciclo", True),
        ("❌ PRECLUSO | SEM PEDIDO DA CONTRATADA", False),
        (PRECLUSO, False),
        (TEMPESTIVO, False),
        ("⚠️ ADIANTADO", False),
        ("", False),
        (None, False),
    ],
)
def test_deteccao_do_marcador_normaliza_emoji_e_caixa(valor, esperado):
    assert tem_sem_pedido(valor) is esperado


def test_vocabulario_do_classificador_permanece_fechado():
    referencia, limite = date(2025, 5, 5), date(2025, 8, 3)
    assert classificar_pedido_por_data_exata(
        date(2025, 5, 1), referencia, limite) == "ADIANTADO"
    assert classificar_pedido_por_data_exata(
        date(2025, 6, 1), referencia, limite) == "TEMPESTIVO"
    assert classificar_pedido_por_data_exata(
        date(2025, 9, 1), referencia, limite) == "PRECLUSO"


def test_classificador_nao_aceita_none_por_contrato():
    """Justifica o desvio nas paginas: sem data nao ha classificacao possivel."""
    with pytest.raises(TypeError):
        classificar_pedido_por_data_exata(None, date(2025, 5, 5), date(2025, 8, 3))


@pytest.mark.parametrize("pagina", [PAGINA_SIMPLES, PAGINA_MULTIPLA])
def test_paginas_desviam_do_classificador_antes_de_chamar(pagina):
    """A chamada tem de estar sob o ramo `else` do teste de ausencia de pedido."""
    arvore = ast.parse(pagina.read_text(encoding="utf-8"))
    chamadas = [
        no for no in ast.walk(arvore)
        if isinstance(no, ast.Call)
        and isinstance(no.func, ast.Name)
        and no.func.id == "classificar_pedido_por_data_exata"
    ]
    assert len(chamadas) == 1
    linha_chamada = chamadas[0].lineno
    protegida = False
    for no in ast.walk(arvore):
        if not (isinstance(no, ast.If)
                and isinstance(no.test, ast.Name)
                and no.test.id == "sem_pedido"):
            continue
        for corpo in no.orelse:
            if any(getattr(filho, "lineno", -1) == linha_chamada
                   for filho in ast.walk(corpo)):
                protegida = True
    assert protegida, "chamada ao classificador nao protegida por if sem_pedido"


# ---------------------------------------------------------------------------
# 4-8  calculo simples
# ---------------------------------------------------------------------------

def test_simples_sem_pedido_desabilita_a_data_e_nao_quebra():
    at = _abrir_pagina(PAGINA_SIMPLES)
    assert at.checkbox(key="sem_pedido_contratada_simples").value is False
    assert at.date_input[1].disabled is False

    at.checkbox(key="sem_pedido_contratada_simples").set_value(True).run()
    assert not at.exception
    assert at.date_input[1].disabled is True


def test_simples_sem_pedido_fica_precluso_sem_efeitos_e_sem_data():
    at = _abrir_pagina(PAGINA_SIMPLES)
    at.checkbox(key="sem_pedido_contratada_simples").set_value(True).run()
    at.button[0].click().run()
    assert not at.exception

    relatorio = _relatorio(at)
    assert "Resultado: " + SITUACAO_SEM_PEDIDO + "." in relatorio
    assert "Efeitos financeiros: não aplicáveis." in relatorio
    # nao afirma pedido inexistente e nao exibe data alguma de pedido
    assert "Não houve pedido da contratada neste ciclo." in relatorio
    assert "Pedido realizado em" not in relatorio
    assert "09/04/2024" not in relatorio


def test_simples_sem_pedido_nao_oferece_acordo_negocial():
    """Nao ha pedido precluso a admitir: o expander de acordo nao aparece."""
    at = _abrir_pagina(PAGINA_SIMPLES)
    at.checkbox(key="sem_pedido_contratada_simples").set_value(True).run()
    at.button[0].click().run()
    assert not _tem_acordo_negocial(at)


def test_simples_com_pedido_precluso_continua_oferecendo_acordo_negocial():
    """Regressao: o caminho atual do PRECLUSO com pedido permanece intacto."""
    at = _abrir_pagina(PAGINA_SIMPLES)
    at.date_input[1].set_value(date(2025, 6, 30)).run()
    at.button[0].click().run()
    assert not at.exception

    relatorio = _relatorio(at)
    assert "Pedido realizado em 30/06/2025." in relatorio
    assert "Resultado: " + PRECLUSO + "." in relatorio
    assert _tem_acordo_negocial(at)


def test_simples_desmarcar_restaura_o_fluxo_normal():
    at = _abrir_pagina(PAGINA_SIMPLES)
    at.checkbox(key="sem_pedido_contratada_simples").set_value(True).run()
    at.button[0].click().run()
    assert SITUACAO_SEM_PEDIDO in _relatorio(at)

    at.checkbox(key="sem_pedido_contratada_simples").set_value(False).run()
    assert at.date_input[1].disabled is False
    at.button[0].click().run()
    assert not at.exception

    relatorio = _relatorio(at)
    assert MARCADOR_SEM_PEDIDO not in relatorio
    assert "Pedido realizado em 09/04/2024." in relatorio
    assert "Início dos efeitos financeiros:" in relatorio


# ---------------------------------------------------------------------------
# 9-14  multiciclo
# ---------------------------------------------------------------------------

def test_multiciclo_tem_um_checkbox_por_ciclo_com_chave_propria():
    at = _preparar_multiplos(_abrir_pagina(PAGINA_MULTIPLA), date(2022, 10, 10))
    chaves = [c.key for c in at.checkbox if str(c.key or "").startswith("sem_pedido_")]
    assert chaves == ["sem_pedido_p1_20231010", "sem_pedido_p2_20241010"]


def test_multiciclo_marcar_um_ciclo_nao_contamina_o_outro():
    """C1 sem pedido, C2 com pedido real: cada ciclo mantem o seu estado."""
    at = _preparar_multiplos(_abrir_pagina(PAGINA_MULTIPLA), date(2022, 10, 10))
    at.checkbox(key="sem_pedido_p1_20231010").set_value(True).run()
    assert not at.exception
    at.date_input(key="p2_20241010").set_value(date(2024, 11, 5)).run()
    assert not at.exception

    assert at.date_input(key="p1_20231010").disabled is True
    assert at.date_input(key="p2_20241010").disabled is False

    resumo = at.dataframe[0].value.to_dict("records")
    assert resumo[0]["Ciclo"] == "C1"
    assert resumo[0]["Data do pedido"] == ""
    assert resumo[0]["Situação preliminar"] == PRECLUSO
    assert resumo[0]["Início financeiro"] == "Sem efeitos financeiros automáticos"

    assert resumo[1]["Ciclo"] == "C2"
    assert resumo[1]["Data do pedido"] == "05/11/2024"
    assert resumo[1]["Início financeiro"] == "11/2024"
    assert "PRECLUSO" not in resumo[1]["Situação preliminar"]


def test_multiciclo_sem_pedido_segue_a_cadeia_teorica_e_nao_inventa_data():
    """C1 sem pedido: o C2 nasce da referencia teorica, nunca de data criada."""
    at = _preparar_multiplos(_abrir_pagina(PAGINA_MULTIPLA), date(2022, 10, 10))
    at.checkbox(key="sem_pedido_p1_20231010").set_value(True).run()
    assert not at.exception

    resumo = at.dataframe[0].value.to_dict("records")
    # C1 apto em 10/10/2023; sem efeito financeiro, o C2 nasce 12 meses depois
    # da propria referencia do C1 — sem hoje, sem data-base artificial e sem
    # herdar a data do ciclo seguinte.
    assert resumo[0]["Referência exata"] == "10/10/2023"
    assert resumo[1]["Referência exata"] == "10/10/2024"
    assert at.date_input(key="p2_20241010").value == date(2024, 10, 10)


def test_multiciclo_sem_pedido_nao_oferece_acordo_negocial():
    at = _preparar_multiplos(_abrir_pagina(PAGINA_MULTIPLA), date(2022, 10, 10))
    at.checkbox(key="sem_pedido_p1_20231010").set_value(True).run()
    assert not _tem_acordo_negocial(at)


def test_multiciclo_precluso_com_pedido_mantem_o_acordo_negocial():
    """Regressao: preclusao por data continua permitindo admissao negocial."""
    at = _preparar_multiplos(_abrir_pagina(PAGINA_MULTIPLA), date(2022, 10, 10))
    at.date_input(key="p1_20231010").set_value(date(2024, 3, 1)).run()
    assert not at.exception
    assert _tem_acordo_negocial(at)


def test_multiciclo_desmarcar_restaura_o_ciclo():
    at = _preparar_multiplos(_abrir_pagina(PAGINA_MULTIPLA), date(2022, 10, 10))
    at.checkbox(key="sem_pedido_p1_20231010").set_value(True).run()
    assert at.dataframe[0].value.to_dict("records")[0]["Data do pedido"] == ""

    at.checkbox(key="sem_pedido_p1_20231010").set_value(False).run()
    assert not at.exception
    assert at.date_input(key="p1_20231010").disabled is False

    resumo = at.dataframe[0].value.to_dict("records")
    assert resumo[0]["Data do pedido"] == "10/10/2023"
    assert resumo[0]["Início financeiro"] == "10/2023"


# ---------------------------------------------------------------------------
# 15-17  persistencia no XLS e triestado na releitura
# ---------------------------------------------------------------------------

def test_estado_2_marcador_em_g_com_u_vazia():
    wb = _planilha_parametros({"C1": SITUACAO_SEM_PEDIDO})
    _registrar_datas_pedido(wb, {"C1": {"data_pedido": None}})
    ws = wb["parametros"]

    assert ws["G3"].value == SITUACAO_SEM_PEDIDO
    assert ws["U1"].value == "DATA_PEDIDO"
    assert ws["U3"].value is None

    lido = _ler_parametros_v10(wb)["por_ciclo"]["C1"]
    assert lido["data_pedido"] is None
    assert tem_sem_pedido(lido["situacao"]) is True


def test_estado_1_pedido_real_permanece_intacto():
    wb = _planilha_parametros({"C1": TEMPESTIVO})
    _registrar_datas_pedido(wb, {"C1": {"data_pedido": date(2025, 3, 20)}})

    lido = _ler_parametros_v10(wb)["por_ciclo"]["C1"]
    assert lido["data_pedido"] == date(2025, 3, 20)
    assert tem_sem_pedido(lido["situacao"]) is False


@pytest.mark.parametrize("situacao", ["", PRECLUSO, "Fora da apuracao"])
def test_estado_3_ausencia_sem_marcador_nao_vira_sem_pedido(situacao):
    """U vazia, sozinha, NUNCA significa que nao houve pedido."""
    wb = _planilha_parametros({"C1": situacao})
    _registrar_datas_pedido(wb, {})

    lido = _ler_parametros_v10(wb)["por_ciclo"]["C1"]
    assert lido["data_pedido"] is None
    assert tem_sem_pedido(lido["situacao"]) is False


def test_arquivo_legado_sem_a_coluna_u_nao_vira_sem_pedido():
    wb = Workbook()
    ws = wb.active
    ws.title = "parametros"
    ws["A1"], ws["B1"], ws["G1"] = "COMPUTAR_NESTA_APURACAO", "CICLO", "SITUACAO"
    ws["B2"], ws["G2"] = "C1", PRECLUSO

    lido = _ler_parametros_v10(wb)["por_ciclo"]["C1"]
    assert lido["data_pedido"] is None
    assert tem_sem_pedido(lido["situacao"]) is False


def test_round_trip_salvar_e_reabrir_preserva_o_estado():
    """Gera, grava em disco, reabre: o marcador continua distinguivel."""
    wb = _planilha_parametros({"C1": SITUACAO_SEM_PEDIDO, "C2": TEMPESTIVO})
    _registrar_datas_pedido(
        wb, {"C1": {"data_pedido": None}, "C2": {"data_pedido": date(2025, 7, 10)}}
    )
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    relido = _ler_parametros_v10(openpyxl.load_workbook(buffer))["por_ciclo"]
    assert tem_sem_pedido(relido["C1"]["situacao"]) is True
    assert relido["C1"]["data_pedido"] is None
    assert tem_sem_pedido(relido["C2"]["situacao"]) is False
    assert relido["C2"]["data_pedido"] == date(2025, 7, 10)


# ---------------------------------------------------------------------------
# 18-20  documentos
# ---------------------------------------------------------------------------

def test_sumario_mostra_nao_houve_pedido_e_nao_nao_informado():
    leitura = _com_situacao(leitura_simples_financeiro(), {"C1": SITUACAO_SEM_PEDIDO})
    dados = montar_dados_sumario_executivo(leitura)
    c1 = next(c for c in dados["ciclos"] if c["ciclo"] == "C1")

    assert c1["data_pedido"] == NAO_HOUVE_PEDIDO
    assert c1["data_pedido"] != NAO_INFORMADO
    assert MARCADOR_SEM_PEDIDO in c1["situacao"].upper()


def test_sumario_preserva_ausencia_de_informacao_no_legado():
    dados = montar_dados_sumario_executivo(leitura_simples_financeiro())
    c1 = next(c for c in dados["ciclos"] if c["ciclo"] == "C1")
    assert c1["data_pedido"] != NAO_HOUVE_PEDIDO


def test_saneador_nao_afirma_pedido_inexistente():
    leitura = _com_situacao(leitura_simples_financeiro(), {"C1": SITUACAO_SEM_PEDIDO})
    texto = _texto(gerar_despacho_saneador(leitura, campos_manuais=CAMPOS_SANEADOR))

    assert "A CONTRATADA apresentou pedido de" not in texto
    assert "A solicitação da CONTRATADA, de" not in texto
    assert "A solicitação tempestiva da CONTRATADA" not in texto
    assert "Não houve pedido da CONTRATADA para o ciclo C1" in texto
    assert "permanece precluso, sem efeitos financeiros" in texto
    assert "[PREENCHER: Data do pedido]" not in texto


def test_saneador_quadro1_declara_a_ausencia_de_pedido():
    leitura = _com_situacao(leitura_simples_financeiro(), {"C1": SITUACAO_SEM_PEDIDO})
    tabela = _quadro1(gerar_despacho_saneador(leitura, campos_manuais=CAMPOS_SANEADOR))
    linha = next(r for r in tabela.rows if r.cells[0].text == "C1")

    assert linha.cells[2].text == NAO_HOUVE_PEDIDO
    assert linha.cells[4].text == "Sem efeitos financeiros"


def test_saneador_caso_misto_preserva_o_ciclo_com_pedido():
    leitura = _com_datas_pedido(
        _com_situacao(leitura_multiciclo_pc(), {"C1": SITUACAO_SEM_PEDIDO}),
        {"C2": date(2025, 7, 10)},
    )
    docx = gerar_despacho_saneador(leitura, campos_manuais=CAMPOS_SANEADOR)
    datas = {r.cells[0].text: r.cells[2].text for r in _quadro1(docx).rows[1:]}

    assert datas["C1"] == NAO_HOUVE_PEDIDO
    assert datas["C2"] == "10/07/2025"
    # havendo pedido em ao menos um ciclo, a redacao ordinaria permanece
    assert "A CONTRATADA apresentou pedido de" in _texto(docx)


def test_termo_nao_afirma_solicitacao_inexistente():
    """Todos os ciclos computados sem pedido: o CONSIDERANDO declara a ausencia."""
    leitura = _com_situacao(
        leitura_multiciclo_pc(),
        {"C1": SITUACAO_SEM_PEDIDO, "C2": SITUACAO_SEM_PEDIDO},
    )
    texto = _texto(gerar_termo_apostila(leitura, campos_manuais=CAMPOS_TERMO))

    assert "A solicitação da CONTRATADA, de" not in texto
    assert "A solicitação tempestiva da CONTRATADA" not in texto
    assert "A inexistência de pedido da CONTRATADA para os ciclos analisados" in texto
    assert "permanecem preclusos, sem efeitos financeiros" in texto
    assert "[PREENCHER: Data da solicitacao da contratada]" not in texto


def test_termo_caso_misto_preserva_a_redacao_ordinaria():
    """Havendo pedido em ao menos um ciclo, a solicitacao existiu e e afirmada.

    A data canonica do CONSIDERANDO vem apenas do ciclo que teve pedido: o
    ciclo sem pedido nao contamina o conjunto com rotulo algum.
    """
    leitura = _com_datas_pedido(
        _com_situacao(leitura_multiciclo_pc(), {"C1": SITUACAO_SEM_PEDIDO}),
        {"C2": date(2025, 7, 10)},
    )
    texto = _texto(gerar_termo_apostila(leitura, campos_manuais=CAMPOS_TERMO))

    assert "A solicitação da CONTRATADA, de 10/07/2025" in texto
    assert "A inexistência de pedido da CONTRATADA" not in texto
    assert NAO_HOUVE_PEDIDO not in texto.split("instruída em")[0]


def test_documentos_nao_produzem_data_alguma_para_ciclo_sem_pedido():
    """Nenhuma data ficticia (hoje, 01/01/1900, data-base) substitui o pedido."""
    leitura = _com_situacao(leitura_simples_financeiro(), {"C1": SITUACAO_SEM_PEDIDO})
    tabela = _quadro1(gerar_despacho_saneador(leitura, campos_manuais=CAMPOS_SANEADOR))
    celula_pedido = next(
        r.cells[2].text for r in tabela.rows if r.cells[0].text == "C1"
    )

    assert not re.search(r"\d{2}/\d{2}/\d{4}", celula_pedido)
    assert celula_pedido == NAO_HOUVE_PEDIDO


# ---------------------------------------------------------------------------
# 21  regressao: Relatorio de Apuracao do multiciclo (POS-processamento)
# ---------------------------------------------------------------------------

def _relatorio_multiciclo_processado(at) -> str:
    """Texto renderizado em "Relatorio de Apuracao" DEPOIS de processar.

    Le a saida pos-processamento — nao o "Resumo antes de processar", nao o
    XLS, nao o payload e nao o codigo-fonte. Foi exatamente essa superficie
    que chegou a producao emitindo "Pedido realizado em ." para ciclo sem
    pedido, justamente porque nenhuma das outras a cobria.
    """
    for botao in at.button:
        if "Processar Análise" in str(botao.label):
            botao.click()
            break
    at.run()
    assert not at.exception
    for bloco in at.info:
        if "C1:" in bloco.value and "Resultado:" in bloco.value:
            return bloco.value
    pytest.skip("relatorio indisponivel (indice sem cobertura no periodo)")


def test_relatorio_multiciclo_nao_afirma_pedido_inexistente():
    """C1 sem pedido + C2 com pedido real: cada ciclo com a sua redacao."""
    at = _preparar_multiplos(
        _abrir_pagina(PAGINA_MULTIPLA, timeout=600), date(2022, 10, 10)
    )
    at.checkbox(key="sem_pedido_p1_20231010").set_value(True).run()
    assert not at.exception

    relatorio = _relatorio_multiciclo_processado(at)

    # C1 — sem pedido: declara a ausencia, nunca afirma pedido inexistente
    assert "**C1:** Não houve pedido da contratada neste ciclo." in relatorio
    assert "Resultado: " + SITUACAO_SEM_PEDIDO + "." in relatorio
    assert "Efeitos financeiros: não aplicáveis." in relatorio

    # C2 — pedido real permanece intacto
    assert "**C2:** Pedido realizado em 10/10/2024." in relatorio
    assert "Início dos efeitos financeiros: 10/2024." in relatorio

    # o defeito exato que chegou a producao
    assert "Pedido realizado em ." not in relatorio


def test_relatorio_multiciclo_preserva_ciclos_com_pedido_real():
    """Sem nenhum ciclo marcado, a redacao ordinaria permanece integral."""
    at = _preparar_multiplos(
        _abrir_pagina(PAGINA_MULTIPLA, timeout=600), date(2022, 10, 10)
    )
    relatorio = _relatorio_multiciclo_processado(at)

    assert "**C1:** Pedido realizado em 10/10/2023." in relatorio
    assert "**C2:** Pedido realizado em 10/10/2024." in relatorio
    assert MARCADOR_SEM_PEDIDO not in relatorio
    assert "Pedido realizado em ." not in relatorio
