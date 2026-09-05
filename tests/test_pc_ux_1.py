"""PC-UX-1: contrato de apresentação sem alteração da metodologia PC."""
from __future__ import annotations

from contextlib import contextmanager
from copy import deepcopy
from datetime import date, datetime
from io import BytesIO
import gc
import os
from pathlib import Path
import shutil

import pytest
from docx import Document
from openpyxl import load_workbook

from _coleta_oficial import gerar_coleta_oficial_preenchida
from _leitor_masterfile_v10 import _totais_canonicos_pc
from _templates_documentos import (
    _extrair_dados,
    _situacao_retroativos_pc,
    _vta_texto_doc,
    gerar_despacho_saneador,
    gerar_termo_apostila,
)
from test_sumario_executivo import (
    leitura_multiciclo_pc,
    leitura_simples_financeiro,
)

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "templates" / "COLETA_REAJUSTE_OFICIAL.xlsx"
CORTE = date(2026, 8, 31)


def _pc(
    numero: str,
    ciclo: str,
    valor: float,
    *,
    pago: bool,
    efeito: bool,
    retro: float = 0.0,
    analise: float = 0.0,
    potencial: float = 0.0,
    dentro: bool = True,
    enquadramento: str = "CICLO",
) -> dict:
    return {
        "numero_pc": numero,
        "ciclo": ciclo,
        "valor_pc": valor,
        "valor_historico_considerado": round(valor + retro, 2),
        "retroativo_reconhecido_a_pagar": retro,
        "valor_atualizado_em_analise": analise,
        "delta_potencial": potencial,
        "pc_pago_a_contratada": "Sim" if pago else "Nao",
        "efeito_financeiro_pc": "Sim" if efeito else "Nao",
        "dentro_do_corte": dentro,
        "enquadramento": enquadramento,
        "entra_no_calculo": "Sim",
        "campos_vta": {"status_consolidacao": "COMPUTADO"},
    }


def _casos_basicos() -> list[dict]:
    return [
        _pc("C0-PAGO", "C0", 100.0, pago=True, efeito=False),
        _pc("C0-ANALISE", "C0", 50.0, pago=False, efeito=False, analise=50.0),
        _pc("C1-PAGO-EFEITO", "C1", 200.0, pago=True, efeito=True, retro=20.0),
        _pc(
            "C1-ANALISE-EFEITO", "C1", 300.0, pago=False, efeito=True,
            analise=330.0, potencial=30.0,
        ),
        _pc("C1-PAGO-SEM-EFEITO", "C1", 80.0, pago=True, efeito=False),
        _pc("C1-FORA", "C1", 70.0, pago=True, efeito=True, retro=7.0, dentro=False),
    ]


def _caso_sintetico_consolidacao() -> list[dict]:
    # Valores intermediarios ja calculados exercitam apenas a consolidacao e a
    # apresentacao. Nao representam prova do calculo originario do XLS-fonte.
    return [
        _pc("C0-PAGO", "C0", 13_559_396.46, pago=True, efeito=False),
        _pc(
            "C0-ANALISE", "C0", 242_807.91, pago=False, efeito=False,
            analise=242_807.91,
        ),
        _pc(
            "C1-PAGO-EFEITO", "C1", 1_697_102.65, pago=True, efeito=True,
            retro=57_701.49,
        ),
        _pc("C1-PAGO-SEM-EFEITO", "C1", 177_217.78, pago=True, efeito=False),
        _pc(
            "C1-ANALISE-EFEITO", "C1", 3_529_897.65, pago=False, efeito=True,
            analise=3_649_914.17, potencial=120_016.52,
        ),
        _pc(
            "C1-ANALISE-SEM-EFEITO", "C1", 377_476.28, pago=False,
            efeito=False, analise=377_476.28,
        ),
    ]


def _texto_docx(conteudo: bytes) -> str:
    doc = Document(BytesIO(conteudo))
    partes = [p.text for p in doc.paragraphs]
    for tabela in doc.tables:
        for linha in tabela.rows:
            partes.extend(celula.text for celula in linha.cells)
    return "\n".join(partes)


def _leitura_documental(itens: list[dict]) -> dict:
    leitura = deepcopy(leitura_multiciclo_pc())
    leitura["controle"]["data_corte"] = CORTE
    leitura["itens_pc_v10"]["itens"] = deepcopy(itens)
    leitura["itens_pc_v10"]["totais_canonicos"] = _totais_canonicos_pc(
        itens, CORTE
    )
    return leitura


def test_a_f_regras_pc_permanecem_separadas_por_estado_e_corte():
    totais = _totais_canonicos_pc(_casos_basicos(), CORTE)
    c0, c1 = totais["por_ciclo"]["C0"], totais["por_ciclo"]["C1"]
    assert c0["quantidade_reconhecida"] == 1
    assert c0["valor_original_reconhecido"] == 100.0
    assert c0["valor_atualizado_reconhecido"] == 100.0
    assert c0["retroativo"] == 0.0
    assert c0["valor_atualizado_em_analise"] == 50.0
    assert c0["delta_potencial"] == 0.0
    assert c1["quantidade_reconhecida"] == 2
    assert c1["valor_original_reconhecido"] == 280.0
    assert c1["valor_atualizado_reconhecido"] == 300.0
    assert c1["retroativo"] == 20.0
    assert c1["valor_atualizado_em_analise"] == 330.0
    assert c1["delta_potencial"] == 30.0
    assert totais["posterior_ao_corte"]["valor_pc"] == 70.0
    assert totais["posterior_ao_corte_por_ciclo"]["C1"]["valor_pc"] == 70.0


def test_b_fora_do_corte_inclui_pc_fora_dos_ciclos_e_fecha_o_total():
    fora = _pc(
        "FORA-999999", "Fora dos ciclos", 999_999.99,
        pago=True, efeito=False, dentro=False,
    )
    totais = _totais_canonicos_pc([fora], CORTE)
    assert totais["posterior_ao_corte"]["valor_pc"] == 999_999.99
    assert totais["informado"]["valor_pc"] - totais["ate_o_corte"]["valor_pc"] == 999_999.99


def test_c_fora_do_corte_soma_c1_posterior_e_fora_dos_ciclos():
    itens = [
        _pc("C1-FORA", "C1", 100.01, pago=True, efeito=True, dentro=False),
        _pc("SEM-CICLO-FORA", "Fora dos ciclos", 999_999.99,
            pago=True, efeito=False, dentro=False),
    ]
    totais = _totais_canonicos_pc(itens, CORTE)
    assert totais["posterior_ao_corte_por_ciclo"]["C1"]["valor_pc"] == 100.01
    assert totais["posterior_ao_corte"]["valor_pc"] == 1_000_100.00


def test_g_ausencia_nao_vira_zero_no_payload_documental():
    ausente = {"controle": {"modo": "PC"}, "totais_canonicos_pc": {
        "ate_o_corte": {"quantidade": 1, "retroativo": None},
    }}
    zero = {"controle": {"modo": "PC"}, "totais_canonicos_pc": {
        "ate_o_corte": {"quantidade": 1, "retroativo": 0.0},
    }}
    assert _situacao_retroativos_pc(ausente)["reconhecido"] is None
    assert _situacao_retroativos_pc(zero)["reconhecido"] == 0.0


def test_h_quadro_itens_pc_le_exclusivamente_colunas_canonicas():
    ws = load_workbook(TEMPLATE, data_only=False)["itens_PC"]
    assert ws["M1"].value == "TODOS OS PCs CADASTRADOS POR CICLO"
    assert ws["M8"].value == "Outras situações / fora dos ciclos"
    assert ws.column_dimensions["M"].width >= 20
    assert ws["O8"].value == "=SUM($D$2:$D$5001)-SUM(O3:O7)"
    assert ws["P8"].value == "=SUM($F$2:$F$5001)-SUM(P3:P7)"
    assert ws["Q8"].value == "=SUM($H$2:$H$5001)-SUM(Q3:Q7)"
    assert ws["R8"].value == "=SUM($I$2:$I$5001)-SUM(R3:R7)"
    assert ws["S8"].value == "=SUM($J$2:$J$5001)-SUM(S3:S7)"
    # A-01: K e coluna de formula ate 5001 e devolve "" nas linhas sem PC. Sem
    # ancorar em A (coluna de entrada), o residual contava as 5.000 vazias.
    assert ws["T8"].value == (
        '=COUNTIFS($A$2:$A$5001,"<>",$K$2:$K$5001,"<>OK",'
        '$K$2:$K$5001,"<>")-SUM(T3:T7)'
    )
    assert "$A$2:$A$5001" in ws["T8"].value
    assert ws["T9"].value == "=SUM(T3:T8)"
    assert ws["M9"].value == "TOTAL"
    assert ws["O9"].value == "=SUM(O3:O8)"
    assert ws["M10"].value == "PCs CONSIDERADOS NA APURAÇÃO ATÉ A DATA DE CORTE"
    assert [ws.cell(11, c).value for c in range(13, 21)] == [
        "Ciclo", "PCs pagos/reconhecidos", "Valor original reconhecido",
        "Valor atualizado reconhecido", "Retroativo reconhecido",
        "Valor em análise — não pagos (regra vigente)",
        "Retroativo potencial", "Fora da data de corte",
    ]
    for linha in range(12, 17):
        assert "$D$2:$D$5001" in ws.cell(linha, 15).value
        assert "$U$2:$U$5001" in ws.cell(linha, 16).value
        assert "$H$2:$H$5001" in ws.cell(linha, 17).value
        assert "$I$2:$I$5001" in ws.cell(linha, 18).value
        assert "$J$2:$J$5001" in ws.cell(linha, 19).value
        assert 'MEMORIA_RESULTADOS!$T$31' in ws.cell(linha, 20).value
    assert ws["M17"].value == "Outras situações / fora dos ciclos"
    assert ws["M17"].alignment.wrap_text is True
    assert "SUMIFS($D$2:$D$5001" in ws["T17"].value
    assert ws["T18"].value == "=SUM(T12:T17)"


def test_pipeline_real_entrega_os_tres_blocos_pc_ux_sem_cabecalho_legado():
    payload = {
        "origem": "Reajustes Múltiplos",
        "indice": "IPCA",
        "data_base_original": "01/01/2025",
        "variacao_acumulada": 0.05,
        "fator_acumulado": 1.05,
        "ciclos": [{
            "ciclo": "C1",
            "data_inicio": "01/01/2026",
            "data_pedido": "15/03/2026",
            "percentual_aplicado": 0.05,
            "financeiro_inicio": "01/03/2026",
            "objeto_analise_atual": True,
            "situacao": "TEMPESTIVO",
        }],
    }
    wb = load_workbook(
        BytesIO(gerar_coleta_oficial_preenchida(payload)), data_only=False
    )
    ws = wb["itens_PC"]
    assert (ws["M1"].value, ws["M10"].value, ws["M20"].value) == (
        "TODOS OS PCs CADASTRADOS POR CICLO",
        "PCs CONSIDERADOS NA APURAÇÃO ATÉ A DATA DE CORTE",
        "COMO OS PCs SÃO TRATADOS",
    )
    assert [ws.cell(2, coluna).value for coluna in range(13, 21)] == [
        "CICLO", "QTD. DE PCs", "VALOR ORIGINAL TOTAL",
        "VALOR DOS PCs COM FATOR DO CICLO", "RETROATIVO RECONHECIDO",
        "VALOR EM ANÁLISE - NÃO PAGOS (ÁREA GEST.)", "RETROATIVO POTENCIAL",
        "QTD. COM ALERTA",
    ]
    assert [ws.cell(11, coluna).value for coluna in range(13, 21)] == [
        "Ciclo", "PCs pagos/reconhecidos", "Valor original reconhecido",
        "Valor atualizado reconhecido", "Retroativo reconhecido",
        "Valor em análise — não pagos (regra vigente)", "Retroativo potencial",
        "Fora da data de corte",
    ]
    assert [ws.cell(linha, 13).value for linha in range(3, 10)] == [
        "C0", "C1", "C2", "C3", "C4",
        "Outras situações / fora dos ciclos", "TOTAL",
    ]
    assert [ws.cell(linha, 13).value for linha in range(12, 19)] == [
        "C0", "C1", "C2", "C3", "C4",
        "Outras situações / fora dos ciclos", "TOTAL",
    ]
    assert ws["M17"].alignment.wrap_text is True
    assert ws.column_dimensions["M"].width >= 20
    # A-01: o XLS que o fiscal baixa nao pode anunciar alerta de linha vazia.
    assert ws["T8"].value == (
        '=COUNTIFS($A$2:$A$5001,"<>",$K$2:$K$5001,"<>OK",'
        '$K$2:$K$5001,"<>")-SUM(T3:T7)'
    )
    assert ws["T9"].value == "=SUM(T3:T8)"
    assert ws["U1"].value == "VALOR_CONSIDERADO"
    assert "M1:T1" in {str(faixa) for faixa in ws.merged_cells.ranges}
    assert ws["U1"].fill.fgColor.rgb != ws["M1"].fill.fgColor.rgb
    assert ws.column_dimensions["V"].hidden is True
    assert not {
        "QTD_PC", "VALOR_PC_TOTAL", "VALOR_ATUALIZADO_TOTAL",
        "VALOR_ATUALIZADO_EM_ANALISE", "QTD_COM_CHECK",
    } & {ws.cell(2, coluna).value for coluna in range(13, 21)}


def test_i_j_resultados_preserva_formulas_e_vta():
    ws = load_workbook(TEMPLATE, data_only=False)["RESULTADOS"]
    assert ws["A9"].value == "1. COMO O VTA FOI CALCULADO"
    assert ws["A15"].value == "2. EXECUÇÃO E RETROATIVO POR CICLO"
    assert ws["B22"].value == '=IF(COUNT(B16:B20)=0,"",ROUND(SUM(B16:B20),2))'
    assert ws["B38"].value == '=IF(OR(B36="",B37=""),"",ROUND(B37-B36,2))'
    assert ws["B86"].value == '=IF(VTA_FINAL="","",VTA_FINAL)'
    assert "Esta conferência não altera o VTA Oficial" in ws["A70"].value
    assert ws["A78"].value is None


def test_o_formulas_e_referencias_foram_reancoradas_sem_tocar_a_l():
    wb = load_workbook(TEMPLATE, data_only=False)
    ws, mem = wb["itens_PC"], wb["MEMORIA_RESULTADOS"]
    assert [ws.cell(1, c).value for c in range(1, 13)] == [
        "NUMERO_PC", "DATA_PC", "CICLO_PC", "VALOR_PC", "FATOR_ACUMULADO",
        "VALOR_ATUALIZADO", "PC_PAGO_A_CONTRATADA",
        "RETROATIVO_RECONHECIDO_A_PAGAR", "VALOR_ATUALIZADO_EM_ANALISE",
        "DELTA_POTENCIAL", "CHECK_PC_FINANCEIRO", "EFEITO_FINANCEIRO_PC",
    ]
    assert "itens_PC!$O$3" in mem["T21"].value
    assert "ROW(itens_PC!$O$3:$O$7)-3" in mem["T22"].value
    assert "itens_PC!$P$3" in mem["T25"].value
    assert "SUM(itens_PC!$O$3:$O$7)" in mem["T35"].value
    assert "ROW(itens_PC!$O$3:$O$7)-3" in mem["W67"].value
    assert wb.defined_names["VTA_FINAL"].value == "MEMORIA_RESULTADOS!$B$26"
    assert wb.defined_names["EXECUTADO_APURADO"].value == "RESULTADOS!$B$83"


def test_consolidacao_sintetica_do_benchmark_fecha_sem_hardcode():
    """Prova a consolidacao; homologacao do XLS-fonte original segue pendente."""
    totais = _totais_canonicos_pc(_caso_sintetico_consolidacao(), CORTE)
    c0, c1 = totais["por_ciclo"]["C0"], totais["por_ciclo"]["C1"]
    assert c0["valor_pc"] == 13_802_204.37
    assert c0["valor_atualizado_reconhecido"] == 13_559_396.46
    assert c0["valor_atualizado_em_analise"] == 242_807.91
    assert c0["retroativo"] == 0.0
    assert c1["valor_original_reconhecido"] == 1_874_320.43
    assert c1["valor_atualizado_reconhecido"] == 1_932_021.92
    assert c1["retroativo"] == 57_701.49
    assert c1["valor_atualizado_em_analise"] == 4_027_390.45
    assert c1["delta_potencial"] == 120_016.52
    reconhecido = round(sum(x["valor_atualizado_reconhecido"] for x in (c0, c1)), 2)
    analise = round(sum(x["valor_atualizado_em_analise"] for x in (c0, c1)), 2)
    assert reconhecido == 15_491_418.38
    assert analise == 4_270_198.36
    assert round(reconhecido + analise, 2) == 19_761_616.74


def test_k_l_q_r_apostila_e_saneador_usam_os_mesmos_valores_e_abrem():
    leitura = _leitura_documental(_caso_sintetico_consolidacao())
    termo = gerar_termo_apostila(leitura)
    saneador = gerar_despacho_saneador(leitura)
    for conteudo in (termo, saneador):
        Document(BytesIO(conteudo))
        texto = _texto_docx(conteudo)
        assert "R$ 57.701,49" in texto
        assert "R$ 120.016,52" in texto
        assert "valor em análise pela área gestora" in texto.lower()
        assert "31/08/2026" in texto
    texto_termo = _texto_docx(termo)
    assert texto_termo.count("SITUAÇÃO DOS VALORES RETROATIVOS") == 1
    assert "Quadro 2 — Execução considerada até a data de corte" in texto_termo
    assert "Valor original" in texto_termo
    assert "Valor atualizado" in texto_termo
    assert "Retroativo reconhecido" in texto_termo
    texto_saneador = _texto_docx(saneador)
    assert "PENDÊNCIA TÉCNICA" in texto_saneador
    assert "PROVIDÊNCIA DA ÁREA GESTORA" in texto_saneador


def test_saneador_pc_restaura_vta_canonico_sem_duplicar_e_outros_metodos_preservam():
    esperado = "R$ 12.345,67"
    for metodo, leitura in (
        ("PC", _leitura_documental(_caso_sintetico_consolidacao())),
        ("SIMPLES", leitura_simples_financeiro()),
        ("D", leitura_simples_financeiro()),
    ):
        leitura = deepcopy(leitura)
        leitura["controle"]["modo"] = metodo
        leitura["resultados_xls"] = {
            "disponivel": True,
            "valores": {"VTA_FINAL": 12_345.67, "RETRO_OFICIAL": 0.0},
        }
        canonico = _vta_texto_doc(_extrair_dados(leitura, None))
        assert canonico == f"{esperado} — PRÉVIA"
        texto = _texto_docx(gerar_despacho_saneador(leitura))
        assert texto.count("Valor Total Atualizado do Contrato") == 1
        assert texto.count(canonico) == 1


@pytest.mark.parametrize("enquadramento", ["INTERVALO_PRECLUSO", "INDETERMINADO"])
def test_documentos_exibem_residual_e_total_fecha(enquadramento):
    residual = _pc(
        "RESIDUAL", "", 25.0, pago=True, efeito=False,
        enquadramento=enquadramento,
    )
    leitura = _leitura_documental(_casos_basicos()[:3] + [residual])
    for gerador in (gerar_termo_apostila, gerar_despacho_saneador):
        doc = Document(BytesIO(gerador(leitura)))
        tabelas = [[c.text for c in linha.cells] for t in doc.tables for linha in t.rows]
        assert any(linha[0] == "Outras situações até a data de corte" for linha in tabelas)
        assert any(linha[0] == "Total" and "R$ 325,00" in linha for linha in tabelas)


def test_resultados_rotulos_sao_condicionais_aos_tres_metodos():
    ws = load_workbook(TEMPLATE, data_only=False)["RESULTADOS"]
    for celula, textos in {
        "B15": ("Financeiro", "Valor pago", "PCs", "Valor original", "Itens", "Valor consumido original"),
        "C15": ("Financeiro", "Valor devido atualizado", "PCs", "Valor atualizado", "Itens", "Valor consumido atualizado"),
    }.items():
        formula = str(ws[celula].value)
        assert formula.startswith("=IF(")
        for texto in textos:
            assert texto in formula
    assert ws["D15"].value == "Diferença"


# --------------------------------------------------------------- A-01 / A-02
# Datas derivadas da regra real de itens_PC!C (ancora = parametros!C2 = 01/2025,
# n = meses desde a ancora, "Fora dos ciclos" quando n > 59) e conferidas no
# Excel antes de virarem expectativa. Nao ha data chutada neste fixture:
#   01/05/2025 -> n=4  -> C0
#   01/03/2026 -> n=14 -> C1
#   01/03/2029 -> n=50 -> C4
#   01/06/2030 -> n=65 -> "Fora dos ciclos"
DATA_C0 = datetime(2025, 5, 1)
DATA_C1 = datetime(2026, 3, 1)
DATA_C1_ANALISE = datetime(2026, 4, 1)
DATA_C1_POS_CORTE = datetime(2026, 9, 15)
DATA_C4 = datetime(2029, 3, 1)
DATA_FORA_DOS_CICLOS = datetime(2030, 6, 1)


def _preparar_cenario_pc(wb):
    """Deixa CONTROLE/parametros no metodo PC com os cinco ciclos abertos."""
    controle = wb.Worksheets("CONTROLE")
    parametros = wb.Worksheets("parametros")
    itens = wb.Worksheets("itens_PC")
    for planilha in (controle, parametros, itens):
        if planilha.ProtectContents:
            planilha.Unprotect()
    controle.Range("B1").Value = "PC (Pedidos de Compra)"
    controle.Range("B2").Value = "C1"
    controle.Range("B3").Value = datetime(2026, 8, 31)
    parametros.Range("A2:A6").Value = (("Sim",),) * 5
    parametros.Range("C2:D2").Value = (
        (datetime(2025, 1, 1), datetime(2025, 12, 31)),
    )
    for linha, ano in enumerate(range(2026, 2030), start=3):
        parametros.Range(f"C{linha}:E{linha}").Value = (
            (datetime(ano, 1, 1), datetime(ano, 12, 31), 0.10),
        )
    # Somente C1 recebe INICIO_EFEITO; C4 fica sem, para gerar alerta real.
    parametros.Range("H3").Value = datetime(2026, 2, 1)
    return controle, parametros, itens


def _limpar_pcs(itens):
    """Zera apenas as colunas de ENTRADA; A:L de formula permanece intacta."""
    for faixa in ("A2:B5001", "D2:D5001", "G2:G5001"):
        itens.Range(faixa).ClearContents()


def _gravar_pc(itens, linha, numero, data, valor, pago):
    itens.Range(f"A{linha}").Value = numero
    itens.Range(f"B{linha}").Value = data
    itens.Range(f"D{linha}").Value = valor
    itens.Range(f"G{linha}").Value = pago


def _alertas(itens):
    """Quadro 1, coluna QTD. COM ALERTA: C0..C4, residual e TOTAL."""
    return [int(itens.Range(f"T{linha}").Value) for linha in range(3, 10)]


@contextmanager
def _excel_com(origem, destino):
    client = pytest.importorskip("win32com.client")
    pythoncom = pytest.importorskip("pythoncom")
    shutil.copy2(origem, destino)
    pythoncom.CoInitialize()
    excel = client.DispatchEx("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False
    wb = None
    try:
        wb = excel.Workbooks.Open(
            str(destino), UpdateLinks=0, ReadOnly=False, CorruptLoad=0
        )
        yield excel, wb
    finally:
        if wb is not None:
            try:
                wb.Close(False)
            except Exception:  # noqa: BLE001 - encerramento best-effort
                pass
        excel.Quit()
        del wb
        del excel
        gc.collect()
        pythoncom.CoUninitialize()


@pytest.mark.skipif(
    os.environ.get("RUN_EXCEL_INTEGRATION") != "1",
    reason="defina RUN_EXCEL_INTEGRATION=1 para executar Excel COM",
)
def test_a01_qtd_com_alerta_conta_apenas_linhas_com_pc(tmp_path):
    """A-01: as 5.000 linhas vazias nao podem virar alerta no Quadro 1."""
    destino = tmp_path / "pc_ux_alertas.xlsx"
    with _excel_com(TEMPLATE, destino) as (excel, wb):
        _, _, itens = _preparar_cenario_pc(wb)

        # CENARIO A - nenhum PC cadastrado.
        excel.CalculateFullRebuild()
        assert _alertas(itens) == [0, 0, 0, 0, 0, 0, 0]
        assert int(itens.Range("N9").Value) == 0
        assert float(itens.Range("O9").Value) == 0.0

        # CENARIO B - PCs validos, nenhum alerta.
        _limpar_pcs(itens)
        _gravar_pc(itens, 2, "C0-PAGO", DATA_C0, 50.0, "Sim")
        _gravar_pc(itens, 3, "C1-PAGO", DATA_C1, 200.0, "Sim")
        excel.CalculateFullRebuild()
        assert (itens.Range("C2").Value, itens.Range("C3").Value) == ("C0", "C1")
        assert itens.Range("K2").Value == "OK"
        assert itens.Range("K3").Value == "OK"
        assert _alertas(itens) == [0, 0, 0, 0, 0, 0, 0]

        # CENARIO C - um alerta dentro de C0:C4 (C4 sem INICIO_EFEITO).
        _limpar_pcs(itens)
        _gravar_pc(itens, 2, "C4-ALERTA", DATA_C4, 400.0, "Sim")
        excel.CalculateFullRebuild()
        assert itens.Range("C2").Value == "C4"
        assert str(itens.Range("K2").Value).startswith("INICIO_EFEITO ausente")
        assert _alertas(itens) == [0, 0, 0, 0, 1, 0, 1]

        # CENARIO D - um alerta fora dos ciclos.
        _limpar_pcs(itens)
        _gravar_pc(
            itens, 2, "FORA-ALERTA", DATA_FORA_DOS_CICLOS, 999_999.99, "Sim"
        )
        excel.CalculateFullRebuild()
        assert itens.Range("C2").Value == "Fora dos ciclos"
        assert itens.Range("K2").Value == "CICLO_PC nao identificado"
        assert _alertas(itens) == [0, 0, 0, 0, 0, 1, 1]

        # CENARIO E - mistura; as 4.995 linhas vazias continuam de fora.
        _gravar_pc(itens, 3, "C1-PAGO", DATA_C1, 200.0, "Sim")
        _gravar_pc(itens, 4, "C0-PAGO", DATA_C0, 50.0, "Sim")
        _gravar_pc(itens, 5, "C4-ALERTA", DATA_C4, 400.0, "Sim")
        _gravar_pc(itens, 6, "C1-ANALISE", DATA_C1_ANALISE, 100.0, "Nao")
        excel.CalculateFullRebuild()
        assert _alertas(itens) == [0, 0, 0, 0, 1, 1, 2]
        assert int(itens.Range("N9").Value) == 5


@pytest.mark.skipif(
    os.environ.get("RUN_EXCEL_INTEGRATION") != "1",
    reason="defina RUN_EXCEL_INTEGRATION=1 para executar Excel COM",
)
def test_a01_pipeline_real_nao_anuncia_alerta_em_linha_vazia(tmp_path):
    """A-01 no artefato que o fiscal baixa, nao apenas no template versionado."""
    payload = {
        "origem": "Reajustes Múltiplos",
        "indice": "IPCA",
        "data_base_original": "01/01/2025",
        "variacao_acumulada": 0.10,
        "fator_acumulado": 1.10,
        "ciclos": [{
            "ciclo": "C1",
            "data_inicio": "01/01/2026",
            "data_pedido": "15/03/2026",
            "percentual_aplicado": 0.10,
            "financeiro_inicio": "01/02/2026",
            "objeto_analise_atual": True,
            "situacao": "TEMPESTIVO",
        }],
    }
    origem = tmp_path / "coleta_pipeline.xlsx"
    origem.write_bytes(gerar_coleta_oficial_preenchida(payload))
    destino = tmp_path / "coleta_pipeline_excel.xlsx"
    with _excel_com(origem, destino) as (excel, wb):
        _, _, itens = _preparar_cenario_pc(wb)
        excel.CalculateFullRebuild()
        assert _alertas(itens) == [0, 0, 0, 0, 0, 0, 0]

        _gravar_pc(itens, 2, "C1-PAGO", DATA_C1, 200.0, "Sim")
        _gravar_pc(itens, 3, "C4-ALERTA", DATA_C4, 400.0, "Sim")
        _gravar_pc(
            itens, 4, "FORA-ALERTA", DATA_FORA_DOS_CICLOS, 999_999.99, "Sim"
        )
        excel.CalculateFullRebuild()
        assert _alertas(itens) == [0, 0, 0, 0, 1, 1, 2]
        assert int(itens.Range("N9").Value) == 3


@pytest.mark.skipif(
    os.environ.get("RUN_EXCEL_INTEGRATION") != "1",
    reason="defina RUN_EXCEL_INTEGRATION=1 para executar Excel COM",
)
def test_calculo_originario_excel_percorre_e_f_h_i_j_u_totais_e_resultados(tmp_path):
    """Prova o motor real sem injetar H/I/J; nao substitui o XLS-fonte ausente."""
    destino = tmp_path / "pc_ux_calculo_originario.xlsx"
    with _excel_com(TEMPLATE, destino) as (excel, wb):
        _, _, itens = _preparar_cenario_pc(wb)
        resultados = wb.Worksheets("RESULTADOS")

        # NOVO-02: C1 + um PC realmente fora dos ciclos fecham o universo.
        _gravar_pc(itens, 2, "C1-PAGO", DATA_C1, 200.0, "Sim")
        _gravar_pc(
            itens, 3, "FORA-CICLOS", DATA_FORA_DOS_CICLOS, 999_999.99, "Sim"
        )
        excel.CalculateFullRebuild()
        assert itens.Range("C3").Value == "Fora dos ciclos"
        assert float(itens.Range("O4").Value) == pytest.approx(200.0)
        assert float(itens.Range("O8").Value) == pytest.approx(999_999.99)
        assert float(itens.Range("O9").Value) == pytest.approx(1_000_199.99)
        assert float(resultados.Range("B55").Value) == pytest.approx(1_000_199.99)

        # Acrescenta C0, C4, C1 em analise e C1 posterior a data de corte.
        _gravar_pc(itens, 4, "C0-PAGO", DATA_C0, 50.0, "Sim")
        _gravar_pc(itens, 5, "C4-ALERTA", DATA_C4, 400.0, "Sim")
        _gravar_pc(itens, 6, "C1-ANALISE", DATA_C1_ANALISE, 100.0, "Nao")
        _gravar_pc(itens, 7, "C1-POS-CORTE", DATA_C1_POS_CORTE, 100.01, "Sim")
        excel.CalculateFullRebuild()

        assert [itens.Range(f"C{linha}").Value for linha in range(2, 8)] == [
            "C1", "Fora dos ciclos", "C0", "C4", "C1", "C1",
        ]
        # C1 em analise: E/F/H/I/J/U saem do motor, sem injecao.
        assert float(itens.Range("E6").Value) == pytest.approx(1.1)
        assert float(itens.Range("F6").Value) == pytest.approx(110.0)
        assert float(itens.Range("H6").Value) == 0.0
        assert float(itens.Range("I6").Value) == pytest.approx(110.0)
        assert float(itens.Range("J6").Value) == pytest.approx(10.0)
        assert float(itens.Range("U6").Value) == pytest.approx(100.0)
        # PC pago com efeito: retroativo reconhecido, nada em analise.
        assert float(itens.Range("H2").Value) == pytest.approx(20.0)
        assert float(itens.Range("U2").Value) == pytest.approx(220.0)
        # Fora dos ciclos nao recebe fator; C4 recebe fator mas nao tem efeito.
        assert str(itens.Range("F3").Value) == ""
        assert float(itens.Range("E5").Value) == pytest.approx(1.4641)
        assert float(itens.Range("F5").Value) == pytest.approx(585.64)

        # Quadro 1 - universo integral, inclusive o posterior ao corte.
        assert float(itens.Range("O3").Value) == pytest.approx(50.0)
        assert float(itens.Range("O4").Value) == pytest.approx(400.01)
        assert float(itens.Range("O7").Value) == pytest.approx(400.0)
        assert float(itens.Range("O8").Value) == pytest.approx(999_999.99)
        assert float(itens.Range("O9").Value) == pytest.approx(1_000_850.00)
        assert float(itens.Range("Q9").Value) == pytest.approx(30.0)
        assert float(itens.Range("R9").Value) == pytest.approx(110.0)
        assert float(itens.Range("S9").Value) == pytest.approx(10.0)
        # A-01: dois PCs com alerta, nunca as linhas vazias.
        assert _alertas(itens) == [0, 0, 0, 0, 1, 1, 2]

        # Quadro 2 - so o que foi pago ate a data de corte.
        assert float(itens.Range("O18").Value) == pytest.approx(250.0)
        assert float(itens.Range("T18").Value) == pytest.approx(1_000_500.00)
        assert float(resultados.Range("B17").Value) == pytest.approx(200.0)
        assert float(resultados.Range("C17").Value) == pytest.approx(220.0)
        assert float(resultados.Range("D17").Value) == pytest.approx(20.0)
        assert float(resultados.Range("B55").Value) == pytest.approx(1_000_850.00)

        wb.Save()
        wb.Close(False)
        reaberto = excel.Workbooks.Open(
            str(destino), UpdateLinks=0, ReadOnly=True, CorruptLoad=0
        )
        try:
            itens_reaberto = reaberto.Worksheets("itens_PC")
            assert float(itens_reaberto.Range("O9").Value) == pytest.approx(
                1_000_850.00
            )
            assert float(itens_reaberto.Range("T18").Value) == pytest.approx(
                1_000_500.00
            )
            assert int(itens_reaberto.Range("T9").Value) == 2
        finally:
            reaberto.Close(False)


@pytest.mark.parametrize("metodo", ["financeiro", "consumido"])
def test_m_n_outros_metodos_nao_recebem_texto_pc(metodo):
    leitura = leitura_simples_financeiro()
    if metodo == "consumido":
        leitura["controle"]["modo"] = "D"
    for gerador in (gerar_termo_apostila, gerar_despacho_saneador):
        texto = _texto_docx(gerador(leitura))
        assert "Metodologia utilizada: Pedidos de Compra (PC)" not in texto
        assert "PROVIDÊNCIA DA ÁREA GESTORA" not in texto
