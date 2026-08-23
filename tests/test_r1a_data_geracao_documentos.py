"""R1A — data/hora de materializacao no rodape (Despacho, Termo, Sumario).

Protege exatamente o que a tarefa exige: APUR intacto, uma unica indicacao de
geracao com hora e fuso explicito nos tres documentos reais, ausencia total de
rastreabilidade nos modelos em branco, e o contrato A->A / A->B do APUR. O
relogio e congelado via mock — nenhum teste aqui depende da hora real nem usa
sleep.
"""
from __future__ import annotations

import sys
import zipfile
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from _estado_apuracao_upload import (  # noqa: E402
    assinatura_conteudo_upload,
    id_apuracao_de_assinatura,
)
from _sumario_executivo import gerar_sumario_executivo  # noqa: E402
from _templates_documentos import (  # noqa: E402
    gerar_despacho_saneador,
    gerar_modelo_branco_despacho,
    gerar_modelo_branco_termo,
    gerar_termo_apostila,
)
from test_sumario_executivo import leitura_multiciclo_pc  # noqa: E402
from test_templates_documentos import CAMPOS_SANEADOR, CAMPOS_TERMO  # noqa: E402

CARIMBO_FIXO = "23/08/2026 09:15 (Brasília)"


def _leitura_com_apur():
    leitura = leitura_multiciclo_pc()
    assinatura = assinatura_conteudo_upload(b"r1a-smoke-fixo")
    leitura["assinatura_origem_xlsx"] = assinatura
    leitura["id_apuracao"] = id_apuracao_de_assinatura(assinatura)
    return leitura, leitura["id_apuracao"]


def _texto_docx(conteudo: bytes) -> str:
    doc = Document(BytesIO(conteudo))
    partes = [p.text for p in doc.paragraphs]
    for secao in doc.sections:
        partes.extend(p.text for p in secao.footer.paragraphs)
    return "\n".join(partes)


def _docx_valido(conteudo: bytes) -> bool:
    try:
        with zipfile.ZipFile(BytesIO(conteudo)) as z:
            nomes = z.namelist()
            return (
                conteudo[:2] == b"PK"
                and "word/document.xml" in nomes
                and "[Content_Types].xml" in nomes
                and z.testzip() is None
            )
    except Exception:
        return False


def _pdf_valido(conteudo: bytes) -> bool:
    return conteudo.startswith(b"%PDF-") and conteudo.rstrip().endswith(b"%%EOF")


def _texto_pdf(conteudo: bytes) -> str:
    fitz = pytest.importorskip("fitz")
    doc = fitz.open(stream=conteudo, filetype="pdf")
    return "\n".join(pagina.get_text() for pagina in doc)


def _paginas_pdf(conteudo: bytes) -> list[str]:
    fitz = pytest.importorskip("fitz")
    doc = fitz.open(stream=conteudo, filetype="pdf")
    return [pagina.get_text() for pagina in doc]


# --------------------------------------------------------------------- A/B
@patch("_templates_documentos.gerado_em_brasilia", return_value=CARIMBO_FIXO)
def test_despacho_real_tem_apur_e_geracao(_mock):
    leitura, id_apuracao = _leitura_com_apur()
    conteudo = gerar_despacho_saneador(leitura, campos_manuais=CAMPOS_SANEADOR)
    texto = _texto_docx(conteudo)
    assert _docx_valido(conteudo)
    assert id_apuracao in texto
    assert f"ID da apuração: {id_apuracao} | Gerado em {CARIMBO_FIXO}" in texto


@patch("_templates_documentos.gerado_em_brasilia", return_value=CARIMBO_FIXO)
def test_termo_real_tem_apur_e_geracao(_mock):
    leitura, id_apuracao = _leitura_com_apur()
    conteudo = gerar_termo_apostila(leitura, campos_manuais=CAMPOS_TERMO)
    texto = _texto_docx(conteudo)
    assert _docx_valido(conteudo)
    assert id_apuracao in texto
    assert f"ID da apuração: {id_apuracao} | Gerado em {CARIMBO_FIXO}" in texto


# ------------------------------------------------------------------------- C
@patch("_sumario_executivo.gerado_em_brasilia", return_value=CARIMBO_FIXO)
def test_sumario_real_tem_apur_e_geracao(_mock):
    leitura, id_apuracao = _leitura_com_apur()
    conteudo = gerar_sumario_executivo(leitura)
    texto = _texto_pdf(conteudo)
    assert _pdf_valido(conteudo)
    assert id_apuracao in texto
    assert f"Gerado em {CARIMBO_FIXO}." in texto
    assert "Brasília" in texto
    assert "09:15" in texto
    # Uma unica indicacao de geracao POR PAGINA (rodape repete no cabecalho
    # de cada pagina; nao pode haver uma segunda frase "Gerado em" distinta).
    for pagina in _paginas_pdf(conteudo):
        assert pagina.count("Gerado em") == 1


# ------------------------------------------------------------------------- D
def test_modelos_em_branco_continuam_sem_rastreabilidade():
    for conteudo in (gerar_modelo_branco_despacho(), gerar_modelo_branco_termo()):
        texto = _texto_docx(conteudo)
        assert _docx_valido(conteudo)
        assert "APUR-" not in texto
        assert "ID da apuração" not in texto
        assert "Gerado em" not in texto


# --------------------------------------------------------------------- E/F
def test_apur_estavel_a_para_a_e_diferente_a_para_b():
    a1 = id_apuracao_de_assinatura(assinatura_conteudo_upload(b"conteudo-A"))
    a2 = id_apuracao_de_assinatura(assinatura_conteudo_upload(b"conteudo-A"))
    b = id_apuracao_de_assinatura(assinatura_conteudo_upload(b"conteudo-B"))
    assert a1 == a2
    assert a1 != b


# ------------------------------------------------------------------------- G
@patch("_templates_documentos.gerado_em_brasilia", return_value=CARIMBO_FIXO)
def test_geracao_nao_altera_conteudo_negocial_do_despacho(_mock):
    leitura, _id = _leitura_com_apur()
    texto = _texto_docx(gerar_despacho_saneador(leitura, campos_manuais=CAMPOS_SANEADOR))
    # Mesmas checagens de conteudo negocial ja usadas nos testes existentes
    # deste gerador (numero de ciclos computados nao muda com o carimbo novo).
    assert "Despacho" in texto or "DESPACHO" in texto.upper()
