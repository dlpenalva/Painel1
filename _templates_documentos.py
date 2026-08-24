"""Geradores de documentos administrativos/juridicos em DOCX.

Gera o Despacho Saneador e a Minuta de Termo de Apostilamento a partir dos
dados canonicos do Objeto Processo de Reajuste. Nao recalcula valores — apenas
apresenta os dados ja consolidados pelos motores oficiais, em LINGUAGEM
ADMINISTRATIVA (nunca expoe vocabulario de implementacao do XLS/Python).

Campos manuais ausentes recebem o marcador [PREENCHER: <descricao>] com
destaque amarelo. Ausencia de dado automatico nunca vira zero.

Nenhum arquivo entregue pode conter emoji/pictograma (sanitizacao no output).

Interface publica:
    gerar_despacho_saneador(leitura_ou_objeto, identificacao, campos_manuais) -> bytes
    gerar_termo_apostila(leitura_ou_objeto, identificacao, campos_manuais) -> bytes
    diagnosticar_campos_manuais(leitura_ou_objeto, identificacao, campos_manuais) -> list[dict]
"""
from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from _sumario_executivo import (
    NAO_INFORMADO,
    formatar_moeda,
    montar_dados_sumario_executivo,
    _num_ou_none,
)
from _objeto_processo_reajuste import obter_objeto_processo_reajuste
from _reajuste_utils import (
    FRASE_SEM_CICLOS_COMPUTADOS,
    expressao_quantidade_ciclos,
    gerado_em_brasilia,
)
from _sanitizacao_documental import remover_emojis_leve

# ---------------------------------------------------------------------------
# Constantes
# ---------------------------------------------------------------------------

PREENCHER_TAG = "[PREENCHER: {}]"
COR_NEGATIVO = RGBColor(0xC0, 0x00, 0x00)
_LETRAS = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"


def _fmt_pct_doc(valor: Any) -> str:
    """Formata percentual documental com exatamente duas casas decimais.

    Entrada no formato decimal canonico (0.0421 -> '4,21%').
    """
    numero = _num_ou_none(valor)
    if numero is None:
        return NAO_INFORMADO
    texto = f"{numero * 100:.2f}".replace(".", ",")
    return f"{texto}%"


def _indice_amigavel_doc(indice: Any) -> str | None:
    """Nome amigavel do indice, sem expor codigo tecnico (SGS-433/189/DIMAC).

    Retorna None quando indefinido, para que o chamador use marcador manual.
    """
    texto = remover_emojis_leve(indice).strip()
    if not texto or texto == NAO_INFORMADO:
        return None
    norm = texto.upper()
    if norm.startswith("IST"):
        # user-facing IST (Anatel); legado "IST (Série Local)" ainda reconhecido.
        return "IST (Anatel)"
    if norm.startswith("ICTI"):
        return "ICTI (Ipeadata)"
    if norm.startswith("IPCA"):
        return "IPCA"
    if norm.startswith("IGP"):
        return "IGP-M"
    if norm.startswith("INPC"):
        return "INPC"
    import re as _re
    limpo = _re.sub(r"\s*\[[^\]]*\]\s*", " ", texto)      # remove "[SGS-433]"
    limpo = _re.sub(r"\s*\(\s*\d+\s*\)\s*$", "", limpo)   # remove "(433)"
    return limpo.strip() or None


CAMPOS_MANUAIS_DESPACHO = [
    ("contrato", "Numero do contrato", "despacho"),
    ("empresa_contratada", "Nome da empresa contratada", "despacho"),
    ("objeto_contrato", "Objeto resumido do contrato", "despacho"),
    ("vigencia_ate", "Data final da vigencia contratual", "despacho"),
    ("tipo_atualizacao", "Tipo da atualizacao contratual", "despacho"),
    ("processo_pleito", "Referencias do pleito da contratada", "despacho"),
    ("referencia_analise", "Referencia onde o resultado da analise consta", "despacho"),
    ("memoria_calculo_ref", "Referencia da memoria de calculo", "despacho"),
    ("adequacao_orcamentaria_ref", "Referencia da adequacao orcamentaria", "despacho"),
    ("adequacao_orcamentaria_valor", "Valor da adequacao orcamentaria", "despacho"),
    ("regularidade_ref", "Referencia das certidoes de regularidade", "despacho"),
    ("regularidade_situacao", "Situacao da regularidade da contratada", "despacho"),
    ("concordancia_ref", "Referencia da manifestacao de concordancia da contratada", "despacho"),
    ("concordancia_situacao", "Situacao da concordancia da contratada", "despacho"),
    ("garantia_situacao", "Situacao da garantia contratual", "despacho"),
    ("docs_desatualizados", "Lista de documentos a desconsiderar (opcional)", "despacho"),
    ("pendencias_complemento", "Complemento manual das pendencias (opcional)", "despacho"),
]

CAMPOS_MANUAIS_TERMO = [
    ("contrato", "Numero do contrato", "termo"),
    ("empresa_contratada", "Nome/qualificacao da empresa contratada", "termo"),
    ("representante_telebras_1_nome", "Nome do 1o representante da Telebras", "termo"),
    ("representante_telebras_1_matricula", "Matricula do 1o representante", "termo"),
    ("representante_telebras_2_cargo", "Cargo do 2o representante da Telebras", "termo"),
    ("representante_telebras_2_matricula", "Matricula do 2o representante", "termo"),
    ("solicitacao_data", "Data da solicitacao da contratada", "termo"),
    ("solicitacao_ref", "Referencia documental da solicitacao da contratada", "termo"),
    ("memoria_calculo_ref", "Referencia da memoria de calculo", "termo"),
    ("concordancia_ref", "Referencia da manifestacao de concordancia da contratada", "termo"),
    ("regularidade_ref", "Referencia das certidoes de regularidade", "termo"),
    ("adequacao_orcamentaria_ref", "Referencia da adequacao orcamentaria", "termo"),
    ("processo_ref", "Numero do processo de instrucao", "termo"),
    ("valor_pago_efetivo", "Valor pago efetivo (quando nao apurado automaticamente)", "termo"),
    ("valor_teorico", "Valor devido apos o reajuste (quando nao apurado automaticamente)", "termo"),
    ("valor_original_contrato", "Valor original do contrato", "termo"),
    ("local_data", "Data (ex.: 20/07/2026)", "termo"),
]

TODOS_CAMPOS_MANUAIS = list(
    {c[0]: c for c in CAMPOS_MANUAIS_DESPACHO + CAMPOS_MANUAIS_TERMO}.values()
)

# Campos que sao opcionais (nao entram como pendencia critica no diagnostico).
_CAMPOS_OPCIONAIS = {
    "docs_desatualizados", "pendencias_complemento", "valor_pago_efetivo",
    "valor_teorico",
}


# ---------------------------------------------------------------------------
# Helpers XML / DOCX
# ---------------------------------------------------------------------------

# Etapa 26H — politica documental da PREVIA: numero XLS oficial sem resultado
# definitivo e exibido como "R$ x — PREVIA", com highlight verde somente na
# palavra PREVIA. Nunca declara VALIDADO nem resolve a divergencia (26C).
ROTULO_PREVIA = "PRÉVIA"
SUFIXO_PREVIA = f" — {ROTULO_PREVIA}"
COR_HIGHLIGHT_PREVIA = "green"


def _vta_texto_doc(dados: dict) -> str:
    """Texto documental do VTA: valor oficial, PREVIA do XLS, ou vazio."""
    vta = dados.get("vta")
    if vta is not None:
        return formatar_moeda(vta)
    vta_previa = dados.get("vta_previa")
    if vta_previa is not None:
        return f"{formatar_moeda(vta_previa)}{SUFIXO_PREVIA}"
    return ""


def _set_highlight(run, cor: str = "yellow") -> None:
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), cor)
    rPr.append(highlight)


def _repetir_cabecalho(tabela) -> None:
    tr = tabela.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _adicionar_run(p, texto: str, negrito: bool = False, tamanho: int = 11,
                   cor: RGBColor | None = None, italico: bool = False) -> Any:
    run = p.add_run(remover_emojis_leve(texto))
    run.bold = negrito
    run.italic = italico
    run.font.name = "Calibri"
    run.font.size = Pt(tamanho)
    if cor:
        run.font.color.rgb = cor
    return run


def _titulo_secao(doc: Document, texto: str, tamanho: int = 11,
                  alinhamento=WD_ALIGN_PARAGRAPH.LEFT) -> Any:
    p = doc.add_paragraph()
    p.alignment = alinhamento
    _adicionar_run(p, texto, negrito=True, tamanho=tamanho)
    return p


def _titulo_quadro(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _adicionar_run(p, texto, negrito=True, tamanho=10)


def _configurar_box_discreto(paragrafo) -> None:
    ppr = paragrafo._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    ppr.append(shd)
    bordas = OxmlElement("w:pBdr")
    for lado in ("top", "left", "bottom", "right"):
        borda = OxmlElement(f"w:{lado}")
        borda.set(qn("w:val"), "single")
        borda.set(qn("w:sz"), "6")
        borda.set(qn("w:space"), "6")
        borda.set(qn("w:color"), "BFBFBF")
        bordas.append(borda)
    ppr.append(bordas)


def _adicionar_box_retroativos(doc: Document, dados: dict, *, saneador: bool) -> None:
    situacao = dados.get("situacao_retroativos_pc") or {}
    reconhecido = _num_ou_none(situacao.get("reconhecido"))
    em_analise = _num_ou_none(situacao.get("em_analise"))
    potencial = _num_ou_none(situacao.get("potencial"))
    if saneador:
        exibir = any(abs(v or 0.0) > 0.004 for v in (reconhecido, em_analise, potencial))
    else:
        exibir = any(abs(v or 0.0) > 0.004 for v in (reconhecido, potencial))
    if not exibir:
        return

    reconhecido = reconhecido or 0.0
    em_analise = em_analise or 0.0
    potencial = potencial or 0.0
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _configurar_box_discreto(p)
    _adicionar_run(p, "SITUAÇÃO DOS VALORES RETROATIVOS", negrito=True, tamanho=10)
    p.add_run().add_break()
    _adicionar_run(p, f"Retroativo reconhecido: {formatar_moeda(reconhecido)}", tamanho=10)
    if saneador:
        p.add_run().add_break()
        _adicionar_run(
            p, f"Valor atualizado em análise: {formatar_moeda(em_analise)}", tamanho=10
        )
    p.add_run().add_break()
    _adicionar_run(p, f"Retroativo potencial: {formatar_moeda(potencial)}", tamanho=10)
    p.add_run().add_break()
    p.add_run().add_break()

    if saneador:
        texto = (
            f"A apuração identificou retroativo reconhecido no valor de "
            f"{formatar_moeda(reconhecido)}. Foram identificados, ainda, "
            "Pedidos de Compra com valores sujeitos à confirmação, "
            f"correspondentes a retroativo potencial de até {formatar_moeda(potencial)}.\n\n"
            "A conversão do retroativo potencial em valor reconhecido dependerá "
            "da aceitação dos respectivos Pedidos de Compra pela área gestora do "
            "contrato, a quem compete conduzir a validação desses eventos e os "
            "procedimentos relacionados ao eventual pagamento.\n\n"
            "Enquanto não houver essa confirmação, o retroativo potencial não "
            "integra o valor reconhecido a pagar."
        )
    else:
        texto = (
            f"No âmbito da apuração, foi reconhecido retroativo de "
            f"{formatar_moeda(reconhecido)}. Adicionalmente, foi identificado "
            f"retroativo potencial de até {formatar_moeda(potencial)}, relacionado "
            "a Pedidos de Compra ainda sujeitos à aceitação pela área gestora do "
            "contrato.\n\n"
            "A confirmação desses valores e a condução do eventual pagamento "
            "competem à área gestora, observados os critérios, o índice, o "
            "percentual e o período de efeitos formalizados nesta Apostila.\n\n"
            "O retroativo potencial não integra, nesta data, o montante "
            "reconhecido a pagar."
        )
    partes = texto.split("\n")
    for indice, parte in enumerate(partes):
        if indice:
            p.add_run().add_break()
        if parte:
            _adicionar_run(p, parte, tamanho=10)


def _run_campo_manual(p, descricao: str, tamanho: int = 11) -> Any:
    run = p.add_run(PREENCHER_TAG.format(descricao))
    run.font.name = "Calibri"
    run.font.size = Pt(tamanho)
    _set_highlight(run, "yellow")
    return run


def _texto_ou_marcador(p, valor: Any, descricao: str, tamanho: int = 11,
                        negrito: bool = False, cor: RGBColor | None = None) -> None:
    if valor is not None and str(valor).strip():
        run = p.add_run(remover_emojis_leve(valor))
        run.bold = negrito
        run.font.name = "Calibri"
        run.font.size = Pt(tamanho)
        if cor:
            run.font.color.rgb = cor
    else:
        _run_campo_manual(p, descricao, tamanho)


def _campo(campos_manuais: dict, chave: str) -> Any:
    if not campos_manuais:
        return None
    v = campos_manuais.get(chave)
    if v is None:
        return None
    if isinstance(v, str) and not v.strip():
        return None
    return v


def _valor_moeda_ou_marcador(p, valor: Any, descricao: str, tamanho: int = 11) -> None:
    numero = _num_ou_none(valor)
    if numero is not None:
        run = p.add_run(formatar_moeda(numero))
        run.font.name = "Calibri"
        run.font.size = Pt(tamanho)
        if numero < 0:
            run.font.color.rgb = COR_NEGATIVO
    else:
        _run_campo_manual(p, descricao, tamanho)


def _configurar_documento() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


def _adicionar_id_apuracao_rodape(doc: Document, dados: dict) -> None:
    """Registra a rastreabilidade sem interferir no corpo juridico do documento.

    So escreve quando ha id_apuracao (nunca em modelo em branco, cujo `dados`
    nao carrega essa chave). A data/hora acompanha o mesmo bloco/run — nao e
    a apuracao, o upload ou o commit: e o instante em que estes bytes
    especificos foram montados (fonte unica: gerado_em_brasilia()).
    """
    id_apuracao = str(dados.get("id_apuracao") or "").strip()
    if not id_apuracao:
        return
    paragrafo = doc.sections[0].footer.paragraphs[0]
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragrafo.add_run(
        f"ID da apuração: {id_apuracao} | Gerado em {gerado_em_brasilia()}"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _adicionar_tabela(
    doc: Document,
    cabecalho: list[str],
    linhas: list[list[str]],
    *,
    repetir_cabecalho: bool = True,
    destacar_placeholders: bool = False,
    destacar_placeholders_embutidos: bool = False,
) -> Any:
    n_cols = len(cabecalho)
    tabela = doc.add_table(rows=1, cols=n_cols)
    tabela.style = "Table Grid"
    tabela.rows[0].height = Pt(16)
    tabela.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    celulas_cab = tabela.rows[0].cells
    for i, texto in enumerate(cabecalho):
        celulas_cab[i].text = ""
        run = celulas_cab[i].paragraphs[0].add_run(remover_emojis_leve(texto))
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
    if repetir_cabecalho:
        _repetir_cabecalho(tabela)
    for linha in linhas:
        row = tabela.add_row()
        row.height = Pt(16)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for i, celula_texto in enumerate(linha):
            row.cells[i].text = ""
            texto = remover_emojis_leve(celula_texto)
            negativo = False
            try:
                val_num = float(
                    str(celula_texto).replace("R$ ", "").replace(".", "").replace(",", ".")
                )
                negativo = val_num < 0
            except (ValueError, AttributeError):
                pass
            paragrafo_celula = row.cells[i].paragraphs[0]
            if texto.endswith(SUFIXO_PREVIA):
                # Etapa 26H: highlight verde SOMENTE na palavra PREVIA.
                run = paragrafo_celula.add_run(texto[: -len(ROTULO_PREVIA)])
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run_previa = paragrafo_celula.add_run(ROTULO_PREVIA)
                run_previa.font.name = "Calibri"
                run_previa.font.size = Pt(10)
                _set_highlight(run_previa, COR_HIGHLIGHT_PREVIA)
                continue
            partes = (
                re.split(r"(\[PREENCHER:[^\]]+\])", texto)
                if destacar_placeholders_embutidos and "[PREENCHER:" in texto
                else [texto]
            )
            for parte in partes:
                if not parte:
                    continue
                run = paragrafo_celula.add_run(parte)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                if destacar_placeholders and parte.startswith("[PREENCHER:"):
                    _set_highlight(run, "yellow")
                elif negativo and "R$" in str(celula_texto):
                    run.font.color.rgb = COR_NEGATIVO
    return tabela


# ---------------------------------------------------------------------------
# Extracao de dados canonicos
# ---------------------------------------------------------------------------

def _extrair_dados(leitura_ou_objeto: dict, identificacao: dict | None) -> dict:
    dados = montar_dados_sumario_executivo(leitura_ou_objeto, identificacao)
    if not dados.get("disponivel"):
        return {
            "disponivel": False,
            "identificacao_externa": dict(identificacao or {}),
            "pendencias": {},
        }

    ciclos = dados.get("ciclos") or []
    ciclos_reajuste = [c for c in ciclos if not c.get("eh_base")]
    ciclos_computados = [c for c in ciclos_reajuste if c.get("computar") == "Sim"]

    financeiro = dados.get("financeiro") or {}
    sintese = dados.get("sintese") or {}
    aditivos_raw = (dados.get("aditivos") or {}).get("itens") or []

    fin_por_ciclo = {r["ciclo"]: r for r in financeiro.get("financeiro_por_ciclo") or []}
    pc_por_ciclo = {r["ciclo"]: r for r in financeiro.get("pc_por_ciclo") or []}

    objeto_proc = obter_objeto_processo_reajuste(leitura_ou_objeto) or {}
    dados_op = objeto_proc.get("dados_operacionais") or {}
    if not dados_op and isinstance(leitura_ou_objeto, dict):
        dados_op = leitura_ou_objeto
    vta_sombra = dados_op.get("vta_sombra") or {}
    parcelas_vta = vta_sombra.get("parcelas_computadas") or []

    aditivos = []
    for ad in aditivos_raw:
        aditivos.append({
            "identificador_interno": ad.get("identificador_interno"),
            "rotulo_documental": (
                ad.get("rotulo_documental") or ad.get("ciclo")
            ),
            "instrumento": ad.get("instrumento"),
            "item": ad.get("item"),
            "tipo_alteracao": ad.get("tipo_alteracao"),
            "ciclo": ad.get("ciclo"),
            "data_alteracao": ad.get("data_alteracao"),
            "quantidade": ad.get("quantidade"),
            "valor_original": ad.get("valor_original"),
            "valor_atualizado": ad.get("valor_atualizado"),
        })

    return {
        "disponivel": True,
        "id_apuracao": dados.get("id_apuracao"),
        "ciclos": ciclos,
        "ciclos_reajuste": ciclos_reajuste,
        "ciclos_computados": ciclos_computados,
        "var_acumulada": sintese.get("variacao_acumulada"),
        "vta": sintese.get("vta"),
        "vta_previa": sintese.get("vta_previa"),
        "vta_execucao_atualizada": sintese.get("vta_execucao_atualizada"),
        "vta_saldo_remanescente_atualizado": sintese.get(
            "vta_saldo_remanescente_atualizado"
        ),
        "fin_por_ciclo": fin_por_ciclo,
        "pc_por_ciclo": pc_por_ciclo,
        "parcelas_vta": parcelas_vta,
        "aditivos": aditivos,
        "financeiro": financeiro,
        "sintese": sintese,
        "identificacao": dados.get("identificacao") or {},
        "identificacao_externa": dict(identificacao or {}),
        "pendencias": objeto_proc.get("pendencias") or {},
        "historico_vu": dados.get("historico_vu") or {},
        "referencias_vta": (
            (leitura_ou_objeto or {}).get("referencias_vta")
            or (dados.get("referencias_vta") if isinstance(dados, dict) else None)
            or {}
        ),
        "situacao_retroativos_pc": _situacao_retroativos_pc(dados_op),
    }


def _situacao_retroativos_pc(dados_operacionais: dict) -> dict[str, float] | None:
    """Expõe nos documentos o mesmo consolidado canônico usado em RESULTADOS."""
    totais_pc = (dados_operacionais.get("itens_pc_v10") or {}).get(
        "totais_canonicos"
    ) or dados_operacionais.get("totais_canonicos_pc") or {}
    ate_o_corte = totais_pc.get("ate_o_corte") or {}
    totais = {
        "reconhecido": _num_ou_none(ate_o_corte.get("retroativo")) or 0.0,
        "em_analise": (
            _num_ou_none(ate_o_corte.get("valor_atualizado_em_analise")) or 0.0
        ),
        "potencial": _num_ou_none(ate_o_corte.get("delta_potencial")) or 0.0,
    }
    if not any(abs(valor) > 0.004 for valor in totais.values()):
        return None
    return {chave: round(valor, 2) for chave, valor in totais.items()}


def _retroativo_total(dados: dict) -> float | None:
    fin = dados.get("financeiro") or {}
    t_fin = fin.get("delta_total_financeiro")
    if t_fin is not None:
        return t_fin
    return fin.get("delta_total_pc")


def _linhas_financeiro(dados: dict) -> list[dict]:
    if dados.get("fin_por_ciclo"):
        return list(dados["fin_por_ciclo"].values())
    if dados.get("pc_por_ciclo"):
        return list(dados["pc_por_ciclo"].values())
    return []


def _valor_pago_total(dados: dict) -> float | None:
    linhas = _linhas_financeiro(dados)
    if not linhas:
        return None
    return round(sum(_num_ou_none(l.get("valor_pago")) or 0.0 for l in linhas), 2)


def _valor_atualizado_total(dados: dict) -> float | None:
    linhas = _linhas_financeiro(dados)
    if not linhas:
        return None
    return round(sum(_num_ou_none(l.get("valor_atualizado")) or 0.0 for l in linhas), 2)


def _indice_doc(dados: dict) -> str | None:
    return _indice_amigavel_doc((dados.get("identificacao") or {}).get("indice"))


def _efeito_financeiro_ciclo(c: dict) -> str:
    """Frase administrativa de efeitos financeiros de um ciclo."""
    situacao = remover_emojis_leve(c.get("situacao") or "").strip().lower()
    inicio = str(c.get("inicio_efeito_financeiro") or "").strip()
    if "preclu" in situacao:
        return "Sem efeitos financeiros"
    if inicio and inicio != NAO_INFORMADO and "/" in inicio:
        return f"A partir de {inicio}"
    return NAO_INFORMADO


_MESES_EXTENSO = (
    "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
)


def _lista_natural(itens: list[str]) -> str:
    if not itens:
        return ""
    if len(itens) == 1:
        return itens[0]
    return ", ".join(itens[:-1]) + " e " + itens[-1]


def _competencias_sem_efeito(c: dict) -> list[str]:
    """Competencias do ciclo que nao produzem efeitos financeiros.

    Fonte unica: o bloco `meses_sem_efeito` que a apuracao ja consolida a
    partir dos marcos canonicos (inicio do ciclo x INICIO_EFEITO_FINANCEIRO).
    O gerador NAO recria temporalidade: sem status "ok" nao ha o que declarar.
    """
    situacao = remover_emojis_leve(c.get("situacao") or "").strip().lower()
    if "preclu" in situacao:
        # Preclusao integral ja e declarada como "Sem efeitos financeiros";
        # nunca vira perda parcial de competencias.
        return []
    bloco = c.get("meses_sem_efeito") or {}
    if str(bloco.get("status") or "") != "ok":
        return []
    return [str(x).strip() for x in (bloco.get("competencias") or []) if str(x).strip()]


def _competencias_por_extenso(competencias: list[str]) -> str:
    """'01/2026', '02/2026' -> 'janeiro e fevereiro de 2026'."""
    grupos: list[tuple[str, list[str]]] = []
    for comp in competencias:
        mes_txt, _, ano = str(comp).partition("/")
        try:
            nome_mes = _MESES_EXTENSO[int(mes_txt) - 1]
        except (ValueError, IndexError):
            nome_mes = mes_txt
        if grupos and grupos[-1][0] == ano:
            grupos[-1][1].append(nome_mes)
        else:
            grupos.append((ano, [nome_mes]))
    return _lista_natural([f"{_lista_natural(m)} de {ano}" for ano, m in grupos])


def _frase_perda_efeitos(c: dict, *, nomear_ciclo: bool) -> str | None:
    competencias = _competencias_sem_efeito(c)
    if not competencias:
        return None
    inicio = str(c.get("inicio_efeito_financeiro") or "").strip()
    if not inicio or inicio == NAO_INFORMADO or "/" not in inicio:
        return None
    ciclo = remover_emojis_leve(c.get("ciclo") or "").strip()
    referencia = f"do ciclo {ciclo}" if (nomear_ciclo and ciclo) else "deste ciclo"
    rotulo = "a competência de" if len(competencias) == 1 else "as competências de"
    return (
        "Em razão da data do pedido, os efeitos financeiros do reajuste "
        f"{referencia} iniciam-se em {inicio}, não alcançando {rotulo} "
        f"{_competencias_por_extenso(competencias)}."
    )


def _paragrafos_perda_efeitos(doc: Document, dados: dict) -> None:
    """Declara, ciclo a ciclo, as competencias sem efeitos financeiros.

    Compartilhada pelo Despacho Saneador e pelo Termo de Apostila para que os
    dois documentos declarem a mesma perda a partir da mesma fonte temporal.
    """
    if dados.get("_modo_branco"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _adicionar_run(
            p,
            "Havendo competências não alcançadas pelos efeitos financeiros do "
            "reajuste em razão da data do pedido, deverão ser expressamente "
            "indicadas neste item.",
        )
        return
    ciclos = dados.get("ciclos_computados") or []
    nomear = len(ciclos) > 1
    for c in ciclos:
        frase = _frase_perda_efeitos(c, nomear_ciclo=nomear)
        if not frase:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _adicionar_run(p, frase)


# ---------------------------------------------------------------------------
# Camada de apresentacao humanizada do VTA (nunca expoe vocabulario do XLS)
# ---------------------------------------------------------------------------

def _descricao_vta_humana(parcela: dict) -> str:
    """Traduz a parcela do VTA para linguagem administrativa."""
    fonte = str(parcela.get("fonte_parcela") or "").strip().lower()
    ciclo = remover_emojis_leve(parcela.get("ciclo") or "").strip().upper()
    if "aditivo" in fonte or "supress" in fonte:
        return f"Aditivo/supressão computável ({ciclo})" if ciclo else "Aditivo/supressão computável"
    if "remanesc" in fonte or "residual" in fonte or "saldo" in fonte:
        return "Saldo remanescente atualizado"
    if ciclo:
        return f"{ciclo} - execução atualizada"
    return "Parcela de composição do Valor Total Atualizado"


def _composicao_didatica_vta(dados: dict) -> list[tuple[str, float | None]]:
    """Agrupa as parcelas em componentes didaticos (execucao por ciclo, saldo,
    aditivos), somando por rubrica. Nunca inventa; apenas soma o que existe.
    """
    executado = _num_ou_none(dados.get("vta_execucao_atualizada"))
    saldo = _num_ou_none(dados.get("vta_saldo_remanescente_atualizado"))
    vta = _num_ou_none(dados.get("vta"))
    if vta is None:
        vta = _num_ou_none(dados.get("vta_previa"))
    if (
        executado is not None
        and saldo is not None
        and vta is not None
        and abs(round(executado + saldo, 2) - round(vta, 2)) <= 0.01
    ):
        return [
            ("Execução atualizada anterior ao corte", round(executado, 2)),
            ("Saldo remanescente atualizado no corte", round(saldo, 2)),
        ]

    parcelas = dados.get("parcelas_vta") or []
    grupos: dict[str, float | None] = {}
    ordem: list[str] = []
    for p in parcelas:
        desc = _descricao_vta_humana(p)
        valor = _num_ou_none(p.get("valor_atualizado"))
        if valor is None:
            valor = _num_ou_none(p.get("valor"))
        if desc not in grupos:
            grupos[desc] = None
            ordem.append(desc)
        if valor is not None:
            grupos[desc] = (grupos[desc] or 0.0) + valor
    return [(d, grupos[d]) for d in ordem]


TITULO_HISTORICO_VU = "HISTÓRICO DOS VALORES UNITÁRIOS POR CICLO"


def montar_historico_vu_documental(dados: dict) -> dict:
    """Estrutura neutra do quadro de VUs, unica para Saneador e Apostila.

    Fonte: dados["historico_vu"] (aba historico_VU via sumario executivo),
    ja truncada em C0..ultimo ciclo da analise SEM filtrar ciclos historicos
    com COMPUTAR=Nao (o quadro e historico contratual). Nunca inventa zeros.
    """
    hvu = dados.get("historico_vu") or {}
    itens = hvu.get("itens") or []
    ciclos = hvu.get("ciclos") or []
    if not itens or not ciclos:
        return {"disponivel": False, "cabecalhos": [], "linhas": [],
                "ciclo_final": None}
    cabecalhos = ["Item"] + [f"VU_{c}" for c in ciclos]
    linhas: list[list[str]] = []
    for reg in itens:
        vus = reg.get("vus") or {}
        linha = [str(reg.get("item") or "")]
        for c in ciclos:
            valor = vus.get(c)
            linha.append(formatar_moeda(valor) if valor is not None else "")
        linhas.append(linha)
    return {"disponivel": True, "cabecalhos": cabecalhos, "linhas": linhas,
            "ciclo_final": hvu.get("ultimo_ciclo")}


def _secao_valores_unitarios_por_ciclo(
    doc: Document, dados: dict, texto_intro: str | None = None
) -> None:
    """Renderiza o HISTORICO DOS VALORES UNITARIOS POR CICLO (C0..ultimo).

    Mesma estrutura para os dois documentos (montar_historico_vu_documental).
    """
    quadro = montar_historico_vu_documental(dados)
    if not quadro["disponivel"]:
        return
    if texto_intro:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _adicionar_run(p, texto_intro)
    _titulo_quadro(doc, TITULO_HISTORICO_VU)
    linhas = quadro["linhas"]
    blocos = [linhas[:10]]
    restante = linhas[10:]
    while restante:
        blocos.append(restante[:12])
        restante = restante[12:]
    for indice, bloco in enumerate(blocos):
        if indice:
            doc.add_page_break()
            _titulo_quadro(doc, f"{TITULO_HISTORICO_VU} (continuação)")
        _adicionar_tabela(
            doc,
            quadro["cabecalhos"],
            bloco,
            repetir_cabecalho=False,
        )
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# TERMO DE APOSTILAMENTO (modelo canonico §6)
# ---------------------------------------------------------------------------

def gerar_termo_apostila(
    leitura_ou_objeto: dict,
    identificacao: dict | None = None,
    campos_manuais: dict | None = None,
    *,
    modo_modelo_em_branco: bool = False,
) -> bytes:
    """Gera a Minuta de Termo de Apostilamento em DOCX e retorna os bytes.

    Etapa 29B: `modo_modelo_em_branco=True` produz um MODELO EM BRANCO — mesma
    estrutura juridica, mas sem afirmar reajuste apurado, formalizacao,
    concordancia, certidoes, adequacao, ausencia de aditivos ou percentual; os
    campos variaveis viram placeholders [PREENCHER: ...] destacados. Com o
    padrao False o documento automatico permanece identico ao main.
    """
    if campos_manuais is None:
        campos_manuais = {}
    dados = _extrair_dados(leitura_ou_objeto, identificacao)
    dados["_modo_branco"] = bool(modo_modelo_em_branco)
    doc = _configurar_documento()

    _ta_titulo(doc, campos_manuais)
    _ta_qualificacao(doc, campos_manuais)
    _ta_considerandos(doc, dados, campos_manuais)
    _ta_abertura(doc)
    _ta_secao1_reajustes(doc, dados, campos_manuais)
    _ta_secao2_retroativo(doc, dados, campos_manuais)
    _ta_secao3_composicao_vta(doc, dados)
    _ta_secao4_valores_unitarios(doc, dados)
    _ta_secao5_aditivos(doc, dados)
    _ta_secoes_finais(doc, campos_manuais)
    _ta_assinaturas(doc, campos_manuais)
    _adicionar_id_apuracao_rodape(doc, dados)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Etapa 29B — modelos em branco (wrappers finos, sem duplicar estrutura).
# Nao leem session_state; produzem bytes deterministicos a partir de {} apenas.
# ---------------------------------------------------------------------------

def gerar_modelo_branco_despacho() -> bytes:
    """Modelo em branco do Despacho Saneador (sem dados de Coleta/sessao)."""
    return gerar_despacho_saneador({}, {}, {}, modo_modelo_em_branco=True)


def gerar_modelo_branco_termo() -> bytes:
    """Modelo em branco do Termo de Apostila (sem dados de Coleta/sessao)."""
    return gerar_termo_apostila({}, {}, {}, modo_modelo_em_branco=True)


def _ta_titulo(doc: Document, cm: dict) -> None:
    _titulo_secao(doc, "MINUTA DE TERMO DE APOSTILAMENTO", tamanho=12,
                  alinhamento=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _adicionar_run(p, "Contrato nº ")
    _texto_ou_marcador(p, _campo(cm, "contrato"), "Numero do contrato")
    doc.add_paragraph()


def _ta_qualificacao(doc: Document, cm: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _adicionar_run(p,
        "A TELECOMUNICAÇÕES BRASILEIRAS S.A. - TELEBRAS, sociedade de economia "
        "mista, vinculada ao Ministério das Comunicações, com sede no SIG, "
        "Quadra 04, Bloco A, Salas 201 a 224, Edifício Capital Financial Center, "
        "CEP nº 70.610-440, inscrita no CNPJ sob o n.º 00.336.701/0001-04, "
        "doravante denominada TELEBRAS, neste ato representada por ")
    _texto_ou_marcador(p, _campo(cm, "representante_telebras_1_nome"), "Nome do 1o representante da Telebras")
    _adicionar_run(p, ", Matrícula ")
    _texto_ou_marcador(p, _campo(cm, "representante_telebras_1_matricula"), "Matricula do 1o representante")
    _adicionar_run(p, ", e por seu ")
    _texto_ou_marcador(p, _campo(cm, "representante_telebras_2_cargo"), "Cargo do 2o representante da Telebras")
    _adicionar_run(p, ", Matrícula ")
    _texto_ou_marcador(p, _campo(cm, "representante_telebras_2_matricula"), "Matricula do 2o representante")
    _adicionar_run(p, ", nos termos da Diretriz nº 229/2018, apostila o Contrato nº ")
    _texto_ou_marcador(p, _campo(cm, "contrato"), "Numero do contrato")
    _adicionar_run(p, ", celebrado com a empresa ")
    _texto_ou_marcador(p, _campo(cm, "empresa_contratada"), "Nome/qualificacao da empresa contratada")
    _adicionar_run(p,
        ", doravante denominada CONTRATADA, com fundamento no parágrafo 7º do "
        "art. 81 da Lei nº 13.303, de 30 de junho de 2016, na legislação "
        "aplicável, no Regulamento de Licitações e Contratos da Telebras e nos "
        "documentos constantes do processo.")
    doc.add_paragraph()


def _ta_considerandos(doc: Document, dados: dict, cm: dict) -> None:
    _titulo_secao(doc, "CONSIDERANDO:")

    def item(numero: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _adicionar_run(p, f"{numero}. ", negrito=True)
        return p

    p1 = item("1")
    _adicionar_run(p1, "A Cláusula Oitava do Contrato nº ")
    _texto_ou_marcador(p1, _campo(cm, "contrato"), "Numero do contrato")
    _adicionar_run(p1,
        ", que disciplina o reajuste contratual, os ciclos de apuração, a "
        "admissibilidade dos pedidos e os respectivos efeitos financeiros;")

    p2 = item("2")
    _adicionar_run(p2,
        "A deliberação da Diretoria Executiva da Telebras, consignada na Ata da "
        "1869ª Reunião Ordinária, de 13 de janeiro de 2026, que revogou a "
        "suspensão anteriormente imposta à tramitação dos reajustes contratuais "
        "e restabeleceu a normalidade do respectivo processamento;")

    p3 = item("3")
    branco = dados.get("_modo_branco")
    if branco:
        _adicionar_run(p3,
            "A solicitação da CONTRATADA a ser identificada pela data ")
        _run_campo_manual(p3, "Data da solicitacao da contratada")
        _adicionar_run(p3, " e pela referência documental ")
        _run_campo_manual(p3, "Referencia documental da solicitacao da contratada")
        _adicionar_run(p3, ";")
    else:
        ciclos = dados.get("ciclos_computados") or []
        situacoes = [
            remover_emojis_leve(c.get("situacao") or "").strip().upper()
            for c in ciclos
        ]
        tempestiva = bool(situacoes) and all(
            situacao.startswith("TEMPESTIVO") for situacao in situacoes
        )
        datas = {
            str(c.get("data_pedido") or "").strip()
            for c in ciclos
            if str(c.get("data_pedido") or "").strip()
            not in ("", NAO_INFORMADO)
        }
        data_canonica = next(iter(datas)) if len(datas) == 1 else None
        _adicionar_run(
            p3,
            "A solicitação tempestiva da CONTRATADA, de " if tempestiva
            else "A solicitação da CONTRATADA, de ",
        )
        _texto_ou_marcador(
            p3,
            data_canonica or _campo(cm, "solicitacao_data"),
            "Data da solicitacao da contratada",
        )
        _adicionar_run(p3, ", instruída em ")
        _texto_ou_marcador(
            p3,
            _campo(cm, "solicitacao_ref"),
            "Referencia documental da solicitacao da contratada",
        )
        _adicionar_run(p3, ";")

    p4 = item("4")
    _adicionar_run(p4,
        "A necessidade de distinguir o histórico já formalizado anteriormente do "
        "objeto da presente análise, evitando duplicidade de contagem ou "
        "sobreposição de efeitos financeiros;")

    p5 = item("5")
    if branco:
        _adicionar_run(p5,
            "As informações da área gestora que vierem a fundamentar a "
            "formalização deverão abranger, quando aplicável, a execução, o "
            "saldo remanescente, os itens contratuais, os aditivos/supressões e "
            "os documentos de suporte da apuração;")
    else:
        _adicionar_run(p5,
            "As informações encaminhadas pela área gestora do contrato quanto à "
            "execução, ao saldo remanescente, aos itens contratuais, aos "
            "aditivos/supressões e aos documentos de suporte da apuração;")

    p6 = item("6")
    if branco:
        # Etapa 29C.1.2: nao afirma memoria existente nem valores apurados.
        _adicionar_run(p6, "A memória de cálculo a ser indicada em ")
        _run_campo_manual(p6, "Referencia da memoria de calculo")
        _adicionar_run(p6,
            " deverá apresentar, quando aplicável, os ciclos de reajuste, os "
            "percentuais aplicáveis, os efeitos financeiros, o eventual "
            "retroativo e a composição do Valor Total Atualizado do Contrato;")
    else:
        _adicionar_run(p6, "A memória de cálculo constante em ")
        _texto_ou_marcador(p6, _campo(cm, "memoria_calculo_ref"), "Referencia da memoria de calculo")
        _adicionar_run(p6,
            ", que apurou os ciclos de reajuste, os percentuais aplicáveis, os "
            "efeitos financeiros, o saldo retroativo a pagar e a composição do Valor "
            "Total Atualizado do Contrato;")

    p7 = item("7")
    if branco:
        # Etapa 29C.1: o considerando em branco nao afirma analise nem apuracao
        # ocorridas — apenas indica os campos a preencher (mesmos 2 placeholders).
        _adicionar_run(p7,
            "O índice contratual e o percentual aplicável deverão ser informados "
            "nos campos a seguir: ")
        _run_campo_manual(p7, "Indice contratual")
        _adicionar_run(p7, " e ")
        _run_campo_manual(p7, "Percentual aplicavel")
        _adicionar_run(p7, ";")
    else:
        _adicionar_run(p7, "O índice contratual utilizado na análise, qual seja ")
        indice = _indice_doc(dados)
        if indice:
            _adicionar_run(p7, indice, negrito=True)
        else:
            _run_campo_manual(p7, "Indice contratual")
        _adicionar_run(p7, ", e o percentual acumulado apurado de ")
        var = dados.get("var_acumulada")
        if var is not None:
            _adicionar_run(p7, _fmt_pct_doc(var), negrito=True)
        else:
            _run_campo_manual(p7, "Percentual acumulado apurado")
        _adicionar_run(p7, ";")

    p8 = item("8")
    if branco:
        _adicionar_run(p8, "A concordância da CONTRATADA, se aplicável, deverá ser registrada em ")
        _run_campo_manual(p8, "Referencia da manifestacao de concordancia da contratada")
        _adicionar_run(p8, ";")
    else:
        _adicionar_run(p8, "A manifestação de concordância da CONTRATADA constante em ")
        _texto_ou_marcador(p8, _campo(cm, "concordancia_ref"), "Referencia da manifestacao de concordancia da contratada")
        _adicionar_run(p8, ";")

    p9 = item("9")
    if branco:
        _adicionar_run(p9, "As certidões de regularidade da CONTRATADA a referenciar em ")
        _run_campo_manual(p9, "Referencia das certidoes de regularidade")
        _adicionar_run(p9, " e a adequação orçamentária a referenciar em ")
        _run_campo_manual(p9, "Referencia da adequacao orcamentaria")
        _adicionar_run(p9, ".")
    else:
        _adicionar_run(p9, "As certidões de regularidade ")
        _texto_ou_marcador(p9, _campo(cm, "regularidade_ref"), "Referencia das certidoes de regularidade")
        _adicionar_run(p9, " da CONTRATADA e a adequação orçamentária ")
        _texto_ou_marcador(p9, _campo(cm, "adequacao_orcamentaria_ref"), "Referencia da adequacao orcamentaria")
        _adicionar_run(p9, ".")
    doc.add_paragraph()


def _ta_abertura(doc: Document) -> None:
    _titulo_secao(doc, "FORMALIZA-SE O PRESENTE TERMO DE APOSTILA:")
    doc.add_paragraph()


def _ta_secao1_reajustes(doc: Document, dados: dict, cm: dict) -> None:
    _titulo_secao(doc, "1. Dos reajustes concedidos")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    branco = dados.get("_modo_branco")
    _adicionar_run(p, "1.1. Ao Contrato nº ")
    _texto_ou_marcador(p, _campo(cm, "contrato"), "Numero do contrato")
    if branco:
        _adicionar_run(p,
            ", os reajustes a serem formalizados deverão ser indicados no "
            "Quadro 1.")
    else:
        _adicionar_run(p,
            ", formalizam-se os reajustes contratuais apurados, conforme Quadro 1.")

    _titulo_quadro(doc, "Quadro 1 — Síntese dos reajustes concedidos")
    cabecalho = ["Ref.", "Ciclo", "Percentual aplicado", "Efeitos financeiros", "Situação"]
    if branco:
        _adicionar_tabela(doc, cabecalho, [[
            "[PREENCHER: Ref.]", "[PREENCHER: Ciclo]",
            "[PREENCHER: Percentual aplicável]", "[PREENCHER: Efeitos financeiros]",
            "[PREENCHER: Situação]",
        ]], destacar_placeholders=True)
        _paragrafos_perda_efeitos(doc, dados)
        doc.add_paragraph()
        return
    linhas: list[list[str]] = []
    ciclos = dados.get("ciclos_computados") or []
    for i, c in enumerate(ciclos):
        pct = c.get("percentual_reajuste")
        linhas.append([
            _LETRAS[i] if i < len(_LETRAS) else str(i + 1),
            remover_emojis_leve(c.get("ciclo") or ""),
            _fmt_pct_doc(pct) if pct is not None else NAO_INFORMADO,
            _efeito_financeiro_ciclo(c),
            remover_emojis_leve(c.get("situacao") or NAO_INFORMADO),
        ])
    ref_acum = _LETRAS[len(ciclos)] if len(ciclos) < len(_LETRAS) else "Acum."
    var = dados.get("var_acumulada")
    linhas.append([
        ref_acum,
        "Acumulado",
        _fmt_pct_doc(var) if var is not None else NAO_INFORMADO,
        "Conforme composição dos ciclos",
        "Percentual acumulado apurado",
    ])
    _adicionar_tabela(doc, cabecalho, linhas)
    _paragrafos_perda_efeitos(doc, dados)
    doc.add_paragraph()


def _ta_secao2_retroativo(doc: Document, dados: dict, cm: dict) -> None:
    _titulo_secao(doc, "2. Da apuração financeira do retroativo")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if dados.get("_modo_branco"):
        # Etapa 29C.1.2: nao afirma apuracao, diferenca nem retroativo; mesmos
        # tres placeholders financeiros, destacados.
        _adicionar_run(p,
            "2.1. Os valores financeiros que eventualmente integrem a "
            "formalização deverão ser informados nos campos e quadros desta "
            "seção, incluindo o valor pago efetivo ")
        _run_campo_manual(p, "Valor pago efetivo")
        _adicionar_run(p, ", o valor devido após o reajuste ")
        _run_campo_manual(p, "Valor devido apos o reajuste")
        _adicionar_run(p, " e a diferença ou retroativo ")
        _run_campo_manual(p, "Valor retroativo a pagar")
        _adicionar_run(p, ", quando aplicável, conforme Quadro 2.")
    else:
        _adicionar_run(p, "2.1. A apuração financeira consolidada indicou valor pago efetivo de ")
        vp = _valor_pago_total(dados)
        if vp is not None:
            _adicionar_run(p, formatar_moeda(vp), negrito=True)
        else:
            _valor_moeda_ou_marcador(p, _campo(cm, "valor_pago_efetivo"), "Valor pago efetivo")
        _adicionar_run(p, " e valor devido após o reajuste de ")
        vat = _valor_atualizado_total(dados)
        if vat is not None:
            _adicionar_run(p, formatar_moeda(vat), negrito=True)
        else:
            _valor_moeda_ou_marcador(p, _campo(cm, "valor_teorico"), "Valor devido apos o reajuste")
        _adicionar_run(p, ", resultando em valor retroativo a pagar de ")
        retro = _retroativo_total(dados)
        if retro is not None:
            _adicionar_run(p, formatar_moeda(retro), negrito=True)
        else:
            _run_campo_manual(p, "Valor retroativo a pagar")
        _adicionar_run(p, ", conforme Quadro 2.")

    _titulo_quadro(doc, "Quadro 2 — Apuração financeira por ciclo")
    cabecalho = ["Ciclo", "Valor pago efetivo", "Valor devido após o reajuste", "Diferença/retroativo"]
    linhas: list[list[str]] = []
    tot_pago = tot_teorico = tot_delta = None
    for lin in _linhas_financeiro(dados):
        vpg = _num_ou_none(lin.get("valor_pago"))
        vtc = _num_ou_none(lin.get("valor_atualizado"))
        vdl = _num_ou_none(lin.get("delta"))
        linhas.append([
            remover_emojis_leve(lin.get("ciclo") or ""),
            formatar_moeda(vpg) if vpg is not None else "",
            formatar_moeda(vtc) if vtc is not None else "",
            formatar_moeda(vdl) if vdl is not None else "",
        ])
        if vpg is not None:
            tot_pago = (tot_pago or 0.0) + vpg
        if vtc is not None:
            tot_teorico = (tot_teorico or 0.0) + vtc
        if vdl is not None:
            tot_delta = (tot_delta or 0.0) + vdl
    if not linhas:
        linhas = [["—", "", "", ""]]
    total_delta = tot_delta if tot_delta is not None else _retroativo_total(dados)
    linhas.append([
        "Total",
        formatar_moeda(tot_pago) if tot_pago is not None else "",
        formatar_moeda(tot_teorico) if tot_teorico is not None else "",
        formatar_moeda(total_delta) if total_delta is not None else "",
    ])
    _adicionar_tabela(doc, cabecalho, linhas)
    doc.add_paragraph()


def _ta_secao3_composicao_vta(doc: Document, dados: dict) -> None:
    _titulo_secao(doc, "3. Da composição do Valor Total Atualizado")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if dados.get("_modo_branco"):
        # Etapa 29C.1.2: descreve a organizacao a ser dada, nao a ja realizada.
        _adicionar_run(p,
            "3.1. Para fins de consolidação contratual, a composição do Valor "
            "Total Atualizado deverá ser organizada de forma evolutiva, "
            "demonstrando, quando aplicável, a execução por ciclo, os "
            "remanescentes intermediários, o saldo remanescente final e os "
            "aditivos/supressões computáveis.")
    else:
        _adicionar_run(p,
            "3.1. Para fins de consolidação contratual, a composição do Valor "
            "Total Atualizado foi organizada de forma evolutiva, demonstrando a "
            "execução por ciclo, os remanescentes intermediários, o saldo "
            "remanescente final e os aditivos/supressões computáveis, quando "
            "aplicáveis.")

    _titulo_quadro(doc, "Quadro 3 — Composição do Valor Total Atualizado")
    cabecalho = ["Ref.", "Descrição", "Valor"]
    linhas: list[list[str]] = []
    for i, (desc, valor) in enumerate(_composicao_didatica_vta(dados)):
        linhas.append([
            _LETRAS[i] if i < len(_LETRAS) else str(i + 1),
            desc,
            formatar_moeda(valor) if valor is not None else "",
        ])
    linhas.append([
        "Total",
        "Valor Total Atualizado do Contrato",
        _vta_texto_doc(dados),
    ])
    _adicionar_tabela(doc, cabecalho, linhas)
    doc.add_paragraph()
    _adicionar_box_retroativos(doc, dados, saneador=False)


# VTA-U2 (regra global): o quadro das "tres referencias do VTA"
# (posicao atual / ultima abertura / contrato integralmente reajustado)
# foi removido dos documentos externos. Era codigo morto — nenhum gerador
# o chamava — e apresentava valores que podiam ser lidos como um segundo
# VTA. Os dados seguem em leitura['referencias_vta'] para auditoria.


def _ta_secao4_valores_unitarios(doc: Document, dados: dict) -> None:
    _titulo_secao(doc, "4. Dos valores unitários")
    _secao_valores_unitarios_por_ciclo(
        doc, dados,
        texto_intro=(
            "4.1. Os valores unitários dos itens, considerados os ciclos de "
            "reajuste aplicáveis até a presente atualização, ficam consolidados "
            "conforme quadro abaixo."
        ),
    )


def _sintese_aditivos_por_ciclo(aditivos: list[dict]) -> list[str]:
    """Rotulos executivos por ciclo, sem expor chaves tecnicas internas."""
    grupos: dict[str, dict[str, Any]] = {}
    for ad in aditivos:
        ciclo = remover_emojis_leve(ad.get("ciclo") or "Sem ciclo").strip()
        grupo = grupos.setdefault(
            ciclo, {"total": 0.0, "tem_valor": False, "tipos": {}}
        )
        tipo = remover_emojis_leve(
            ad.get("tipo_alteracao") or "Alteração"
        ).strip()
        chave_tipo = "supressao" if "supr" in tipo.lower() else (
            "acrescimo" if "acresc" in tipo.lower()
            or "acrésc" in tipo.lower() else "alteracao"
        )
        grupo["tipos"][chave_tipo] = grupo["tipos"].get(chave_tipo, 0) + 1
        valor = _num_ou_none(ad.get("valor_atualizado"))
        if valor is not None:
            grupo["total"] += valor
            grupo["tem_valor"] = True

    def _ordem(ciclo: str) -> tuple[int, str]:
        texto = ciclo.upper()
        return (
            int(texto[1]) if len(texto) == 2 and texto[0] == "C"
            and texto[1].isdigit() else 99,
            texto,
        )

    saida: list[str] = []
    for ciclo in sorted(grupos, key=_ordem):
        grupo = grupos[ciclo]
        tipos = grupo["tipos"]
        partes_tipo = []
        if tipos.get("acrescimo"):
            n = tipos["acrescimo"]
            partes_tipo.append(f"{n} acréscimo" + ("" if n == 1 else "s"))
        if tipos.get("supressao"):
            n = tipos["supressao"]
            partes_tipo.append(f"{n} supressão" if n == 1 else f"{n} supressões")
        if tipos.get("alteracao"):
            n = tipos["alteracao"]
            partes_tipo.append(f"{n} alteração" if n == 1 else f"{n} alterações")
        tipos_txt = " e ".join(partes_tipo)
        impacto = (
            f"impacto atualizado total {formatar_moeda(grupo['total'])}"
            if grupo["tem_valor"] else "impacto a confirmar"
        )
        saida.append(f"{ciclo} — {tipos_txt} — {impacto}")
    return saida


def _ta_secao5_aditivos(doc: Document, dados: dict) -> None:
    _titulo_secao(doc, "5. Dos aditivos e supressões considerados")
    aditivos = dados.get("aditivos") or []
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if dados.get("_modo_branco"):
        _adicionar_run(p1, "5.1. Registrar os aditivos e supressões considerados: ")
        _run_campo_manual(p1, "Aditivos e supressões considerados")
        _adicionar_run(p1, ".")
    elif not aditivos:
        _adicionar_run(p1,
            "5.1. Não foram identificados aditivos ou supressões específicos na "
            "base processada, sem prejuízo da conferência dos instrumentos já "
            "formalizados no processo.")
    else:
        _adicionar_run(
            p1,
            "5.1. Foram consideradas as alterações contratuais registradas "
            "na apuração: " + "; ".join(_sintese_aditivos_por_ciclo(aditivos))
            + ".",
        )

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _adicionar_run(p2,
        "5.2. Os aditivos e supressões computáveis integram o Valor Total "
        "Atualizado quando não estiverem refletidos na execução atualizada, no "
        "saldo remanescente ou no valor formalizado anterior, vedada a dupla "
        "contagem.")
    espaco_apos_capitulo = doc.add_paragraph()
    espaco_apos_capitulo.paragraph_format.space_after = Pt(6)


def _ta_secoes_finais(doc: Document, cm: dict) -> None:
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _adicionar_run(p6, "6. ", negrito=True)
    _adicionar_run(p6,
        "Permanecem inalteradas e em pleno vigor as demais cláusulas e condições "
        "do Contrato e de seus instrumentos posteriores não modificadas por este "
        "Termo de Apostila.")

    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _adicionar_run(p7, "7. ", negrito=True)
    _adicionar_run(p7,
        "A CONTRATADA deverá atualizar a garantia contratual, prevista na "
        "cláusula própria do Contrato, no prazo contratualmente estabelecido, "
        "observado o novo valor após a formalização deste Termo de Apostila.")

    p8 = doc.add_paragraph()
    p8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _adicionar_run(p8, "8. ", negrito=True)
    _adicionar_run(p8,
        "O presente apostilamento vincula-se, para todos os fins, aos documentos "
        "instruídos no Processo ")
    _texto_ou_marcador(p8, _campo(cm, "processo_ref"), "Numero do processo de instrucao")
    _adicionar_run(p8, ".")
    doc.add_paragraph()


def _ta_assinaturas(doc: Document, cm: dict) -> None:
    p_local = doc.add_paragraph()
    p_local.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _adicionar_run(p_local, "Brasília/DF, ")
    _texto_ou_marcador(p_local, _campo(cm, "local_data"), "Data")
    _adicionar_run(p_local, ".")
    doc.add_paragraph()
    doc.add_paragraph()

    # Dois representantes da TELEBRAS (nenhuma assinatura da CONTRATADA).
    for chave_nome, desc_nome in (
        ("representante_telebras_1_nome", "Nome do 1o representante da Telebras"),
        ("representante_telebras_2_cargo", "Cargo do 2o representante da Telebras"),
    ):
        p_ent = doc.add_paragraph()
        p_ent.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _adicionar_run(p_ent, "TELECOMUNICAÇÕES BRASILEIRAS S.A. - TELEBRAS")
        p_rep = doc.add_paragraph()
        p_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _texto_ou_marcador(p_rep, _campo(cm, chave_nome), desc_nome)
        doc.add_paragraph()


# ---------------------------------------------------------------------------
# DESPACHO SANEADOR (modelo canonico §7)
# ---------------------------------------------------------------------------

def gerar_despacho_saneador(
    leitura_ou_objeto: dict,
    identificacao: dict | None = None,
    campos_manuais: dict | None = None,
    *,
    modo_modelo_em_branco: bool = False,
) -> bytes:
    """Gera o Despacho Saneador em DOCX e retorna os bytes.

    `modo_modelo_em_branco=True` reutiliza a mesma estrutura enxuta sem afirmar
    fatos não comprovados. Nos dois modos, o gerador apenas apresenta dados já
    consolidados; não recalcula valores nem cria classificação processual.
    """
    if campos_manuais is None:
        campos_manuais = {}
    dados = _extrair_dados(leitura_ou_objeto, identificacao)
    dados["_modo_branco"] = bool(modo_modelo_em_branco)
    doc = _configurar_documento()

    _ds_assunto_enxuto(doc, dados, campos_manuais)
    _ds_secao1_identificacao(doc, dados, campos_manuais)
    _ds_secao2_pedido_parametros(doc, dados, campos_manuais)
    _ds_secao3_resultado(doc, dados, campos_manuais)
    _ds_secao4_documentos(doc, dados, campos_manuais)
    _ds_secao5_pendencias(doc, dados, campos_manuais)
    _ds_secao6_conclusao(doc, dados, campos_manuais)
    _adicionar_id_apuracao_rodape(doc, dados)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _ds_valor_identificacao(dados: dict, cm: dict, chave_manual: str,
                            *aliases: str) -> Any:
    externa = dados.get("identificacao_externa") or {}
    for alias in aliases:
        valor = externa.get(alias)
        if valor is not None and str(valor).strip():
            return valor
    return _campo(cm, chave_manual)


def _ds_texto_ou_tag(valor: Any, descricao: str) -> str:
    if valor is None or not str(valor).strip() or str(valor).strip() == NAO_INFORMADO:
        return PREENCHER_TAG.format(descricao)
    if hasattr(valor, "strftime"):
        try:
            return valor.strftime("%d/%m/%Y")
        except (TypeError, ValueError):
            pass
    return remover_emojis_leve(valor).strip()


def _ds_tipo_atualizacao(dados: dict, cm: dict) -> str | None:
    valor = _ds_valor_identificacao(
        dados, cm, "tipo_atualizacao",
        "tipo_atualizacao", "tipo_instrumento", "tipo_analise",
    )
    if valor is not None and str(valor).strip():
        return remover_emojis_leve(valor).strip()
    if dados.get("_modo_branco"):
        return None
    return "atualização contratual"


def _ds_assunto_enxuto(doc: Document, dados: dict, cm: dict) -> None:
    _titulo_secao(doc, "DESPACHO SANEADOR", tamanho=12,
                  alinhamento=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _adicionar_run(p, "Assunto: ", negrito=True)
    _adicionar_run(p, "Saneamento para formalização de ")
    tipo = _ds_tipo_atualizacao(dados, cm)
    if tipo:
        _adicionar_run(p, tipo)
    else:
        _run_campo_manual(p, "Tipo ou instrumento")
    _adicionar_run(p, " — ")
    contrato = _ds_valor_identificacao(
        dados, cm, "contrato", "contrato", "numero_contrato"
    )
    _texto_ou_marcador(p, contrato, "Numero do contrato")
    _adicionar_run(p, ".")

    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _adicionar_run(p_ref, "Referência(s): ", negrito=True)
    _texto_ou_marcador(
        p_ref, _campo(cm, "processo_pleito"), "Referencia do pedido"
    )
    doc.add_paragraph()


def _ds_titulo(doc: Document, numero: int, texto: str) -> Any:
    p = _titulo_secao(doc, f"{numero}. {texto.upper()}", tamanho=11)
    p.paragraph_format.keep_with_next = True
    return p


def _ds_titulo_quadro(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    _adicionar_run(p, texto, negrito=True, tamanho=10)


def _ds_secao1_identificacao(doc: Document, dados: dict, cm: dict) -> None:
    _ds_titulo(doc, 1, "Identificação")
    contrato = _ds_valor_identificacao(
        dados, cm, "contrato", "contrato", "numero_contrato"
    )
    contratada = _ds_valor_identificacao(
        dados, cm, "empresa_contratada",
        "empresa_contratada", "contratada",
    )
    objeto = _ds_valor_identificacao(
        dados, cm, "objeto_contrato", "objeto_contrato", "objeto"
    )
    vigencia = _ds_valor_identificacao(
        dados, cm, "vigencia_ate",
        "vigencia_ate", "fim_vigencia", "data_fim_vigencia",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _adicionar_run(p, "Realiza-se o saneamento processual do Contrato nº ")
    _texto_ou_marcador(p, contrato, "Numero do contrato")
    _adicionar_run(p, ", celebrado com ")
    _texto_ou_marcador(p, contratada, "Nome da empresa contratada")
    _adicionar_run(p, ", cujo objeto é ")
    _texto_ou_marcador(p, objeto, "Objeto resumido do contrato")
    _adicionar_run(p, ", com vigência até ")
    if vigencia is not None:
        _adicionar_run(p, _ds_texto_ou_tag(vigencia, "Data final da vigencia contratual"))
    else:
        _run_campo_manual(p, "Data final da vigencia contratual")
    _adicionar_run(p, ".")


def _ds_ciclos_relevantes(dados: dict) -> list[dict]:
    return list(dados.get("ciclos_computados") or [])


def _ds_secao2_pedido_parametros(doc: Document, dados: dict, cm: dict) -> None:
    _ds_titulo(doc, 2, "Pedido e parâmetros da análise")
    branco = bool(dados.get("_modo_branco"))
    ciclos = _ds_ciclos_relevantes(dados)
    tipo = _ds_tipo_atualizacao(dados, cm)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if branco:
        _adicionar_run(p, "Deverão ser registrados o pedido de ")
        _run_campo_manual(p, "Tipo da atualizacao contratual")
        _adicionar_run(
            p,
            ", a data correspondente e a respectiva referência documental. "
            "Os parâmetros da análise deverão constar do Quadro 1.",
        )
    else:
        _adicionar_run(p, "A CONTRATADA apresentou pedido de ")
        _adicionar_run(p, tipo or "atualização contratual")
        if len(ciclos) == 1:
            _adicionar_run(p, " em ")
            pedido = ciclos[0].get("data_pedido")
            if pedido and pedido != NAO_INFORMADO:
                _adicionar_run(p, pedido)
            else:
                _run_campo_manual(p, "Data do pedido")
        elif ciclos:
            _adicionar_run(p, " nas datas indicadas no Quadro 1")
        else:
            _adicionar_run(p, " em ")
            _run_campo_manual(p, "Data do pedido")
        _adicionar_run(p, ", conforme ")
        _texto_ou_marcador(p, _campo(cm, "processo_pleito"), "Referencia do pedido")
        if not ciclos:
            _adicionar_run(
                p,
                ". Os parâmetros de ciclo não foram disponibilizados pela fonte "
                "canônica; o Quadro 1 deve ser complementado.",
            )
        elif len(ciclos) == 1:
            _adicionar_run(p, ". A análise considerou ")
            ciclo = ciclos[0]
            _adicionar_run(p, "a data-base ")
            _adicionar_run(p, _ds_texto_ou_tag(ciclo.get("data_inicio"), "Data-base"))
            _adicionar_run(p, ", a referência econômica ")
            indice = _indice_doc(dados)
            if indice:
                _adicionar_run(p, indice)
            else:
                _run_campo_manual(p, "Indice ou referencia economica")
            _adicionar_run(p, " e classificou o pedido como ")
            situacao = ciclo.get("situacao")
            if situacao and situacao != NAO_INFORMADO:
                _adicionar_run(p, remover_emojis_leve(situacao))
            else:
                _run_campo_manual(p, "Situacao do pedido")
            _adicionar_run(p, ", com efeitos financeiros ")
            efeito = _efeito_financeiro_ciclo(ciclo)
            if efeito != NAO_INFORMADO:
                _adicionar_run(p, efeito.lower())
            else:
                _run_campo_manual(p, "Data ou competencia do efeito financeiro")
            _adicionar_run(p, ".")
        else:
            _adicionar_run(p, ". A análise considerou ")
            _adicionar_run(
                p,
                "as datas-base, situações e efeitos financeiros indicados no "
                "Quadro 1, com referência econômica ",
            )
            indice = _indice_doc(dados)
            if indice:
                _adicionar_run(p, indice)
            else:
                _run_campo_manual(p, "Indice ou referencia economica")
            _adicionar_run(p, ".")

    _paragrafos_perda_efeitos(doc, dados)

    _ds_titulo_quadro(doc, "Quadro 1 - Síntese da análise")
    cabecalho = [
        "Ciclo", "Data-base", "Data do pedido", "Situação",
        "Efeito financeiro", "Percentual",
    ]
    if branco:
        linhas = [[PREENCHER_TAG.format(rotulo) for rotulo in cabecalho]]
    else:
        linhas = []
        for ciclo in ciclos:
            pct = ciclo.get("percentual_reajuste")
            linhas.append([
                _ds_texto_ou_tag(ciclo.get("ciclo"), "Ciclo"),
                _ds_texto_ou_tag(ciclo.get("data_inicio"), "Data-base"),
                _ds_texto_ou_tag(ciclo.get("data_pedido"), "Data do pedido"),
                _ds_texto_ou_tag(
                    remover_emojis_leve(ciclo.get("situacao") or ""), "Situacao"
                ),
                _ds_texto_ou_tag(
                    _efeito_financeiro_ciclo(ciclo), "Efeito financeiro"
                ),
                _fmt_pct_doc(pct) if pct is not None
                else PREENCHER_TAG.format("Percentual"),
            ])
        if not linhas:
            linhas = [[PREENCHER_TAG.format(rotulo) for rotulo in cabecalho]]
    _adicionar_tabela(
        doc, cabecalho, linhas,
        destacar_placeholders=True,
        destacar_placeholders_embutidos=True,
    )
    doc.add_paragraph()


def _ds_total_presente(dados: dict, chave: str) -> float | None:
    linhas = _linhas_financeiro(dados)
    if not linhas:
        return None
    valores = []
    for linha in linhas:
        valor = _num_ou_none(linha.get(chave))
        if valor is None:
            return None
        valores.append(valor)
    return round(sum(valores), 2)


def _ds_secao3_resultado(doc: Document, dados: dict, cm: dict) -> None:
    _ds_titulo(doc, 3, "Resultado essencial")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _adicionar_run(p, "Conforme memória de cálculo ")
    memoria_ref = _campo(cm, "memoria_calculo_ref") or _campo(cm, "referencia_analise")
    if dados.get("_modo_branco"):
        _run_campo_manual(p, "Referencia da memoria de calculo")
        _adicionar_run(
            p, ", os resultados essenciais deverão ser preenchidos no Quadro 2."
        )
    else:
        _texto_ou_marcador(p, memoria_ref, "Referencia da memoria de calculo")
        _adicionar_run(p, ", a apuração apresentou os resultados abaixo.")

    tipo = (_ds_tipo_atualizacao(dados, cm) or "").strip().lower()
    rotulo_devido = (
        "Valor devido após o reajuste"
        if tipo == "reajuste"
        else "Valor devido após a atualização contratual"
    )
    pago = _ds_total_presente(dados, "valor_pago")
    devido = _ds_total_presente(dados, "valor_atualizado")
    retro = _retroativo_total(dados)
    vta = _vta_texto_doc(dados)
    linhas = [
        [
            "Valor pago no período analisado",
            formatar_moeda(pago) if pago is not None
            else PREENCHER_TAG.format("Valor pago no periodo analisado"),
        ],
        [
            rotulo_devido,
            formatar_moeda(devido) if devido is not None
            else PREENCHER_TAG.format("Valor devido apos a atualizacao contratual"),
        ],
        [
            "Retroativo a pagar",
            formatar_moeda(retro) if retro is not None
            else PREENCHER_TAG.format("Valor retroativo a pagar"),
        ],
        [
            "Valor Total Atualizado do Contrato",
            vta or PREENCHER_TAG.format("Valor Total Atualizado do Contrato"),
        ],
    ]
    _ds_titulo_quadro(doc, "Quadro 2 - Síntese financeira")
    _adicionar_tabela(
        doc, ["Resultado", "Valor"], linhas,
        destacar_placeholders=True,
        destacar_placeholders_embutidos=True,
    )
    doc.add_paragraph()
    _adicionar_box_retroativos(doc, dados, saneador=True)


def _ds_juntar_campos(*valores: tuple[Any, str]) -> str:
    partes = []
    for valor, descricao in valores:
        if valor is None or not str(valor).strip():
            partes.append(PREENCHER_TAG.format(descricao))
        elif isinstance(valor, (int, float)) and not isinstance(valor, bool):
            partes.append(formatar_moeda(valor))
        else:
            partes.append(remover_emojis_leve(valor).strip())
    return " / ".join(partes)


def _ds_secao4_documentos(doc: Document, dados: dict, cm: dict) -> None:
    _ds_titulo(doc, 4, "Documentos e verificações")
    memoria_ref = _campo(cm, "memoria_calculo_ref") or _campo(cm, "referencia_analise")
    linhas = [
        ["Memória de cálculo", _ds_juntar_campos(
            (memoria_ref, "Referencia da memoria de calculo"),
        )],
        ["Adequação orçamentária", _ds_juntar_campos(
            (_campo(cm, "adequacao_orcamentaria_ref"),
             "Referencia da adequacao orcamentaria"),
            (_campo(cm, "adequacao_orcamentaria_valor"),
             "Valor ou situacao da adequacao orcamentaria"),
        )],
        ["Regularidade da contratada", _ds_juntar_campos(
            (_campo(cm, "regularidade_ref"),
             "Referencia da regularidade da contratada"),
            (_campo(cm, "regularidade_situacao"),
             "Situacao da regularidade da contratada"),
        )],
        ["Concordância da contratada", _ds_juntar_campos(
            (_campo(cm, "concordancia_ref"),
             "Referencia da concordancia da contratada"),
            (_campo(cm, "concordancia_situacao"),
             "Situacao da concordancia da contratada"),
        )],
        ["Garantia contratual", _ds_juntar_campos(
            (_campo(cm, "garantia_situacao"),
             "Situacao da garantia contratual"),
        )],
    ]
    _ds_titulo_quadro(doc, "Quadro 3 - Documentos e verificações")
    _adicionar_tabela(
        doc, ["Documento ou verificação", "Referência ou situação"], linhas,
        destacar_placeholders=True,
        destacar_placeholders_embutidos=True,
    )
    doc.add_paragraph()


def _ds_pendencias_tecnicas(dados: dict, cm: dict) -> list[str]:
    resultado: list[str] = []
    pendencias = dados.get("pendencias") or {}
    for chave in ("bloqueantes", "advertencias"):
        for item in pendencias.get(chave) or []:
            texto = remover_emojis_leve(item).strip()
            if texto and texto not in resultado:
                resultado.append(texto)
    docs = _campo(cm, "docs_desatualizados")
    if docs:
        itens = docs if isinstance(docs, (list, tuple)) else [docs]
        texto = "Documentos desatualizados: " + ", ".join(str(item) for item in itens)
        resultado.append(remover_emojis_leve(texto))
    complemento = _campo(cm, "pendencias_complemento")
    if complemento:
        resultado.append(remover_emojis_leve(complemento).strip())
    if cm.get("pendencia_critica") and not resultado:
        resultado.append("Pendência impeditiva indicada para complementação.")
    return resultado


def _ds_secao5_pendencias(doc: Document, dados: dict, cm: dict) -> None:
    _ds_titulo(doc, 5, "Pendências")
    pendencias = _ds_pendencias_tecnicas(dados, cm)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if dados.get("_modo_branco"):
        _adicionar_run(p, "Registrar as pendências relevantes para o prosseguimento: ")
        _run_campo_manual(p, "Pendencias relevantes")
        _adicionar_run(p, ".")
    elif pendencias:
        _adicionar_run(p, "Pendências: " + "; ".join(pendencias) + ".")
    else:
        _adicionar_run(p, "Não foram identificadas pendências técnicas na apuração.")
    if _ds_ha_pendencia_documental(dados, cm):
        _adicionar_run(
            p,
            " Os campos documentais destacados permanecem sujeitos a "
            "preenchimento e conferência.",
        )


def _ds_ha_pendencia_documental(dados: dict, cm: dict) -> bool:
    if dados.get("_modo_branco"):
        return True
    memoria_ref = _campo(cm, "memoria_calculo_ref") or _campo(cm, "referencia_analise")
    obrigatorios = (
        memoria_ref,
        _campo(cm, "adequacao_orcamentaria_ref"),
        _campo(cm, "adequacao_orcamentaria_valor"),
        _campo(cm, "regularidade_ref"),
        _campo(cm, "regularidade_situacao"),
        _campo(cm, "concordancia_ref"),
        _campo(cm, "concordancia_situacao"),
        _campo(cm, "garantia_situacao"),
    )
    return any(
        valor is None or (isinstance(valor, str) and not valor.strip())
        for valor in obrigatorios
    )


def _ds_tem_pendencia_impeditiva(dados: dict, cm: dict) -> bool:
    pendencias = dados.get("pendencias") or {}
    if pendencias.get("bloqueantes"):
        return True
    flag = cm.get("pendencia_critica")
    if isinstance(flag, str):
        return flag.strip().lower() in ("sim", "true", "1", "critica", "critico")
    return bool(flag)


def _ds_secao6_conclusao(doc: Document, dados: dict, cm: dict) -> None:
    _ds_titulo(doc, 6, "Conclusão")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if dados.get("_modo_branco"):
        _adicionar_run(
            p,
            "Após o preenchimento e a conferência das informações, deverá ser "
            "avaliado se a instrução reúne condições para prosseguir à "
            "formalização.",
        )
    elif _ds_tem_pendencia_impeditiva(dados, cm):
        _adicionar_run(
            p,
            "A instrução deverá ser complementada quanto às pendências acima "
            "antes do prosseguimento para formalização.",
        )
    else:
        _adicionar_run(
            p,
            "Após a complementação e conferência das informações documentais "
            "indicadas, deverá ser avaliado o prosseguimento da instrução para "
            "formalização.",
        )


def _ds_assunto(doc: Document, cm: dict) -> None:
    _titulo_secao(doc, "DESPACHO SANEADOR", tamanho=12,
                  alinhamento=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _adicionar_run(p, "Assunto: ", negrito=True)
    _adicionar_run(p, "Saneamento para formalização de Termo de Apostila de Reajuste - ")
    _texto_ou_marcador(p, _campo(cm, "contrato"), "Numero do contrato")
    _adicionar_run(p, ".")
    doc.add_paragraph()


def _ds_par(doc: Document, numero: str) -> Any:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _adicionar_run(p, f"{numero}. ", negrito=True)
    return p


def _ds_par1(doc: Document, dados: dict) -> None:
    p = _ds_par(doc, "1")
    if dados.get("_modo_branco"):
        # Etapa 29C.1.2: sentido de finalidade, sem afirmar consolidacao feita.
        _adicionar_run(p,
            "Este modelo de despacho saneador destina-se à consolidação dos "
            "elementos documentais, financeiros e formais necessários à "
            "instrução de eventual Termo de Apostila de reajuste contratual. "
            "Os campos destacados devem ser revisados e preenchidos.")
        return
    _adicionar_run(p,
        "Este despacho saneador consolida os elementos documentais, financeiros "
        "e formais necessários à instrução do Termo de Apostila destinado ao "
        "registro de reajuste contratual, com a finalidade de demonstrar a "
        "regularidade mínima da instrução antes da formalização.")


def _ds_par2(doc: Document, dados: dict, cm: dict) -> None:
    p = _ds_par(doc, "2")
    if dados.get("_modo_branco"):
        # Etapa 29C.1.2: nao afirma pleito apresentado nem datas consideradas.
        _adicionar_run(p, "Registrar a referência do eventual pleito da contratada: ")
        _run_campo_manual(p, "Referencias do pleito da contratada")
        _adicionar_run(p,
            ". Para a verificação da anualidade, deverão ser informados a data "
            "da proposta em ")
        _run_campo_manual(p, "Data da proposta")
        _adicionar_run(p, ", o índice contratual ")
        _run_campo_manual(p, "Indice contratual")
        _adicionar_run(p, " e as datas de pedido aplicáveis: ")
        _run_campo_manual(p, "Datas de pedido por ciclo")
        _adicionar_run(p, ".")
        return
    _adicionar_run(p, "A contratada apresentou pleito de reajuste por meio de ")
    _texto_ou_marcador(p, _campo(cm, "processo_pleito"), "Referencias do pleito da contratada")
    _adicionar_run(p,
        ". Para fins de verificação da anualidade, foram consideradas a data da "
        "proposta em ")
    _texto_ou_marcador(p, _campo(cm, "data_proposta"), "Data da proposta")
    _adicionar_run(p, ", o índice contratual ")
    indice = _indice_doc(dados)
    if indice:
        _adicionar_run(p, indice, negrito=True)
    else:
        _run_campo_manual(p, "Indice contratual")
    _adicionar_run(p, " e as datas de pedido registradas na análise: ")
    datas = []
    for c in dados.get("ciclos_reajuste") or []:
        dp = str(c.get("data_pedido") or "").strip()
        if dp and dp != NAO_INFORMADO:
            datas.append(f"{remover_emojis_leve(c.get('ciclo') or '')} em {dp}")
    if datas:
        _adicionar_run(p, "; ".join(datas) + ".")
    else:
        _run_campo_manual(p, "Datas de pedido por ciclo")
        _adicionar_run(p, ".")


def _ds_par3(doc: Document, dados: dict, cm: dict) -> None:
    p = _ds_par(doc, "3")
    if dados.get("_modo_branco"):
        # Etapa 29C.1.2: nao afirma acordo, concessao nem apuracao.
        _adicionar_run(p,
            "Os ciclos que poderão integrar a formalização deverão ser "
            "conferidos nos quadros deste modelo e na seguinte referência: ")
        _run_campo_manual(p, "Referencia onde o resultado da analise consta")
        _adicionar_run(p, ".")
        return
    ciclos = dados.get("ciclos_computados") or []
    if ciclos:
        nomes = ", ".join(remover_emojis_leve(c.get("ciclo") or "") for c in ciclos)
        _adicionar_run(p, f"Acordou-se na concessão de {nomes}, conforme exposto em ")
    else:
        _adicionar_run(p, "Acordou-se na concessão dos ciclos apurados, conforme exposto em ")
    _texto_ou_marcador(p, _campo(cm, "referencia_analise"), "Referencia onde o resultado da analise consta")
    _adicionar_run(p, ".")


def _ds_par4_quadro1(doc: Document, dados: dict, cm: dict) -> None:
    branco = dados.get("_modo_branco")
    p = _ds_par(doc, "4")
    if branco:
        # Etapa 29C.1.2: valor original como instrucao ("deverá ser informado"),
        # sem afirmar que a informacao ja foi prestada.
        _adicionar_run(p, "A análise deverá registrar os ciclos indicados no Quadro 1 abaixo. Quantidade de ciclos: ")
        _run_campo_manual(p, "Quantidade de ciclos objeto da análise")
        _adicionar_run(p, ", com variação acumulada de ")
        _run_campo_manual(p, "Variacao acumulada")
        _adicionar_run(p, ". O valor original do contrato deverá ser informado: ")
        _run_campo_manual(p, "Valor original do contrato")
        _adicionar_run(p, ".")
    else:
        # Conta os ciclos efetivamente considerados, nao as linhas do Quadro 1:
        # ciclos "Fora da apuracao" continuam no quadro para rastreabilidade,
        # mas afirmar que foram considerados contradiz o proprio quadro.
        # Mesma fonte canonica ja usada no item 3 e no Quadro 1 do Termo.
        expressao = expressao_quantidade_ciclos(len(dados.get("ciclos_computados") or []))
        if expressao is None:
            _adicionar_run(p, FRASE_SEM_CICLOS_COMPUTADOS)
        else:
            _adicionar_run(p, f"A análise de reajuste considerou {expressao}, com variação acumulada de ")
            var = dados.get("var_acumulada")
            if var is not None:
                _adicionar_run(p, _fmt_pct_doc(var), negrito=True)
            else:
                _run_campo_manual(p, "Variacao acumulada")
            _adicionar_run(p, ".")
        _adicionar_run(p, " O valor original do contrato informado foi de ")
        _valor_moeda_ou_marcador(p, _campo(cm, "valor_original_contrato"), "Valor original do contrato")
        _adicionar_run(p, ".")

    _titulo_quadro(doc, "Quadro 1 - Síntese dos ciclos de reajuste")
    cabecalho = ["Ciclo", "Data-base", "Data do pedido", "Início financeiro",
                 "Fim financeiro", "Situação", "Percentual aplicado"]
    linhas: list[list[str]] = []
    for c in dados.get("ciclos_reajuste") or []:
        pct = c.get("percentual_reajuste")
        linhas.append([
            remover_emojis_leve(c.get("ciclo") or ""),
            c.get("data_inicio") or NAO_INFORMADO,
            c.get("data_pedido") or NAO_INFORMADO,
            c.get("inicio_efeito_financeiro") or NAO_INFORMADO,
            c.get("data_fim") or NAO_INFORMADO,
            remover_emojis_leve(c.get("situacao") or NAO_INFORMADO),
            _fmt_pct_doc(pct) if pct is not None else NAO_INFORMADO,
        ])
    if branco:
        _adicionar_tabela(doc, cabecalho, [[
            "[PREENCHER: Ciclo]", "[PREENCHER: Data-base]",
            "[PREENCHER: Data do pedido]", "[PREENCHER: Início financeiro]",
            "[PREENCHER: Fim financeiro]", "[PREENCHER: Situação]",
            "[PREENCHER: Percentual aplicável]",
        ]], destacar_placeholders=True)
        doc.add_paragraph()
        return
    if not linhas:
        linhas = [["—"] * 7]
    _adicionar_tabela(doc, cabecalho, linhas)
    doc.add_paragraph()


def _ds_par5_quadro2(doc: Document, dados: dict) -> None:
    p = _ds_par(doc, "5")
    if dados.get("_modo_branco"):
        # Etapa 29C.1.1: o item 5 em branco nao afirma apuracao/analise/
        # consolidacao realizadas — apenas instrui o preenchimento, com os
        # mesmos tres placeholders financeiros destacados.
        _adicionar_run(p,
            "Os valores da apuração financeira deverão ser preenchidos no "
            "quadro abaixo, incluindo, quando aplicável, o valor pago efetivo, "
            "o valor teórico calculado e a diferença ou retroativo "
            "correspondente: ")
        _run_campo_manual(p, "Valor pago efetivo")
        _adicionar_run(p, ", ")
        _run_campo_manual(p, "Valor teorico calculado")
        _adicionar_run(p, " e ")
        _run_campo_manual(p, "Valor retroativo a pagar")
        _adicionar_run(p, ".")
    else:
        _adicionar_run(p, "A apuração financeira consolidada indicou valor pago efetivo de ")
        vp = _valor_pago_total(dados)
        if vp is not None:
            _adicionar_run(p, formatar_moeda(vp), negrito=True)
        else:
            _run_campo_manual(p, "Valor pago efetivo")
        _adicionar_run(p, " e valor teórico calculado de ")
        vat = _valor_atualizado_total(dados)
        if vat is not None:
            _adicionar_run(p, formatar_moeda(vat), negrito=True)
        else:
            _run_campo_manual(p, "Valor teorico calculado")
        _adicionar_run(p, ", resultando em valor retroativo a pagar de ")
        retro = _retroativo_total(dados)
        if retro is not None:
            _adicionar_run(p, formatar_moeda(retro), negrito=True)
        else:
            _run_campo_manual(p, "Valor retroativo a pagar")
        _adicionar_run(p, ".")

    _titulo_quadro(doc, "Quadro 2 - Apuração financeira por ciclo")
    cabecalho = ["Ciclo", "Valor pago efetivo", "Valor teórico calculado", "Diferença/retroativo"]
    if dados.get("_modo_branco"):
        _adicionar_tabela(doc, cabecalho, [[
            "[PREENCHER: Ciclo]", "[PREENCHER: Valor pago efetivo]",
            "[PREENCHER: Valor teórico calculado]", "[PREENCHER: Valor retroativo]",
        ]], destacar_placeholders=True)
        doc.add_paragraph()
        return
    linhas: list[list[str]] = []
    tot_pago = tot_teorico = tot_delta = None
    for lin in _linhas_financeiro(dados):
        vpg = _num_ou_none(lin.get("valor_pago"))
        vtc = _num_ou_none(lin.get("valor_atualizado"))
        vdl = _num_ou_none(lin.get("delta"))
        linhas.append([
            remover_emojis_leve(lin.get("ciclo") or ""),
            formatar_moeda(vpg) if vpg is not None else "",
            formatar_moeda(vtc) if vtc is not None else "",
            formatar_moeda(vdl) if vdl is not None else "",
        ])
        if vpg is not None:
            tot_pago = (tot_pago or 0.0) + vpg
        if vtc is not None:
            tot_teorico = (tot_teorico or 0.0) + vtc
        if vdl is not None:
            tot_delta = (tot_delta or 0.0) + vdl
    if not linhas:
        linhas = [["—", "", "", ""]]
    linhas.append([
        "Total",
        formatar_moeda(tot_pago) if tot_pago is not None else "",
        formatar_moeda(tot_teorico) if tot_teorico is not None else "",
        formatar_moeda(tot_delta) if tot_delta is not None else "",
    ])
    _adicionar_tabela(doc, cabecalho, linhas)
    doc.add_paragraph()


def _ds_par6_quadro3(doc: Document, dados: dict, cm: dict) -> None:
    p = _ds_par(doc, "6")
    if dados.get("_modo_branco"):
        # Etapa 29C.1.2: a premissa de corte como obrigacao de informar, sem
        # afirmar que ja foi definida ou utilizada.
        _adicionar_run(p,
            "Deverá ser informada a premissa de corte a ser adotada para o "
            "cálculo do eventual retroativo e do valor remanescente do "
            "contrato: ")
        _run_campo_manual(p, "Descricao da data/posicao de corte adotada")
        _adicionar_run(p, ".")
    else:
        _adicionar_run(p,
            "Para fins de consolidação contratual, foi adotada a premissa de "
            "considerar, para fins de cálculo do retroativo e consequente cálculo do "
            "valor remanescente do contrato, ")
        _texto_ou_marcador(p, _campo(cm, "data_corte_descricao"), "Descricao da data/posicao de corte adotada")
        _adicionar_run(p, ".")

    _titulo_quadro(doc, "Quadro 3 - Memória fiscal do Valor Total Atualizado Estimado")
    cabecalho = ["Descrição", "Valor"]
    if dados.get("_modo_branco"):
        _adicionar_tabela(doc, cabecalho, [
            ["[PREENCHER: Descrição da parcela]", "[PREENCHER: Valor]"],
            ["Valor total do contrato estimado", "[PREENCHER: Valor Total Atualizado]"],
        ], destacar_placeholders=True)
        doc.add_paragraph()
        return
    linhas: list[list[str]] = []
    for desc, valor in _composicao_didatica_vta(dados):
        linhas.append([desc, formatar_moeda(valor) if valor is not None else ""])
    linhas.append(["Valor total do contrato estimado", _vta_texto_doc(dados)])
    _adicionar_tabela(doc, cabecalho, linhas)
    doc.add_paragraph()


def _ds_par7_composicao(doc: Document, dados: dict) -> None:
    p = _ds_par(doc, "7")
    _adicionar_run(p,
        "De forma didática, o Valor Total Atualizado Estimado do Contrato pode "
        "ser lido pela seguinte composição:")
    cabecalho = ["Parcela", "Valor"]
    if dados.get("_modo_branco"):
        _adicionar_tabela(doc, cabecalho, [
            ["[PREENCHER: Parcela]", "[PREENCHER: Valor]"],
            ["Valor Total Atualizado Estimado do Contrato",
             "[PREENCHER: Valor Total Atualizado]"],
        ], destacar_placeholders=True)
        doc.add_paragraph()
        return
    componentes = _composicao_didatica_vta(dados)
    linhas = [[desc, formatar_moeda(valor) if valor is not None else ""]
              for desc, valor in componentes]
    linhas.append(["Valor Total Atualizado Estimado do Contrato",
                   _vta_texto_doc(dados)])
    _adicionar_tabela(doc, cabecalho, linhas)
    doc.add_paragraph()


def _ds_bloco_historico_vu(doc: Document, dados: dict) -> None:
    """Bloco sem numeracao propria (entre os paragrafos 7 e 8): historico de
    VUs vinculado a consolidacao dos valores apurados. Preserva a numeracao
    juridica existente do despacho."""
    _secao_valores_unitarios_por_ciclo(
        doc, dados,
        texto_intro=(
            "Para fins de consolidação da evolução dos preços contratuais, "
            "apresenta-se abaixo o histórico dos valores unitários dos itens "
            "até o último ciclo considerado nesta análise."
        ),
    )


def _ds_par8_aditivos(doc: Document, dados: dict) -> None:
    p = _ds_par(doc, "8")
    aditivos = dados.get("aditivos") or []
    if dados.get("_modo_branco"):
        _adicionar_run(p, "Registrar os aditivos ou supressões considerados: ")
        _run_campo_manual(p, "Aditivos e supressões aplicáveis")
        _adicionar_run(p, ".")
        return
    if not aditivos:
        _adicionar_run(p,
            "Quanto às alterações contratuais consideradas, não foram "
            "identificados eventos específicos na base processada, sem prejuízo "
            "da conferência dos instrumentos já formalizados no processo.")
        return
    _adicionar_run(
        p,
        "Quanto às alterações contratuais consideradas, registra-se: "
        + "; ".join(_sintese_aditivos_por_ciclo(aditivos)) + ".",
    )


def _ds_par9_adequacao(doc: Document, dados: dict, cm: dict) -> None:
    p = _ds_par(doc, "9")
    if dados.get("_modo_branco"):
        _adicionar_run(p, "Registrar a adequação orçamentária, quando aplicável: ")
        _run_campo_manual(p, "Referência e valor da adequação orçamentária")
        _adicionar_run(p, ".")
        return
    _adicionar_run(p,
        "Foi realizada a adequação orçamentária necessária ao prosseguimento da "
        "instrução, no valor de ")
    _valor_moeda_ou_marcador(p, _campo(cm, "adequacao_orcamentaria_valor"), "Valor da adequacao orcamentaria")
    _adicionar_run(p, ", conforme documento ")
    _texto_ou_marcador(p, _campo(cm, "adequacao_orcamentaria_ref"), "Referencia da adequacao orcamentaria")
    _adicionar_run(p, ".")


def _ds_par10_regularidade(doc: Document, dados: dict, cm: dict) -> None:
    p = _ds_par(doc, "10")
    if dados.get("_modo_branco"):
        _adicionar_run(p, "Registrar as certidões de regularidade, quando aplicável: ")
        _run_campo_manual(p, "Referência das certidões de regularidade")
        _adicionar_run(p, ".")
        return
    _adicionar_run(p, "As certidões de regularidade estão presentes em ")
    _texto_ou_marcador(p, _campo(cm, "regularidade_ref"), "Referencia das certidoes de regularidade")
    _adicionar_run(p, ".")


def _ds_par11_concordancia(doc: Document, dados: dict, cm: dict) -> None:
    p = _ds_par(doc, "11")
    if dados.get("_modo_branco"):
        _adicionar_run(p, "Registrar a manifestação da contratada, quando aplicável: ")
        _run_campo_manual(p, "Referência da manifestação da contratada")
        _adicionar_run(p, ".")
        return
    _adicionar_run(p, "A contratada manifestou concordância com os valores propostos conforme registrado em ")
    _texto_ou_marcador(p, _campo(cm, "concordancia_ref"), "Referencia da manifestacao de concordancia da contratada")
    _adicionar_run(p, ".")


def _ds_par12_garantia(doc: Document, dados: dict) -> None:
    p = _ds_par(doc, "12")
    if dados.get("_modo_branco"):
        # Etapa 29C.1.2: nao afirma comunicacao ja realizada a contratada.
        _adicionar_run(p,
            "Registrar, quando aplicável, a comunicação à contratada sobre a "
            "necessidade de atualização ou endosso da garantia contratual, "
            "observados o prazo e as condições previstos no contrato.")
        return
    _adicionar_run(p,
        "A contratada foi informada da necessidade de apresentação do endosso da "
        "garantia contratual, quando aplicável, observando-se o prazo e as "
        "condições previstos no contrato.")


def _ds_par13_docs(doc: Document, cm: dict) -> None:
    docs = _campo(cm, "docs_desatualizados")
    if not docs:
        return
    p = _ds_par(doc, "13")
    _adicionar_run(p,
        "Após atualizações e alinhamentos internos, alguns documentos instruídos "
        "mostram-se desatualizados, devendo ser desconsiderados: ")
    if isinstance(docs, (list, tuple)):
        _adicionar_run(p, ", ".join(str(d) for d in docs))
    else:
        _adicionar_run(p, str(docs))
    _adicionar_run(p, ".")


def _tem_pendencia_critica(dados: dict, cm: dict) -> bool:
    """Soft-block: nao afirmar 'inexiste pendencia critica' se houver pendencia."""
    if not dados.get("disponivel"):
        return True
    flag = cm.get("pendencia_critica")
    if isinstance(flag, str):
        return flag.strip().lower() in ("sim", "true", "1", "critica", "critico")
    return bool(flag)


def _ds_conclusao(doc: Document, dados: dict, cm: dict) -> None:
    # Numeracao final: o item de documentos desatualizados (13) so existe quando
    # ha docs_desatualizados. A conclusao vem logo apos — 14 nesse caso, 13 caso
    # contrario. A logica de soft-block do texto permanece inalterada.
    numero = "14" if _campo(cm, "docs_desatualizados") else "13"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _adicionar_run(p, f"{numero}. ", negrito=True)
    if dados.get("_modo_branco"):
        # Etapa 29C.1: a conclusao do modelo em branco nao afirma consolidacao,
        # saneamento nem aptidao — apenas orienta a avaliacao apos o preenchimento.
        _adicionar_run(p,
            "Após o preenchimento e a conferência dos campos deste modelo, "
            "deverá ser avaliado se a instrução reúne condições para prosseguir "
            "à formalização do Termo de Apostila, observadas as alçadas "
            "competentes e os procedimentos internos aplicáveis.")
    elif _tem_pendencia_critica(dados, cm):
        _adicionar_run(p,
            "Diante do exposto, os elementos disponíveis encontram-se "
            "consolidados para análise, permanecendo pendentes as complementações "
            "ou validações indicadas antes da formalização do Termo de Apostila.")
    else:
        _adicionar_run(p,
            "Diante do exposto, estando conferidos os elementos documentais, "
            "financeiros e formais acima indicados, e inexistindo pendência "
            "crítica impeditiva, a instrução poderá prosseguir para formalização "
            "do Termo de Apostila, observadas as alçadas competentes e os "
            "procedimentos internos aplicáveis.")
    doc.add_paragraph()


def _ds_quadro4(doc: Document, dados: dict, cm: dict) -> None:
    _titulo_quadro(doc, "Quadro 4 - Síntese dos principais valores")
    cabecalho = ["Parcela", "Valor"]
    linhas: list[list[str]] = []

    val_orig = _num_ou_none(_campo(cm, "valor_original_contrato"))
    linhas.append(["Valor original do contrato",
                   formatar_moeda(val_orig) if val_orig is not None
                   else PREENCHER_TAG.format("Valor original do contrato")])

    var = dados.get("var_acumulada")
    if var is not None:
        linhas.append(["Variação acumulada do reajuste", _fmt_pct_doc(var)])

    for lin in _linhas_financeiro(dados):
        delta = _num_ou_none(lin.get("delta"))
        if delta is not None:
            ciclo = remover_emojis_leve(lin.get("ciclo") or "")
            linhas.append([f"Retroativo {ciclo}".strip(), formatar_moeda(delta)])

    retro = _retroativo_total(dados)
    if retro is not None:
        linhas.append(["Valor retroativo/represado a pagar", formatar_moeda(retro)])

    vta_texto = _vta_texto_doc(dados)
    if vta_texto:
        linhas.append(["Valor Total Atualizado Estimado do Contrato", vta_texto])

    adeq = _num_ou_none(_campo(cm, "adequacao_orcamentaria_valor"))
    if adeq is not None:
        linhas.append(["Adequação orçamentária registrada", formatar_moeda(adeq)])

    _adicionar_tabela(doc, cabecalho, linhas,
                      destacar_placeholders=bool(dados.get("_modo_branco")))


# ---------------------------------------------------------------------------
# Diagnostico de campos manuais
# ---------------------------------------------------------------------------

def diagnosticar_campos_manuais(
    leitura_ou_objeto: dict,
    identificacao: dict | None = None,
    campos_manuais: dict | None = None,
) -> list[dict]:
    """Retorna lista de campos manuais pendentes: {campo, descricao, documento}."""
    if campos_manuais is None:
        campos_manuais = {}
    pendentes = []
    vistos: set[str] = set()
    for chave, descricao, documento in TODOS_CAMPOS_MANUAIS:
        if chave in _CAMPOS_OPCIONAIS:
            continue
        if _campo(campos_manuais, chave) is None and chave not in vistos:
            vistos.add(chave)
            pendentes.append({"campo": chave, "descricao": descricao, "documento": documento})
    return pendentes
